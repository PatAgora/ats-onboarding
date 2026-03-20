# public.py
from __future__ import annotations
import os
import importlib
import json
from datetime import datetime, timedelta
from uuid import uuid4
from typing import Optional

from flask import (
    Blueprint, render_template, request, redirect, url_for,
    session, flash, current_app, abort
)
from werkzeug.utils import secure_filename
from sqlalchemy.orm import Session as SASession
from sqlalchemy import select

public_bp = Blueprint(
    "public",
    __name__,
    template_folder="templates_public",
    static_folder="static_public",
    static_url_path="/public_static",
)

# ----------------- Helpers -----------------
def _models():
    """Lazy import to avoid circular import with app.py."""
    from app import engine, Application, Candidate, Job  # noqa: WPS433
    return engine, Application, Candidate, Job

def _load_from_app(name: str):
    """Safely load a callable/attr from app at runtime (avoid top-level imports)."""
    try:
        app_mod = importlib.import_module("app")
        return getattr(app_mod, name, None)
    except Exception:
        return None

def _get_user_id() -> Optional[int]:
    return session.get("associate_user_id")

def _require_login(next_url: str):
    """If not logged in as an associate user, redirect to portal login with ?next=..."""
    if not _get_user_id():
        return redirect(url_for("associate.login", next=next_url))
    return None

def _static_upload_dir(*parts: str) -> str:
    root = current_app.root_path  # app.py folder
    return os.path.join(root, "static", *parts)

def _save_cv(file_storage) -> Optional[str]:
    """Save uploaded CV to /static/uploads/cvs and return relative path 'uploads/cvs/<file>'."""
    if not file_storage or not (file_storage.filename or "").strip():
        return None
    fname = secure_filename(file_storage.filename)
    ext = os.path.splitext(fname)[1].lower() or ".pdf"
    new_name = f"{uuid4().hex}{ext}"
    upload_dir = _static_upload_dir("uploads", "cvs")
    os.makedirs(upload_dir, exist_ok=True)
    file_storage.save(os.path.join(upload_dir, new_name))
    return f"uploads/cvs/{new_name}"

def _maybe_set(obj, name: str, value):
    """Set obj.name = value if the mapped attribute exists and value is truthy."""
    if value and hasattr(obj, name):
        try:
            setattr(obj, name, value)
        except Exception:
            pass

def _mirror_cv_to_candidate(cand, relpath: Optional[str]):
    """Mirror CV presence to Candidate for internal views."""
    if not cand or not relpath:
        return
    for attr in ("cv_path", "cv_url", "resume_path", "cv", "cv_filename", "cv_file_path"):
        if hasattr(cand, attr):
            _maybe_set(cand, attr, relpath)
            break
    if hasattr(cand, "has_cv"):
        try:
            cand.has_cv = True
        except Exception:
            pass

def _abs_from_static_rel(rel_path: str) -> str:
    root = current_app.root_path
    return os.path.join(root, "static", rel_path)

def _safe_extract_text(file_path: str, original_name: str) -> str:
    """Basic extractor for PDF/DOCX/TXT."""
    name = (original_name or file_path or "").lower()
    # PDF
    if name.endswith(".pdf"):
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                bits = [(p.extract_text() or "") for p in pdf.pages]
            txt = "\n".join(bits).strip()
            if txt:
                return txt
        except Exception:
            pass
    # DOCX
    if name.endswith(".docx"):
        try:
            import docx
            d = docx.Document(file_path)
            txt = "\n".join([p.text for p in d.paragraphs]).strip()
            if txt:
                return txt
        except Exception:
            pass
    # Fallback: TXT
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception:
        return ""

def _job_text(job) -> str:
    for attr in ("description", "desc", "requirements", "summary", "details", "spec"):
        if hasattr(job, attr):
            v = getattr(job, attr) or ""
            v = v.strip()
            if v:
                return v
    return (getattr(job, "title", "") or "").strip()

def _first(d: dict, keys: list[str], default=None):
    """Return the first non-None value for any of the provided keys in dict d."""
    if not isinstance(d, dict):
        return default
    for k in keys:
        if k in d and d[k] is not None:
            return d[k]
    return default

def _normalise_score(raw) -> Optional[int]:
    """
    Accepts many shapes:
      - number (0–1 or 0–100)
      - "78" or "78%" strings
      - (score, explanation) tuple/list
    Returns an int 0..100, or None.
    """
    if raw is None:
        return None

    # Tuple like (score, explanation)
    if isinstance(raw, (tuple, list)) and raw:
        raw = raw[0]

    # Strings (possibly with %)
    if isinstance(raw, str):
        s = raw.strip()
        if s.endswith("%"):
            s = s[:-1].strip()
        try:
            val = float(s)
        except Exception:
            return None
    else:
        # Numbers
        try:
            val = float(raw)
        except Exception:
            return None

    # Scale if 0–1
    if 0.0 <= val <= 1.0:
        val *= 100.0

    # Clamp and round
    val = max(0.0, min(100.0, val))
    return int(round(val))

# ----------------- Magic Link Helpers -----------------
def _portal_signer():
    """Create a URL-safe timed serializer for magic links."""
    from itsdangerous import URLSafeTimedSerializer
    from flask import current_app
    return URLSafeTimedSerializer(current_app.config["SECRET_KEY"], salt="portal-magic-link")

def _send_magic_link(email: str, name: str, next_url: str, is_signup: bool = False):
    """Generate and send a magic link email for account setup."""
    from app import send_email, APP_BASE_URL
    
    token = _portal_signer().dumps({
        "email": email,
        "name": name,
        "next": next_url,
        "is_signup": is_signup
    })
    
    # Build the magic link URL - directs to password setup
    verify_url = f"{APP_BASE_URL}/auth/verify?token={token}"
    
    if is_signup:
        subject = "Complete Your Registration - Optimus Solutions Careers"
        html_body = f"""
        <div style="font-family: 'Inter', Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
            <div style="background: #0a1628; padding: 30px; border-radius: 12px 12px 0 0; text-align: center;">
                <h1 style="color: #00d4ff; margin: 0; font-size: 24px;">Optimus Solutions</h1>
                <p style="color: #94a3b8; margin: 10px 0 0; font-size: 14px;">Careers Portal</p>
            </div>
            <div style="background: #ffffff; padding: 40px 30px; border: 1px solid #e5e7eb; border-top: none; border-radius: 0 0 12px 12px;">
                <h2 style="color: #0a1628; margin: 0 0 20px; font-size: 20px;">Welcome, {name}!</h2>
                <p style="color: #334155; margin: 0 0 20px; line-height: 1.6;">
                    Thank you for registering with Optimus Solutions. Click the button below to verify your email and set up your secure account credentials.
                </p>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="{verify_url}" style="background: #0066cc; color: #ffffff; padding: 14px 28px; border-radius: 8px; text-decoration: none; font-weight: 600; display: inline-block;">
                        Complete Registration
                    </a>
                </div>
                <p style="color: #64748b; font-size: 14px; margin: 20px 0 0;">
                    This link will expire in 30 minutes. If you didn't request this email, you can safely ignore it.
                </p>
            </div>
            <div style="text-align: center; padding: 20px; color: #94a3b8; font-size: 12px;">
                <p>&copy; 2026 Optimus Solutions. All rights reserved.</p>
            </div>
        </div>
        """
    else:
        subject = "Password Reset - Optimus Solutions Careers"
        html_body = f"""
        <div style="font-family: 'Inter', Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
            <div style="background: #0a1628; padding: 30px; border-radius: 12px 12px 0 0; text-align: center;">
                <h1 style="color: #00d4ff; margin: 0; font-size: 24px;">Optimus Solutions</h1>
                <p style="color: #94a3b8; margin: 10px 0 0; font-size: 14px;">Careers Portal</p>
            </div>
            <div style="background: #ffffff; padding: 40px 30px; border: 1px solid #e5e7eb; border-top: none; border-radius: 0 0 12px 12px;">
                <h2 style="color: #0a1628; margin: 0 0 20px; font-size: 20px;">Password Reset Request</h2>
                <p style="color: #334155; margin: 0 0 20px; line-height: 1.6;">
                    Click the button below to reset your password.
                </p>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="{verify_url}" style="background: #0066cc; color: #ffffff; padding: 14px 28px; border-radius: 8px; text-decoration: none; font-weight: 600; display: inline-block;">
                        Reset Password
                    </a>
                </div>
                <p style="color: #64748b; font-size: 14px; margin: 20px 0 0;">
                    This link will expire in 30 minutes. If you didn't request this email, you can safely ignore it.
                </p>
            </div>
            <div style="text-align: center; padding: 20px; color: #94a3b8; font-size: 12px;">
                <p>&copy; 2026 Optimus Solutions. All rights reserved.</p>
            </div>
        </div>
        """
    
    # This will raise an exception if email fails - caller must handle it
    send_email(to_email=email, subject=subject, html_body=html_body)
    current_app.logger.info(f"Magic link email sent successfully to {email}")

def _validate_password_strength(password: str) -> tuple[bool, str]:
    """Validate password meets security requirements."""
    if len(password) < 8:
        return False, "Password must be at least 8 characters long"
    if not any(c.isupper() for c in password):
        return False, "Password must contain at least one uppercase letter"
    if not any(c.islower() for c in password):
        return False, "Password must contain at least one lowercase letter"
    if not any(c.isdigit() for c in password):
        return False, "Password must contain at least one number"
    if not any(c in "!@#$%^&*()_+-=[]{}|;':\",./<>?" for c in password):
        return False, "Password must contain at least one special character"
    return True, ""

# ----------------- Auth Routes -----------------

@public_bp.get("/auth/login")
def login():
    """Show login page - email + password + 2FA."""
    next_url = request.args.get("next") or url_for("public.jobs_index")
    return render_template("auth_login_public.html", next=next_url)

@public_bp.post("/auth/login")
def login_post():
    """Handle login with email + password, then redirect to 2FA if enabled."""
    from app import engine, Candidate
    next_url = request.form.get("next") or request.args.get("next") or url_for("public.jobs_index")
    email = (request.form.get("email") or "").strip().lower()
    password = request.form.get("password") or ""

    if not email:
        flash("Please enter your email address", "danger")
        return redirect(url_for("public.login", next=next_url))

    with SASession(engine) as s:
        cand = s.query(Candidate).filter(Candidate.email.ilike(email)).first()
        
        if not cand:
            flash("No account found with that email. Please register first.", "info")
            return redirect(url_for("public.signup", next=next_url))
        
        # Check if account is locked
        if hasattr(cand, 'is_locked') and cand.is_locked():
            flash("Account is temporarily locked. Please try again later.", "danger")
            return redirect(url_for("public.login", next=next_url))
        
        # Check if password is set up
        if not cand.password_hash:
            flash("Please complete your account setup first. Check your email for a registration link.", "info")
            return redirect(url_for("public.login", next=next_url))
        
        # Verify password
        if not cand.check_password(password):
            # Increment failed attempts
            cand.failed_login_attempts = (cand.failed_login_attempts or 0) + 1
            if cand.failed_login_attempts >= 5:
                cand.locked_until = datetime.utcnow() + timedelta(minutes=15)
                flash("Too many failed attempts. Account locked for 15 minutes.", "danger")
            else:
                flash("Invalid email or password.", "danger")
            s.commit()
            return redirect(url_for("public.login", next=next_url))
        
        # Password correct - reset failed attempts
        cand.failed_login_attempts = 0
        cand.locked_until = None
        
        # Check if 2FA is enabled
        if cand.totp_enabled and cand.totp_secret:
            # Store candidate ID temporarily for 2FA verification
            session["portal_2fa_pending_id"] = cand.id
            session["portal_2fa_next"] = next_url
            s.commit()
            return redirect(url_for("public.verify_2fa"))
        
        # Check if 2FA setup is required (account complete but no 2FA)
        if not cand.totp_enabled:
            session["portal_setup_2fa_id"] = cand.id
            session["portal_setup_2fa_next"] = next_url
            s.commit()
            flash("Please set up two-factor authentication to secure your account.", "info")
            return redirect(url_for("public.setup_2fa"))
        
        # Fully authenticated - log in
        cand.last_login_at = datetime.utcnow()
        session["associate_user_id"] = cand.id
        s.commit()
        
    flash("Signed in successfully!", "success")
    return redirect(next_url)

@public_bp.get("/auth/signup")
def signup():
    """Show signup page."""
    next_url = request.args.get("next") or url_for("public.jobs_index")
    return render_template("auth_signup_public.html", next=next_url)

@public_bp.post("/auth/signup")
def signup_post():
    """Create pending registration and send magic link to set up credentials."""
    from app import engine, Candidate
    next_url = request.form.get("next") or request.args.get("next") or url_for("public.jobs_index")

    email = (request.form.get("email") or "").strip().lower()
    name = (request.form.get("name") or "").strip()
    phone = (request.form.get("phone") or "").strip()
    about = (request.form.get("about") or "").strip()

    if not email or not name:
        flash("Please provide your name and email address", "danger")
        return redirect(url_for("public.signup", next=next_url))

    with SASession(engine) as s:
        existing = s.query(Candidate).filter(Candidate.email.ilike(email)).first()
        if existing:
            if existing.password_hash and existing.totp_enabled:
                flash("An account with this email already exists. Please sign in.", "info")
                return redirect(url_for("public.login", next=next_url))
            else:
                # Account exists but not fully set up - resend magic link
                flash("We've sent you a new link to complete your registration.", "info")
                try:
                    _send_magic_link(email, existing.name or name, next_url, is_signup=True)
                except Exception as e:
                    current_app.logger.exception("Failed to send magic link: %s", e)
                    flash("We were unable to send your verification email. Please contact administration at enquiries@agoraconsulting.ai for assistance.", "danger")
                    return redirect(url_for("public.signup", next=next_url))
                return render_template("auth_check_email_public.html", email=email, next=next_url, is_signup=True)
        
    # Try to send email FIRST before creating the user
    try:
        _send_magic_link(email, name, next_url, is_signup=True)
    except Exception as e:
        current_app.logger.exception("Failed to send verification email: %s", e)
        flash("We were unable to send your verification email. Please contact administration at enquiries@agoraconsulting.ai for assistance.", "danger")
        return redirect(url_for("public.signup", next=next_url))
    
    # Email sent successfully - now create the candidate
    with SASession(engine) as s:
        # Double-check no one created the account in the meantime
        existing = s.query(Candidate).filter(Candidate.email.ilike(email)).first()
        if existing:
            flash("Check your email to complete your registration.", "success")
            return render_template("auth_check_email_public.html", email=email, next=next_url, is_signup=True)
        
        cand = Candidate(
            name=name,
            email=email,
            phone=phone,
            email_verified=False,
            source="portal"
        )
        if hasattr(cand, "about"):
            cand.about = about
        s.add(cand)
        s.commit()

    flash("Check your email to complete your registration.", "success")
    return render_template("auth_check_email_public.html", email=email, next=next_url, is_signup=True)

@public_bp.get("/auth/verify")
def verify_magic_link():
    """Verify magic link and redirect to password setup."""
    from app import engine, Candidate
    from itsdangerous import SignatureExpired, BadSignature
    
    token = request.args.get("token", "")
    
    if not token:
        flash("Invalid verification link.", "danger")
        return redirect(url_for("public.jobs_index"))
    
    try:
        data = _portal_signer().loads(token, max_age=30*60)  # 30 minutes expiry
    except SignatureExpired:
        flash("This link has expired. Please request a new one.", "warning")
        return redirect(url_for("public.signup"))
    except BadSignature:
        flash("Invalid verification link.", "danger")
        return redirect(url_for("public.jobs_index"))
    
    email = (data.get("email") or "").strip().lower()
    next_url = data.get("next") or url_for("public.jobs_index")
    
    if not email:
        flash("Invalid verification link.", "danger")
        return redirect(url_for("public.jobs_index"))
    
    with SASession(engine) as s:
        cand = s.query(Candidate).filter(Candidate.email.ilike(email)).first()
        
        if not cand:
            flash("Account not found. Please register again.", "warning")
            return redirect(url_for("public.signup", next=next_url))
        
        # Mark email as verified
        cand.email_verified = True
        cand.email_verified_at = datetime.utcnow()
        s.commit()
        
        # Store candidate ID for password setup
        session["portal_setup_password_id"] = cand.id
        session["portal_setup_next"] = next_url
    
    # Redirect to password setup
    return redirect(url_for("public.setup_password"))

@public_bp.get("/auth/setup-password")
def setup_password():
    """Show password setup form."""
    cand_id = session.get("portal_setup_password_id")
    if not cand_id:
        flash("Please use the link from your email.", "warning")
        return redirect(url_for("public.login"))
    
    from app import engine, Candidate
    with SASession(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            session.pop("portal_setup_password_id", None)
            flash("Account not found.", "danger")
            return redirect(url_for("public.signup"))
        cand_name = cand.name
        cand_email = cand.email
    
    return render_template("auth_setup_password_public.html", 
                         name=cand_name, email=cand_email)

@public_bp.post("/auth/setup-password")
def setup_password_post():
    """Handle password setup."""
    cand_id = session.get("portal_setup_password_id")
    if not cand_id:
        flash("Please use the link from your email.", "warning")
        return redirect(url_for("public.login"))
    
    password = request.form.get("password") or ""
    confirm_password = request.form.get("confirm_password") or ""
    
    if password != confirm_password:
        flash("Passwords do not match.", "danger")
        return redirect(url_for("public.setup_password"))
    
    is_valid, error_msg = _validate_password_strength(password)
    if not is_valid:
        flash(error_msg, "danger")
        return redirect(url_for("public.setup_password"))
    
    from app import engine, Candidate
    with SASession(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            session.pop("portal_setup_password_id", None)
            flash("Account not found.", "danger")
            return redirect(url_for("public.signup"))
        
        cand.set_password(password)
        s.commit()
        
        # Move to 2FA setup
        session["portal_setup_2fa_id"] = cand.id
        session.pop("portal_setup_password_id", None)
    
    flash("Password set successfully! Now set up two-factor authentication.", "success")
    return redirect(url_for("public.setup_2fa"))

@public_bp.get("/auth/setup-2fa")
def setup_2fa():
    """Show 2FA setup page with QR code."""
    cand_id = session.get("portal_setup_2fa_id")
    if not cand_id:
        flash("Please sign in first.", "warning")
        return redirect(url_for("public.login"))
    
    import pyotp
    import io
    import base64
    
    from app import engine, Candidate
    with SASession(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            session.pop("portal_setup_2fa_id", None)
            flash("Account not found.", "danger")
            return redirect(url_for("public.signup"))
        
        # Generate TOTP secret if not exists
        if not cand.totp_secret:
            cand.totp_secret = cand.generate_totp_secret()
            s.commit()
        
        secret = cand.totp_secret
        email = cand.email
    
    # Generate provisioning URI
    totp = pyotp.TOTP(secret)
    provisioning_uri = totp.provisioning_uri(name=email, issuer_name="Optimus Solutions Portal")
    
    # Generate QR code
    qr_data_uri = ""
    try:
        import qrcode
        qr = qrcode.QRCode(version=1, box_size=10, border=4)
        qr.add_data(provisioning_uri)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        qr_data_uri = f"data:image/png;base64,{base64.b64encode(buffer.getvalue()).decode()}"
    except Exception as e:
        # Fallback to Google Charts API
        print(f"QR code generation error: {e}")
        import urllib.parse
        encoded_uri = urllib.parse.quote(provisioning_uri)
        qr_data_uri = f"https://api.qrserver.com/v1/create-qr-code/?size=200x200&data={encoded_uri}"
    
    return render_template("auth_setup_2fa_public.html",
                         qr_code=qr_data_uri,
                         secret=secret,
                         email=email)

@public_bp.post("/auth/setup-2fa")
def setup_2fa_post():
    """Verify 2FA token and complete setup."""
    cand_id = session.get("portal_setup_2fa_id")
    if not cand_id:
        flash("Please sign in first.", "warning")
        return redirect(url_for("public.login"))
    
    token = (request.form.get("token") or "").strip().replace(" ", "")
    
    if not token or len(token) != 6 or not token.isdigit():
        flash("Please enter a valid 6-digit code.", "danger")
        return redirect(url_for("public.setup_2fa"))
    
    from app import engine, Candidate
    with SASession(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            session.pop("portal_setup_2fa_id", None)
            flash("Account not found.", "danger")
            return redirect(url_for("public.signup"))
        
        if not cand.verify_totp(token):
            flash("Invalid code. Please try again.", "danger")
            return redirect(url_for("public.setup_2fa"))
        
        # Enable 2FA
        cand.totp_enabled = True
        
        # Generate backup codes
        backup_codes = cand.generate_backup_codes(count=10)
        
        cand.last_login_at = datetime.utcnow()
        s.commit()
        
        # Fully authenticated
        session["associate_user_id"] = cand.id
        session.pop("portal_setup_2fa_id", None)
        session.pop("portal_setup_2fa_next", None)
    
    # Show backup codes
    return render_template("auth_backup_codes_public.html", backup_codes=backup_codes)

@public_bp.get("/auth/verify-2fa")
def verify_2fa():
    """Show 2FA verification page."""
    cand_id = session.get("portal_2fa_pending_id")
    if not cand_id:
        flash("Please sign in first.", "warning")
        return redirect(url_for("public.login"))
    
    return render_template("auth_verify_2fa_public.html")

@public_bp.post("/auth/verify-2fa")
def verify_2fa_post():
    """Verify 2FA token for login."""
    cand_id = session.get("portal_2fa_pending_id")
    next_url = session.get("portal_2fa_next") or url_for("public.jobs_index")
    
    if not cand_id:
        flash("Please sign in first.", "warning")
        return redirect(url_for("public.login"))
    
    token = (request.form.get("token") or "").strip().replace(" ", "")
    use_backup = request.form.get("use_backup") == "1"
    
    from app import engine, Candidate
    with SASession(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            session.pop("portal_2fa_pending_id", None)
            flash("Account not found.", "danger")
            return redirect(url_for("public.login"))
        
        verified = False
        if use_backup:
            # Verify backup code
            if cand.verify_backup_code(token):
                verified = True
        else:
            # Verify TOTP
            if len(token) == 6 and token.isdigit() and cand.verify_totp(token):
                verified = True
        
        if not verified:
            flash("Invalid code. Please try again.", "danger")
            return redirect(url_for("public.verify_2fa"))
        
        # Fully authenticated
        cand.last_login_at = datetime.utcnow()
        session["associate_user_id"] = cand.id
        session.pop("portal_2fa_pending_id", None)
        session.pop("portal_2fa_next", None)
        s.commit()
    
    flash("Signed in successfully!", "success")
    return redirect(next_url)

@public_bp.get("/auth/forgot-password")
def forgot_password():
    """Show forgot password page."""
    return render_template("auth_forgot_password_public.html")

@public_bp.post("/auth/forgot-password")
def forgot_password_post():
    """Send password reset link."""
    from app import engine, Candidate
    email = (request.form.get("email") or "").strip().lower()
    
    if not email:
        flash("Please enter your email address.", "danger")
        return redirect(url_for("public.forgot_password"))
    
    with SASession(engine) as s:
        cand = s.query(Candidate).filter(Candidate.email.ilike(email)).first()
        if cand:
            try:
                _send_magic_link(email, cand.name or email.split("@")[0], 
                               url_for("public.jobs_index"), is_signup=False)
            except Exception as e:
                current_app.logger.exception("Failed to send reset email: %s", e)
    
    # Always show success to prevent email enumeration
    flash("If an account exists with that email, you'll receive a password reset link.", "success")
    return render_template("auth_check_email_public.html", email=email, 
                         next=url_for("public.jobs_index"), is_signup=False)

@public_bp.get("/auth/logout")
def logout():
    """Log out and clear session."""
    session.pop("associate_user_id", None)
    session.pop("portal_2fa_pending_id", None)
    session.pop("portal_2fa_next", None)
    session.pop("portal_setup_2fa_id", None)
    session.pop("portal_setup_password_id", None)
    flash("You have been signed out.", "success")
    return redirect(url_for("public.jobs_index"))

# ----------------- Public Jobs -----------------
@public_bp.get("/jobs")
def jobs_index():
    engine, _, _, Job = _models()
    with SASession(engine) as s:
        jobs = (
            s.query(Job)
             .filter((Job.status == "Open") | (Job.status == None))  # noqa: E711
             .order_by(Job.created_at.desc())
             .all()
        )
    return render_template("jobs_index.html", jobs=jobs)

@public_bp.get("/jobs/<int:job_id>")
def job_public(job_id: int):
    engine, _, Candidate, Job = _models()
    with SASession(engine) as s:
        job = s.get(Job, job_id)
        if not job or (job.status or "Open") != "Open":
            flash("Job not found or not open", "warning")
            return redirect(url_for("public.jobs_index"))
        cand = s.get(Candidate, _get_user_id()) if _get_user_id() else None
    return render_template("job_public.html", job=job, cand=cand)

# ----------------- Apply Flow -----------------
@public_bp.get("/jobs/<int:job_id>/apply")
def job_apply(job_id: int):
    need_login = _require_login(url_for("public.job_apply", job_id=job_id))
    if need_login:
        return need_login

    engine, _, Candidate, Job = _models()
    user_id = _get_user_id()

    with SASession(engine) as s:
        job = s.get(Job, job_id)
        if not job or (job.status or "Open") != "Open":
            flash("Job not found or not open", "warning")
            return redirect(url_for("public.jobs_index"))
        cand = s.get(Candidate, user_id) if user_id else None

    return render_template("apply_form_public.html", job=job, cand=cand)

@public_bp.post("/jobs/<int:job_id>/apply")
def job_apply_post(job_id: int):
    """
    Create/update an application, persist CV file, create a Document row for the CV,
    mirror CV to Candidate, then run AI summary + tagging + scoring immediately.
    """
    need_login = _require_login(url_for("public.job_apply", job_id=job_id))
    if need_login:
        return need_login

    engine, Application, Candidate, Job = _models()
    user_id = _get_user_id()

    cover_note_val = (request.form.get("cover_note") or "").strip()
    cv_file = request.files.get("cv_file")
    relpath = _save_cv(cv_file)
    original_name = (cv_file.filename if cv_file else "") or ""

    with SASession(engine) as s:
        job = s.get(Job, job_id)
        if not job or (job.status or "Open") != "Open":
            flash("Job not found or not open", "warning")
            return redirect(url_for("public.jobs_index"))

        cand = s.get(Candidate, user_id)
        if not cand:
            session.pop("associate_user_id", None)
            return redirect(url_for("associate.login", next=url_for("public.job_apply", job_id=job_id)))

        # Latest application per candidate/job
        existing = (
            s.query(Application)
             .filter(Application.candidate_id == cand.id, Application.job_id == job.id)
             .order_by(Application.created_at.desc())
             .first()
        )

        if existing:
            try:
                existing.status = existing.status or "Declared"
            except Exception:
                pass
            for attr in ("cv_file_path", "cv_path", "cv_url", "resume_path", "cv", "cv_filename"):
                _maybe_set(existing, attr, relpath)
            for attr in ("cover_note", "cover_letter", "note", "notes", "summary"):
                _maybe_set(existing, attr, cover_note_val)
            _mirror_cv_to_candidate(cand, relpath)
            s.commit()
            app_obj = existing
            app_id = existing.id
        else:
            app_obj = Application(
                candidate_id=cand.id,
                job_id=job.id,
                status="Declared",
                created_at=datetime.utcnow(),
            )
            s.add(app_obj)
            s.commit()
            for attr in ("cv_file_path", "cv_path", "cv_url", "resume_path", "cv", "cv_filename"):
                _maybe_set(app_obj, attr, relpath)
            for attr in ("cover_note", "cover_letter", "note", "notes", "summary"):
                _maybe_set(app_obj, attr, cover_note_val)
            _mirror_cv_to_candidate(cand, relpath)
            s.commit()
            app_id = app_obj.id

        # ---- Ensure a Document row for this CV ----
        Document = _load_from_app("Document")
        doc_row = None
        if Document and relpath:
            try:
                order_col = getattr(Document, "uploaded_at", Document.id)
                doc_row = s.execute(
                    select(Document)
                    .where(Document.candidate_id == cand.id)
                    .where(getattr(Document, "doc_type", Document.doc_type) == "cv")
                    .order_by(order_col.desc())
                    .limit(1)
                ).scalar_one_or_none()
                if not doc_row or (getattr(doc_row, "filename", None) != relpath):
                    payload = dict(candidate_id=cand.id, filename=relpath)
                    if hasattr(Document, "application_id"):
                        payload["application_id"] = app_id
                    if hasattr(Document, "doc_type"):
                        payload["doc_type"] = "cv"
                    if hasattr(Document, "original_name"):
                        payload["original_name"] = original_name
                    if hasattr(Document, "uploaded_at"):
                        payload["uploaded_at"] = datetime.utcnow()
                    doc_row = Document(**payload)
                    s.add(doc_row)
                    s.commit()
            except Exception as e:
                current_app.logger.exception("Failed to upsert Document row for CV: %s", e)

        # ---- AI: extract text -> summary -> tags/skills -> score ----
        extract_cv_text = _load_from_app("extract_cv_text")
        ai_summarise = _load_from_app("ai_summarise")
        ai_score_with_explanation = _load_from_app("ai_score_with_explanation")

        cv_text = ""
        if relpath:
            abs_path = _abs_from_static_rel(relpath)
            try:
                if callable(extract_cv_text):
                    try:
                        if doc_row is not None:
                            cv_text = extract_cv_text(doc_row) or ""
                            if not cv_text:
                                cv_text = extract_cv_text(abs_path) or ""
                        else:
                            cv_text = extract_cv_text(abs_path) or ""
                    except TypeError:
                        cv_text = extract_cv_text(abs_path) or ""
                if not cv_text:
                    cv_text = _safe_extract_text(abs_path, original_name)
            except Exception:
                cv_text = _safe_extract_text(abs_path, original_name)

        # 1) AI Summary
        summary = ""
        try:
            if callable(ai_summarise):
                summary = ai_summarise(cv_text or "") or ""
        except Exception as e:
            current_app.logger.exception("ai_summarise failed: %s", e)

        if hasattr(app_obj, "ai_summary"):
            try:
                app_obj.ai_summary = summary
            except Exception:
                pass
        if hasattr(cand, "ai_summary"):
            try:
                cand.ai_summary = summary
            except Exception:
                pass

        # 2) Tags/skills
        try:
            TaxonomyTag = _load_from_app("TaxonomyTag")
            CandidateTag = _load_from_app("CandidateTag")
            if TaxonomyTag and CandidateTag:
                text_lc = (cv_text or "").lower()
                matched_ids = set()
                all_tags = s.execute(select(TaxonomyTag)).scalars().all()
                for tg in all_tags:
                    token = (tg.tag or "").strip()
                    if token and token.lower() in text_lc:
                        matched_ids.add(tg.id)

                if matched_ids:
                    existing_ids = {
                        tid for (tid,) in s.execute(
                            select(CandidateTag.tag_id).where(CandidateTag.candidate_id == cand.id)
                        ).all()
                    }
                    for tid in matched_ids:
                        if tid not in existing_ids:
                            s.add(CandidateTag(candidate_id=cand.id, tag_id=tid))

                    # Mirror into Candidate.skills
                    all_tag_rows = s.execute(
                        select(TaxonomyTag)
                        .join(CandidateTag, CandidateTag.tag_id == TaxonomyTag.id)
                        .where(CandidateTag.candidate_id == cand.id)
                        .order_by(TaxonomyTag.tag.asc())
                    ).scalars().all()
                    tag_names = [t.tag for t in all_tag_rows if (t.tag or "").strip()]
                    cand.skills = ", ".join(dict.fromkeys(tag_names))
        except Exception as e:
            current_app.logger.debug("Tagging skipped: %s", e)

        # 3) AI Score vs Job (robust normalisation + multiple payload shapes)
        if (cv_text or "").strip():
            try:
                score_val: Optional[int] = None
                explanation_text: Optional[str] = None

                if callable(ai_score_with_explanation):
                    payload = ai_score_with_explanation(_job_text(job), cv_text)

                    if isinstance(payload, dict):
                        raw_score = _first(
                            payload,
                            ["blended_score", "score", "match_score", "overall", "total", "similarity", "similarity_score"],
                        )
                        score_val = _normalise_score(raw_score)
                        explanation_text = _first(
                            payload,
                            ["explanation", "reason", "rationale", "details", "why"],
                        )
                    elif isinstance(payload, (tuple, list)):
                        score_val = _normalise_score(payload[0] if payload else None)
                        if len(payload) > 1 and isinstance(payload[1], str):
                            explanation_text = payload[1]
                    else:
                        # plain number or string
                        score_val = _normalise_score(payload)

                # optional fallback scorer
                if score_val is None:
                    ai_score_simple = _load_from_app("ai_score_simple")
                    if callable(ai_score_simple):
                        score_val = _normalise_score(ai_score_simple(_job_text(job), cv_text))

                # persist score
                if score_val is not None:
                    if hasattr(app_obj, "ai_score"):
                        try:
                            app_obj.ai_score = score_val
                        except Exception:
                            pass
                    if hasattr(cand, "ai_score"):
                        try:
                            cand.ai_score = score_val
                        except Exception:
                            pass
                # persist explanation
                if explanation_text and hasattr(app_obj, "ai_explanation"):
                    try:
                        app_obj.ai_explanation = explanation_text
                    except Exception:
                        pass

            except Exception as e:
                current_app.logger.exception("AI scoring failed: %s", e)

        # 4) Ensure Shortlist (optional)
        Shortlist = _load_from_app("Shortlist")
        if Shortlist:
            try:
                exists_short = s.execute(
                    select(Shortlist).where(
                        Shortlist.candidate_id == cand.id,
                        Shortlist.job_id == job.id,
                    )
                ).scalar_one_or_none()
                if not exists_short:
                    payload = dict(candidate_id=cand.id, job_id=job.id)
                    if hasattr(Shortlist, "created_at"):
                        payload["created_at"] = datetime.utcnow()
                    s.add(Shortlist(**payload))
            except Exception:
                pass

        s.commit()

    # Optional legacy hook
    run_cv_scoring = _load_from_app("run_cv_scoring")
    try:
        if callable(run_cv_scoring):
            run_cv_scoring(app_id)
    except Exception:
        current_app.logger.exception("run_cv_scoring failed for app_id=%s", app_id)

    return redirect(url_for("public.apply_done", job_id=job_id, app_id=app_id))

@public_bp.get("/jobs/<int:job_id>/apply/done")
def apply_done(job_id: int):
    return render_template("apply_done_public.html", job_id=job_id)

# ----------------- Optional token-based route -----------------
@public_bp.get("/apply/<token>")
def public_apply(token: str):
    engine, _, _, Job = _models()
    with SASession(engine) as s:
        job = s.query(Job).filter(Job.public_token == token).first()
        if not job:
            abort(404)
        return redirect(url_for("public.job_apply", job_id=job.id))
