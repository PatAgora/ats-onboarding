import os, uuid, datetime, json, mimetypes, re, smtplib, ssl, base64, hashlib, hmac, uuid, json, re, time, requests
from email.message import EmailMessage
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, jsonify, abort
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
from flask import Response
from wtforms import StringField, TextAreaField, BooleanField, SelectField, IntegerField, FileField, DateTimeLocalField, SubmitField
from typing import Optional, List, Dict, Tuple
from sqlalchemy import or_, and_
import csv
from io import StringIO
from flask import current_app
from flask import Blueprint
from sqlalchemy import exists
from flask import redirect, url_for
from flask import make_response

from wtforms.validators import DataRequired, Email, Optional as WTOptional
from werkzeug.utils import secure_filename
from sqlalchemy import create_engine, select, ForeignKey, func, Column, text, Float
from sqlalchemy.orm import declarative_base, relationship, Session, selectinload
from sqlalchemy.types import JSON as SA_JSON
from sqlalchemy import text
from sqlalchemy.exc import IntegrityError
from dotenv import load_dotenv
load_dotenv()
import requests
from dateutil import parser as dtparser
from sqlalchemy import literal_column
from sqlalchemy import delete
from sqlalchemy import Column, Integer, String, Text, DateTime, ForeignKey
from sqlalchemy.orm import relationship, backref
from wtforms import Form
from wtforms.validators import DataRequired
from wtforms import StringField, SelectField
from flask_login import LoginManager, UserMixin, login_user, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from flask import session
from openai import OpenAI
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

# reCAPTCHA configuration
RECAPTCHA_SITE_KEY = os.getenv("RECAPTCHA_SITE_KEY", "")
RECAPTCHA_SECRET_KEY = os.getenv("RECAPTCHA_SECRET_KEY", "")
RECAPTCHA_ENABLED = bool(RECAPTCHA_SITE_KEY and RECAPTCHA_SECRET_KEY)

import datetime
from flask import abort, render_template, request
from sqlalchemy import or_, false
from sqlalchemy.orm import selectinload

from wtforms import FileField
from wtforms.validators import DataRequired as WTDataRequired

class UploadCVForm(FlaskForm):
    cv = FileField("Upload CV (PDF/DOC/DOCX)", validators=[WTDataRequired()])
    submit = SubmitField("Upload")

# Optional e-sign SDKs
try:
    from dropbox_sign import ApiClient, Configuration
    from dropbox_sign.apis import SignatureRequestApi
    from dropbox_sign.models import SignatureRequestSendRequest, SubSignatureRequestSigner
except Exception:
    ApiClient = None
    Configuration = None
    SignatureRequestApi = None
    SignatureRequestSendRequest = None
    SubSignatureRequestSigner = None

try:
    import docusign_esign as dse
except Exception:
    dse = None

SECRET_KEY = os.getenv("FLASK_SECRET_KEY", "dev-secret")
DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///ats.db")
UPLOAD_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), "uploads"))
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

SECRET_KEY   = os.getenv("FLASK_SECRET_KEY", "dev-secret")
DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///ats.db")

SUMSUB_APP_TOKEN = os.getenv("SUMSUB_APP_TOKEN")
SUMSUB_SECRET_KEY = os.getenv("SUMSUB_SECRET_KEY")
SUMSUB_BASE_URL = os.getenv("SUMSUB_BASE_URL", "https://api.sumsub.com")
SUMSUB_LEVEL_NAME = os.getenv("SUMSUB_LEVEL_NAME", "basic-kyc-level")

SUMSUB_SDK_URLS = [
    "https://static.sumsub.com/idensic/static/sns-websdk-builder.js",  # legacy path
    "https://static.sumsub.com/idensic/assets/websdk-builder.js",       # newer path
]

engine = create_engine(DATABASE_URL, future=True)
Base = declarative_base()

class RoleType(Base):
    __tablename__ = "role_types"
    id = Column(Integer, primary_key=True, autoincrement=True)
    name = Column(String(100), unique=True, nullable=False)
    default_rate = Column(Integer, default=0)

class EngagementPlan(Base):
    __tablename__ = "engagement_plans"
    id = Column(Integer, primary_key=True, autoincrement=True)
    engagement_id = Column(Integer, ForeignKey("engagements.id"), index=True)
    role_type = Column(String(50), default="")
    planned_count = Column(Integer, default=0)

    # cost to us per head per day
    pay_rate = Column(Integer, default=0)

    # what we bill client per head per day
    charge_rate = Column(Integer, default=0)

    # legacy; keep in sync with charge_rate when saving for backwards reports
    rate = Column(Integer, default=0)

    # version snapshot / plan iteration number
    version_int = Column(Integer, default=1)
    
    # intake date for 4-week lookahead feature
    intake_date = Column(DateTime, nullable=True)

TRUSTID_BASE_URL = os.getenv("TRUSTID_BASE_URL", "https://api.trustid.co.uk")
TRUSTID_API_KEY = os.getenv("TRUSTID_API_KEY", "")
TRUSTID_WEBHOOK_SECRET = os.getenv("TRUSTID_WEBHOOK_SECRET", "")

ESIGN_PROVIDER = os.getenv("ESIGN_PROVIDER", "dropbox_sign").lower()

HELLOSIGN_API_KEY = os.getenv("HELLOSIGN_API_KEY", "")

DOCUSIGN_BASE_PATH = os.getenv("DOCUSIGN_BASE_PATH", "https://demo.docusign.net/restapi")
DOCUSIGN_ACCESS_TOKEN = os.getenv("DOCUSIGN_ACCESS_TOKEN", "")
DOCUSIGN_ACCOUNT_ID = os.getenv("DOCUSIGN_ACCOUNT_ID", "")
DOCUSIGN_FROM_NAME = os.getenv("DOCUSIGN_FROM_NAME", "Talent Ops")
DOCUSIGN_FROM_EMAIL = os.getenv("DOCUSIGN_FROM_EMAIL", "talent@example.com")

SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASS = os.getenv("SMTP_PASS", "")
SMTP_FROM = os.getenv("SMTP_FROM", "Talent <noreply@example.com>")

_client = None
def get_openai_client():
    """
    Return a clean OpenAI client compatible with the new SDK.

    We:
    - pass api_key directly
    - explicitly build a bare httpx client with trust_env=False so it doesn't
      try to forward local proxy env vars that the SDK then treats as `proxies`
    
    Returns None if OPENAI_API_KEY is not configured or is a placeholder.
    """
    # Check if API key is configured and not a placeholder
    if not OPENAI_API_KEY or OPENAI_API_KEY.startswith("sk-your-") or OPENAI_API_KEY == "your_openai_api_key_here":
        return None

    from openai import OpenAI
    import httpx

    httpx_client = httpx.Client(trust_env=False, timeout=30.0)

    return OpenAI(
        api_key=OPENAI_API_KEY,
        http_client=httpx_client,
    )

APP_BASE_URL = os.getenv("APP_BASE_URL", "http://127.0.0.1:5000")

INTERVIEWER_EMAIL = os.getenv("INTERVIEWER_EMAIL", "interviewer@example.com")
TIMEZONE = os.getenv("TIMEZONE", "Europe/London")

ALLOWED_EXTENSIONS = {"pdf", "doc", "docx"}

# Role taxonomy for planning and reporting
ROLE_TYPES = [
    "Project Director",
    "Project Manager",
    "Ops Manager",
    "Team Leader",
    "Case Handler",
    "Admin",
]

# Opportunity stages (for new Opportunities feature)
OPPORTUNITY_STAGES = [
    "Lead",
    "Closed Won",
    "Closed Lost",
]

app = Flask(__name__)
app.config["SECRET_KEY"] = SECRET_KEY
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Session timeout configuration (30 minutes for regular sessions)
from datetime import timedelta
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)
app.config['SESSION_REFRESH_EACH_REQUEST'] = True  # Refresh timeout on each request

# Initialize CSRF protection
csrf = CSRFProtect(app)

# Initialize rate limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

# Security headers
@app.after_request
def set_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    response.headers['Content-Security-Policy'] = "default-src 'self' https://cdn.tailwindcss.com https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; script-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com https://cdn.jsdelivr.net; style-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; font-src 'self' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com data:"
    return response

# Health check endpoint for Railway
@app.route("/health")
def health():
    return jsonify({"status": "healthy", "timestamp": datetime.datetime.utcnow().isoformat()}), 200

@app.route("/system/status")
def system_status():
    """
    Diagnostic endpoint to check system configuration
    """
    openai_client = get_openai_client()
    
    status = {
        "timestamp": datetime.datetime.utcnow().isoformat(),
        "openai": {
            "configured": openai_client is not None,
            "key_present": bool(OPENAI_API_KEY),
            "key_valid": OPENAI_API_KEY and not OPENAI_API_KEY.startswith("sk-your-") and OPENAI_API_KEY != "your_openai_api_key_here" if OPENAI_API_KEY else False
        },
        "database": {
            "connected": True  # If we get here, DB is working
        },
        "features": {
            "ai_scoring": openai_client is not None,
            "ai_summarization": openai_client is not None,
            "authentication": True
        }
    }
    
    # Test OpenAI connection if configured
    if openai_client:
        try:
            # Simple test to verify API key works
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )
            status["openai"]["api_test"] = "success"
        except Exception as e:
            status["openai"]["api_test"] = f"failed: {str(e)}"
    
    return jsonify(status), 200


# SECURITY: Diagnostic endpoints removed to prevent information disclosure
# These were temporary debugging endpoints that exposed:
# - Database schema structure (/system/db-schema)
# - User email addresses (/system/list-users)
# Removed for CREST compliance


@app.context_processor
def inject_template_helpers():
    def view_exists(name: str) -> bool:
        try:
            return name in current_app.view_functions
        except Exception:
            return False
    
    def csp_nonce():
        """Generate a CSP nonce for inline scripts."""
        import secrets
        return secrets.token_urlsafe(16)
    
    return {"view_exists": view_exists, "csp_nonce": csp_nonce}

# --- Auth / sessions ---
login_manager = LoginManager(app)
login_manager.login_view = "login"  # default guard for worker-only pages

@login_manager.unauthorized_handler
def unauthorized_callback():
    """Handle unauthorized access - return JSON for API routes, redirect for others"""
    if request.path.startswith('/api/'):
        return jsonify({"ok": False, "error": "Unauthorized. Please log in."}), 401
    return redirect(url_for('login', next=request.url))

@login_manager.user_loader
def load_user(user_id):
    with Session(engine) as s:
        user = s.get(User, int(user_id))
        if user:
            s.expunge(user)
            return user
    return None

# ========== Authentication Routes ==========

@app.route("/setup-first-user", methods=["GET", "POST"])
@limiter.limit("5 per hour")  # Rate limit: 5 setup attempts per hour per IP
def setup_first_user():
    """
    One-time setup route to create the first admin user
    Only works if NO users exist in the database
    """
    with Session(engine) as s:
        # Check if any users exist
        user_count = s.execute(select(func.count(User.id))).scalar()
        
        if user_count > 0:
            flash("Setup already completed. Users already exist.", "error")
            return redirect(url_for('login'))
    
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        
        if not name or not email or not password:
            flash("All fields are required", "error")
            return render_template("setup_first_user.html")
        
        # Validate password length
        if len(password) < 8:
            flash("Password must be at least 8 characters", "error")
            return render_template("setup_first_user.html")
        
        with Session(engine) as s:
            # Double-check no users exist
            user_count = s.execute(select(func.count(User.id))).scalar()
            if user_count > 0:
                flash("Setup already completed", "error")
                return redirect(url_for('login'))
            
            # Create first user
            new_user = User(
                name=name,
                email=email,
                password_hash=generate_password_hash(password, method='pbkdf2:sha256')
            )
            s.add(new_user)
            s.commit()
            
            flash(f"Account created successfully! You can now login.", "success")
            return redirect(url_for('login'))
    
    return render_template("setup_first_user.html")

@app.route("/run-migration-secret-xyz123")
def run_migration():
    """
    One-time migration route to add security columns
    Visit this URL once to run the migration, then it self-disables
    """
    try:
        from sqlalchemy import text
        
        results = []
        
        with engine.begin() as conn:
            # Check which columns already exist
            result = conn.execute(text("""
                SELECT column_name FROM information_schema.columns 
                WHERE table_name = 'users'
            """))
            existing_cols = {row[0] for row in result}
            results.append(f"Existing columns: {', '.join(existing_cols)}")
            
            # Add security columns
            security_columns = {
                "role": "VARCHAR(50) DEFAULT 'employee'",
                "is_active": "BOOLEAN DEFAULT TRUE",
                "last_login": "TIMESTAMP",
                "failed_login_attempts": "INTEGER DEFAULT 0",
                "locked_until": "TIMESTAMP",
                "totp_secret": "VARCHAR(32)",
                "totp_enabled": "BOOLEAN DEFAULT FALSE",
                "backup_codes": "TEXT",
                "session_token": "VARCHAR(255)",
                "last_ip": "VARCHAR(45)",
                "last_user_agent": "TEXT",
                "magic_token": "VARCHAR(255)",
                "magic_token_expires": "TIMESTAMP"
            }
            
            for col_name, col_def in security_columns.items():
                if col_name not in existing_cols:
                    try:
                        conn.execute(text(f"ALTER TABLE users ADD COLUMN {col_name} {col_def}"))
                        results.append(f"✅ Added column: {col_name}")
                    except Exception as e:
                        results.append(f"❌ Failed to add {col_name}: {str(e)}")
                else:
                    results.append(f"⏭️  Column {col_name} already exists")
            
            # Create audit_logs table
            try:
                conn.execute(text("""
                    CREATE TABLE IF NOT EXISTS audit_logs (
                        id SERIAL PRIMARY KEY,
                        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL,
                        user_id INTEGER,
                        user_email VARCHAR(255),
                        event_type VARCHAR(50) NOT NULL,
                        event_category VARCHAR(50) NOT NULL,
                        resource_type VARCHAR(50),
                        resource_id INTEGER,
                        action VARCHAR(255) NOT NULL,
                        details TEXT,
                        status VARCHAR(20) DEFAULT 'success',
                        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE SET NULL
                    )
                """))
                results.append("✅ Created audit_logs table")
            except Exception as e:
                results.append(f"⚠️  audit_logs: {str(e)}")
            
            # Create password_history table
            try:
                conn.execute(text("""
                    CREATE TABLE IF NOT EXISTS password_history (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER NOT NULL,
                        password_hash VARCHAR(255) NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL,
                        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
                    )
                """))
                results.append("✅ Created password_history table")
            except Exception as e:
                results.append(f"⚠️  password_history: {str(e)}")
            
            # Create indexes
            indexes = [
                "CREATE INDEX IF NOT EXISTS idx_audit_logs_timestamp ON audit_logs(timestamp)",
                "CREATE INDEX IF NOT EXISTS idx_audit_logs_user_id ON audit_logs(user_id)",
                "CREATE INDEX IF NOT EXISTS idx_password_history_user_id ON password_history(user_id)",
                "CREATE INDEX IF NOT EXISTS idx_password_history_created_at ON password_history(created_at)",
                "CREATE INDEX IF NOT EXISTS idx_users_email_idx ON users(email)",
                "CREATE INDEX IF NOT EXISTS idx_users_role_idx ON users(role)",
                "CREATE INDEX IF NOT EXISTS idx_users_totp_enabled ON users(totp_enabled)"
            ]
            
            for idx_stmt in indexes:
                try:
                    conn.execute(text(idx_stmt))
                except Exception:
                    pass
            
            results.append("✅ Created indexes")
            
            # Set first user as admin if exists
            try:
                conn.execute(text("""
                    UPDATE users 
                    SET role = 'admin' 
                    WHERE id = (SELECT MIN(id) FROM users)
                    AND role IS NULL
                """))
                results.append("✅ Set first user as admin")
            except Exception as e:
                results.append(f"⚠️  Admin role: {str(e)}")
        
        html = "<html><body style='font-family: monospace; padding: 20px;'>"
        html += "<h1>✅ Migration Complete!</h1>"
        html += "<pre>" + "\n".join(results) + "</pre>"
        html += "<p><strong>Next steps:</strong></p>"
        html += "<ol>"
        html += "<li>I will now redeploy the app with security features enabled</li>"
        html += "<li>Account lockout will be active</li>"
        html += "<li>Audit logging will work</li>"
        html += "</ol>"
        html += "<p><a href='/login'>Go to Login</a></p>"
        html += "</body></html>"
        
        return html
        
    except Exception as e:
        import traceback
        return f"<pre>❌ Migration failed:\n\n{str(e)}\n\n{traceback.format_exc()}</pre>", 500

@app.route("/login", methods=["GET", "POST"])
@limiter.limit("10 per minute")  # Rate limit: 10 login attempts per minute per IP
def login():
    """Staff login page with full security features"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    form = WorkerLoginForm()
    
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        
        if not email or not password:
            flash("Please enter both email and password", "error")
            return render_template("auth_login.html", form=form)
        
        try:
            with Session(engine) as s:
                # Query user
                user = s.scalar(select(User).where(User.email == email))
                
                if not user:
                    # Audit log failed attempt
                    try:
                        log_audit_event('login', 'auth', f'Failed login attempt for {email}', 
                                      details={'reason': 'user_not_found'}, status='failure')
                    except:
                        pass
                    flash("Invalid email or password", "error")
                    return render_template("auth_login.html", form=form)
                
                # Check if account is locked
                try:
                    if user.is_locked():
                        remaining_time = user.locked_until - datetime.datetime.utcnow()
                        remaining_minutes = int(remaining_time.total_seconds() / 60) + 1
                        try:
                            log_audit_event('login', 'auth', 
                                          f'Locked account login attempt for {email}',
                                          details={'remaining_minutes': remaining_minutes},
                                          status='failure')
                        except:
                            pass
                        flash(f"Account locked. Try again in {remaining_minutes} minutes.", "error")
                        return render_template("auth_login.html", form=form)
                except:
                    pass  # If lockout check fails, continue with login
                
                # Verify password
                if not check_password_hash(user.password_hash, password):
                    # Failed login - increment attempts
                    try:
                        user.failed_login_attempts = (user.failed_login_attempts or 0) + 1
                        
                        if user.failed_login_attempts >= 5:
                            user.locked_until = datetime.datetime.utcnow() + timedelta(minutes=30)
                            s.commit()
                            try:
                                log_audit_event('login', 'auth', 
                                              f'Account locked for {email} after 5 failed attempts',
                                              status='warning')
                            except:
                                pass
                            flash("Too many failed attempts. Account locked for 30 minutes.", "error")
                        else:
                            s.commit()
                            remaining = 5 - user.failed_login_attempts
                            try:
                                log_audit_event('login', 'auth', 
                                              f'Failed login attempt for {email}',
                                              details={'attempts': user.failed_login_attempts, 'remaining': remaining},
                                              status='failure')
                            except:
                                pass
                            flash(f"Invalid email or password. {remaining} attempts remaining.", "error")
                    except:
                        # If security features fail, show generic message
                        flash("Invalid email or password", "error")
                    return render_template("auth_login.html", form=form)
                
                # Password is correct - reset security counters
                try:
                    user.failed_login_attempts = 0
                    user.locked_until = None
                    user.last_login = datetime.datetime.utcnow()
                    s.commit()
                except:
                    pass  # Continue even if security updates fail
                
                # Check if 2FA is enabled
                if user.totp_enabled:
                    # Store user ID and email in session for 2FA verification
                    session['pending_user_id'] = user.id
                    session['pending_user_email'] = user.email
                    session['remember_me'] = request.form.get('remember_me') == 'on'
                    session['next_page'] = request.args.get('next', '/')
                    return redirect(url_for('verify_2fa'))
                
                # 2FA NOT enabled - require setup before allowing access
                # Store user ID in session and redirect to 2FA setup
                session['setup_2fa_user_id'] = user.id
                session['setup_2fa_remember'] = request.form.get('remember_me') == 'on'
                session['setup_2fa_next'] = request.args.get('next', '/')
                flash("Two-Factor Authentication is required. Please set up 2FA to continue.", "warning")
                return redirect(url_for('mandatory_2fa_setup'))
        
        except Exception as e:
            # Log the actual error for debugging
            print(f"Login error: {str(e)}")
            import traceback
            traceback.print_exc()
            flash(f"Login error: {str(e)}", "error")
            return render_template("auth_login.html", form=form)
    
    return render_template("auth_login.html", form=form)

@app.route("/forgot-password", methods=["GET", "POST"])
@limiter.limit("5 per minute")
def forgot_password():
    """Forgot password page - sends password reset link"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        
        if not email:
            flash("Please enter your email address", "error")
            return render_template("forgot_password.html")
        
        with Session(engine) as s:
            user = s.scalar(select(User).where(User.email == email))
            
            # Always show success message (security: don't reveal if email exists)
            # In production, you would send an actual email here
            if user:
                # Generate reset token (in production, send via email)
                import secrets
                reset_token = secrets.token_urlsafe(32)
                # Store token in session temporarily (in production, store in DB with expiry)
                session['reset_token'] = reset_token
                session['reset_email'] = email
                log_audit_event('password_reset', 'auth', f'Password reset requested for {email}', status='info')
        
        flash("If an account exists with that email, you will receive a password reset link shortly.", "success")
        return redirect(url_for('login'))
    
    return render_template("forgot_password.html")


@app.route("/request-access", methods=["GET", "POST"])
@limiter.limit("3 per minute")
def request_access():
    """Request access page - new user registration request"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        company = request.form.get("company", "").strip()
        reason = request.form.get("reason", "").strip()
        
        if not name or not email:
            flash("Name and email are required", "error")
            return render_template("request_access.html")
        
        with Session(engine) as s:
            # Check if user already exists
            existing = s.scalar(select(User).where(User.email == email))
            if existing:
                flash("An account with this email already exists. Please login or reset your password.", "info")
                return redirect(url_for('login'))
        
        # Log the access request (in production, store in DB and notify admin)
        log_audit_event('access_request', 'auth', f'Access requested by {name} ({email})', 
                       details={'company': company, 'reason': reason}, status='info')
        
        flash("Your access request has been submitted. An administrator will review your request and contact you shortly.", "success")
        return redirect(url_for('login'))
    
    return render_template("request_access.html")


@app.route("/logout")
@login_required
def logout():
    """Staff logout"""
    # Log logout before clearing session
    log_audit_event('logout', 'auth', 'User logged out')
    logout_user()
    flash("You have been logged out", "info")
    return redirect(url_for('login'))

# ========== Magic Link Authentication ==========
@app.route("/magic-link/<token>")
def magic_link_login(token):
    """
    Magic link login - allows user to set password on first login.
    Token is single-use and expires after 48 hours.
    """
    with Session(engine) as s:
        user = s.scalar(select(User).where(User.magic_token == token))
        
        if not user:
            flash("Invalid or expired magic link. Please contact your administrator.", "error")
            return redirect(url_for('login'))
        
        # Check if token has expired
        if user.magic_token_expires and user.magic_token_expires < datetime.datetime.utcnow():
            flash("This magic link has expired. Please contact your administrator for a new one.", "error")
            return redirect(url_for('login'))
        
        # Store user info in session for password setup
        session['magic_link_user_id'] = user.id
        session['magic_link_user_email'] = user.email
        session['magic_link_user_name'] = user.name
        
        return redirect(url_for('magic_link_set_password'))

@app.route("/magic-link/set-password", methods=["GET", "POST"])
def magic_link_set_password():
    """
    Allow user to set their password after clicking magic link.
    """
    user_id = session.get('magic_link_user_id')
    user_email = session.get('magic_link_user_email')
    user_name = session.get('magic_link_user_name')
    
    if not user_id:
        flash("Invalid session. Please use your magic link again.", "error")
        return redirect(url_for('login'))
    
    if request.method == "POST":
        password = request.form.get("password", "")
        confirm_password = request.form.get("confirm_password", "")
        
        if not password or not confirm_password:
            flash("Both password fields are required", "error")
            return render_template("magic_link_set_password.html", user_name=user_name, user_email=user_email)
        
        if password != confirm_password:
            flash("Passwords do not match", "error")
            return render_template("magic_link_set_password.html", user_name=user_name, user_email=user_email)
        
        # Validate password strength
        is_valid, message, score = validate_password_strength(password, user_email)
        if not is_valid:
            flash(message, "error")
            return render_template("magic_link_set_password.html", user_name=user_name, user_email=user_email)
        
        with Session(engine) as s:
            user = s.get(User, user_id)
            if not user:
                flash("User not found", "error")
                return redirect(url_for('login'))
            
            # Set the new password
            user.password_hash = generate_password_hash(password, method='pbkdf2:sha256')
            # Clear the magic token (single use)
            user.magic_token = None
            user.magic_token_expires = None
            s.commit()
            
            # Log the password setup
            log_audit_event('update', 'auth', f'User set password via magic link: {user_email}',
                          'user', user_id)
            
            # Clear session
            session.pop('magic_link_user_id', None)
            session.pop('magic_link_user_email', None)
            session.pop('magic_link_user_name', None)
            
            flash("Your password has been set successfully! Please log in.", "success")
            return redirect(url_for('login'))
    
    return render_template("magic_link_set_password.html", user_name=user_name, user_email=user_email)

def generate_magic_link(user_id: int) -> str:
    """
    Generate a magic link token for a user.
    Returns the full URL for the magic link.
    """
    import secrets
    token = secrets.token_urlsafe(32)
    expires = datetime.datetime.utcnow() + datetime.timedelta(hours=48)
    
    with Session(engine) as s:
        user = s.get(User, user_id)
        if user:
            user.magic_token = token
            user.magic_token_expires = expires
            s.commit()
    
    return token

# ========== One-Time Admin Setup ==========
@app.route("/setup/admin", methods=["GET", "POST"])
def setup_admin():
    """
    One-time admin setup route (no authentication required)
    Only works if no super_admin users exist in the database
    """
    with Session(engine) as s:
        # Check if any super_admin already exists
        existing_admin = s.scalar(
            select(User).where(User.role == 'super_admin')
        )
        
        if existing_admin:
            flash("Admin user already exists. Please login or contact support.", "error")
            return redirect(url_for('login'))
        
        if request.method == "POST":
            name = request.form.get("name", "").strip()
            email = request.form.get("email", "").strip().lower()
            password = request.form.get("password", "")
            confirm_password = request.form.get("confirm_password", "")
            
            if not name or not email or not password or not confirm_password:
                flash("All fields are required", "error")
                return render_template("setup_admin.html")
            
            if password != confirm_password:
                flash("Passwords do not match", "error")
                return render_template("setup_admin.html")
            
            # Validate password strength
            is_valid, message, score = validate_password_strength(password, email)
            if not is_valid:
                flash(message, "error")
                return render_template("setup_admin.html")
            
            try:
                # Create super admin user
                new_admin = User(
                    name=name,
                    email=email,
                    password_hash=generate_password_hash(password, method='pbkdf2:sha256'),
                    role='super_admin',
                    is_active=True,
                    created_at=datetime.datetime.utcnow()
                )
                s.add(new_admin)
                s.commit()
                
                # Log the admin creation
                log_audit_event('create', 'user_mgmt', f'Initial super admin created: {email}',
                              'user', new_admin.id, details={'setup_route': True})
                
                flash("✅ Super Admin created successfully! You can now login.", "success")
                return redirect(url_for('login'))
            except Exception as e:
                flash(f"Error creating admin: {str(e)}", "error")
                return render_template("setup_admin.html")
    
    return render_template("setup_admin.html")

# ========== Admin User Management (Super Admin Only) ==========
def super_admin_required(f):
    """Decorator to require super_admin role for access"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            flash("Please log in to access this page.", "error")
            return redirect(url_for('login'))
        if current_user.role != 'super_admin':
            flash("Access denied. Super Admin privileges required.", "error")
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

@app.route("/admin/create-user", methods=["GET", "POST"])
@login_required
@super_admin_required
def admin_create_user():
    """Create a new user with magic link for password setup"""
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        role = request.form.get("role", "employee")
        send_magic_link = request.form.get("send_magic_link") == "1"
        
        if not name or not email:
            flash("Name and email are required", "error")
            return render_template("admin_create_user.html")
        
        try:
            with Session(engine) as s:
                # Check if user already exists
                existing = s.scalar(select(User).where(User.email == email))
                if existing:
                    flash(f"User with email {email} already exists!", "error")
                    return render_template("admin_create_user.html")
                
                # Generate magic link token
                import secrets
                magic_token = secrets.token_urlsafe(32)
                magic_expires = datetime.datetime.utcnow() + datetime.timedelta(hours=48)
                
                # Create new user with placeholder password (user will set via magic link)
                new_user = User(
                    name=name,
                    email=email,
                    password_hash=generate_password_hash(secrets.token_urlsafe(32), method='pbkdf2:sha256'),
                    role=role,
                    is_active=True,
                    created_at=datetime.datetime.utcnow(),
                    magic_token=magic_token,
                    magic_token_expires=magic_expires
                )
                s.add(new_user)
                s.commit()
                
                # Generate the magic link URL
                magic_link_url = request.url_root.rstrip('/') + url_for('magic_link_login', token=magic_token)
                
                # Log user creation
                log_audit_event('create', 'user_mgmt', f'Created new user: {email}',
                              'user', new_user.id, {'role': role, 'created_by': current_user.email})
                
                # Send welcome email with magic link if requested
                if send_magic_link:
                    try:
                        html_body = f"""
                        <html>
                        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
                            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                                <h2 style="color: #1e3a8a;">Welcome to Optimus OS1</h2>
                                <p>Hi {name},</p>
                                <p>Your account has been created. Click the button below to set your password and get started:</p>
                                <div style="text-align: center; margin: 30px 0;">
                                    <a href="{magic_link_url}" style="background: #1e3a8a; color: white; padding: 15px 30px; text-decoration: none; border-radius: 8px; display: inline-block; font-weight: bold; font-size: 16px;">
                                        Set Your Password
                                    </a>
                                </div>
                                <div style="background: #f3f4f6; padding: 15px; border-radius: 8px; margin: 20px 0;">
                                    <p style="margin: 5px 0;"><strong>Email:</strong> {email}</p>
                                    <p style="margin: 5px 0;"><strong>Role:</strong> {role.replace('_', ' ').title()}</p>
                                </div>
                                <p style="color: #666; font-size: 14px;">
                                    <i>This link expires in 48 hours. If you didn't request this account, please ignore this email.</i>
                                </p>
                                <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 20px 0;">
                                <p style="color: #999; font-size: 12px;">This is an automated message from Optimus OS1.</p>
                            </div>
                        </body>
                        </html>
                        """
                        send_email(email, "Welcome to Optimus OS1 - Set Your Password", html_body)
                        flash(f"✅ User created! Magic link sent to {email}", "success")
                    except Exception as e:
                        # Show the magic link directly if email fails
                        flash(f"✅ User created but email failed. Share this link with {name}: {magic_link_url}", "warning")
                else:
                    # Show the magic link directly
                    flash(f"✅ User created! Share this link with {name} to set their password:", "success")
                    flash(f"🔗 {magic_link_url}", "info")
                
                return redirect(url_for('admin_list_users'))
        except Exception as e:
            flash(f"Error creating user: {str(e)}", "error")
            return render_template("admin_create_user.html")
    
    return render_template("admin_create_user.html")

@app.route("/admin/list-users")
@login_required
@super_admin_required
def admin_list_users():
    """Temporary page to list all users"""
    try:
        with Session(engine) as s:
            users = s.scalars(select(User).order_by(User.created_at.desc())).all()
            # Detach from session
            for user in users:
                s.expunge(user)
            return render_template("admin_list_users.html", users=users)
    except Exception as e:
        flash(f"Error loading users: {str(e)}", "error")
        return render_template("admin_list_users.html", users=[])

@app.route("/admin/unlock-user/<int:user_id>", methods=["POST"])
@login_required
@super_admin_required
def admin_unlock_user(user_id):
    """Admin route to unlock a locked user account"""
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_list_users'))
        
        user.failed_login_attempts = 0
        user.locked_until = None
        s.commit()
        
        # Log unlock action
        log_audit_event('unlock', 'user_mgmt', f'Unlocked user account: {user.email}',
                      'user', user.id)
        
        flash(f"User {user.email} unlocked successfully", "success")
    
    return redirect(url_for('admin_list_users'))

@app.route("/admin/update-user", methods=["POST"])
@login_required
@super_admin_required
def admin_update_user():
    """Admin route to update user details"""
    user_id = request.form.get('user_id', type=int)
    name = request.form.get('name', '').strip()
    email = request.form.get('email', '').strip().lower()
    role = request.form.get('role', 'employee')
    
    if not user_id or not name or not email:
        flash("Missing required fields", "error")
        return redirect(url_for('admin_list_users'))
    
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_list_users'))
        
        # Check if email is being changed and if new email already exists
        if email != user.email:
            existing = s.scalar(select(User).where(User.email == email))
            if existing:
                flash(f"Email {email} is already in use", "error")
                return redirect(url_for('admin_list_users'))
        
        old_values = {'name': user.name, 'email': user.email, 'role': user.role}
        user.name = name
        user.email = email
        user.role = role
        s.commit()
        
        # Log update action
        log_audit_event('update', 'user_mgmt', f'Updated user: {email}',
                      'user', user.id,
                      {'old_values': old_values, 'new_values': {'name': name, 'email': email, 'role': role}})
        
        flash(f"User {email} updated successfully", "success")
    
    return redirect(url_for('admin_list_users'))

@app.route("/admin/send-magic-link/<int:user_id>", methods=["POST"])
@login_required
@super_admin_required
def admin_send_magic_link(user_id):
    """Admin route to send a magic link for password reset"""
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_list_users'))
        
        # Generate magic link token
        import secrets
        magic_token = secrets.token_urlsafe(32)
        magic_expires = datetime.datetime.utcnow() + datetime.timedelta(hours=48)
        
        user.magic_token = magic_token
        user.magic_token_expires = magic_expires
        s.commit()
        
        # Generate the magic link URL
        magic_link_url = request.url_root.rstrip('/') + url_for('magic_link_login', token=magic_token)
        
        # Log the action
        log_audit_event('update', 'user_mgmt', f'Sent magic link to: {user.email}',
                      'user', user.id)
        
        # Try to send email
        try:
            html_body = f"""
            <html>
            <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
                <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                    <h2 style="color: #1e3a8a;">Reset Your Password</h2>
                    <p>Hi {user.name},</p>
                    <p>A password reset has been requested for your account. Click the button below to set a new password:</p>
                    <div style="text-align: center; margin: 30px 0;">
                        <a href="{magic_link_url}" style="background: #1e3a8a; color: white; padding: 15px 30px; text-decoration: none; border-radius: 8px; display: inline-block; font-weight: bold; font-size: 16px;">
                            Reset Password
                        </a>
                    </div>
                    <p style="color: #666; font-size: 14px;">
                        <i>This link expires in 48 hours. If you didn't request this, please ignore this email.</i>
                    </p>
                    <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 20px 0;">
                    <p style="color: #999; font-size: 12px;">This is an automated message from Optimus OS1.</p>
                </div>
            </body>
            </html>
            """
            send_email(user.email, "Reset Your Password - Optimus OS1", html_body)
            flash(f"✅ Magic link sent to {user.email}", "success")
        except Exception as e:
            # Show the magic link directly if email fails
            flash(f"Email failed. Share this link with {user.name}:", "warning")
            flash(f"🔗 {magic_link_url}", "info")
    
    return redirect(url_for('admin_list_users'))

@app.route("/admin/disable-user-2fa/<int:user_id>", methods=["POST"])
@login_required
@super_admin_required
def admin_disable_user_2fa(user_id):
    """Admin route to disable 2FA for a user"""
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_list_users'))
        
        user.totp_secret = None
        user.totp_enabled = False
        user.backup_codes = None
        s.commit()
        
        # Log 2FA disable action
        log_audit_event('update', 'security', f'Admin disabled 2FA for: {user.email}',
                      'user', user.id)
        
        flash(f"2FA disabled for {user.email}. They will need to set it up again.", "success")
    
    return redirect(url_for('admin_list_users'))

@app.route("/admin/deactivate-user/<int:user_id>", methods=["POST"])
@login_required
@super_admin_required
def admin_deactivate_user(user_id):
    """Admin route to deactivate a user account"""
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_list_users'))
        
        if user.id == current_user.id:
            flash("You cannot deactivate your own account", "error")
            return redirect(url_for('admin_list_users'))
        
        user.is_active = False
        s.commit()
        
        # Log deactivation action
        log_audit_event('update', 'user_mgmt', f'Deactivated user: {user.email}',
                      'user', user.id)
        
        flash(f"User {user.email} has been deactivated", "success")
    
    return redirect(url_for('admin_list_users'))

@app.route("/admin/activate-user/<int:user_id>", methods=["POST"])
@login_required
@super_admin_required
def admin_activate_user(user_id):
    """Admin route to activate a user account"""
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_list_users'))
        
        user.is_active = True
        s.commit()
        
        # Log activation action
        log_audit_event('update', 'user_mgmt', f'Activated user: {user.email}',
                      'user', user.id)
        
        flash(f"User {user.email} has been activated", "success")
    
    return redirect(url_for('admin_list_users'))

@app.route("/admin/delete-user/<int:user_id>", methods=["POST"])
@login_required
@super_admin_required
def admin_delete_user(user_id):
    """Admin route to permanently delete a user account"""
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_list_users'))
        
        if user.id == current_user.id:
            flash("You cannot delete your own account", "error")
            return redirect(url_for('admin_list_users'))
        
        user_email = user.email
        user_name = user.name
        
        # Delete the user
        s.delete(user)
        s.commit()
        
        # Log deletion action
        log_audit_event('delete', 'user_mgmt', f'Deleted user: {user_email}',
                      'user', user_id, {'deleted_user_name': user_name, 'deleted_user_email': user_email})
        
        flash(f"User {user_email} has been permanently deleted", "success")
    
    return redirect(url_for('admin_list_users'))


# ========== Password Security Functions ==========

def validate_password_strength(password: str, email: str = None) -> Tuple[bool, str, int]:
    """
    Validate password strength according to security policy.
    
    Returns: (is_valid, error_message, strength_score)
    - is_valid: True if password meets all requirements
    - error_message: Description of what's missing or "Strong password"
    - strength_score: 0-5 based on requirements met
    """
    errors = []
    score = 0
    
    # Minimum 12 characters
    if len(password) < 12:
        errors.append("at least 12 characters")
    else:
        score += 1
    
    # At least one uppercase letter
    if not re.search(r'[A-Z]', password):
        errors.append("an uppercase letter")
    else:
        score += 1
    
    # At least one lowercase letter
    if not re.search(r'[a-z]', password):
        errors.append("a lowercase letter")
    else:
        score += 1
    
    # At least one number
    if not re.search(r'\d', password):
        errors.append("a number")
    else:
        score += 1
    
    # At least one special character
    if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
        errors.append("a special character (!@#$%^&*() etc.)")
    else:
        score += 1
    
    # Cannot contain email username
    if email:
        username = email.split('@')[0].lower()
        if username in password.lower():
            errors.append("cannot contain your email username")
            score = 0
    
    if errors:
        return False, f"Password must contain {', '.join(errors)}", score
    
    return True, "Strong password", score

# ========== Security Enhancement Functions ==========

def verify_recaptcha(token: str) -> bool:
    """
    Verify Google reCAPTCHA v3 token.
    Returns True if valid, False otherwise.
    """
    if not RECAPTCHA_ENABLED:
        return True  # Skip verification if not configured
    
    try:
        response = requests.post(
            'https://www.google.com/recaptcha/api/siteverify',
            data={
                'secret': RECAPTCHA_SECRET_KEY,
                'response': token
            },
            timeout=5
        )
        result = response.json()
        
        # reCAPTCHA v3 returns a score from 0.0 to 1.0
        # 1.0 is very likely a good interaction, 0.0 is very likely a bot
        # We'll accept scores >= 0.5
        return result.get('success', False) and result.get('score', 0) >= 0.5
    except Exception as e:
        current_app.logger.error(f"reCAPTCHA verification failed: {str(e)}")
        return True  # Fail open to avoid blocking legitimate users if reCAPTCHA is down

def sanitize_input(text: str, allow_html: bool = False) -> str:
    """
    Sanitize user input to prevent XSS attacks.
    
    Args:
        text: Input text to sanitize
        allow_html: If True, allows safe HTML tags (for rich text)
    
    Returns:
        Sanitized text safe for display
    """
    if not text:
        return text
    
    import bleach
    
    if allow_html:
        # Allow safe HTML tags for rich text content
        allowed_tags = [
            'p', 'br', 'strong', 'em', 'u', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
            'ul', 'ol', 'li', 'a', 'blockquote', 'code', 'pre'
        ]
        allowed_attrs = {
            'a': ['href', 'title'],
            'img': ['src', 'alt', 'title']
        }
        return bleach.clean(text, tags=allowed_tags, attributes=allowed_attrs, strip=True)
    else:
        # Strip all HTML for plain text fields
        return bleach.clean(text, tags=[], strip=True)

def check_password_history(user_id: int, new_password: str, history_limit: int = 5) -> bool:
    """
    Check if password has been used recently.
    
    Args:
        user_id: User ID to check
        new_password: New password to check against history
        history_limit: Number of previous passwords to check (default 5)
    
    Returns:
        True if password is OK to use, False if it's been used recently
    """
    try:
        with Session(engine) as s:
            # Get last N password hashes for this user
            recent_passwords = s.execute(
                text("""
                    SELECT password_hash FROM password_history
                    WHERE user_id = :user_id
                    ORDER BY created_at DESC
                    LIMIT :limit
                """),
                {"user_id": user_id, "limit": history_limit}
            ).fetchall()
            
            # Check if new password matches any recent passwords
            for (old_hash,) in recent_passwords:
                if check_password_hash(old_hash, new_password):
                    return False  # Password has been used recently
            
            return True  # Password is OK to use
    except Exception as e:
        current_app.logger.error(f"Password history check failed: {str(e)}")
        return True  # Fail open if check fails

def save_password_to_history(user_id: int, password_hash: str):
    """
    Save password hash to history for reuse prevention.
    
    Args:
        user_id: User ID
        password_hash: Hashed password to save
    """
    try:
        with Session(engine) as s:
            s.execute(
                text("""
                    INSERT INTO password_history (user_id, password_hash, created_at)
                    VALUES (:user_id, :password_hash, CURRENT_TIMESTAMP)
                """),
                {"user_id": user_id, "password_hash": password_hash}
            )
            s.commit()
    except Exception as e:
        current_app.logger.error(f"Failed to save password history: {str(e)}")

# ========== Audit Logging Functions ==========

def log_audit_event(event_type: str, event_category: str, action: str, 
                   resource_type: str = None, resource_id: int = None,
                   details: dict = None, status: str = 'success'):
    """
    Log an audit event for security and compliance tracking.
    
    Args:
        event_type: Type of event (login, logout, create, update, delete, view, export)
        event_category: Category (auth, user_mgmt, data_access, security)
        action: Description of the action taken
        resource_type: Type of resource affected (candidate, job, user, etc.)
        resource_id: ID of the affected resource
        details: Additional details as dictionary (will be JSON encoded)
        status: Result status (success, failure, warning)
    
    Examples:
        log_audit_event('login', 'auth', 'User logged in successfully')
        log_audit_event('create', 'user_mgmt', 'Created new user', 'user', new_user.id)
        log_audit_event('login', 'auth', 'Failed login attempt', status='failure')
    """
    try:
        log = AuditLog(
            user_id=current_user.id if current_user.is_authenticated else None,
            user_email=current_user.email if current_user.is_authenticated else None,
            event_type=event_type,
            event_category=event_category,
            resource_type=resource_type,
            resource_id=resource_id,
            action=action,
            details=json.dumps(details) if details else None,
            ip_address=request.remote_addr if request else None,
            user_agent=request.headers.get('User-Agent') if request else None,
            status=status
        )
        
        with Session(engine) as s:
            s.add(log)
            s.commit()
    except Exception as e:
        # Don't let audit logging failures break the application
        current_app.logger.error(f"Audit logging failed: {str(e)}")

@login_required
@app.route("/admin/system-diagnostics")
def admin_system_diagnostics():
    """System diagnostics and configuration check"""
    openai_client = get_openai_client()
    
    # Test OpenAI connection
    openai_test_result = None
    if openai_client:
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )
            openai_test_result = "✓ Connection successful"
        except Exception as e:
            openai_test_result = f"✗ Connection failed: {str(e)}"
    
    diagnostics = {
        "openai": {
            "configured": openai_client is not None,
            "key_present": bool(OPENAI_API_KEY),
            "key_prefix": OPENAI_API_KEY[:10] + "..." if OPENAI_API_KEY and len(OPENAI_API_KEY) > 10 else "Not set",
            "test_result": openai_test_result,
            "features_enabled": [
                "AI Match Scoring" if openai_client else "AI Match Scoring (disabled - no API key)",
                "AI CV Summarization" if openai_client else "AI CV Summarization (disabled - no API key)",
                "GPT-powered Candidate Analysis" if openai_client else "GPT-powered Analysis (disabled - no API key)"
            ]
        },
        "database": {
            "status": "✓ Connected",
            "engine": str(engine.url).split('@')[0] + "@[hidden]"  # Hide credentials
        },
        "authentication": {
            "status": "✓ Enabled",
            "login_manager": "Flask-Login",
            "password_hashing": "pbkdf2:sha256"
        }
    }
    
    return render_template("admin_system_diagnostics.html", diagnostics=diagnostics)

# ========== Admin Routes for Sidebar Navigation ==========
@login_required
@app.route("/admin/users")
def admin_users():
    """User management page - redirects to existing list users"""
    return redirect(url_for('admin_list_users'))

@login_required
@app.route("/admin/approvals")
def admin_approvals():
    """Approvals management page - timesheets, expenses, etc."""
    return render_template("admin_approvals.html")

@login_required
@app.route("/admin/invoices")
def admin_invoices():
    """Invoices management page"""
    import json
    with Session(engine) as s:
        # Get all invoices with counts
        invoices = s.scalars(select(Invoice).order_by(Invoice.invoice_date.desc())).all()
        
        # Calculate summary stats
        draft_invoices = [i for i in invoices if i.status == "Draft"]
        pending_invoices = [i for i in invoices if i.status == "Pending"]
        paid_invoices = [i for i in invoices if i.status == "Paid"]
        overdue_invoices = [i for i in invoices if i.status == "Overdue"]
        
        # Get unique clients for filter dropdown
        clients = s.execute(text("SELECT DISTINCT client FROM engagements WHERE client IS NOT NULL")).fetchall()
        client_list = [{"id": c[0], "name": c[0]} for c in clients if c[0]]
        
        return render_template("admin_invoices.html",
            invoices=invoices,
            draft_count=len(draft_invoices),
            draft_amount=sum(i.total_amount or 0 for i in draft_invoices),
            pending_count=len(pending_invoices),
            pending_amount=sum(i.total_amount or 0 for i in pending_invoices),
            paid_count=len(paid_invoices),
            paid_amount=sum(i.total_amount or 0 for i in paid_invoices),
            overdue_count=len(overdue_invoices),
            overdue_amount=sum(i.total_amount or 0 for i in overdue_invoices),
            clients=client_list
        )

@app.route("/admin/invoices/create", methods=["GET", "POST"])
@login_required
def admin_create_invoice():
    """Create a new invoice"""
    import json
    
    if request.method == "POST":
        with Session(engine) as s:
            # Generate invoice number
            year = datetime.datetime.utcnow().year
            month = datetime.datetime.utcnow().month
            count = s.scalar(select(func.count(Invoice.id)).where(
                Invoice.invoice_number.like(f"INV-{year}{month:02d}%")
            )) or 0
            invoice_number = f"INV-{year}{month:02d}-{count + 1:04d}"
            
            # Get form data
            engagement_id = request.form.get("engagement_id")
            client_name = request.form.get("client_name", "").strip()
            engagement_name = request.form.get("engagement_name", "").strip()
            
            # Parse dates
            invoice_date = datetime.datetime.utcnow()
            due_date_str = request.form.get("due_date", "")
            due_date = None
            if due_date_str:
                try:
                    due_date = datetime.datetime.strptime(due_date_str, "%Y-%m-%d")
                except:
                    pass
            
            # Get line items from form
            line_items = []
            descriptions = request.form.getlist("item_description[]")
            quantities = request.form.getlist("item_quantity[]")
            rates = request.form.getlist("item_rate[]")
            
            subtotal = 0
            for i, desc in enumerate(descriptions):
                if desc.strip():
                    qty = float(quantities[i]) if i < len(quantities) and quantities[i] else 1
                    rate = float(rates[i]) if i < len(rates) and rates[i] else 0
                    amount = qty * rate
                    subtotal += amount
                    line_items.append({
                        "description": desc.strip(),
                        "quantity": qty,
                        "rate": rate,
                        "amount": amount
                    })
            
            # Calculate VAT
            vat_rate = float(request.form.get("vat_rate", 20))
            vat_amount = subtotal * (vat_rate / 100)
            total_amount = subtotal + vat_amount
            
            # Create invoice
            invoice = Invoice(
                invoice_number=invoice_number,
                engagement_id=int(engagement_id) if engagement_id else None,
                client_name=client_name,
                engagement_name=engagement_name,
                invoice_date=invoice_date,
                due_date=due_date,
                subtotal=subtotal,
                vat_rate=vat_rate,
                vat_amount=vat_amount,
                total_amount=total_amount,
                status="Draft",
                line_items=json.dumps(line_items),
                notes=request.form.get("notes", ""),
                payment_terms=request.form.get("payment_terms", "Net 30"),
                created_by=current_user.id
            )
            s.add(invoice)
            s.commit()
            
            log_audit_event('create', 'billing', f'Created invoice {invoice_number}',
                          'invoice', invoice.id, {'client': client_name, 'amount': total_amount})
            
            flash(f"✅ Invoice {invoice_number} created successfully!", "success")
            return redirect(url_for('admin_invoices'))
    
    # GET - show create form
    with Session(engine) as s:
        engagements = s.scalars(select(Engagement).order_by(Engagement.name)).all()
        return render_template("admin_invoice_create.html", engagements=engagements)

@app.route("/admin/invoices/<int:invoice_id>")
@login_required
def admin_view_invoice(invoice_id):
    """View a single invoice"""
    import json
    with Session(engine) as s:
        invoice = s.get(Invoice, invoice_id)
        if not invoice:
            flash("Invoice not found", "error")
            return redirect(url_for('admin_invoices'))
        
        line_items = json.loads(invoice.line_items) if invoice.line_items else []
        return render_template("admin_invoice_view.html", invoice=invoice, line_items=line_items)

@app.route("/admin/invoices/<int:invoice_id>/edit", methods=["GET", "POST"])
@login_required
def admin_edit_invoice(invoice_id):
    """Edit an invoice"""
    import json
    with Session(engine) as s:
        invoice = s.get(Invoice, invoice_id)
        if not invoice:
            flash("Invoice not found", "error")
            return redirect(url_for('admin_invoices'))
        
        if request.method == "POST":
            # Update invoice fields
            invoice.client_name = request.form.get("client_name", invoice.client_name)
            invoice.engagement_name = request.form.get("engagement_name", invoice.engagement_name)
            invoice.notes = request.form.get("notes", "")
            invoice.payment_terms = request.form.get("payment_terms", "Net 30")
            
            # Parse due date
            due_date_str = request.form.get("due_date", "")
            if due_date_str:
                try:
                    invoice.due_date = datetime.datetime.strptime(due_date_str, "%Y-%m-%d")
                except:
                    pass
            
            # Get line items
            line_items = []
            descriptions = request.form.getlist("item_description[]")
            quantities = request.form.getlist("item_quantity[]")
            rates = request.form.getlist("item_rate[]")
            
            subtotal = 0
            for i, desc in enumerate(descriptions):
                if desc.strip():
                    qty = float(quantities[i]) if i < len(quantities) and quantities[i] else 1
                    rate = float(rates[i]) if i < len(rates) and rates[i] else 0
                    amount = qty * rate
                    subtotal += amount
                    line_items.append({
                        "description": desc.strip(),
                        "quantity": qty,
                        "rate": rate,
                        "amount": amount
                    })
            
            invoice.line_items = json.dumps(line_items)
            invoice.subtotal = subtotal
            invoice.vat_rate = float(request.form.get("vat_rate", 20))
            invoice.vat_amount = subtotal * (invoice.vat_rate / 100)
            invoice.total_amount = subtotal + invoice.vat_amount
            
            s.commit()
            
            log_audit_event('update', 'billing', f'Updated invoice {invoice.invoice_number}',
                          'invoice', invoice.id)
            
            flash(f"✅ Invoice updated successfully!", "success")
            return redirect(url_for('admin_view_invoice', invoice_id=invoice.id))
        
        # GET - show edit form
        line_items = json.loads(invoice.line_items) if invoice.line_items else []
        engagements = s.scalars(select(Engagement).order_by(Engagement.name)).all()
        return render_template("admin_invoice_edit.html", invoice=invoice, 
                             line_items=line_items, engagements=engagements)

@app.route("/admin/invoices/<int:invoice_id>/status", methods=["POST"])
@login_required
def admin_update_invoice_status(invoice_id):
    """Update invoice status"""
    with Session(engine) as s:
        invoice = s.get(Invoice, invoice_id)
        if not invoice:
            flash("Invoice not found", "error")
            return redirect(url_for('admin_invoices'))
        
        new_status = request.form.get("status", invoice.status)
        old_status = invoice.status
        invoice.status = new_status
        
        if new_status == "Paid" and not invoice.paid_date:
            invoice.paid_date = datetime.datetime.utcnow()
        elif new_status == "Pending" and not invoice.sent_at:
            invoice.sent_at = datetime.datetime.utcnow()
        
        s.commit()
        
        log_audit_event('update', 'billing', f'Invoice {invoice.invoice_number} status: {old_status} → {new_status}',
                      'invoice', invoice.id)
        
        flash(f"✅ Invoice status updated to {new_status}", "success")
    
    return redirect(url_for('admin_invoices'))

@app.route("/admin/invoices/<int:invoice_id>/delete", methods=["POST"])
@login_required
def admin_delete_invoice(invoice_id):
    """Delete an invoice"""
    with Session(engine) as s:
        invoice = s.get(Invoice, invoice_id)
        if not invoice:
            flash("Invoice not found", "error")
            return redirect(url_for('admin_invoices'))
        
        invoice_number = invoice.invoice_number
        s.delete(invoice)
        s.commit()
        
        log_audit_event('delete', 'billing', f'Deleted invoice {invoice_number}',
                      'invoice', invoice_id)
        
        flash(f"✅ Invoice {invoice_number} deleted", "success")
    
    return redirect(url_for('admin_invoices'))

@app.route("/admin/invoices/export")
@login_required
def admin_export_invoices():
    """Export invoices to CSV"""
    import csv
    import io
    
    with Session(engine) as s:
        invoices = s.scalars(select(Invoice).order_by(Invoice.invoice_date.desc())).all()
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Header row
        writer.writerow([
            "Invoice Number", "Client", "Engagement", "Invoice Date", "Due Date",
            "Subtotal", "VAT Rate", "VAT Amount", "Total", "Status", "Paid Date"
        ])
        
        # Data rows
        for inv in invoices:
            writer.writerow([
                inv.invoice_number,
                inv.client_name,
                inv.engagement_name or "",
                inv.invoice_date.strftime("%Y-%m-%d") if inv.invoice_date else "",
                inv.due_date.strftime("%Y-%m-%d") if inv.due_date else "",
                f"{inv.subtotal:.2f}",
                f"{inv.vat_rate:.1f}%",
                f"{inv.vat_amount:.2f}",
                f"{inv.total_amount:.2f}",
                inv.status,
                inv.paid_date.strftime("%Y-%m-%d") if inv.paid_date else ""
            ])
        
        output.seek(0)
        
        log_audit_event('export', 'billing', f'Exported {len(invoices)} invoices to CSV')
        
        return Response(
            output.getvalue(),
            mimetype="text/csv",
            headers={"Content-Disposition": f"attachment;filename=invoices_export_{datetime.datetime.utcnow().strftime('%Y%m%d')}.csv"}
        )

@app.route("/admin/invoices/settings", methods=["GET", "POST"])
@login_required
def admin_invoice_settings():
    """Invoice settings page"""
    # For now, just render a settings page
    # In a full implementation, these would be stored in a settings table
    return render_template("admin_invoice_settings.html")

@login_required
@app.route("/admin/audit-log")
def admin_audit_log():
    """Audit log page - shows all system activity"""
    # Audit log: viewing audit log (meta, but useful)
    
    with Session(engine) as s:
        # Get real audit log entries
        try:
            audit_entries = s.scalars(
                select(AuditLog)
                .order_by(AuditLog.timestamp.desc())
                .limit(500)
            ).all()
        except:
            audit_entries = []
    return render_template("admin_audit_log.html", audit_entries=audit_entries)

# ========== Portal Admin (Public Portal User Management) ==========

@login_required
@app.route("/admin/portal-users")
def admin_portal_users():
    """Admin page for managing public portal users (candidates who signed up via the portal)."""
    search = request.args.get("search", "").strip()
    status_filter = request.args.get("status", "")  # all, verified, unverified
    source_filter = request.args.get("source", "")  # all, portal, manual, import
    page = int(request.args.get("page", 1))
    per_page = 25
    
    with Session(engine) as s:
        query = select(Candidate).order_by(Candidate.created_at.desc())
        
        # Apply search filter
        if search:
            search_term = f"%{search}%"
            query = query.where(
                or_(
                    Candidate.name.ilike(search_term),
                    Candidate.email.ilike(search_term),
                    Candidate.phone.ilike(search_term),
                    Candidate.skills.ilike(search_term)
                )
            )
        
        # Apply status filter
        if status_filter == "verified":
            query = query.where(Candidate.email_verified == True)
        elif status_filter == "unverified":
            query = query.where(or_(Candidate.email_verified == False, Candidate.email_verified == None))
        
        # Apply source filter
        if source_filter and source_filter != "all":
            query = query.where(Candidate.source == source_filter)
        
        # Get total count
        count_query = select(func.count()).select_from(query.subquery())
        total = s.scalar(count_query) or 0
        
        # Paginate
        offset = (page - 1) * per_page
        query = query.offset(offset).limit(per_page)
        
        candidates = s.scalars(query).all()
        
        # Get stats
        total_candidates = s.scalar(select(func.count(Candidate.id))) or 0
        verified_count = s.scalar(select(func.count(Candidate.id)).where(Candidate.email_verified == True)) or 0
        portal_signups = s.scalar(select(func.count(Candidate.id)).where(Candidate.source == "portal")) or 0
        recent_signups = s.scalar(
            select(func.count(Candidate.id)).where(
                Candidate.created_at >= datetime.datetime.utcnow() - datetime.timedelta(days=7)
            )
        ) or 0
    
    total_pages = (total + per_page - 1) // per_page
    
    return render_template(
        "admin_portal_users.html",
        candidates=candidates,
        search=search,
        status_filter=status_filter,
        source_filter=source_filter,
        page=page,
        total_pages=total_pages,
        total=total,
        total_candidates=total_candidates,
        verified_count=verified_count,
        portal_signups=portal_signups,
        recent_signups=recent_signups
    )

@login_required
@app.route("/admin/portal-users/<int:cand_id>")
def admin_portal_user_detail(cand_id: int):
    """View/edit a portal user's details."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found", "danger")
            return redirect(url_for("admin_portal_users"))
        
        # Get applications for this candidate (eagerly load job relationship)
        applications = s.scalars(
            select(Application)
            .options(selectinload(Application.job))
            .where(Application.candidate_id == cand_id)
            .order_by(Application.created_at.desc())
        ).all()
        
        # Get documents
        documents = s.scalars(
            select(Document)
            .where(Document.candidate_id == cand_id)
            .order_by(Document.uploaded_at.desc())
        ).all()
        
        # Force load all attributes before session closes
        _ = cand.name, cand.email, cand.phone, cand.status, cand.source
        _ = cand.email_verified, cand.email_verified_at, cand.created_at
        _ = cand.last_login_at, cand.last_activity_at
        if hasattr(cand, 'about'):
            _ = cand.about
        
        for app in applications:
            if app.job:
                _ = app.job.title
        
        for doc in documents:
            _ = doc.filename, doc.original_name, doc.doc_type, doc.uploaded_at
    
    return render_template(
        "admin_portal_user_detail.html",
        cand=cand,
        applications=applications,
        documents=documents
    )

@login_required
@app.route("/admin/portal-users/<int:cand_id>/verify", methods=["POST"])
def admin_portal_user_verify(cand_id: int):
    """Manually verify a portal user's email."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found", "danger")
            return redirect(url_for("admin_portal_users"))
        
        cand.email_verified = True
        cand.email_verified_at = datetime.datetime.utcnow()
        s.commit()
        
        flash(f"Email verified for {cand.name}", "success")
    
    return redirect(url_for("admin_portal_user_detail", cand_id=cand_id))

@login_required
@app.route("/admin/portal-users/<int:cand_id>/send-magic-link", methods=["POST"])
def admin_portal_user_send_magic_link(cand_id: int):
    """Send a magic link to a portal user."""
    from itsdangerous import URLSafeTimedSerializer
    
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found", "danger")
            return redirect(url_for("admin_portal_users"))
        
        if not cand.email:
            flash("Candidate has no email address", "danger")
            return redirect(url_for("admin_portal_user_detail", cand_id=cand_id))
        
        # Generate magic link
        signer = URLSafeTimedSerializer(app.config["SECRET_KEY"], salt="portal-magic-link")
        token = signer.dumps({
            "email": cand.email,
            "name": cand.name,
            "next": "/jobs",
            "is_signup": False
        })
        
        verify_url = f"{APP_BASE_URL}/auth/verify?token={token}"
        
        html_body = f"""
        <div style="font-family: 'Inter', Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
            <div style="background: #0a1628; padding: 30px; border-radius: 12px 12px 0 0; text-align: center;">
                <h1 style="color: #00d4ff; margin: 0; font-size: 24px;">Optimus Solutions</h1>
                <p style="color: #94a3b8; margin: 10px 0 0; font-size: 14px;">Careers Portal</p>
            </div>
            <div style="background: #ffffff; padding: 40px 30px; border: 1px solid #e5e7eb; border-top: none; border-radius: 0 0 12px 12px;">
                <h2 style="color: #0a1628; margin: 0 0 20px; font-size: 20px;">Sign In Request</h2>
                <p style="color: #334155; margin: 0 0 20px; line-height: 1.6;">
                    An admin has sent you a sign-in link for the Optimus Solutions Careers Portal.
                </p>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="{verify_url}" style="background: #0066cc; color: #ffffff; padding: 14px 28px; border-radius: 8px; text-decoration: none; font-weight: 600; display: inline-block;">
                        Sign In Securely
                    </a>
                </div>
                <p style="color: #64748b; font-size: 14px; margin: 20px 0 0;">
                    This link will expire in 30 minutes.
                </p>
            </div>
        </div>
        """
        
        try:
            send_email(
                to_email=cand.email,
                subject="Sign in to Optimus Solutions Careers",
                html_body=html_body
            )
            flash(f"Magic link sent to {cand.email}", "success")
        except Exception as e:
            flash(f"Failed to send email: {e}", "danger")
    
    return redirect(url_for("admin_portal_user_detail", cand_id=cand_id))

@login_required
@app.route("/admin/portal-users/<int:cand_id>/update", methods=["POST"])
def admin_portal_user_update(cand_id: int):
    """Update a portal user's details."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found", "danger")
            return redirect(url_for("admin_portal_users"))
        
        # Update fields
        cand.name = request.form.get("name", cand.name)
        cand.email = request.form.get("email", cand.email)
        cand.phone = request.form.get("phone", cand.phone)
        cand.status = request.form.get("status", cand.status)
        cand.source = request.form.get("source", cand.source)
        
        if hasattr(cand, "about"):
            cand.about = request.form.get("about", cand.about)
        
        s.commit()
        flash("Portal user updated", "success")
    
    return redirect(url_for("admin_portal_user_detail", cand_id=cand_id))

@login_required
@app.route("/admin/portal-users/<int:cand_id>/delete", methods=["POST"])
def admin_portal_user_delete(cand_id: int):
    """Delete a portal user."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found", "danger")
            return redirect(url_for("admin_portal_users"))
        
        name = cand.name
        s.delete(cand)
        s.commit()
        flash(f"Portal user {name} deleted", "success")
    
    return redirect(url_for("admin_portal_users"))

@login_required
@app.route("/change-password", methods=["GET", "POST"])
def change_password():
    """Allow users to change their password"""
    if request.method == "POST":
        current_password = request.form.get("current_password", "")
        new_password = request.form.get("new_password", "")
        confirm_password = request.form.get("confirm_password", "")
        
        if not current_password or not new_password or not confirm_password:
            flash("All fields are required", "error")
            return render_template("change_password.html")
        
        if new_password != confirm_password:
            flash("New passwords do not match", "error")
            return render_template("change_password.html")
        
        # Validate password strength
        is_valid, message, score = validate_password_strength(new_password, current_user.email)
        if not is_valid:
            flash(message, "error")
            return render_template("change_password.html")
        
        with Session(engine) as s:
            user = s.get(User, current_user.id)
            if not user or not check_password_hash(user.password_hash, current_password):
                flash("Current password is incorrect", "error")
                return render_template("change_password.html")
            
            # Check password history (prevent reuse of last 5 passwords)
            if not check_password_history(current_user.id, new_password, history_limit=5):
                flash("You cannot reuse any of your last 5 passwords. Please choose a different password.", "error")
                return render_template("change_password.html")
            
            # Save old password to history before changing
            save_password_to_history(current_user.id, user.password_hash)
            
            # Update password
            user.password_hash = generate_password_hash(new_password, method='pbkdf2:sha256')
            s.commit()
            
            # Log password change
            log_audit_event('update', 'user_mgmt', 'Changed password', 'user', current_user.id)
            
            flash("Password changed successfully", "success")
            return redirect(url_for('index'))
    
    return render_template("change_password.html")

# ========== Two-Factor Authentication (2FA) Routes ==========

@app.route("/security/2fa/setup", methods=["GET", "POST"])
@login_required
def setup_2fa():
    """Setup Two-Factor Authentication"""
    if request.method == "POST":
        action = request.form.get("action")
        
        if action == "generate":
            # Generate new TOTP secret
            with Session(engine) as s:
                user = s.get(User, current_user.id)
                user.totp_secret = user.generate_totp_secret()
                s.commit()
                
                try:
                    # Generate provisioning URI
                    import pyotp
                    import qrcode
                    import io
                    import base64
                    
                    totp = pyotp.TOTP(user.totp_secret)
                    provisioning_uri = totp.provisioning_uri(
                        name=user.email,
                        issuer_name="Optimus OS1"
                    )
                    
                    # Generate QR code locally
                    qr = qrcode.QRCode(version=1, box_size=10, border=4)
                    qr.add_data(provisioning_uri)
                    qr.make(fit=True)
                    qr_img = qr.make_image(fill_color="black", back_color="white")
                    
                    # Convert to base64
                    buffer = io.BytesIO()
                    qr_img.save(buffer, format='PNG')
                    qr_base64 = base64.b64encode(buffer.getvalue()).decode()
                    qr_url = f"data:image/png;base64,{qr_base64}"
                    
                    return render_template("setup_2fa.html", 
                                         qr_code=None,
                                         qr_url=qr_url,
                                         secret=user.totp_secret,
                                         step="verify")
                    
                except Exception as e:
                    import traceback
                    error_details = traceback.format_exc()
                    current_app.logger.error(f"2FA setup failed: {str(e)}\n{error_details}")
                    flash("Error setting up 2FA. Please use manual entry.", "error")
                    
                    # Return with manual entry only
                    return render_template("setup_2fa.html", 
                                         qr_code=None,
                                         qr_url=None,
                                         secret=user.totp_secret,
                                         step="verify")
        
        elif action == "verify":
            token = request.form.get("token", "").strip()
            
            with Session(engine) as s:
                user = s.get(User, current_user.id)
                
                if user.verify_totp(token):
                    # Enable 2FA
                    user.totp_enabled = True
                    
                    # Generate backup codes
                    backup_codes = user.generate_backup_codes()
                    s.commit()
                    
                    log_audit_event('update', 'security', '2FA enabled', 'user', current_user.id)
                    
                    return render_template("setup_2fa.html",
                                         step="complete",
                                         backup_codes=backup_codes)
                else:
                    flash("Invalid verification code. Please try again.", "error")
                    return render_template("setup_2fa.html", step="start")
    
    # GET request - show initial setup page
    with Session(engine) as s:
        user = s.get(User, current_user.id)
        if user.totp_enabled:
            return render_template("setup_2fa.html", step="already_enabled")
    
    return render_template("setup_2fa.html", step="start")

@app.route("/security/2fa/disable", methods=["POST"])
@login_required
def disable_2fa():
    """Disable Two-Factor Authentication"""
    password = request.form.get("password", "")
    
    with Session(engine) as s:
        user = s.get(User, current_user.id)
        
        # Verify password before disabling
        if not user or not check_password_hash(user.password_hash, password):
            flash("Incorrect password", "error")
            return redirect(url_for('setup_2fa'))
        
        user.totp_enabled = False
        user.totp_secret = None
        user.backup_codes = None
        s.commit()
        
        log_audit_event('update', 'security', '2FA disabled', 'user', current_user.id, status='warning')
        
        flash("Two-Factor Authentication has been disabled", "success")
        return redirect(url_for('index'))

@app.route("/security/2fa/verify", methods=["GET", "POST"])
def verify_2fa():
    """Verify 2FA token during login"""
    # This should be called after password verification but before login_user()
    # User ID should be in session from the initial login attempt
    
    if 'pending_user_id' not in session:
        return redirect(url_for('login'))
    
    user_email = session.get('pending_user_email', '')
    
    if request.method == "POST":
        token = request.form.get("token", "").strip()
        use_backup = request.form.get("use_backup") == "1"
        
        with Session(engine) as s:
            user = s.get(User, session['pending_user_id'])
            
            if not user:
                session.pop('pending_user_id', None)
                session.pop('pending_user_email', None)
                flash("Session expired. Please login again.", "error")
                return redirect(url_for('login'))
            
            verified = False
            if use_backup:
                verified = user.verify_backup_code(token)
                if verified:
                    s.commit()  # Save updated backup codes
                    log_audit_event('login', 'auth', 'Logged in with backup code', 
                                  'user', user.id, status='warning')
            else:
                verified = user.verify_totp(token)
            
            if verified:
                # 2FA successful - complete login
                session.pop('pending_user_id', None)
                session.pop('pending_user_email', None)
                
                remember = session.get('remember_me', False)
                session.pop('remember_me', None)
                
                # Refresh user to ensure all attributes are loaded before login_user
                s.refresh(user)
                
                if remember:
                    login_user(user, remember=True, duration=timedelta(days=30))
                else:
                    login_user(user, remember=False)
                    session.permanent = True
                
                log_audit_event('login', 'auth', 'Successful 2FA verification', 'user', user.id)
                
                next_page = session.get('next_page', '/')
                session.pop('next_page', None)
                return redirect(next_page)
            else:
                flash("Invalid code. Please try again.", "error")
    
    return render_template("verify_2fa.html", user_email=user_email)

@app.route("/security/2fa/mandatory-setup", methods=["GET", "POST"])
def mandatory_2fa_setup():
    """
    Mandatory 2FA setup for users who haven't enabled it yet.
    Called after successful password verification but before granting access.
    """
    # Check if user is in the setup flow
    if 'setup_2fa_user_id' not in session:
        return redirect(url_for('login'))
    
    user_id = session['setup_2fa_user_id']
    
    if request.method == "POST":
        action = request.form.get("action")
        
        if action == "generate":
            # Generate new TOTP secret
            with Session(engine) as s:
                user = s.get(User, user_id)
                if not user:
                    session.pop('setup_2fa_user_id', None)
                    flash("Session expired. Please login again.", "error")
                    return redirect(url_for('login'))
                
                # Generate and store the secret
                totp_secret = user.generate_totp_secret()
                user.totp_secret = totp_secret
                user_email = user.email
                s.commit()
            
            # Store the secret in session for QR code endpoint
            session['qr_totp_secret'] = totp_secret
            session['qr_user_email'] = user_email
            
            app.logger.info(f"TOTP secret generated for {user_email}, redirecting to verify step")
            
            # Use a dedicated endpoint for the QR code image
            qr_url = url_for('get_2fa_qr_code')
            
            return render_template("mandatory_2fa_setup.html", 
                                 qr_url=qr_url,
                                 secret=totp_secret,
                                 step="verify",
                                 user_email=user_email)
        
        elif action == "verify":
            token = request.form.get("token", "").strip()
            
            with Session(engine) as s:
                user = s.get(User, user_id)
                
                if not user:
                    session.pop('setup_2fa_user_id', None)
                    flash("Session expired. Please login again.", "error")
                    return redirect(url_for('login'))
                
                # Store values before any session operations
                user_email = user.email
                user_totp_secret = user.totp_secret
                
                if user.verify_totp(token):
                    # Enable 2FA
                    user.totp_enabled = True
                    backup_codes = user.generate_backup_codes()
                    s.commit()
                    
                    remember = session.get('setup_2fa_remember', False)
                    next_page = session.get('setup_2fa_next', '/')
                    
                    # Clear setup session data
                    session.pop('setup_2fa_user_id', None)
                    session.pop('setup_2fa_remember', None)
                    session.pop('setup_2fa_next', None)
                    session.pop('qr_totp_secret', None)
                    session.pop('qr_user_email', None)
                    
                    # Log in the user while still in session context
                    # Make sure user object has all attributes loaded
                    s.refresh(user)
                    if remember:
                        login_user(user, remember=True, duration=timedelta(days=30))
                    else:
                        login_user(user, remember=False)
                        session.permanent = True
                    
                    log_audit_event('login', 'auth', f'2FA setup completed and logged in: {user_email}', status='success')
                    
                    # Show backup codes before redirecting
                    return render_template("mandatory_2fa_setup.html",
                                         step="complete",
                                         backup_codes=backup_codes,
                                         next_page=next_page)
                else:
                    flash("Invalid verification code. Please try again.", "error")
                    # Regenerate QR URL for retry
                    qr_url = url_for('get_2fa_qr_code')
                    return render_template("mandatory_2fa_setup.html", 
                                         step="verify",
                                         secret=user_totp_secret,
                                         qr_url=qr_url,
                                         user_email=user_email)
    
    # GET request - show initial setup page
    with Session(engine) as s:
        user = s.get(User, user_id)
        if not user:
            session.pop('setup_2fa_user_id', None)
            return redirect(url_for('login'))
        user_email = user.email
    
    return render_template("mandatory_2fa_setup.html", step="start", user_email=user_email)

@app.route("/security/2fa/qr-code.png")
def get_2fa_qr_code():
    """
    Generate and serve the 2FA QR code as a PNG image.
    This endpoint is used during 2FA setup to display the QR code.
    """
    import pyotp
    import qrcode
    import io
    from flask import Response
    
    # Get the TOTP secret from session
    totp_secret = session.get('qr_totp_secret')
    user_email = session.get('qr_user_email')
    
    if not totp_secret or not user_email:
        app.logger.error("No TOTP secret or email in session for QR code generation")
        # Return a 1x1 transparent PNG as fallback
        return Response(b'', mimetype='image/png', status=404)
    
    try:
        totp = pyotp.TOTP(totp_secret)
        provisioning_uri = totp.provisioning_uri(
            name=user_email,
            issuer_name="Optimus OS1"
        )
        
        # Generate QR code
        qr = qrcode.QRCode(version=1, box_size=10, border=4)
        qr.add_data(provisioning_uri)
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="black", back_color="white")
        
        # Convert to bytes
        buffer = io.BytesIO()
        qr_img.save(buffer, format='PNG')
        buffer.seek(0)
        
        app.logger.info(f"QR code image served for {user_email}")
        
        return Response(buffer.getvalue(), mimetype='image/png')
    except Exception as e:
        app.logger.error(f"QR code image generation failed: {str(e)}")
        return Response(b'', mimetype='image/png', status=500)

def _signer() -> URLSafeTimedSerializer:
    # used for candidate magic links
    return URLSafeTimedSerializer(app.config["SECRET_KEY"], salt="candidate-login")

def get_role_category_names(session: Session) -> List[str]:
    """
    Return role names from taxonomy_categories where type='role',
    ordered by name. Falls back to the hard-coded ROLE_TYPES if empty.
    """
    rows = session.scalars(
        select(TaxonomyCategory)
        .where(TaxonomyCategory.type == "role")
        .order_by(TaxonomyCategory.name.asc())
    ).all()
    names = [r.name for r in rows]
    # Safe fallback for fresh installs (keeps app usable if taxonomy is empty)
    return names or list(ROLE_TYPES)

# -------------- DB Models --------------
from sqlalchemy import String, Integer, DateTime, Boolean

# --- Taxonomy models (subject groups/tags) ---
# --- User model for authentication ---
class User(Base, UserMixin):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, autoincrement=True)
    name = Column(String(200), nullable=False)
    email = Column(String(200), unique=True, nullable=False, index=True)
    # Map to actual database column name
    password_hash = Column('pw_hash', String(255), nullable=False)
    created_at = Column(DateTime, nullable=True)
    
    # Security columns - NOW ENABLED (migration completed)
    role = Column(String(50), default='employee', nullable=True)
    is_active = Column(Boolean, default=True, nullable=True)
    last_login = Column(DateTime, nullable=True)
    failed_login_attempts = Column(Integer, default=0, nullable=True)
    locked_until = Column(DateTime, nullable=True)
    
    # 2FA/MFA columns
    totp_secret = Column(String(32), nullable=True)
    totp_enabled = Column(Boolean, default=False, nullable=True)
    backup_codes = Column(Text, nullable=True)  # JSON array of backup codes
    
    # Session security columns
    session_token = Column(String(255), nullable=True)
    last_ip = Column(String(45), nullable=True)
    last_user_agent = Column(Text, nullable=True)
    
    # Magic link columns for passwordless login/setup
    magic_token = Column(String(255), nullable=True)
    magic_token_expires = Column(DateTime, nullable=True)
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password, method='pbkdf2:sha256')
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    def is_locked(self):
        """Check if account is currently locked"""
        if self.locked_until is None:
            return False
        return self.locked_until > datetime.datetime.utcnow()
    
    def generate_totp_secret(self):
        """Generate a new TOTP secret for 2FA"""
        import pyotp
        return pyotp.random_base32()
    
    def verify_totp(self, token):
        """Verify a TOTP token"""
        if not self.totp_secret:
            return False
        import pyotp
        totp = pyotp.TOTP(self.totp_secret)
        return totp.verify(token, valid_window=1)  # Allow 1 window before/after
    
    def generate_backup_codes(self, count=10):
        """Generate backup codes for account recovery"""
        import secrets
        codes = [secrets.token_hex(4).upper() for _ in range(count)]
        self.backup_codes = json.dumps(codes)
        return codes
    
    def verify_backup_code(self, code):
        """Verify and consume a backup code"""
        if not self.backup_codes:
            return False
        codes = json.loads(self.backup_codes)
        code = code.upper().replace('-', '')
        if code in codes:
            codes.remove(code)
            self.backup_codes = json.dumps(codes)
            return True
        return False

class PasswordHistory(Base):
    """Track password history to prevent reuse"""
    __tablename__ = "password_history"
    
    id = Column(Integer, primary_key=True, autoincrement=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False, index=True)
    password_hash = Column(String(255), nullable=False)
    created_at = Column(DateTime, default=datetime.datetime.utcnow, nullable=False, index=True)

class AuditLog(Base):
    """Comprehensive audit logging for security and compliance"""
    __tablename__ = "audit_logs"
    
    id = Column(Integer, primary_key=True, autoincrement=True)
    timestamp = Column(DateTime, default=datetime.datetime.utcnow, nullable=False, index=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=True, index=True)
    user_email = Column(String(255), nullable=True, index=True)
    event_type = Column(String(50), nullable=False, index=True)  # login, logout, create, update, delete, view, export
    event_category = Column(String(50), nullable=False, index=True)  # auth, user_mgmt, data_access, security
    resource_type = Column(String(50), nullable=True)  # candidate, job, engagement, user
    resource_id = Column(Integer, nullable=True)
    action = Column(String(255), nullable=False)
    details = Column(Text, nullable=True)  # JSON string with additional info
    ip_address = Column(String(45), nullable=True)
    user_agent = Column(String(500), nullable=True)
    status = Column(String(20), nullable=False, default='success')  # success, failure, warning

class CandidateTag(Base):
    __tablename__ = "candidate_tags"
    id = Column(Integer, primary_key=True, autoincrement=True)
    candidate_id = Column(ForeignKey("candidates.id"), index=True, nullable=False)
    tag_id = Column(ForeignKey("taxonomy_tags.id"), index=True, nullable=False)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

# --- DB Models ---
class Engagement(Base):
    __tablename__ = "engagements"
    id = Column(Integer, primary_key=True)
    name = Column(String)
    client = Column(String)
    status = Column(String, default="Active")
    ref = Column(String, unique=True)
    start_date = Column(DateTime, nullable=True)
    end_date = Column(DateTime, nullable=True)
    sow_signed_at = Column(DateTime)
    description = Column(Text, default="")
    opportunity_id = Column(Integer, ForeignKey("opportunities.id"), nullable=True, unique=True)

    # <--- add this
    plan_version = Column(Integer, default=1)

    opportunity = relationship(
        "Opportunity",
        backref=backref("engagement", uselist=False)
    )
    jobs = relationship("Job", back_populates="engagement", cascade="all, delete-orphan")

class Job(Base):
    __tablename__ = "jobs"
    id = Column(Integer, primary_key=True, autoincrement=True)
    engagement_id = Column(ForeignKey("engagements.id"))
    title = Column(String(200))
    description = Column(String(4000))
    role_type = Column(String(50), default="")
    location = Column(String(200), default="")
    salary_range = Column(String(200), default="")
    status = Column(String(50), default="Open")
    public_token = Column(String(36), unique=True, default=lambda: str(uuid.uuid4()))
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

    engagement = relationship("Engagement", back_populates="jobs")
    applications = relationship("Application", back_populates="job", cascade="all, delete-orphan")

class RoleAliasForm(FlaskForm):
    canonical = SelectField("Canonical role", choices=[(r, r) for r in ROLE_TYPES], validators=[DataRequired()])
    alias = StringField("Alias (case-insensitive)", validators=[DataRequired()])

class WorkerUser(UserMixin):
    """
    Lightweight adapter around your Users table row to satisfy Flask-Login.
    We'll just carry id + email + name in memory.
    """
    def __init__(self, row):
        self.id = str(row.id)
        self.email = row.email
        self.name = row.name

# NOTE: user_loader is defined at line 274, this duplicate has been removed

def sign_request(method, path, body=None):
    ts = str(int(time.time()))
    payload = (ts + method.upper() + path + (body or "")).encode()
    signature = hmac.new(SUMSUB_SECRET_KEY.encode(), payload, hashlib.sha256).hexdigest()
    return {"X-App-Token": SUMSUB_APP_TOKEN, "X-App-Access-Ts": ts, "X-App-Access-Sig": signature}

# --- Role/session helpers ---

def worker_required(fn):
    """Guard routes for internal workers (Flask-Login)."""
    from functools import wraps
    @wraps(fn)
    def _wrap(*a, **kw):
        if not current_user.is_authenticated:
            return redirect(url_for("login", next=request.path))
        return fn(*a, **kw)
    return _wrap

def candidate_required(fn):
    """Guard routes for candidate self-service (magic-link)."""
    from functools import wraps
    @wraps(fn)
    def _wrap(*a, **kw):
        if not session.get("candidate_id"):
            return redirect(url_for("candidate_login", next=request.path))
        return fn(*a, **kw)
    return _wrap

class WorkerLoginForm(FlaskForm):
    email = StringField("Email", validators=[DataRequired(), Email()])
    password = StringField("Password", validators=[DataRequired()])
    submit = SubmitField("Sign in")

class WorkerSignupForm(FlaskForm):
    name = StringField("Name", validators=[DataRequired()])
    email = StringField("Email", validators=[DataRequired(), Email()])
    password = StringField("Password", validators=[DataRequired()])
    submit = SubmitField("Create account")

class CandidateLoginForm(FlaskForm):
    email = StringField("Email", validators=[DataRequired(), Email()])
    submit = SubmitField("Send me a login link")

class Timesheet(Base):
    __tablename__ = "timesheets"
    id = Column(Integer, primary_key=True, autoincrement=True)
    user_id = Column(Integer, nullable=False)
    engagement_id = Column(Integer, nullable=False)
    period_start = Column(DateTime, nullable=False)
    period_end = Column(DateTime, nullable=False)
    hours = Column(Integer, default=0)
    notes = Column(String(4000), default="")
    status = Column(String(50), default="Draft")
    submitted_at = Column(DateTime, nullable=True)

class Invoice(Base):
    """Invoice model for client billing"""
    __tablename__ = "invoices"
    id = Column(Integer, primary_key=True, autoincrement=True)
    invoice_number = Column(String(50), unique=True, nullable=False)
    engagement_id = Column(Integer, ForeignKey("engagements.id"), nullable=True)
    client_name = Column(String(200), nullable=False)
    engagement_name = Column(String(200), nullable=True)
    
    # Dates
    invoice_date = Column(DateTime, default=datetime.datetime.utcnow)
    due_date = Column(DateTime, nullable=True)
    paid_date = Column(DateTime, nullable=True)
    
    # Financial
    subtotal = Column(Float, default=0.0)
    vat_rate = Column(Float, default=20.0)  # VAT percentage
    vat_amount = Column(Float, default=0.0)
    total_amount = Column(Float, default=0.0)
    
    # Status: Draft, Pending, Paid, Overdue, Cancelled
    status = Column(String(50), default="Draft")
    
    # Details
    line_items = Column(Text, default="[]")  # JSON array of line items
    notes = Column(Text, default="")
    payment_terms = Column(String(200), default="Net 30")
    
    # Tracking
    created_by = Column(Integer, ForeignKey("users.id"), nullable=True)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)
    sent_at = Column(DateTime, nullable=True)
    
    # Relationship
    engagement = relationship("Engagement", backref="invoices")

class Candidate(Base):
    __tablename__ = "candidates"
    id = Column(Integer, primary_key=True, autoincrement=True)
    name = Column(String(200))
    email = Column(String(200))
    phone = Column(String(50), default="")
    skills = Column(String(2000), default="")
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

    # card fields…
    onboarded_at = Column(DateTime, nullable=True)
    esign_status = Column(String(50), default=None)
    trustid_rtw_date = Column(DateTime, nullable=True)
    trustid_idv_date = Column(DateTime, nullable=True)
    trustid_dbs_date = Column(DateTime, nullable=True)
    updated_cv_requested_at = Column(DateTime, nullable=True)

    # NEW
    ai_summary = Column(String(6000), default="")
    
    # Resource Pool enhancement fields
    postcode = Column(String(20), nullable=True)  # UK postcode for radius search
    last_login_at = Column(DateTime, nullable=True)  # Last login to candidate portal
    last_activity_at = Column(DateTime, nullable=True)  # Last activity (application, profile update, etc.)
    optimus_interview_result = Column(String(50), nullable=True)  # Pass, Fail, Pending, Not Required
    optimus_assessment_result = Column(String(50), nullable=True)  # Pass, Fail, Pending, Not Required
    previously_vetted = Column(Boolean, default=False)  # Has been previously vetted
    min_day_rate = Column(Integer, nullable=True)  # Minimum acceptable day rate
    max_day_rate = Column(Integer, nullable=True)  # Maximum acceptable day rate
    
    # Status and availability fields per wireframe requirements
    status = Column(String(50), default="Available")  # Available, On Assignment, On Hold, Unavailable, etc.
    availability = Column(String(100), default="Immediately available")  # Immediately available, 2 weeks notice, etc.
    day_rate = Column(Integer, nullable=True)  # Day rate in pounds
    location = Column(String(200), nullable=True)  # Location/city
    clearance_level = Column(String(100), nullable=True)  # Security clearance level
    
    # GAP 3.1 & 3.3: Contract end/leaving tracking
    leaving_confirmed = Column(Boolean, default=False)  # Confirmed leaving when contract ends
    leaving_confirmed_at = Column(DateTime, nullable=True)  # When leaving was confirmed
    
    # Portal authentication fields
    email_verified = Column(Boolean, default=False)  # Email verified via magic link
    email_verified_at = Column(DateTime, nullable=True)  # When email was verified
    about = Column(Text, nullable=True)  # "Tell us about yourself" field from signup
    source = Column(String(100), default="portal")  # Where candidate came from (portal, manual, import)

    documents = relationship("Document", back_populates="candidate", cascade="all, delete-orphan")

class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, autoincrement=True)
    candidate_id = Column(ForeignKey("candidates.id"))
    doc_type = Column(String(50), default="cv")
    filename = Column(String(500))
    original_name = Column(String(500))
    uploaded_at = Column(DateTime, default=datetime.datetime.utcnow)
    candidate = relationship("Candidate", back_populates="documents")

class Application(Base):
    __tablename__ = "applications"
    id = Column(Integer, primary_key=True, autoincrement=True)
    job_id = Column(ForeignKey("jobs.id"))
    candidate_id = Column(ForeignKey("candidates.id"))
    cover_note = Column(String(4000), default="")
    status = Column(String(50), default="New")  # Pipeline status
    ai_score = Column(Integer, default=0)
    ai_summary = Column(String(6000), default="")
    ai_explanation = Column(String(8000), default="")  # <-- NEW: store "why this score"
    interview_scheduled_at = Column(DateTime, nullable=True)
    interview_completed_at = Column(DateTime, nullable=True)

    # ✅ make sure THIS is present:
    interview_notes = Column(Text, default="")

    onboarding_email_sent = Column(Boolean, default=False)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

    job = relationship("Job", back_populates="applications")
    candidate = relationship("Candidate")

class TrustIDCheck(Base):
    __tablename__ = "trustid_checks"
    id = Column(Integer, primary_key=True, autoincrement=True)
    application_id = Column(ForeignKey("applications.id"))
    rtw = Column(Boolean, default=False)
    idv = Column(Boolean, default=True)
    dbs = Column(Boolean, default=False)
    trustid_application_id = Column(String(100), default="")
    status = Column(String(50), default="Created")
    result_json = Column(String(20000), default="")

class Shortlist(Base):
    __tablename__ = "shortlists"
    id = Column(Integer, primary_key=True, autoincrement=True)
    job_id = Column(ForeignKey("jobs.id"), index=True, nullable=False)
    candidate_id = Column(ForeignKey("candidates.id"), index=True, nullable=False)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

class ESigRequest(Base):
    __tablename__ = "esign_requests"
    id = Column(Integer, primary_key=True, autoincrement=True)
    application_id = Column(ForeignKey("applications.id"))
    candidate_id = Column(Integer, ForeignKey("candidates.id"), nullable=True)
    engagement_id = Column(Integer, ForeignKey("engagements.id"), nullable=True)
    provider = Column(String(50), default="dropbox_sign")
    request_id = Column(String(100), default="")
    status = Column(String(50), default="Draft")  # Draft, Sent, Signed, Declined, Error
    sent_at = Column(DateTime, nullable=True)
    signed_at = Column(DateTime, nullable=True)
    created_at = Column(DateTime, nullable=True)
    end_date = Column(DateTime, nullable=True)

class WebhookEvent(Base):
    __tablename__ = "webhook_events"
    id = Column(Integer, primary_key=True, autoincrement=True)
    source = Column(String(50))
    event_type = Column(String(100))
    payload = Column(String(40000))
    received_at = Column(DateTime, default=datetime.datetime.utcnow)

# --- Taxonomy (DB-driven) ---
class TaxonomyCategory(Base):
    __tablename__ = "taxonomy_categories"
    id = Column(Integer, primary_key=True, autoincrement=True)
    # 'role' or 'subject'
    type = Column(String(20), nullable=False)
    name = Column(String(200), nullable=False, unique=False)

    tags = relationship("TaxonomyTag", back_populates="category", cascade="all, delete-orphan")

class TaxonomyTag(Base):
    __tablename__ = "taxonomy_tags"
    id = Column(Integer, primary_key=True, autoincrement=True)
    category_id = Column(ForeignKey("taxonomy_categories.id"), index=True, nullable=False)
    tag = Column(String(200), nullable=False)

    category = relationship("TaxonomyCategory", back_populates="tags")

# ---- Opportunities model ----
class Opportunity(Base):
    __tablename__ = "opportunities"
    id = Column(Integer, primary_key=True)
    name = Column(String)
    client = Column(String)
    stage = Column(String)
    owner = Column(String)
    est_start = Column(DateTime)
    est_value = Column(Integer)
    probability = Column(Integer, default=0)
    notes = Column(Text, default="")
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

    _engagement_id = Column(Integer)
    _engagement_ref = Column(String)

    # NEW client contact fields
    client_contact_name  = Column(String, default="")
    client_contact_role  = Column(String, default="")
    client_contact_phone = Column(String, default="")
    client_contact_email = Column(String, default="")

# ---- VettingCheck model (per wireframe requirements) ----
class VettingCheck(Base):
    """
    Pre-employment vetting check items for each candidate.
    12 check types as per wireframe:
    - Right to Work, Identity Verification, Address History, DBS Check,
    - Employment History, References, Qualifications, Professional Registration,
    - Credit Check, Directorship/Disqualification, Sanctions/PEP, Social Media Review
    """
    __tablename__ = "vetting_check"
    id = Column(Integer, primary_key=True)
    candidate_id = Column(Integer, ForeignKey("candidates.id"), nullable=False)
    check_type = Column(String(100), nullable=False)
    status = Column(String(50), default="NOT STARTED")  # NOT STARTED, In Progress, Complete, N/A, Failed
    notes = Column(Text, default="")
    completed_at = Column(DateTime)
    updated_at = Column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

# ---- CandidateNote model (Notes & Activity panel) ----
class CandidateNote(Base):
    """Notes and activity log for candidates"""
    __tablename__ = "candidate_notes"
    id = Column(Integer, primary_key=True)
    candidate_id = Column(Integer, ForeignKey("candidates.id"), nullable=False)
    user_email = Column(String)  # Who added the note
    note_type = Column(String(50), default="note")  # note, email, activity, system
    content = Column(Text, default="")
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

from public import public_bp
app.register_blueprint(public_bp)
app.secret_key = os.environ.get("FLASK_SECRET", "dev-secret")  # for session

# ---- Lightweight schema patches for existing SQLite DBs ----
def ensure_schema():
    with engine.begin() as conn:
        # ===== jobs table =====
        try:
            conn.execute(text("ALTER TABLE jobs ADD COLUMN role_type TEXT DEFAULT ''"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE jobs ADD COLUMN status TEXT DEFAULT 'Open'"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE jobs ADD COLUMN location TEXT DEFAULT ''"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE jobs ADD COLUMN salary_range TEXT DEFAULT ''"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE jobs ADD COLUMN public_token TEXT UNIQUE"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE jobs ADD COLUMN created_at DATETIME DEFAULT CURRENT_TIMESTAMP"))
        except Exception:
            pass
        try:
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_jobs_status ON jobs(status)"))
        except Exception:
            pass
        try:
            conn.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_jobs_public_token ON jobs(public_token)"))
        except Exception:
            pass

        # ===== applications table =====
        for coldef in [
            "interview_completed_at DATETIME",
            "interview_notes TEXT DEFAULT ''",
            "ai_explanation TEXT DEFAULT ''",
            "ai_score INTEGER DEFAULT 0",
            "ai_summary TEXT DEFAULT ''",
            "onboarding_email_sent BOOLEAN DEFAULT 0",
            "interview_scheduled_at DATETIME",
            "status TEXT DEFAULT 'New'",
            "created_at DATETIME DEFAULT CURRENT_TIMESTAMP",
        ]:
            try:
                conn.execute(text(f"ALTER TABLE applications ADD COLUMN {coldef}"))
            except Exception:
                pass

        # ===== candidates table =====
        try:
            conn.execute(text("ALTER TABLE candidates ADD COLUMN ai_summary TEXT DEFAULT ''"))
        except Exception:
            pass
        for coldef in [
            "onboarded_at DATETIME",
            "esign_status TEXT",
            "trustid_rtw_date DATETIME",
            "trustid_idv_date DATETIME",
            "trustid_dbs_date DATETIME",
            "updated_cv_requested_at DATETIME",
        ]:
            try:
                conn.execute(text(f"ALTER TABLE candidates ADD COLUMN {coldef}"))
            except Exception:
                pass
        try:
            conn.execute(text("ALTER TABLE candidates ADD COLUMN phone TEXT DEFAULT ''"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE candidates ADD COLUMN skills TEXT DEFAULT ''"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE candidates ADD COLUMN created_at DATETIME DEFAULT CURRENT_TIMESTAMP"))
        except Exception:
            pass
        
        # Resource Pool enhancement columns
        for coldef in [
            "postcode TEXT",
            "last_login_at DATETIME",
            "last_activity_at DATETIME",
            "optimus_interview_result TEXT",
            "optimus_assessment_result TEXT",
            "previously_vetted BOOLEAN DEFAULT 0",
            "min_day_rate INTEGER",
            "max_day_rate INTEGER",
        ]:
            try:
                conn.execute(text(f"ALTER TABLE candidates ADD COLUMN {coldef}"))
            except Exception:
                pass
        
        for stmt in [
            "CREATE INDEX IF NOT EXISTS idx_candidates_created_at ON candidates(created_at)",
            "CREATE INDEX IF NOT EXISTS idx_candidates_email ON candidates(email)",
            "CREATE INDEX IF NOT EXISTS idx_candidates_name ON candidates(name)",
        ]:
            try:
                conn.execute(text(stmt))
            except Exception:
                pass

        # ===== documents table =====
        for stmt in [
            "CREATE INDEX IF NOT EXISTS idx_documents_candidate ON documents(candidate_id)",
            "CREATE INDEX IF NOT EXISTS idx_documents_uploaded_at ON documents(uploaded_at)",
        ]:
            try:
                conn.execute(text(stmt))
            except Exception:
                pass

        # ===== shortlists table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS shortlists (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              job_id INTEGER NOT NULL,
              candidate_id INTEGER NOT NULL,
              created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """))
        except Exception:
            pass
        for stmt in [
            "CREATE INDEX IF NOT EXISTS idx_shortlists_job ON shortlists(job_id)",
            "CREATE INDEX IF NOT EXISTS idx_shortlists_candidate ON shortlists(candidate_id)",
        ]:
            try:
                conn.execute(text(stmt))
            except Exception:
                pass

        # ===== taxonomy tables =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS taxonomy_categories (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              type TEXT NOT NULL,
              name TEXT NOT NULL
            )
            """))
        except Exception:
            pass
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS taxonomy_tags (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              category_id INTEGER NOT NULL REFERENCES taxonomy_categories(id) ON DELETE CASCADE,
              tag TEXT NOT NULL
            )
            """))
        except Exception:
            pass
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS candidate_tags (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              candidate_id INTEGER NOT NULL,
              tag_id INTEGER NOT NULL,
              created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """))
        except Exception:
            pass
        for stmt in [
            "CREATE INDEX IF NOT EXISTS idx_tax_cat_type ON taxonomy_categories(type)",
            "CREATE INDEX IF NOT EXISTS idx_tax_tags_cat ON taxonomy_tags(category_id)",
            "CREATE INDEX IF NOT EXISTS idx_candidate_tags_cand ON candidate_tags(candidate_id)",
        ]:
            try:
                conn.execute(text(stmt))
            except Exception:
                pass

        # ===== role_aliases table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS role_aliases (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              canonical TEXT NOT NULL,
              alias TEXT NOT NULL
            )
            """))
        except Exception:
            pass
        try:
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_role_alias_canon ON role_aliases(canonical)"))
        except Exception:
            pass
        try:
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_role_alias_alias ON role_aliases(alias)"))
        except Exception:
            pass

        # ===== users table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS users (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              name TEXT,
              email TEXT UNIQUE,
              pw_hash TEXT,
              created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """))
        except Exception:
            pass
        try:
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_users_email ON users(email)"))
        except Exception:
            pass

        # ===== timesheets table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS timesheets (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              user_id INTEGER NOT NULL,
              engagement_id INTEGER NOT NULL,
              period_start DATETIME NOT NULL,
              period_end DATETIME NOT NULL,
              hours INTEGER DEFAULT 0,
              notes TEXT,
              status TEXT DEFAULT 'Draft',
              submitted_at DATETIME
            )
            """))
        except Exception:
            pass
        for stmt in [
            "CREATE INDEX IF NOT EXISTS idx_timesheets_eng_period ON timesheets(engagement_id, period_start)",
            "CREATE INDEX IF NOT EXISTS idx_timesheets_user_period ON timesheets(user_id, period_start)",
        ]:
            try:
                conn.execute(text(stmt))
            except Exception:
                pass

        # ===== engagements table =====
        try:
            conn.execute(text("ALTER TABLE engagements ADD COLUMN ref TEXT"))
        except Exception:
            pass
        try:
            conn.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_engagements_ref ON engagements(ref)"))
        except Exception:
            pass

        for coldef in [
            "sow_signed_at DATETIME",
            "description TEXT DEFAULT ''",
            "start_date DATETIME",
            "end_date DATETIME",
            "opportunity_id INTEGER UNIQUE",
            "plan_version INTEGER DEFAULT 1"
        ]:
            try:
                conn.execute(text(f"ALTER TABLE engagements ADD COLUMN {coldef}"))
            except Exception:
                pass

        try:
            conn.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_engagements_opp ON engagements(opportunity_id)"))
        except Exception:
            pass

        # ===== opportunities table =====
        for coldef in [
            "_engagement_id INTEGER",
            "_engagement_ref TEXT",
            "created_at DATETIME DEFAULT CURRENT_TIMESTAMP",
            "probability INTEGER DEFAULT 0",
            "notes TEXT DEFAULT ''",
            "client_contact_name TEXT",
            "client_contact_role TEXT",
            "client_contact_phone TEXT",
            "client_contact_email TEXT"
        ]:
            try:
                conn.execute(text(f"ALTER TABLE opportunities ADD COLUMN {coldef}"))
            except Exception:
                pass
        try:
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_opps_engid ON opportunities(_engagement_id)"))
        except Exception:
            pass

        # ===== engagement_plans table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS engagement_plans (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              engagement_id INTEGER,
              role_type TEXT DEFAULT '',
              planned_count INTEGER DEFAULT 0,
              pay_rate INTEGER DEFAULT 0,
              charge_rate INTEGER DEFAULT 0,
              -- legacy 'rate' kept for backwards compat; mirrors charge_rate
              rate INTEGER DEFAULT 0,
              version_int INTEGER DEFAULT 1
            )
            """))
        except Exception:
            pass

        # backfill new columns if table pre-existed
        for coldef in [
            "pay_rate INTEGER DEFAULT 0",
            "charge_rate INTEGER DEFAULT 0",
            "version_int INTEGER DEFAULT 1",
            "rate INTEGER DEFAULT 0"
        ]:
            try:
                conn.execute(text(f"ALTER TABLE engagement_plans ADD COLUMN {coldef}"))
            except Exception:
                pass

        # ===== trustid_checks table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS trustid_checks (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              application_id INTEGER,
              rtw BOOLEAN DEFAULT 0,
              idv BOOLEAN DEFAULT 1,
              dbs BOOLEAN DEFAULT 0,
              trustid_application_id TEXT DEFAULT '',
              status TEXT DEFAULT 'Created',
              result_json TEXT DEFAULT ''
            )
            """))
        except Exception:
            pass

        # ===== esign_requests table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS esign_requests (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              application_id INTEGER,
              provider TEXT DEFAULT 'dropbox_sign',
              request_id TEXT DEFAULT '',
              status TEXT DEFAULT 'Draft',
              sent_at DATETIME,
              signed_at DATETIME
            )
            """))
        except Exception:
            pass

        # ===== webhook_events table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS webhook_events (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              source TEXT,
              event_type TEXT,
              payload TEXT,
              received_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """))
        except Exception:
            pass

        # ===== role_types table =====
        try:
            conn.execute(text("""
            CREATE TABLE IF NOT EXISTS role_types (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              name TEXT UNIQUE NOT NULL,
              default_rate INTEGER DEFAULT 0
            )
            """))
        except Exception:
            pass

        # ===== Helpful engagement ref backfill =====
        try:
            existing_refs = [r[0] for r in conn.execute(
                text("SELECT ref FROM engagements WHERE ref IS NOT NULL AND ref <> ''")
            ).all()]
            import re as _re
            prefix = "EG"
            width = 3
            max_n = 0
            for ref in existing_refs:
                m = _re.fullmatch(rf"{_re.escape(prefix)}(\d+)", ref or "")
                if m:
                    try:
                        max_n = max(max_n, int(m.group(1)))
                    except Exception:
                        pass
            rows_to_fill = conn.execute(
                text("SELECT id FROM engagements WHERE ref IS NULL OR ref = '' ORDER BY id ASC")
            ).all()
            for (eid,) in rows_to_fill:
                max_n += 1
                new_ref = f"{prefix}{str(max_n).zfill(width)}"
                conn.execute(text("UPDATE engagements SET ref=:r WHERE id=:i"), {"r": new_ref, "i": eid})
        except Exception:
            pass

        # ===== Seed taxonomy & role_aliases if empty =====
        try:
            existing = conn.execute(text("SELECT COUNT(1) FROM taxonomy_categories")).scalar() or 0
            if existing == 0:
                role_names = [
                    'Project Director','Project Manager','Ops Manager',
                    'Team Leader','Case Handler','Admin'
                ]
                for rn in role_names:
                    conn.execute(
                        text("INSERT INTO taxonomy_categories(type,name) VALUES('role', :n)"),
                        {"n": rn},
                    )

                subjects = {
                    "Financial Crime / Compliance": [
                        "AML","KYC","CDD","EDD","Sanctions","Screening",
                        "TM (Transaction Monitoring)","FRA","SAR/STR","PEP","Adverse Media"
                    ],
                    "Case Handling / QA / Ops": [
                        "Case Handling","Quality Assurance","MI/Reporting",
                        "Playbooks","Workflows","SLA Management"
                    ],
                    "Domains": [
                        "Retail Banking","Corporate Banking","FinTech",
                        "Payments","Crypto","Wealth/Private Banking"
                    ],
                    "Tools & Platforms": [
                        "Actimize","Pega","Mantas","Nice","Salesforce",
                        "Dynamics","Oracle","Workday","Excel","SQL"
                    ],
                }
                for cat, tags in subjects.items():
                    conn.execute(
                        text("INSERT INTO taxonomy_categories(type,name) VALUES('subject', :n)"),
                        {"n": cat},
                    )
                    cat_id = conn.execute(
                        text("SELECT id FROM taxonomy_categories WHERE type='subject' AND name=:n ORDER BY id DESC LIMIT 1"),
                        {"n": cat},
                    ).scalar()
                    for t in tags:
                        conn.execute(
                            text("INSERT INTO taxonomy_tags(category_id, tag) VALUES(:cid, :tag)"),
                            {"cid": cat_id, "tag": t},
                        )
        except Exception:
            pass

        try:
            n_alias = conn.execute(text("SELECT COUNT(1) FROM role_aliases")).scalar() or 0
            if n_alias == 0:
                seeds = [
                    ("Project Director", "programme director"),
                    ("Project Director", "program director"),
                    ("Project Director", "pd"),
                    ("Project Manager",  "programme manager"),
                    ("Project Manager",  "program manager"),
                    ("Project Manager",  "pm"),
                    ("Project Manager",  "project lead"),
                    ("Ops Manager",      "operations manager"),
                    ("Ops Manager",      "service delivery manager"),
                    ("Team Leader",      "team lead"),
                    ("Team Leader",      "supervisor"),
                    ("Case Handler",     "kyc analyst"),
                    ("Case Handler",     "aml analyst"),
                    ("Case Handler",     "investigator"),
                    ("Admin",            "administrator"),
                    ("Admin",            "office admin"),
                ]
                for canon, alias in seeds:
                    conn.execute(
                        text("INSERT INTO role_aliases(canonical, alias) VALUES(:c, :a)"),
                        {"c": canon, "a": alias},
                    )
        except Exception:
            pass

        # ===== Seed role_types table =====
        try:
            count_roles = conn.execute(text("SELECT COUNT(1) FROM role_types")).scalar() or 0
            if count_roles == 0:
                for r in ROLE_TYPES:
                    conn.execute(
                        text("INSERT INTO role_types (name, default_rate) VALUES (:n, 0)"),
                        {"n": r},
                    )
        except Exception:
            pass

        # ===== SECURITY: Add security columns to users table (CREST compliance) =====
        # TEMPORARILY DISABLED - Causing deployment issues
        # Will be re-enabled after testing
        pass

        # ===== Seed admin user if no users exist =====
        try:
            user_count = conn.execute(text("SELECT COUNT(*) FROM users")).scalar() or 0
            if user_count == 0:
                from werkzeug.security import generate_password_hash
                import secrets
                
                # Generate a secure default password
                default_password = secrets.token_urlsafe(16)
                password_hash_value = generate_password_hash(default_password, method='pbkdf2:sha256')
                
                # Create admin user
                conn.execute(
                    text("""
                        INSERT INTO users (name, email, password_hash, role, is_active, created_at)
                        VALUES (:name, :email, :password_hash, :role, :is_active, CURRENT_TIMESTAMP)
                    """),
                    {
                        "name": "Admin",
                        "email": "admin@example.com",
                        "password_hash": password_hash_value,
                        "role": "super_admin",
                        "is_active": 1
                    }
                )
                
                # Log the credentials (WARNING: In production, send this via secure channel)
                print("\n" + "="*80)
                print("🔐 ADMIN USER CREATED")
                print("="*80)
                print(f"Email: admin@example.com")
                print(f"Password: {default_password}")
                print("="*80)
                print("⚠️  IMPORTANT: Change this password immediately after first login!")
                print("="*80 + "\n")
        except Exception as e:
            print(f"Warning: Could not seed admin user: {e}")

# Run on import - wrapped in try/catch to prevent startup failure
try:
    ensure_schema()
except Exception as e:
    print(f"⚠️  WARNING: Schema migration failed: {e}")
    print("⚠️  App will continue but some features may not work correctly")
    import traceback
    traceback.print_exc()
Base.metadata.create_all(engine)

# ---------- Taxonomy tagging helpers ----------
WORD = r"[A-Za-z][A-Za-z\-/&\.\(\) ]+[A-Za-z]"

def _collect_text_for_candidate(session: Session, cand: Candidate) -> str:
    """Return searchable text: latest CV + name + existing skills."""
    doc = session.scalar(
        select(Document)
        .where(Document.candidate_id == cand.id)
        .order_by(Document.uploaded_at.desc())
    )
    parts = [cand.name or "", cand.email or "", cand.skills or ""]
    if doc:
        parts.append(extract_cv_text(doc) or "")
    text = " \n ".join(p for p in parts if p).lower()
    # compress whitespace
    text = re.sub(r"\s+", " ", text)
    return text

def _job_is_open(job) -> bool:
    if job is None:
        return False
    status = (getattr(job, "status", "") or "").strip().lower()
    if status in {"withdrawn","closed","filled","archived","cancelled","canceled"}:
        return False
    # also treat explicit timestamps if you ever add them
    if getattr(job, "closed_at", None) is not None: return False
    if getattr(job, "archived_at", None) is not None: return False
    return True

def _load_role_aliases(session: Session) -> dict:
    """
    Merge built-in aliases with DB-managed aliases.
    Returns: {canonical_lower: set(alias_lower, ...)}
    """
    # start from your existing dict (if you want to keep the built-ins)
    base = {
        "project director": {"programme director","program director","pd","delivery director"},
        "project manager":  {"project lead","pm","programme manager","program manager","scrum master"},
        "ops manager":      {"operations manager","operations lead","ops lead","service delivery manager"},
        "team leader":      {"supervisor","team lead","shift lead","coach"},
        "case handler":     {"investigator","analyst","kyc analyst","aml analyst","qa analyst","caseworker"},
        "admin":            {"coordinator","assistant","administrator","pa","ea","office admin"},
    }
    rows = session.execute(text("SELECT canonical, alias FROM role_aliases")).all()
    for canon, alias in rows:
        if not canon or not alias:
            continue
        base.setdefault(canon.strip().lower(), set()).add(alias.strip().lower())
    return base

def _retag_candidate_skills(s, cand, overwrite: bool = False):
    """
    Derive tags/role from candidate text and update:
      - Candidate.skills (human-readable)
      - CandidateTag (normalized join rows)
    Uses the taxonomy helpers already in this codebase.
    """
    # Collect text (CV text preferred; falls back to freeform)
    text_lc = _collect_text_for_candidate(s, cand)
    if not text_lc:
        return

    # Load taxonomy terms once
    groups = _get_subject_term_set(s)
    term_list = [t for t in groups.get("__all__", [])]

    # Derive role + subject tags
    role = _normalise_role(text_lc)
    tags = _derive_subject_tags(text_lc, term_list)

    # ---------- Update Candidate.skills (readable) ----------
    current = (cand.skills or "").strip()
    existing_tokens = {w.strip().lower() for w in re.split(r"[,/|;]", current) if w.strip()}

    # Ensure role token is first (if found)
    parts = []
    if role:
        parts.append(role)

    # Add new tags (avoid duplicates, case-insensitive)
    new_tags = []
    for t in sorted(tags, key=str.lower):
        if t.lower() not in existing_tokens:
            new_tags.append(t)

    if overwrite:
        merged = ", ".join(parts + new_tags)
        if current:
            # keep any freeform skills too
            merged = (merged + (" | " if merged and current else "") + current).strip()
        cand.skills = merged or current
    else:
        # append tags that aren't already there
        if parts:
            for rt in parts:
                if rt.lower() not in existing_tokens:
                    current = (f"{rt} | " + current).strip(" |")
                    existing_tokens.add(rt.lower())
        if new_tags:
            add_str = ", ".join(new_tags)
            cand.skills = (current + (" | " if current else "") + add_str).strip()
        else:
            cand.skills = current

    # ---------- Sync CandidateTag (normalized) ----------
    # Map existing tag ids for quick checks
    existing_rels = s.scalars(
        select(CandidateTag).where(CandidateTag.candidate_id == cand.id)
    ).all()
    existing_tag_ids = {r.tag_id for r in existing_rels}

    # Lookup tag ids by name (case-insensitive)
    if tags:
        tag_rows = s.scalars(
            select(TaxonomyTag).where(func.lower(TaxonomyTag.tag).in_([t.lower() for t in tags]))
        ).all()
        for tg in tag_rows:
            if tg.id not in existing_tag_ids:
                s.add(CandidateTag(candidate_id=cand.id, tag_id=tg.id))

# --- Tag matching helper (place near other helpers) ---
def _match_tags_from_text(s, text: str) -> List[TaxonomyTag]:
    """Return TaxonomyTag rows whose tag string appears (case-insensitive) in text."""
    text_lc = (text or "").lower()
    if not text_lc.strip():
        return []
    all_tags = s.scalars(select(TaxonomyTag)).all()
    hits: List[TaxonomyTag] = []
    for tg in all_tags:
        token = (tg.tag or "").strip()
        if token and token.lower() in text_lc:
            hits.append(tg)
    return hits

def _create_engagement_from_opportunity(s: Session, opp: Opportunity) -> Engagement:
    existing = s.scalar(select(Engagement).where(Engagement.opportunity_id == opp.id))
    if existing:
        return existing

    new_ref = _next_engagement_ref(s)

    e = Engagement(
        ref=new_ref,
        name=opp.name,
        client=opp.client or "",
        status="Active",
        start_date=opp.est_start,              # now valid
        description=(opp.notes or ""),
        opportunity_id=opp.id,
    )
    s.add(e)
    s.flush()
    return e

def create_engagement_for_opportunity(s, opp: Opportunity) -> Engagement:
    # create the engagement with sensible defaults
    e = Engagement(
        name=opp.name or f"Engagement for {opp.client or 'Unknown Client'}",
        client=opp.client or "",
        status="Active",
        description=opp.notes or "",
        opportunity=opp,               # sets e.opportunity_id automatically
    )
    s.add(e)
    s.flush()  # get e.id

    # ensure engagement has a ref (EG00X...) if your ref backfill didn’t set it yet
    if not e.ref:
        # find next numeric
        prefix, width = "EG", 3
        max_n = s.execute(text("""
            SELECT COALESCE(MAX(CAST(SUBSTR(ref, 3) AS INTEGER)), 0)
            FROM engagements
            WHERE ref LIKE :patt
        """), {"patt": f"{prefix}%"}).scalar() or 0
        e.ref = f"{prefix}{str(max_n + 1).zfill(width)}"

    # cache back onto opportunity for quick linking in UI
    opp._engagement_id = e.id
    opp._engagement_ref = e.ref

    return e

def _get_subject_term_set(session: Session) -> Dict[str, List[str]]:
    """
    Build {category_name: [tag, ...]} from DB taxonomy.
    Also returns a flat set of all tags under key '__all__'.
    """
    cats = session.execute(
        text("SELECT c.id, c.name FROM taxonomy_categories c WHERE c.type='subject'")
    ).all()
    result: Dict[str, List[str]] = {}
    all_terms: List[str] = []
    for cid, cname in cats:
        tags = [t for (t,) in session.execute(
            text("SELECT tag FROM taxonomy_tags WHERE category_id=:cid"),
            {"cid": cid}
        ).all()]
        result[cname] = tags
        all_terms.extend(tags)
    result["__all__"] = sorted(set(all_terms), key=lambda s: s.lower())
    return result

def _rebuild_ai_summary_and_tags(s, cand, doc=None, job=None, appn=None):
    """
    Extract CV text → generate AI summary → write to latest Application and Candidate
    → derive & write skills/tags. If a Job is provided, also compute AI score + explanation.

    Args:
        s: SQLAlchemy Session
        cand: Candidate ORM instance
        doc:  (optional) Document instance for the CV. If None, pick the latest CV.
        job:  (optional) Job ORM instance — if provided, we will compute ai_score + ai_explanation
        appn: (optional) Application ORM instance to write into; if None we fetch latest for candidate
    """
    # pick latest CV if not provided (prefer uploaded_at; fallback to id)
    if doc is None:
        doc = s.scalar(
            select(Document)
            .where(Document.candidate_id == cand.id, getattr(Document, "doc_type", literal_column("'cv'")) == "cv")
            .order_by(getattr(Document, "uploaded_at", Document.id).desc())
        )

    # Safely extract text (prefer project helper)
    def _safe_extract_text(file_path: str, original_name: str) -> str:
        name = (original_name or file_path).lower()
        # PDF
        if name.endswith(".pdf"):
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    bits = [(p.extract_text() or "") for p in pdf.pages]
                txt = "\n".join(bits).strip()
                if txt:
                    return txt
            except Exception:
                pass
        # DOCX
        if name.endswith(".docx"):
            try:
                import docx
                d = docx.Document(file_path)
                txt = "\n".join([p.text for p in d.paragraphs]).strip()
                if txt:
                    return txt
            except Exception:
                pass
        # Plain text
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception:
            return ""

    cv_text = ""
    if doc:
        try:
            if "extract_cv_text" in globals() and callable(globals()["extract_cv_text"]):
                cv_text = extract_cv_text(doc) or ""
            else:
                file_path = _doc_file_path(doc)
                cv_text = _safe_extract_text(file_path, getattr(doc, "original_name", "") or doc.filename)
        except Exception:
            cv_text = ""

    # Generate summary (fallback to skills if no text at all)
    try:
        summary = ai_summarise(cv_text or (cand.skills or "")) if (cv_text or cand.skills) else ""
    except Exception as e:
        summary = ""
        current_app.logger.exception("ai_summarise failed: %s", e)

    # Choose target application
    target_app = appn
    if target_app is None:
        target_app = s.scalar(
            select(Application)
            .where(Application.candidate_id == cand.id)
            .order_by(Application.created_at.desc())
        )

    if target_app:
        target_app.ai_summary = summary

    # Mirror on Candidate if that column exists
    if hasattr(cand, "ai_summary"):
        cand.ai_summary = summary

    # Skills/Tags (prefer project helper if present)
    try:
        if "_retag_candidate_skills" in globals() and callable(globals()["_retag_candidate_skills"]):
            _retag_candidate_skills(s, cand, overwrite=False)
        else:
            # Minimal fallback: match taxonomy tags in text and store
            text_lc = (cv_text or "").lower()
            matched_ids = set()
            matched_tags = []
            all_tags = s.scalars(select(TaxonomyTag)).all()
            for tg in all_tags:
                token = (tg.tag or "").strip()
                if token and token.lower() in text_lc:
                    matched_ids.add(tg.id)
                    matched_tags.append(token)

            if matched_ids:
                existing_ids = {
                    tid for (tid,) in s.execute(
                        select(CandidateTag.tag_id).where(CandidateTag.candidate_id == cand.id)
                    ).all()
                }
                for tid in matched_ids:
                    if tid not in existing_ids:
                        s.add(CandidateTag(candidate_id=cand.id, tag_id=tid))

                # Mirror tag names into Candidate.skills for search/filter UX
                all_tag_rows = s.scalars(
                    select(TaxonomyTag)
                    .join(CandidateTag, CandidateTag.tag_id == TaxonomyTag.id)
                    .where(CandidateTag.candidate_id == cand.id)
                    .order_by(TaxonomyTag.tag.asc())
                ).all()
                tag_names = [t.tag for t in all_tag_rows if (t.tag or "").strip()]
                cand.skills = ", ".join(dict.fromkeys(tag_names))
    except Exception as e:
        current_app.logger.warning("retagging failed: %s", e)

    # Optional AI score if a Job provided
    if job and target_app:
        client = get_openai_client()
        if not client:
            current_app.logger.warning("⚠️  OpenAI API key not configured - skipping AI scoring for job %s", job.id)
            # Use heuristic scoring as fallback
            try:
                heur = _heuristic_components(job.description or "", cv_text or (cand.skills or ""))
                target_app.ai_score = int(heur.get("score", 0))
                overlaps = heur.get("overlap", [])
                top_matches = ", ".join(sorted(overlaps)[:5]) if overlaps else "none"
                target_app.ai_explanation = f"Keyword matching score (Configure OPENAI_API_KEY for AI-powered scoring). Top matches: {top_matches}"
                current_app.logger.info("✓ Using heuristic scoring: %d/100", target_app.ai_score)
            except Exception as e:
                current_app.logger.warning("Heuristic scoring failed: %s", e)
                target_app.ai_score = 0
                target_app.ai_explanation = "Scoring unavailable"
        else:
            try:
                current_app.logger.info("🤖 Computing AI score for candidate %s on job %s", cand.id, job.id)
                result = ai_score_with_explanation(job.description or "", cv_text or (cand.skills or ""))
                target_app.ai_score = int(result.get("final", 0) or 0)
                target_app.ai_explanation = (result.get("explanation") or "")[:7999]
                current_app.logger.info("✓ AI Score: %d/100 (GPT: %d, Heuristic: %d)", 
                                      target_app.ai_score, 
                                      result.get("gpt", 0),
                                      result.get("heuristic", 0))
            except Exception as e:
                current_app.logger.exception("❌ ai_score_with_explanation failed: %s", e)
                # Fallback to heuristic
                try:
                    heur = _heuristic_components(job.description or "", cv_text or (cand.skills or ""))
                    target_app.ai_score = int(heur.get("score", 0))
                    target_app.ai_explanation = f"AI scoring failed, using keyword matching: {heur.get('score', 0)}/100"
                except Exception:
                    target_app.ai_score = 0
                    target_app.ai_explanation = "Scoring failed"

def parse_date_dmy(s: Optional[str]):
    if not s:
        return None
    try:
        # dayfirst=True so "09-10-2025" is 9th Oct (not Sept 10)
        return dtparser.parse(s, dayfirst=True).replace(tzinfo=None)
    except Exception:
        return None

def format_date_dmy(dt: Optional[datetime.datetime]) -> str:
    if not dt:
        return ""
    return dt.strftime("%d-%m-%Y")

# --- Sequential engagement ref (EG001, EG002, ...) ---
def _next_engagement_ref(session: Session, prefix: str = "EG", width: int = 3) -> str:
    import re as _re
    # pull all refs that match pattern and take max
    rows = session.execute(text(
        "SELECT ref FROM engagements WHERE ref IS NOT NULL AND ref <> ''"
    )).all()
    max_n = 0
    for (ref,) in rows:
        m = _re.fullmatch(rf"{_re.escape(prefix)}(\d+)", ref or "")
        if m:
            try:
                max_n = max(max_n, int(m.group(1)))
            except ValueError:
                pass
    return f"{prefix}{str(max_n + 1).zfill(width)}"

# --- WTForm for Create Engagement (vertical layout fields) ---
class CreateEngagementForm(FlaskForm):
    name = StringField("Engagement Name", validators=[DataRequired()])
    client = StringField("Client", validators=[WTOptional()])
    status = SelectField("Status",
                         choices=[("Active","Active"), ("On Hold","On Hold"), ("Finished","Finished")],
                         default="Active")
    start_date = StringField("Start Date (DD-MM-YYYY)", validators=[WTOptional()])
    end_date = StringField("End Date (DD-MM-YYYY)", validators=[WTOptional()])
    sow_signed_on = StringField("Statement of Work Signed (DD-MM-YYYY)", validators=[WTOptional()])
    description = TextAreaField("Description of Work", validators=[WTOptional()])

def _derive_subject_tags(text_lc: str, all_terms: List[str]) -> List[str]:
    """Match whole-word-ish terms (case-insensitive) within candidate text."""
    found = []
    for term in all_terms:
        t = term.strip()
        if not t:
            continue
        # allow terms with / or () etc; use simple contains with word guards when safe
        # If simple letters/spaces, enforce word boundaries; otherwise fallback to 'in'
        if re.fullmatch(r"[A-Za-z0-9 ]+", t):
            pat = r"(?:^|[^A-Za-z0-9])" + re.escape(t.lower()) + r"(?:[^A-Za-z0-9]|$)"
            if re.search(pat, text_lc):
                found.append(term)
        else:
            if t.lower() in text_lc:
                found.append(term)
    # de-dup, stable
    seen, out = set(), []
    for x in found:
        k = x.lower()
        if k not in seen:
            seen.add(k); out.append(x)
    return out

# Simple role normaliser: map common aliases -> canonical
_ROLE_ALIASES = {
    "project director": ["programme director","program director","pd","delivery director"],
    "project manager": ["project lead","pm","programme manager","program manager","scrum master"],
    "ops manager":     ["operations manager","operations lead","ops lead","service delivery manager"],
    "team leader":     ["supervisor","team lead","shift lead","coach"],
    "case handler":    ["investigator","analyst","kyc analyst","aml analyst","qa analyst","caseworker"],
    "admin":           ["coordinator","assistant","administrator","pa","ea","office admin"],
}

def load_subject_groups_with_tags(session: Session):
    cats = session.scalars(
        select(TaxonomyCategory).where(TaxonomyCategory.type.in_(["role","subject"]))
        .order_by(TaxonomyCategory.type.asc(), TaxonomyCategory.name.asc())
    ).all()
    out = []
    for c in cats:
        tags = session.scalars(
            select(TaxonomyTag).where(TaxonomyTag.category_id == c.id).order_by(TaxonomyTag.tag.asc())
        ).all()
        if tags:
            out.append((c, tags))
    return out

def _normalise_role(text_lc: str) -> str:
    # direct hit
    for canonical in ROLE_TYPES:
        if canonical.lower() in text_lc:
            return canonical

    # DB + built-ins
    with Session(engine) as s:
        aliases = _load_role_aliases(s)

    for canonical_lc, alias_set in aliases.items():
        for a in alias_set:
            if re.search(r"(?:^|[^A-Za-z0-9])"+re.escape(a)+r"(?:[^A-Za-z0-9]|$)", text_lc):
                # return with ROLE_TYPES' proper case if present; else title-case the canonical
                for r in ROLE_TYPES:
                    if r.lower() == canonical_lc:
                        return r
                return canonical_lc.title()
    return ""

def _latest_doc_for_candidate(session: Session, cand_id: int) -> Optional[Document]:
    return session.scalar(
        select(Document)
        .where(Document.candidate_id == cand_id)
        .order_by(Document.uploaded_at.desc())
    )

def sumsub_sign_request(method: str, path: str, body_str: str = "") -> dict:
    """
    Create Sumsub HMAC signature headers for the given request.
    IMPORTANT: body_str must match the EXACT bytes you send on the wire.
    We therefore JSON-dump with compact separators and send via data=body_str.
    """
    if not SUMSUB_APP_TOKEN or not SUMSUB_SECRET_KEY:
        # Make it obvious why a 500 would occur
        raise RuntimeError("SUMSUB_APP_TOKEN/SUMSUB_SECRET_KEY are not set")

    ts = str(int(time.time()))
    to_sign = f"{ts}{method.upper()}{path}{body_str}".encode("utf-8")
    sig = hmac.new(SUMSUB_SECRET_KEY.encode("utf-8"), to_sign, hashlib.sha256).hexdigest()

    return {
        "X-App-Token": SUMSUB_APP_TOKEN,
        "X-App-Access-Ts": ts,
        "X-App-Access-Sig": sig,
    }

# --- add near other small helpers ---
def _latest_application_for_candidate(session: Session, cand_id: int) -> Optional[Application]:
    return session.scalar(
        select(Application)
        .where(Application.candidate_id == cand_id)
        .order_by(Application.created_at.desc())
    )

# ---------- AI scoring & summarisation ----------

def _tokenize_words(s: str) -> List[str]:
    return re.findall(r"[a-zA-Z]{3,}", (s or "").lower())

def _heuristic_components(job_desc: str, cv_text: str):
    jd_words = set(_tokenize_words(job_desc))
    cv_words = set(_tokenize_words(cv_text))
    if not jd_words or not cv_words:
        return {"score": 0, "overlap": [], "jd_words": 0}
    overlap = sorted(jd_words & cv_words)
    base = min(100, int(100 * len(overlap) / max(10, len(jd_words))))
    return {"score": base, "overlap": overlap[:20], "jd_words": len(jd_words)}

def _jd_completeness_words(job_desc: str) -> Tuple[int, float]:
    wc = len(_tokenize_words(job_desc))
    completeness = max(0.0, min(1.0, wc / 60.0))  # ~1.0 at 60+ words
    return wc, completeness

def _gpt_score_and_rationale(job_desc: str, cv_text: str) -> Tuple[int, List[str]]:
    client = get_openai_client()
    if not client:
        return 0, []
    prompt = f"""
You are an ATS scoring assistant.

Task:
1) Rate the candidate's CV match to the job description on a 0–100 integer scale.
2) Provide exactly 3 short bullets (≤ ~12 words) explaining the score.

Return strictly this JSON (no prose):
{{
  "score": <int 0-100>,
  "bullets": ["...", "...", "..."]
}}

JOB DESCRIPTION:
{job_desc}

CANDIDATE CV:
{cv_text}
"""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        raw = resp.choices[0].message.content.strip()
        import json as _json
        cleaned = re.sub(r"^```json|```$", "", raw).strip()
        data = _json.loads(cleaned)
        g = int(data.get("score", 0))
        bullets = [str(b) for b in (data.get("bullets") or [])][:3]
        return max(0, min(100, g)), bullets
    except Exception as e:
        print("GPT scoring error:", e)
        return 0, []

def ai_score_with_explanation(job_desc: str, cv_text: str) -> Dict:
    heur = _heuristic_components(job_desc, cv_text)
    jd_wc, completeness = _jd_completeness_words(job_desc)
    gpt_score, bullets = _gpt_score_and_rationale(job_desc, cv_text)

    gpt_w = 0.2 + 0.8 * completeness  # rely more on GPT when JD is richer
    heur_w = 1.0 - gpt_w
    blended = int(round(gpt_w * gpt_score + heur_w * heur["score"]))
    if completeness < 0.25:
        blended = min(blended, 88)  # cap if JD is extremely short

    exp_lines = [
        f"JD length: {jd_wc} words (completeness {int(completeness*100)}%).",
        f"Weights → GPT {gpt_w:.2f}, Heuristic {heur_w:.2f}.",
        f"Raw scores → GPT {gpt_score}, Overlap {heur['score']}."
    ]
    if heur["overlap"]:
        exp_lines.append("Top overlaps: " + ", ".join(sorted(heur["overlap"])[:8]) + ".")
    if bullets:
        exp_lines.append("Why: " + " • ".join(bullets))
    explanation = " ".join(exp_lines)

    return {
        "final": blended,
        "gpt": gpt_score,
        "heuristic": heur["score"],
        "jd_words": jd_wc,
        "completeness": completeness,
        "weights": {"gpt": gpt_w, "heuristic": heur_w},
        "overlap": heur["overlap"],
        "bullets": bullets,
        "explanation": explanation,
    }

def _smart_truncate(txt: str, limit: int) -> str:
    """
    Truncate text gracefully at the nearest sentence or word boundary.
    Ensures summaries end cleanly instead of being cut mid-word.
    """
    if len(txt) <= limit:
        return txt
    cut = txt[:limit]
    # Prefer to end on sentence, bullet, or whitespace
    for sep in [". ", "•", "\n", " "]:
        i = cut.rfind(sep)
        if i >= 60:  # ensure we don’t truncate too early
            return cut[:i + 1].rstrip() + " …"
    return cut.rstrip() + " …"

def ai_summarise(text: str, max_chars: int = 1400) -> str:
    text = _truncate_for_ai(text or "", 12000)
    if not text:
        return ""
    client = get_openai_client()
    if client:
        try:
            prompt = (
                "Summarize the candidate’s experience for recruiters in 5–7 bullets. "
                "Keep bullets concise; avoid fluff; focus on impact, domains, tools.\n\n"
                f"TEXT:\n{text}"
            )
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content": prompt}],
                temperature=0.2,
            )
            out = (resp.choices[0].message.content or "").strip()
            if out:
                return _smart_truncate(out, max_chars)
        except Exception as e:
            try:
                current_app.logger.exception("ai_summarise OpenAI failed: %s", e)
            except Exception:
                pass
    # Fallback (your existing tidy version)
    import re
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    joined = re.sub(r"\s+", " ", " ".join(lines[:40]))
    bullets = []
    roles = re.findall(
        r"\b[A-Z][a-z]+(?: [A-Z][a-z]+)* (?:Manager|Officer|Analyst|Director|Consultant|Engineer|Lead|Head|Specialist|Advisor|Associate)\b",
        joined,
    )
    roles = list(dict.fromkeys(roles))
    if roles:
        bullets.append("• Experience includes: " + ", ".join(roles[:4]) + (", ..." if len(roles) > 4 else ""))
    bullets.append("• " + joined[:800])
    return _smart_truncate("\n".join(bullets), max_chars)

# --- Optional text extraction deps ---
try:
    from PyPDF2 import PdfReader            # pip install PyPDF2
except Exception:
    PdfReader = None

try:
    import docx                              # python-docx  -> pip install python-docx
except Exception:
    docx = None

try:
    import textract                          # for legacy .doc  -> pip install textract
except Exception:
    textract = None

class RoleTypeForm(FlaskForm):
    name = StringField("Role name", validators=[DataRequired()])
    default_rate = IntegerField("Default £/day", validators=[WTOptional()])
    submit = SubmitField("Add role")

# -------------- Forms --------------
class EngagementForm(FlaskForm):
    name = StringField("Engagement Name", validators=[DataRequired()])
    client = StringField("Client", validators=[WTOptional()])
    # Use text inputs so we can accept DD-MM-YYYY; we’ll parse server-side.
    start_date = StringField("Start Date (DD-MM-YYYY)", validators=[WTOptional()])
    end_date = StringField("End Date (DD-MM-YYYY)", validators=[WTOptional()])
    sow_signed_at = StringField("SOW Signed (DD-MM-YYYY)", validators=[WTOptional()])
    status = SelectField("Status",
                         choices=[("Active","Active"), ("Paused","Paused"), ("Completed","Completed")],
                         validators=[DataRequired()],
                         default="Active")
    description = TextAreaField("Description", validators=[WTOptional()])

class JobForm(FlaskForm):
    engagement_id = SelectField("Engagement", coerce=int, validators=[DataRequired()])
    title = StringField("Job Title", validators=[DataRequired()])
    description = TextAreaField("Description", validators=[DataRequired()])
    role_type = SelectField("Role Type", choices=[(r, r) for r in ROLE_TYPES], validators=[WTOptional()])
    location = StringField("Location", validators=[WTOptional()])
    salary_range = StringField("Salary Range", validators=[WTOptional()])

class ApplyForm(FlaskForm):
    name = StringField("Your Name", validators=[DataRequired()])
    email = StringField("Email", validators=[DataRequired(), Email()])
    phone = StringField("Phone", validators=[WTOptional()])
    cover_note = TextAreaField("Cover Note", validators=[WTOptional()])
    cv = FileField("Upload CV (PDF/DOC/DOCX)", validators=[DataRequired()])

class InterviewForm(FlaskForm):
    scheduled_at = DateTimeLocalField("Interview Date & Time", format="%Y-%m-%dT%H:%M", validators=[DataRequired()])
    interviewer_email = StringField("Interviewer Email", validators=[WTOptional()])

class TaxCategoryForm(FlaskForm):
    type = SelectField("Type", choices=[("role","Role"), ("subject","Subject")], validators=[DataRequired()])
    name = StringField("Category name", validators=[DataRequired()])

class TaxCategoryRenameForm(FlaskForm):
    name = StringField("New name", validators=[DataRequired()])

class TaxTagForm(FlaskForm):
    category_id = SelectField("Category", coerce=int, validators=[DataRequired()])
    tag = StringField("Tag", validators=[DataRequired()])

# ---- Opportunities form ----
class OpportunityForm(FlaskForm):
    name = StringField("Opportunity Name", validators=[DataRequired()])
    client = StringField("Client", validators=[WTOptional()])

    stage = SelectField(
        "Stage",
        choices=[
            ("Qualified Lead", "Qualified Lead"),
            ("Proposal", "Proposal"),
            ("Procurement", "Procurement"),
            ("Closed Won", "Closed Won"),
            ("Closed Lost", "Closed Lost"),
        ],
        validators=[DataRequired()],
        default="Qualified Lead",
    )

    owner = SelectField("Owner", coerce=int, validators=[WTOptional()])

    # Display as DD-MM-YYYY to the user
    est_start = StringField("Est. Start (DD-MM-YYYY)", validators=[WTOptional()])

    est_value = IntegerField("Est. Value (£)", validators=[WTOptional()])

    notes = TextAreaField("Notes", validators=[WTOptional()])

    # New client contact fields
    client_contact_name  = StringField("Client contact name", validators=[WTOptional()])
    client_contact_role  = StringField("Client contact role", validators=[WTOptional()])
    client_contact_phone = StringField("Client contact phone", validators=[WTOptional()])
    client_contact_email = StringField("Client contact email", validators=[WTOptional(), Email()])

# -------------- Helpers --------------
def allowed_file(fname:str)->bool:
    return "." in fname and fname.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def save_upload(file_storage, subdir="cvs"):
    fname = secure_filename(file_storage.filename)
    if not allowed_file(fname):
        raise ValueError("Unsupported file type")
    new_name = f"{uuid.uuid4()}_{fname}"
    target_dir = os.path.join(app.config["UPLOAD_FOLDER"], subdir)
    os.makedirs(target_dir, exist_ok=True)
    path = os.path.join(target_dir, new_name)
    file_storage.save(path)
    return new_name, path, fname

def parse_date(s: Optional[str]):
    if not s:
        return None
    try:
        return dtparser.parse(s).replace(tzinfo=None)
    except Exception:
        return None

def parse_int(s, default=0):
    try:
        return int(str(s).replace(",", "").strip())
    except Exception:
        return default

def currency_fmt(n: int) -> str:
    try:
        return f"£{int(n):,}"
    except Exception:
        return "£0"

def _doc_file_path(doc: "Document") -> str:
    """Absolute path to a stored candidate document.
    
    Handles multiple storage patterns:
    1. filename = "uploads/cvs/xxx.pdf" (stored in static/uploads/cvs/)
    2. filename = "xxx.pdf" (stored in uploads/cvs/)
    """
    filename = doc.filename or ""
    
    # If filename contains path prefix like "uploads/cvs/", look in static folder
    if filename.startswith("uploads/"):
        static_path = os.path.join(os.path.dirname(__file__), "static", filename)
        if os.path.exists(static_path):
            return static_path
    
    # Try the standard uploads folder path
    # Handle case where filename already has "uploads/cvs/" prefix
    if filename.startswith("uploads/cvs/"):
        clean_filename = filename.replace("uploads/cvs/", "")
    else:
        clean_filename = filename
    
    standard_path = os.path.join(app.config["UPLOAD_FOLDER"], "cvs", clean_filename)
    if os.path.exists(standard_path):
        return standard_path
    
    # Fallback to static/uploads/cvs with clean filename
    static_fallback = os.path.join(os.path.dirname(__file__), "static", "uploads", "cvs", clean_filename)
    if os.path.exists(static_fallback):
        return static_fallback
    
    # Return the standard path even if it doesn't exist (for error handling upstream)
    return standard_path

def _truncate_for_ai(text: str, limit: int = 12000) -> str:
    """Light truncation to keep prompts reasonable."""
    if not text:
        return ""
    t = re.sub(r"\s+\n", "\n", text)
    t = re.sub(r"[ \t]+", " ", t)
    return t[:limit]

def _extract_text_from_pdf(path: str) -> str:
    if PdfReader is None:
        return ""
    try:
        with open(path, "rb") as f:
            reader = PdfReader(f)
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n".join(pages)
    except Exception as e:
        print("PDF extract error:", e)
        return ""

def _extract_text_from_docx(path: str) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(path)
        paras = [p.text for p in d.paragraphs]
        return "\n".join(paras)
    except Exception as e:
        print("DOCX extract error:", e)
        return ""

def _extract_text_from_doc(path: str) -> str:
    """Legacy .doc via textract if available."""
    if textract is None:
        return ""
    try:
        raw = textract.process(path)
        return raw.decode("utf-8", errors="ignore")
    except Exception as e:
        print("DOC extract error:", e)
        return ""

def extract_cv_text(doc: "Document") -> str:
    """
    Extracts best-effort text from a stored CV Document.
    Supports: .pdf, .docx, .doc (optional via textract).
    """
    if not doc:
        return ""
    path = _doc_file_path(doc)
    ext = os.path.splitext(doc.original_name or doc.filename)[1].lower()

    text = ""
    if ext == ".pdf":
        text = _extract_text_from_pdf(path)
    elif ext == ".docx":
        text = _extract_text_from_docx(path)
    elif ext == ".doc":
        text = _extract_text_from_doc(path)

    # Fallback: try reading as plain text (some users upload .txt with wrong ext)
    if not text:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception:
            pass

    return _truncate_for_ai(text)

def send_email(
    to_email,
    subject,
    html_body,
    attachments: Optional[List[Tuple[str, bytes, str]]] = None,
):
    # Dev / offline mode: if SMTP_HOST is not configured (or just whitespace),
    # don't try to actually send. Just log what would've gone out.
    if not (SMTP_HOST or "").strip():
        print("[DEV] SMTP not configured; would email:")
        print(" To:", to_email)
        print(" Subj:", subject)
        print(" Body:", html_body[:200], "...")
        if attachments:
            for filename, _, mime in attachments:
                print(f" Attachment: {filename} ({mime})")
        return

    msg = EmailMessage()
    msg["From"] = SMTP_FROM
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.set_content("This email contains HTML and possibly an ICS invite.")
    msg.add_alternative(html_body, subtype="html")

    if attachments:
        for filename, content, mime in attachments:
            maintype, subtype = mime.split("/", 1)
            msg.add_attachment(
                content,
                maintype=maintype,
                subtype=subtype,
                filename=filename,
            )

    context = ssl.create_default_context()
    with smtplib.SMTP(SMTP_HOST.strip(), SMTP_PORT) as server:
        server.starttls(context=context)
        if SMTP_USER and SMTP_PASS:
            server.login(SMTP_USER, SMTP_PASS)
        server.send_message(msg)

# Lightweight stopword set for JD term selection
_STOP = {
    "the","and","for","with","from","into","your","our","you","are","this","that","will","have","has",
    "per","role","job","team","work","working","experience","years","year","etc","such","across","able",
    "strong","good","great","excellent","within","including","include","on","in","of","to","a","an","as",
    "by","is","be","we","they","their","them","it","its","or","not","at","over","under"
}

def _jd_top_terms(job_desc: str, n: int = 8) -> List[str]:
    tokens = _tokenize_words(job_desc)
    freq: Dict[str, int] = {}
    for t in tokens:
        if t in _STOP or len(t) < 3:
            continue
        freq[t] = freq.get(t, 0) + 1
    # rank by frequency then alphabetically for stability
    top = sorted(freq.items(), key=lambda kv: (-kv[1], kv[0]))[:n]
    return [w for w, _ in top]

def _cheap_overlap_score(jd_terms: List[str], skills: str) -> int:
    """Very cheap overlap against normalised skills text (no CV parsing)."""
    if not skills:
        return 0
    s = skills.lower()
    hits = sum(1 for t in jd_terms if t and t in s)
    # weight a tiny bit by string length to avoid ties
    return hits * 10 + min(5, len(s) // 2000)

def ics_invite(summary, description, start_dt: datetime.datetime, end_dt: datetime.datetime, location, organizer_email, attendee_email):
    """
    Build a standard .ics invite attachment for email.
    Automatically converts datetimes to UTC if tzinfo is present.
    """
    from email.utils import parseaddr

    # Extract pure email address from “Name <email>” if needed
    organizer_email = parseaddr(organizer_email)[1] or organizer_email
    uid = f"{uuid.uuid4()}@{request.host}".replace(":", "")
    dtfmt = "%Y%m%dT%H%M%SZ"

    def _fmt(dt):
        if not dt:
            return ""
        if dt.tzinfo:
            dt = dt.astimezone(datetime.timezone.utc)
        return dt.strftime(dtfmt)

    start_utc = _fmt(start_dt)
    end_utc = _fmt(end_dt)

    ics = f"""BEGIN:VCALENDAR
PRODID:-//ATS//Interview//EN
VERSION:2.0
CALSCALE:GREGORIAN
METHOD:REQUEST
BEGIN:VEVENT
DTSTAMP:{datetime.datetime.utcnow().strftime(dtfmt)}
DTSTART:{start_utc}
DTEND:{end_utc}
SUMMARY:{summary}
UID:{uid}
DESCRIPTION:{description}
ORGANIZER;CN=Talent Ops:MAILTO:{organizer_email}
ATTENDEE;CN=Candidate;RSVP=TRUE:MAILTO:{attendee_email}
LOCATION:{location}
END:VEVENT
END:VCALENDAR
"""
    return ics.encode("utf-8")

# ---- Jinja helpers ----
@app.template_filter("slugify")
def slugify(value: str) -> str:
    s = (value or "").lower().replace(" ", "_")
    return re.sub(r"[^a-z0-9_]", "", s)

def slugify_role(name: str) -> str:
    s = (name or "").lower().replace(" ", "_")
    return re.sub(r"[^a-z0-9_]", "", s)

# -------------- Routes --------------
from sqlalchemy.orm import selectinload  # make sure this import exists

@app.route("/")
@login_required
def index():
    """
    Dashboard page per wireframe requirements.
    
    Filters: Client, Engagement, Role, Intake
    
    KPIs (Row 1):
    - Live Engagements (count)
    - Delivery Target (total target headcount)
    - Associates Offered (total offered)
    - Vetting Completed (count)
    - Contracted (count)
    
    Features:
    - Delivery Summary Chart (horizontal bar: Associates Offered vs Target per client)
    - Recent Jobs Widget with Manage button and Public Link
    """
    # Audit log: dashboard view
    
    now = datetime.datetime.utcnow()
    d7  = now - datetime.timedelta(days=7)
    d30 = now - datetime.timedelta(days=30)
    d3  = now - datetime.timedelta(days=3)
    
    # Get filter params (now multi-select, returns lists)
    client_filter = request.args.getlist("client") or []
    engagement_filter = request.args.getlist("engagement") or []
    role_filter = request.args.getlist("role") or []
    intake_filter = request.args.getlist("intake") or []

    with Session(engine) as s:
        # Get filter dropdown data
        all_clients = s.scalars(
            select(Engagement.client)
            .distinct()
            .where(Engagement.client.isnot(None), Engagement.client != "")
            .order_by(Engagement.client)
        ).all()
        
        all_engagements_list = s.scalars(
            select(Engagement)
            .where(Engagement.status == "Active")
            .order_by(Engagement.name)
        ).all()
        
        all_roles = s.scalars(
            select(Job.title)
            .distinct()
            .where(Job.title.isnot(None), Job.title != "")
            .order_by(Job.title)
        ).all()
        
        all_intakes = []
        try:
            intake_dates = s.scalars(
                select(EngagementPlan.intake_date)
                .distinct()
                .where(EngagementPlan.intake_date.isnot(None))
                .order_by(EngagementPlan.intake_date)
            ).all()
            all_intakes = [d.strftime("%Y-%m-%d") if d else None for d in intake_dates if d]
        except Exception:
            pass
        
        # Build base engagement query with filters
        eng_query = select(Engagement).where(Engagement.status == "Active")
        if client_filter:
            eng_query = eng_query.where(Engagement.client.in_(client_filter))
        if engagement_filter:
            eng_ids_int = [int(e) for e in engagement_filter if e.isdigit()]
            if eng_ids_int:
                eng_query = eng_query.where(Engagement.id.in_(eng_ids_int))
        
        active_engagements = s.scalars(eng_query).all()
        
        # === KPIs per wireframe ===
        
        # 1. Live Engagements
        live_engagements = len(active_engagements)
        
        # 2. Delivery Target (sum of planned headcount across filtered engagements)
        eng_ids = [e.id for e in active_engagements]
        delivery_target = 0
        if eng_ids:
            delivery_target = s.scalar(
                select(func.coalesce(func.sum(EngagementPlan.planned_count), 0))
                .where(EngagementPlan.engagement_id.in_(eng_ids))
            ) or 0
        
        # 3. Associates Offered (applications in "Offer" or later stage)
        offer_stages = ["Offer", "Vetting", "Vetting In-Flight", "Contract Issued", "Contract Signed", "Onboarding", "Hired"]
        associates_offered_query = (
            select(func.count(func.distinct(Application.candidate_id)))
            .select_from(Application)
            .join(Job, Job.id == Application.job_id)
            .where(Application.status.in_(offer_stages))
        )
        if eng_ids:
            associates_offered_query = associates_offered_query.where(Job.engagement_id.in_(eng_ids))
        associates_offered = s.scalar(associates_offered_query) or 0
        
        # 4. Vetting Completed (applications that have passed vetting)
        vetting_completed_query = (
            select(func.count(Application.id))
            .select_from(Application)
            .join(Job, Job.id == Application.job_id)
            .where(Application.status.in_(["Contract Issued", "Contract Signed", "Hired", "Onboarding"]))
        )
        if eng_ids:
            vetting_completed_query = vetting_completed_query.where(Job.engagement_id.in_(eng_ids))
        vetting_completed = s.scalar(vetting_completed_query) or 0
        
        # 5. Contracted (signed contracts)
        contracted_query = (
            select(func.count(ESigRequest.id))
            .select_from(ESigRequest)
            .outerjoin(Application, Application.id == ESigRequest.application_id)
            .outerjoin(Job, Job.id == Application.job_id)
            .where(func.lower(ESigRequest.status).in_(["signed", "completed"]))
        )
        if eng_ids:
            contracted_query = contracted_query.where(Job.engagement_id.in_(eng_ids))
        contracted = s.scalar(contracted_query) or 0
        
        # === Delivery Summary Chart Data ===
        # For each client/engagement, get Associates Offered vs Delivery Target
        delivery_chart_data = []
        for eng in active_engagements:
            target = s.scalar(
                select(func.coalesce(func.sum(EngagementPlan.planned_count), 0))
                .where(EngagementPlan.engagement_id == eng.id)
            ) or 0
            
            offered = s.scalar(
                select(func.count(func.distinct(Application.candidate_id)))
                .select_from(Application)
                .join(Job, Job.id == Application.job_id)
                .where(Application.status.in_(offer_stages))
                .where(Job.engagement_id == eng.id)
            ) or 0
            
            if target > 0 or offered > 0:
                delivery_chart_data.append({
                    "name": eng.name,
                    "client": eng.client or "Unknown",
                    "target": target,
                    "offered": offered,
                    "pct": int(100 * offered / target) if target else 0,
                    "eng_id": eng.id,  # GAP 2.2: Add engagement ID for clickable chart
                })
        
        # Sort by percentage (lowest first to show who needs attention)
        delivery_chart_data.sort(key=lambda x: x["pct"])
        
        # === 4-Week Lookahead Data ===
        # Show engagements with upcoming intake dates (next 4 weeks)
        # Format: horizontal bar chart with engagement names on Y-axis (like Delivery Summary)
        # IMPORTANT: Applies the same filters as other dashboard sections
        lookahead_data = []
        four_weeks_ahead = now + datetime.timedelta(weeks=4)
        
        try:
            # Build query for engagements with intake dates in the next 4 weeks
            lookahead_query = (
                select(
                    Engagement.id,
                    Engagement.name,
                    Engagement.client,
                    EngagementPlan.intake_date,
                    func.sum(EngagementPlan.planned_count).label("total_target")
                )
                .join(EngagementPlan, EngagementPlan.engagement_id == Engagement.id)
                .where(EngagementPlan.intake_date.isnot(None))
                .where(EngagementPlan.intake_date >= now.date())
                .where(EngagementPlan.intake_date <= four_weeks_ahead.date())
            )
            
            # Apply filters to 4-week lookahead
            if client_filter:
                lookahead_query = lookahead_query.where(Engagement.client.in_(client_filter))
            if eng_ids:  # engagement_filter already processed
                lookahead_query = lookahead_query.where(Engagement.id.in_(eng_ids))
            if intake_filter:
                # Convert intake filter strings to dates for comparison
                intake_dates_filter = []
                for intake_str in intake_filter:
                    try:
                        intake_dates_filter.append(datetime.datetime.strptime(intake_str, "%Y-%m-%d").date())
                    except:
                        pass
                if intake_dates_filter:
                    lookahead_query = lookahead_query.where(EngagementPlan.intake_date.in_(intake_dates_filter))
            
            lookahead_query = lookahead_query.group_by(Engagement.id, EngagementPlan.intake_date).order_by(EngagementPlan.intake_date)
            
            upcoming_intakes = s.execute(lookahead_query).all()
            
            for row in upcoming_intakes:
                # Count associates offered for this engagement
                eng_offered = s.scalar(
                    select(func.count(func.distinct(Application.candidate_id)))
                    .select_from(Application)
                    .join(Job, Job.id == Application.job_id)
                    .where(Application.status.in_(offer_stages))
                    .where(Job.engagement_id == row.id)
                ) or 0
                
                target = row.total_target or 0
                intake_label = row.intake_date.strftime("%d %b") if row.intake_date else "TBD"
                
                lookahead_data.append({
                    "name": row.name,
                    "label": f"{row.name} ({intake_label})",
                    "client": row.client or "Unknown",
                    "intake_date": intake_label,
                    "target": target,
                    "offered": eng_offered,
                    "pct": int(100 * eng_offered / target) if target else 0,
                    "eng_id": row.id,
                })
        except Exception as e:
            pass
        
        # === Open Jobs List ===
        open_jobs_query = (
            select(Job)
            .options(selectinload(Job.engagement))
            .where(Job.status == "Open")
        )
        # Apply filters to open jobs
        if eng_ids:
            open_jobs_query = open_jobs_query.where(Job.engagement_id.in_(eng_ids))
        if role_filter:
            open_jobs_query = open_jobs_query.where(Job.title.in_(role_filter))
        open_jobs = s.scalars(
            open_jobs_query.order_by(Job.created_at.desc())
        ).all()
        
        # === Recent Jobs Widget ===
        recent_jobs = s.scalars(
            select(Job)
            .options(selectinload(Job.engagement))
            .order_by(Job.created_at.desc())
            .limit(5)
        ).all()
        
        # === Legacy data for backward compatibility ===
        total_candidates = s.scalar(select(func.count(Candidate.id))) or 0
        total_apps = s.scalar(select(func.count(Application.id))) or 0
        interviews = s.scalar(
            select(func.count()).select_from(Application).where(Application.status.in_(["Interview", "Interview Scheduled", "Interview Completed"]))
        ) or 0
        onboarding = s.scalar(
            select(func.count()).select_from(Application).where(Application.status == "Onboarding")
        ) or 0
        
        total_active_engagements = len(active_engagements)
        
        total_new_apps_7d = s.scalar(
            select(func.count(Application.id))
            .where(Application.created_at >= d7)
        ) or 0
        
        active_vacancies = s.scalar(
            select(func.count(Job.id)).where(Job.status == "Open")
        ) or 0
        
        new_candidates_7 = s.scalar(select(func.count(Candidate.id)).where(Candidate.created_at >= d7)) or 0
        new_candidates_30 = s.scalar(select(func.count(Candidate.id)).where(Candidate.created_at >= d30)) or 0
        new_docs_7 = s.scalar(select(func.count(Document.id)).where(Document.uploaded_at >= d7)) or 0
        new_docs_30 = s.scalar(select(func.count(Document.id)).where(Document.uploaded_at >= d30)) or 0

        # Upcoming interviews (next 7 days)
        upcoming_interviews = s.execute(
            select(Application, Candidate, Job)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .where(Application.interview_scheduled_at != None)
            .where(Application.interview_scheduled_at >= now)
            .where(Application.interview_scheduled_at <= now + datetime.timedelta(days=7))
            .order_by(Application.interview_scheduled_at.asc())
            .limit(10)
        ).all()

        # Unsigned contracts > 3 days
        unsigned_contracts = s.execute(
            select(ESigRequest, Application, Candidate, Job)
            .join(Application, Application.id == ESigRequest.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .where(func.lower(ESigRequest.status).in_(["sent", "delivered"]))
            .where(ESigRequest.sent_at != None)
            .where(ESigRequest.sent_at < d3)
            .order_by(ESigRequest.sent_at.asc())
            .limit(10)
        ).all()

        # TrustID in progress > 3 days
        vetting_in_progress = s.execute(
            select(TrustIDCheck, Application, Candidate, Job)
            .join(Application, Application.id == TrustIDCheck.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .where(TrustIDCheck.status.in_(["Created", "InProgress"]))
            .order_by(TrustIDCheck.id.desc())
        ).all()

        # Stuck apps
        stuck_apps = s.execute(
            select(Application, Candidate, Job)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .where(Application.created_at < d7)
            .where(Application.status.in_(["New", "Screening", "Interview", "Applications Pending Review"]))
            .order_by(Application.created_at.asc())
            .limit(10)
        ).all()

        # Engagement progress tiles
        plan_rows = s.execute(
            select(EngagementPlan.engagement_id, func.sum(EngagementPlan.planned_count))
            .group_by(EngagementPlan.engagement_id)
        ).all()
        planned_map = {eng_id: int(cnt or 0) for eng_id, cnt in plan_rows}

        signed_rows = s.execute(
            select(Engagement.id, func.count(ESigRequest.id))
            .select_from(Engagement)
            .join(Job, Job.engagement_id == Engagement.id, isouter=True)
            .join(Application, Application.job_id == Job.id, isouter=True)
            .join(ESigRequest, ESigRequest.application_id == Application.id, isouter=True)
            .where(func.lower(ESigRequest.status).in_(["signed", "completed"]))
            .group_by(Engagement.id)
        ).all()
        signed_map = {eng_id: int(cnt or 0) for eng_id, cnt in signed_rows}

        engagement_progress = []
        for e in active_engagements:
            planned = planned_map.get(e.id, 0)
            signed = signed_map.get(e.id, 0)
            pct = int(100 * signed / planned) if planned else 0
            engagement_progress.append((e, planned, signed, pct))
        engagement_progress.sort(key=lambda t: t[3])
        
        # Top engagements by app count
        top_engagements = []

    return render_template(
        "dashboard.html",
        # New wireframe KPIs
        live_engagements=live_engagements,
        delivery_target=delivery_target,
        associates_offered=associates_offered,
        vetting_completed=vetting_completed,
        contracted=contracted,
        delivery_chart_data=delivery_chart_data,
        lookahead_data=lookahead_data,
        open_jobs=open_jobs,
        # Filters
        all_clients=all_clients,
        all_engagements_list=all_engagements_list,
        all_roles=all_roles,
        all_intakes=all_intakes,
        client_filter=client_filter,
        engagement_filter=engagement_filter,
        role_filter=role_filter,
        intake_filter=intake_filter,
        # Legacy data
        total_candidates=total_candidates,
        total_apps=total_apps,
        interviews=interviews,
        onboarding=onboarding,
        active_engagements=active_engagements,
        recent_jobs=recent_jobs,
        new_candidates_7=new_candidates_7,
        new_candidates_30=new_candidates_30,
        new_docs_7=new_docs_7,
        new_docs_30=new_docs_30,
        upcoming_interviews=upcoming_interviews,
        unsigned_contracts=unsigned_contracts,
        vetting_in_progress=vetting_in_progress,
        stuck_apps=stuck_apps,
        engagement_progress=engagement_progress,
        total_active_engagements=total_active_engagements,
        total_new_apps_7d=total_new_apps_7d,
        top_engagements=top_engagements,
        active_vacancies=active_vacancies,
    )

@login_required
@app.route("/action/candidate/regenerate_summary", methods=["POST"])
def candidate_regenerate_summary():
    """
    Rebuild AI summary from latest CV, tag the candidate, mirror tag names into Candidate.skills,
    and (if possible) compute/update an AI match score against the latest or provided job.
    """
    # --- Inputs
    try:
        cand_id = int(request.form.get("candidate_id") or 0)
    except Exception:
        cand_id = 0
    job_id = request.form.get("job_id")
    try:
        job_id = int(job_id) if job_id else None
    except Exception:
        job_id = None

    if not cand_id:
        flash("Missing candidate id.", "warning")
        return redirect(request.referrer or url_for("resource_pool"))

    # --- Imports kept inside to avoid circulars
    from sqlalchemy import select
    from sqlalchemy.orm import Session as SASession

    try:
        from app import (
            engine, Candidate, Application, Document, TaxonomyTag, CandidateTag, Job, Shortlist
        )
    except Exception:
        # If Shortlist doesn't exist in your app, handle later.
        Shortlist = None  # type: ignore

    # --- Fallback file-text extractor
    def _safe_extract_text(file_path: str, original_name: str) -> str:
        name = (original_name or file_path or "").lower()
        # PDF
        if name.endswith(".pdf"):
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    bits = [(p.extract_text() or "") for p in pdf.pages]
                txt = "\n".join(bits).strip()
                if txt:
                    return txt
            except Exception:
                pass
        # DOCX
        if name.endswith(".docx"):
            try:
                import docx
                d = docx.Document(file_path)
                txt = "\n".join([p.text for p in d.paragraphs]).strip()
                if txt:
                    return txt
            except Exception:
                pass
        # Plain text
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception:
            return ""

    # Resolve absolute file path from Document
    def _doc_abs_path(doc) -> str:
        # Prefer helper if your app defines it
        try:
            from app import _doc_file_path  # type: ignore
            return _doc_file_path(doc)
        except Exception:
            # Otherwise assume documents are stored beneath /static
            root = current_app.root_path
            rel = getattr(doc, "filename", "") or ""
            return os.path.join(root, "static", rel)

    # Choose the best job text for scoring
    def _job_text(job) -> str:
        for attr in ("description", "desc", "requirements", "summary", "details", "spec"):
            if hasattr(job, attr):
                v = getattr(job, attr) or ""
                v = v.strip()
                if v:
                    return v
        return (getattr(job, "title", "") or "").strip()

    # Try to use your app's AI utilities if available
    try:
        from app import extract_cv_text
    except Exception:
        extract_cv_text = None  # type: ignore
    try:
        from app import ai_summarise
    except Exception:
        ai_summarise = None  # type: ignore
    try:
        from app import ai_score_with_explanation
    except Exception:
        ai_score_with_explanation = None  # type: ignore

    with SASession(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found.", "warning")
            return redirect(url_for("resource_pool"))

        # Latest CV doc (prefer uploaded_at if present)
        order_col = getattr(Document, "uploaded_at", Document.id)
        doc = s.execute(
            select(Document)
            .where(Document.candidate_id == cand_id)
            .where(getattr(Document, "doc_type", Document.doc_type) == "cv")
            .order_by(order_col.desc())
            .limit(1)
        ).scalar_one_or_none()
        if not doc:
            flash("No CV found for this candidate.", "warning")
            return redirect(url_for("candidate_profile", cand_id=cand_id))

        # Extract text from CV
        cv_text = ""
        try:
            file_path = _doc_abs_path(doc)
            if extract_cv_text and callable(extract_cv_text):
                # Some apps accept a Document; some a file path
                try:
                    cv_text = extract_cv_text(doc) or ""
                except TypeError:
                    cv_text = extract_cv_text(file_path) or ""
            else:
                cv_text = _safe_extract_text(file_path, getattr(doc, "original_name", "") or doc.filename)
        except Exception:
            cv_text = ""

        # Build summary
        summary = ""
        try:
            if ai_summarise and callable(ai_summarise):
                summary = ai_summarise(cv_text or "") or ""
        except Exception as e:
            current_app.logger.exception("ai_summarise failed: %s", e)

        # Persist summary on latest application (and mirror to Candidate.ai_summary if present)
        latest_app = s.execute(
            select(Application)
            .where(Application.candidate_id == cand_id)
            .order_by(Application.created_at.desc())
            .limit(1)
        ).scalar_one_or_none()
        if latest_app and hasattr(latest_app, "ai_summary"):
            latest_app.ai_summary = summary
        if hasattr(cand, "ai_summary"):
            cand.ai_summary = summary

        # --- TAGS: single source of truth ---
        text_lc = (cv_text or "").lower()
        matched_ids = set()
        all_tags = s.execute(select(TaxonomyTag)).scalars().all()
        for tg in all_tags:
            token = (tg.tag or "").strip()
            if token and token.lower() in text_lc:
                matched_ids.add(tg.id)

        if matched_ids:
            existing_ids = {
                tid for (tid,) in s.execute(
                    select(CandidateTag.tag_id).where(CandidateTag.candidate_id == cand.id)
                ).all()
            }
            for tid in matched_ids:
                if tid not in existing_ids:
                    s.add(CandidateTag(candidate_id=cand.id, tag_id=tid))

            # Mirror tags into Candidate.skills (for RP search)
            all_tag_rows = s.execute(
                select(TaxonomyTag)
                .join(CandidateTag, CandidateTag.tag_id == TaxonomyTag.id)
                .where(CandidateTag.candidate_id == cand.id)
                .order_by(TaxonomyTag.tag.asc())
            ).scalars().all()
            tag_names = [t.tag for t in all_tag_rows if (t.tag or "").strip()]
            # preserve order & uniqueness
            cand.skills = ", ".join(dict.fromkeys(tag_names))

        # --- Scoring against job (if any)
        # Prefer explicit job_id from form, else use job on latest application
        target_job = None
        if job_id:
            target_job = s.get(Job, job_id)
        elif latest_app:
            target_job = s.get(Job, getattr(latest_app, "job_id", None))

        if target_job:
            job_txt = _job_text(target_job)
            if ai_score_with_explanation and job_txt and (cv_text or "").strip():
                try:
                    payload = ai_score_with_explanation(job_txt, cv_text) or {}
                    # Write to Application and Candidate if fields exist
                    score = payload.get("blended_score")
                    if latest_app is not None and hasattr(latest_app, "ai_score"):
                        latest_app.ai_score = int(score) if score is not None else None
                    if hasattr(cand, "ai_score"):
                        cand.ai_score = int(score) if score is not None else None
                    if hasattr(latest_app, "ai_explanation"):
                        latest_app.ai_explanation = payload.get("explanation")
                except Exception as e:
                    current_app.logger.exception("ai_score_with_explanation failed: %s", e)

            # Ensure shortlist if model/table exists
            if Shortlist is not None:
                try:
                    exists_short = s.execute(
                        select(Shortlist).where(
                            Shortlist.candidate_id == cand.id,
                            Shortlist.job_id == target_job.id,
                        )
                    ).scalar_one_or_none()
                    if not exists_short:
                        s.add(Shortlist(candidate_id=cand.id, job_id=target_job.id, created_at=datetime.utcnow()))
                except Exception as e:
                    current_app.logger.exception("Shortlist upsert failed: %s", e)

        s.commit()

    flash("AI summary regenerated, tags updated, and match score recalculated (when possible).", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=job_id))


@app.post("/action/score_candidate/<int:cand_id>")
def action_score_candidate(cand_id: int):
    """
    Recalculate AI match score for a candidate against a specific job (job_id required).
    Also ensures candidate is shortlisted for that job (if Shortlist table exists).
    """
    # job_id can come from form or query
    j = request.form.get("job_id") or request.args.get("job_id")
    try:
        job_id = int(j) if j else 0
    except Exception:
        job_id = 0
    if not job_id:
        flash("Missing job id for scoring.", "warning")
        return redirect(url_for("candidate_profile", cand_id=cand_id))

    from sqlalchemy import select
    from sqlalchemy.orm import Session as SASession
    try:
        from app import (
            engine, Candidate, Application, Document, Job, Shortlist
        )
    except Exception:
        Shortlist = None  # type: ignore

    # optional utilities
    try:
        from app import extract_cv_text, ai_score_with_explanation
    except Exception:
        extract_cv_text = None  # type: ignore
        ai_score_with_explanation = None  # type: ignore

    def _doc_abs_path(doc) -> str:
        try:
            from app import _doc_file_path  # type: ignore
            return _doc_file_path(doc)
        except Exception:
            root = current_app.root_path
            rel = getattr(doc, "filename", "") or ""
            return os.path.join(root, "static", rel)

    def _safe_extract_text(file_path: str, original_name: str) -> str:
        name = (original_name or file_path or "").lower()
        if name.endswith(".pdf"):
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    bits = [(p.extract_text() or "") for p in pdf.pages]
                txt = "\n".join(bits).strip()
                if txt:
                    return txt
            except Exception:
                pass
        if name.endswith(".docx"):
            try:
                import docx
                d = docx.Document(file_path)
                txt = "\n".join([p.text for p in d.paragraphs]).strip()
                if txt:
                    return txt
            except Exception:
                pass
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception:
            return ""

    def _job_text(job) -> str:
        for attr in ("description", "desc", "requirements", "summary", "details", "spec"):
            if hasattr(job, attr):
                v = getattr(job, attr) or ""
                v = v.strip()
                if v:
                    return v
        return (getattr(job, "title", "") or "").strip()

    with SASession(engine) as s:
        cand = s.get(Candidate, cand_id)
        job = s.get(Job, job_id)
        if not cand or not job:
            flash("Candidate or job not found.", "warning")
            return redirect(url_for("candidate_profile", cand_id=cand_id))

        # Latest app for this candidate (prefer same job if available)
        latest_app = s.execute(
            select(Application)
            .where(Application.candidate_id == cand_id)
            .order_by(Application.created_at.desc())
            .limit(1)
        ).scalar_one_or_none()

        # Latest CV doc
        order_col = getattr(Document, "uploaded_at", Document.id)
        doc = s.execute(
            select(Document)
            .where(Document.candidate_id == cand_id)
            .where(getattr(Document, "doc_type", Document.doc_type) == "cv")
            .order_by(order_col.desc())
            .limit(1)
        ).scalar_one_or_none()
        if not doc:
            flash("No CV found for this candidate.", "warning")
            return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=job_id))

        # Extract CV text
        cv_text = ""
        try:
            file_path = _doc_abs_path(doc)
            if extract_cv_text and callable(extract_cv_text):
                try:
                    cv_text = extract_cv_text(doc) or ""
                except TypeError:
                    cv_text = extract_cv_text(file_path) or ""
            else:
                cv_text = _safe_extract_text(file_path, getattr(doc, "original_name", "") or doc.filename)
        except Exception:
            cv_text = ""

        # Score
        wrote = False
        if ai_score_with_explanation and cv_text.strip():
            try:
                payload = ai_score_with_explanation(_job_text(job), cv_text) or {}
                score = payload.get("blended_score")
                if latest_app is not None and hasattr(latest_app, "ai_score"):
                    latest_app.ai_score = int(score) if score is not None else None
                    wrote = True
                if hasattr(cand, "ai_score"):
                    cand.ai_score = int(score) if score is not None else None
                    wrote = True
                if latest_app is not None and hasattr(latest_app, "ai_explanation"):
                    latest_app.ai_explanation = payload.get("explanation")
            except Exception as e:
                current_app.logger.exception("ai_score_with_explanation failed: %s", e)

        # Ensure shortlist if your app has it
        if Shortlist is not None:
            try:
                exists_short = s.execute(
                    select(Shortlist).where(
                        Shortlist.candidate_id == cand.id,
                        Shortlist.job_id == job.id,
                    )
                ).scalar_one_or_none()
                if not exists_short:
                    s.add(Shortlist(candidate_id=cand.id, job_id=job.id, created_at=datetime.utcnow()))
                    wrote = True
            except Exception as e:
                current_app.logger.exception("Shortlist upsert failed: %s", e)

        if wrote:
            s.commit()

    flash("AI score recalculated.", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=job_id))

# ---- Opportunities: list + create ----
@login_required
@app.route("/opportunities", methods=["GET", "POST"])
def opportunities_():
    with Session(engine) as s:
        # pull users for the "Owner" dropdown
        users_rows = s.execute(
            text("SELECT id, name FROM users ORDER BY name ASC")
        ).all()

    # build the form *after* we have users
    form = OpportunityForm()
    form.owner.choices = [(u[0], u[1]) for u in users_rows]  # (id, display)

    if form.validate_on_submit():
        # Parse DD-MM-YYYY to datetime (UK style)
        raw_start = form.est_start.data or ""
        est_start_dt = parse_date_dmy(raw_start)

        # est_value is already numeric (IntegerField), but could be None
        val = form.est_value.data if form.est_value.data is not None else 0

        # map owner_id -> friendly owner name
        owner_name = ""
        if form.owner.data:
            # form.owner.data is the user_id (int)
            with Session(engine) as s:
                row = s.execute(
                    text("SELECT name FROM users WHERE id=:i"),
                    {"i": form.owner.data},
                ).first()
                if row:
                    owner_name = row[0] or ""

        with Session(engine) as s:
            opp = Opportunity(
                name=form.name.data,
                client=form.client.data or "",
                stage=form.stage.data,
                owner=owner_name,  # storing name for now
                est_start=est_start_dt,
                est_value=int(val or 0),
                # we keep probability in db but we no longer ask user for it,
                # so default to 0 for new Leads, 100 for Closed Won, 0 for Closed Lost
                probability=100 if form.stage.data.lower() == "closed won" else 0,
                notes=form.notes.data or "",
                client_contact_name=form.client_contact_name.data or "",
                client_contact_role=form.client_contact_role.data or "",
                client_contact_phone=form.client_contact_phone.data or "",
                client_contact_email=form.client_contact_email.data or "",
            )
            s.add(opp)
            s.flush()  # now opp.id exists

            # if they immediately saved it as Closed Won, auto-create engagement
            if (opp.stage or "").strip().lower() == "closed won":
                e = create_engagement_for_opportunity(s, opp)
                if e:
                    flash(f"Engagement {e.ref} created from Closed Won opportunity.", "success")

            s.commit()

        flash("Opportunity created", "success")
        return redirect(url_for("opportunities_"))

    # GET render
    with Session(engine) as s:
        rows = s.scalars(
            select(Opportunity).order_by(Opportunity.created_at.desc())
        ).all()

        # Attach engagement links just like before
        eng_links = {
            oid: (eid, ref)
            for (eid, ref, oid) in s.execute(
                text("SELECT id, ref, opportunity_id FROM engagements WHERE opportunity_id IS NOT NULL")
            ).all()
        }
        for o in rows:
            tup = eng_links.get(o.id)
            if tup:
                o._engagement_id, o._engagement_ref = tup[0], tup[1]
            else:
                o._engagement_id = None
                o._engagement_ref = None

    # We are intentionally NOT calculating pipeline tiles anymore
    return render_template(
        "opportunities_.html",
        form=form,
        items=rows,
    )

@login_required
@app.route("/admin/roles", methods=["GET", "POST"])
def admin_roles():
    form = RoleTypeForm()
    with Session(engine) as s:
        if form.validate_on_submit():
            rt = RoleType(
                name=form.name.data.strip(),
                default_rate=int(form.default_rate.data or 0),
            )
            try:
                s.add(rt)
                s.commit()
                flash("Role added", "success")
            except IntegrityError:
                s.rollback()
                flash("That role already exists", "warning")

        roles = s.scalars(
            select(RoleType).order_by(RoleType.name.asc())
        ).all()

    return render_template("admin_roles.html", form=form, roles=roles)

@login_required
@app.route("/sumsub/access_token/<external_user_id>", methods=["POST", "GET"])
def sumsub_access_token(external_user_id: str):
    """
    Generate a temporary access token for Sumsub WebSDK Next.
    Works with both GET and POST for easy testing.

    Query params:
      - ttlInSecs (default 600)
      - levelName (optional but recommended; e.g. 'id-and-liveness' or 'basic-kyb-level')
    """
    try:
        ttl = request.args.get("ttlInSecs", "600")
        level_name = (request.args.get("levelName") or "").strip()

        from urllib.parse import quote_plus as q
        path = f"/resources/accessTokens?userId={q(external_user_id)}&ttlInSecs={q(ttl)}"
        if level_name:
            path += f"&levelName={q(level_name)}"

        headers = sumsub_sign_request("POST", path, "")
        resp = requests.post(f"{SUMSUB_BASE_URL}{path}", headers=headers, timeout=15)

        content_type = (resp.headers.get("Content-Type") or "")
        if "application/json" in content_type.lower():
            return jsonify(resp.json()), resp.status_code
        return (resp.text, resp.status_code, {"Content-Type": "text/plain; charset=utf-8"})
    except Exception as e:
        current_app.logger.exception("Sumsub access token failed")
        return jsonify({"error": "sumsub_access_token_failed", "detail": str(e)}), 500

@login_required
@app.route("/sumsub/create_applicant", methods=["POST"])
def sumsub_create_applicant():
    """
    Create a new Sumsub applicant.

    Send JSON:
      {
        "externalUserId": "test_user_001",
        "levelName": "id-and-liveness"   # optional; falls back to env SUMSUB_LEVEL_NAME or first available
      }

    Notes:
    - Sumsub requires `levelName` in the QUERY STRING (and included in the signature).
    - We validate levelName against your account’s levels and set the correct applicant type.
    """
    from urllib.parse import quote_plus

    try:
        data = request.get_json(silent=True) or {}
        external_user_id = (data.get("externalUserId") or f"user_{int(time.time())}").strip()
        requested_level = (data.get("levelName") or os.getenv("SUMSUB_LEVEL_NAME", "")).strip()

        # ---- 1) Fetch levels from Sumsub and validate ----
        levels_path = "/resources/applicants/-/levels"
        lv_headers = sumsub_sign_request("GET", levels_path, "")
        lv_resp = requests.get(f"{SUMSUB_BASE_URL}{levels_path}", headers=lv_headers, timeout=15)

        levels_json = {}
        try:
            levels_json = lv_resp.json() if lv_resp.headers.get("Content-Type","").lower().startswith("application/json") else {}
        except Exception:
            levels_json = {}

        # Normalise to a simple list of dicts with 'name' & 'applicantType'
        items = []
        if isinstance(levels_json, dict) and "list" in levels_json and isinstance(levels_json["list"], dict):
            items = levels_json["list"].get("items", []) or []
        elif isinstance(levels_json, dict) and "items" in levels_json:
            items = levels_json.get("items", []) or []

        available = [(it.get("name"), (it.get("applicantType") or it.get("type") or "individual")) for it in items if it.get("name")]
        available_names = [n for n, _ in available]

        # If no requested level, pick first available
        if not requested_level:
            if not available_names:
                return jsonify({
                    "error": "sumsub_create_applicant_failed",
                    "detail": "No levels available in your Sumsub project. Create a level in the Sumsub dashboard first."
                }), 400
            requested_level = available_names[0]

        # Validate requested level
        match = None
        for n, a_type in available:
            if n == requested_level:
                match = (n, a_type)
                break

        if not match:
            return jsonify({
                "error": "invalid_level_name",
                "detail": f"Level '{requested_level}' not found.",
                "availableLevels": available_names
            }), 400

        level_name, applicant_type = match  # applicant_type: 'individual' or 'company'

        # ---- 2) Build payload with the correct applicant type ----
        payload = {
            "externalUserId": external_user_id,
            "type": "company" if str(applicant_type).lower() == "company" else "individual"
        }
        body_str = json.dumps(payload, separators=(",", ":"))

        # ---- 3) Call Sumsub create with levelName in query (and sign including the query) ----
        q_level = quote_plus(level_name)
        path = f"/resources/applicants?levelName={q_level}"

        headers = sumsub_sign_request("POST", path, body_str)
        headers["Content-Type"] = "application/json"

        resp = requests.post(
            f"{SUMSUB_BASE_URL}{path}",
            data=body_str,  # send EXACT bytes we signed
            headers=headers,
            timeout=20,
        )

        ctype = (resp.headers.get("Content-Type") or "").lower()
        if "application/json" in ctype and (resp.text or "").strip():
            try:
                return jsonify(resp.json()), resp.status_code
            except Exception:
                pass

        # Always return a JSON summary for non-JSON/empty bodies
        return jsonify({
            "status": resp.status_code,
            "reason": resp.reason,
            "headers": dict(resp.headers or {}),
            "text": (resp.text or "").strip()
        }), resp.status_code

    except Exception as e:
        current_app.logger.exception("Sumsub create_applicant failed")
        return jsonify({"error": "sumsub_create_applicant_failed", "detail": str(e)}), 500

# Debug routes disabled in production demo
# @app.route("/__ping")
# def __ping():
#     return jsonify({
#         "ok": True,
#         "loaded_from": __file__,
#         "has_sumsub_create": "sumsub_create_applicant" in current_app.view_functions,
#         "routes_count": len(list(current_app.url_map.iter_rules())),
#     })

# @app.route("/__routes")
# def __routes():
#     routes = []
#     for r in current_app.url_map.iter_rules():
#         routes.append({
#             "rule": str(r),
#             "endpoint": r.endpoint,
#             "methods": sorted(m for m in r.methods if m not in {"HEAD", "OPTIONS"}),
#         })
#     # sort for readability
#     routes.sort(key=lambda x: x["rule"])
    return jsonify({"routes": routes})

@login_required
@app.route("/sumsub/levels", methods=["GET"])
def sumsub_list_levels():
    """
    Lists all verification levels available in your Sumsub project.
    """
    try:
        path = "/resources/applicants/-/levels"
        headers = sumsub_sign_request("GET", path, "")
        resp = requests.get(f"{SUMSUB_BASE_URL}{path}", headers=headers, timeout=15)

        ctype = (resp.headers.get("Content-Type") or "").lower()
        if "application/json" in ctype and (resp.text or "").strip():
            return jsonify(resp.json()), resp.status_code

        # Always return JSON, even if Sumsub sends text/empty
        return jsonify({
            "status": resp.status_code,
            "reason": resp.reason,
            "headers": dict(resp.headers or {}),
            "text": (resp.text or "").strip()
        }), resp.status_code

    except Exception as e:
        current_app.logger.exception("Sumsub list levels failed")
        return jsonify({"error": "sumsub_list_levels_failed", "detail": str(e)}), 500

# --- Sumsub WebSDK launcher page ---
@login_required
@app.route("/sumsub/start/<external_user_id>", methods=["GET"])
def sumsub_start(external_user_id: str):
    """
    Simple page that hosts Sumsub WebSDK for a given externalUserId.
    Query params (all optional):
      ttlInSecs   -> defaults 600
      levelName   -> if omitted, we’ll pass nothing (Sumsub falls back to applicant's level)
      theme       -> 'light' or 'dark' (default 'light')
      lang        -> 'en' by default
    """
    try:
        ttl = int(request.args.get("ttlInSecs") or 600)
    except Exception:
        ttl = 600

    level_name = (request.args.get("levelName") or os.getenv("SUMSUB_LEVEL_NAME", "")).strip()
    theme      = (request.args.get("theme") or "light").strip() or "light"
    lang       = (request.args.get("lang") or "en").strip() or "en"

    # Nice label in the header so you know you’re on sandbox
    label = "Sumsub Sandbox" if "sandbox" in (SUMSUB_BASE_URL or "").lower() else "Sumsub"

    return render_template(
        "sumsub_start.html",
        external_user_id=external_user_id,
        ttl=ttl,
        level_name=level_name,
        theme=theme,
        lang=lang,
        sumsub_label=label,
    )

@login_required
@app.route("/sumsub/applicant/<applicant_id>", methods=["GET"])
def sumsub_get_applicant(applicant_id: str):
    """Fetch applicant object from Sumsub."""
    try:
        path = f"/resources/applicants/{applicant_id}"
        headers = sumsub_sign_request("GET", path, "")
        resp = requests.get(f"{SUMSUB_BASE_URL}{path}", headers=headers, timeout=15)

        content_type = (resp.headers.get("Content-Type") or "")
        if "application/json" in content_type.lower():
            return jsonify(resp.json()), resp.status_code
        else:
            return (resp.text, resp.status_code, {"Content-Type": "text/plain; charset=utf-8"})
    except Exception as e:
        current_app.logger.exception("Sumsub get applicant failed")
        return jsonify({"error": "sumsub_get_applicant_failed", "detail": str(e)}), 500

# --- Convert Opportunity → Engagement ---
@login_required
@app.route("/opportunities/<int:opp_id>/convert", methods=["POST"], endpoint="opportunity_convert_v2")
def opportunity_convert_v2(opp_id):
    with Session(engine) as s:
        opp = s.get(Opportunity, opp_id)
        if not opp:
            abort(404)

        # only Closed Won can convert
        if (opp.stage or "").lower() != "closed won":
            flash("Only Closed Won opportunities can be converted.", "warning")
            return redirect(url_for("opportunities_"))

        # already converted?
        existing = s.scalar(select(Engagement).where(Engagement.opportunity_id == opp.id))
        if existing:
            flash(f"Engagement {existing.ref or existing.id} already exists.", "info")
            return redirect(url_for("engagement_dashboard", eng_id=existing.id))

        e = create_engagement_for_opportunity(s, opp)
        s.commit()
        flash(f"Engagement {e.ref or e.id} created.", "success")
        return redirect(url_for("engagement_dashboard", eng_id=e.id))

# GAP 4.1: API endpoint for converting opportunity to engagement
@login_required
@app.route("/api/opportunity/<int:opp_id>/convert", methods=["POST"])
def api_opportunity_convert(opp_id):
    """
    GAP 4.1: Convert opportunity to engagement via API.
    Allows conversion when stage is Proposal, Negotiation, or Won.
    """
    with Session(engine) as s:
        opp = s.get(Opportunity, opp_id)
        if not opp:
            return jsonify({"ok": False, "error": "Opportunity not found"}), 404
        
        # Allow conversion at Proposal, Negotiation, or Won stages
        allowed_stages = ["proposal", "negotiation", "won", "closed won"]
        current_stage = (opp.stage or "").lower()
        
        if current_stage not in allowed_stages:
            return jsonify({
                "ok": False, 
                "error": f"Cannot convert at stage '{opp.stage}'. Allowed stages: Proposal, Negotiation, Won"
            }), 400
        
        # Check if already converted
        existing = s.scalar(select(Engagement).where(Engagement.opportunity_id == opp.id))
        if existing:
            return jsonify({
                "ok": True,
                "message": f"Engagement {existing.ref or existing.id} already exists",
                "engagement_id": existing.id,
                "engagement_ref": existing.ref
            })
        
        # Create the engagement
        e = create_engagement_for_opportunity(s, opp)
        
        # Update opportunity stage to Won if converting from Proposal/Negotiation
        if current_stage in ["proposal", "negotiation"]:
            opp.stage = "Won"
        
        # Audit log
        log_audit_event(
            'create', 'engagement',
            f'Converted opportunity "{opp.name}" to engagement "{e.ref}"',
            'engagement', e.id,
            {'opportunity_id': opp.id, 'opportunity_name': opp.name}
        )
        
        s.commit()
        
        return jsonify({
            "ok": True,
            "message": f"Engagement {e.ref or e.id} created successfully",
            "engagement_id": e.id,
            "engagement_ref": e.ref
        })

# API endpoint for moving opportunity between pipeline stages (drag-and-drop)
@csrf.exempt
@login_required
@app.route("/api/opportunity/move-stage", methods=["POST"])
def api_opportunity_move_stage():
    """
    Move an opportunity to a different pipeline stage.
    Used by the drag-and-drop Kanban on the Projects page.
    """
    data = request.get_json()
    if not data:
        return jsonify({"ok": False, "error": "No data provided"}), 400
    
    opp_id = data.get("opportunity_id")
    from_stage = data.get("from_stage")
    to_stage = data.get("to_stage")
    
    if not opp_id or not to_stage:
        return jsonify({"ok": False, "error": "Missing required fields"}), 400
    
    with Session(engine) as s:
        opp = s.get(Opportunity, int(opp_id))
        if not opp:
            return jsonify({"ok": False, "error": "Opportunity not found"}), 404
        
        # Valid pipeline stages
        valid_stages = ["Qualified Lead", "Proposal", "Procurement", "Closed Won", "Closed Lost"]
        if to_stage not in valid_stages:
            return jsonify({"ok": False, "error": f"Invalid stage: {to_stage}"}), 400
        
        old_stage = opp.stage
        opp.stage = to_stage
        
        # Audit log: opportunity stage move
        log_audit_event('update', 'workflow',
                       f'Pipeline stage changed: {old_stage} → {to_stage}',
                       'opportunity', opp.id,
                       {'old_stage': old_stage, 'new_stage': to_stage,
                        'opportunity_name': opp.name, 'client': opp.client})
        
        # If moved to Closed Won, optionally create engagement
        if to_stage == "Closed Won":
            # Check if engagement already exists
            existing_eng = s.scalar(select(Engagement).where(Engagement.opportunity_id == opp.id))
            if not existing_eng:
                # Auto-create engagement for Closed Won
                eng = _create_engagement_from_opportunity(s, opp)
                s.commit()
                return jsonify({
                    "ok": True, 
                    "message": f"Moved to {to_stage} and created engagement {eng.ref}",
                    "engagement_created": True,
                    "engagement_id": eng.id,
                    "engagement_ref": eng.ref
                })
        
        s.commit()
        return jsonify({
            "ok": True, 
            "message": f"Opportunity moved from {old_stage} to {to_stage}",
            "old_stage": old_stage,
            "new_stage": to_stage
        })

@login_required
@app.route("/opportunity/<int:opp_id>/convert_to_engagement", methods=["POST"])
def opportunity_convert_to_engagement(opp_id):
    with Session(engine) as s:
        opp = s.get(Opportunity, opp_id)
        if not opp:
            abort(404)
        if (opp.stage or "").lower() != "closed won":
            flash("Only Closed Won opportunities can be converted.", "warning")
            return redirect(url_for("opportunity_edit", opp_id=opp_id))

        eng = _create_engagement_from_opportunity(s, opp)
        s.commit()
        flash(f"Engagement {eng.ref} created.", "success")
        return redirect(url_for("engagement_dashboard", eng_id=eng.id))

@login_required
@app.route("/admin/opportunities/backfill_engagements", methods=["POST"])
def backfill_engagements_for_won():
    created = 0
    with Session(engine) as s:
        won = s.scalars(select(Opportunity).where(Opportunity.stage=="Closed Won")).all()
        for opp in won:
            e = s.scalar(select(Engagement).where(Engagement.opportunity_id==opp.id))
            if not e:
                _create_engagement_from_opportunity(s, opp)
                created += 1
        s.commit()
    flash(f"Created {created} engagements from Closed Won opportunities.", "success")
    return redirect(url_for("opportunities_"))

# ---- Opportunities: edit ----
@login_required
@app.route("/opportunity/<int:opp_id>", methods=["GET", "POST"])
def opportunity_edit(opp_id):
    # Load the opportunity row
    with Session(engine) as s:
        opp = s.scalar(select(Opportunity).where(Opportunity.id == opp_id))
        if not opp:
            abort(404)

        # Load owners for dropdown
        users_rows = s.execute(
            text("SELECT id, name FROM users ORDER BY name ASC")
        ).all()
        choices = [(u[0], u[1]) for u in users_rows]

        # We'll need to map owner name -> id for initial form data
        owner_id_for_form = None
        if opp.owner:
            for uid, uname in choices:
                if (uname or "").strip() == (opp.owner or "").strip():
                    owner_id_for_form = uid
                    break

        # Prefill form
        form = OpportunityForm(
            name=opp.name,
            client=opp.client,
            stage=opp.stage,
            owner=owner_id_for_form,
            est_start=format_date_dmy(opp.est_start),  # "DD-MM-YYYY"
            est_value=opp.est_value or 0,
            notes=opp.notes or "",
            client_contact_name = opp.client_contact_name or "",
            client_contact_role = opp.client_contact_role or "",
            client_contact_phone = opp.client_contact_phone or "",
            client_contact_email = opp.client_contact_email or "",
        )
        form.owner.choices = choices

        if form.validate_on_submit():
            prev_stage = (opp.stage or "").strip().lower()

            # Map owner dropdown (user id) back to owner name string for storage
            new_owner_name = ""
            if form.owner.data:
                row = s.execute(
                    text("SELECT name FROM users WHERE id=:i"),
                    {"i": form.owner.data},
                ).first()
                if row:
                    new_owner_name = row[0] or ""

            opp.name = form.name.data
            opp.client = form.client.data or ""
            opp.stage = form.stage.data
            opp.owner = new_owner_name
            opp.est_start = parse_date_dmy(form.est_start.data)
            opp.est_value = int(form.est_value.data or 0)

            # probability is now implicit; we keep writing it for downstream logic
            opp.probability = 100 if (opp.stage or "").strip().lower() == "closed won" else 0

            opp.notes = form.notes.data or ""

            opp.client_contact_name  = form.client_contact_name.data  or ""
            opp.client_contact_role  = form.client_contact_role.data  or ""
            opp.client_contact_phone = form.client_contact_phone.data or ""
            opp.client_contact_email = form.client_contact_email.data or ""

            # If it becomes Closed Won, create/link an engagement (idempotent)
            new_stage = (opp.stage or "").strip().lower()
            e = None
            if new_stage == "closed won" and prev_stage != "closed won":
                e = create_engagement_for_opportunity(s, opp)

            # Also handle the case where it was already Closed Won but not linked yet
            if new_stage == "closed won" and not e:
                e = create_engagement_for_opportunity(s, opp)

            s.commit()
            flash("Opportunity updated", "success")

            if e:
                flash(f"Engagement {e.ref} created.", "success")
                return redirect(url_for("engagement_dashboard", eng_id=e.id))

            return redirect(url_for("opportunities_"))

    return render_template("opportunity_edit.html", form=form, opp=opp)

@app.route(
    "/action/onboarding_email/<int:app_id>",
    methods=["POST"],
    endpoint="action_onboarding_email"
)
def action_onboarding_email(app_id):
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)

        cand = s.scalar(select(Candidate).where(Candidate.id == appn.candidate_id))
        job  = s.scalar(select(Job).where(Job.id == appn.job_id))

        link = f"{APP_BASE_URL}/application/{appn.id}"
        html = f"""
        <h3>Welcome to onboarding</h3>
        <p>Hi {cand.name}, thanks for applying for <b>{job.title}</b>. We'll guide you through vetting next.</p>
        <p>You can check status here: <a href="{link}">{link}</a></p>
        """

        # Send email (no-op in dev if SMTP_* not configured)
        send_email(cand.email, f"Onboarding for {job.title}", html)

        # Update flags
        appn.onboarding_email_sent = True
        appn.status = "Accepted"  # Candidate has accepted, onboarding/vetting begins
        s.commit()

    flash("Onboarding email sent", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@app.route("/kanban")
@login_required
def kanban():
    columns = ["New","Screening","Interview","Offer","Onboarding","Hired","Rejected"]
    with Session(engine) as s:
        data = {col: s.execute(select(Application, Candidate, Job)
                               .join(Candidate, Candidate.id==Application.candidate_id)
                               .join(Job, Job.id==Application.job_id)
                               .where(Application.status==col)
                               .order_by(Application.created_at.desc())
                               ).all() for col in columns}
    return render_template("kanban.html", columns=columns, data=data)

@login_required
@app.route("/kanban/move", methods=["POST"])
def kanban_move():
    payload = request.json or {}
    app_id = int(payload.get("app_id"))
    new_status = payload.get("new_status")
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id==app_id))
        if not appn:
            abort(404)
        old_status = appn.status
        appn.status = new_status
        
        # Get candidate name for audit log
        cand = s.scalar(select(Candidate).where(Candidate.id == appn.candidate_id))
        candidate_name = cand.name if cand else None
        candidate_id = appn.candidate_id
        
        s.commit()
    
    # Audit log: kanban move (after commit to avoid db lock)
    log_audit_event('update', 'workflow', 
                   f'Kanban move: {old_status} → {new_status}',
                   'application', app_id,
                   {'old_status': old_status, 'new_status': new_status,
                    'candidate_id': candidate_id, 'candidate_name': candidate_name})
    
    return jsonify({"ok": True})

# ============== WORKFLOW PAGE ==============
@login_required
@app.route("/workflow")
def workflow():
    """
    Workflow page showing headline figures for each stage and draggable associate cards.
    Filters: owner, client, engagement, role, intake
    
    9 Workflow Stages:
    1) Pipeline - Initial applications received
    2) Shortlist - CV/App marked as passed
    3) I&A - Tele-screen marked as passed (Interview & Assessment)
    4) Client Review - Internal interviews/assessments passed
    5) Offered - Client interview/assessment passed, offer made
    6) Accepted - Offer accepted by candidate
    7) Ready to Contract - Vetting marked as complete
    8) Contract Sent - Placement created, contract sent via eSignature
    9) Placed - Contract signed
    """
    # Query params for filters
    eng_status_filter = request.args.get("eng_status", "active")  # 'active' or 'all'
    owner_filter = request.args.get("owner", "all")
    client_filter = request.args.get("client", "all")
    engagement_filter = request.args.get("engagement", "all")
    role_filter = request.args.get("role", "all")
    intake_filter = request.args.get("intake", "all")
    
    # New 10-stage workflow (including Rejected/Withdrawn as 6th stage)
    stages = [
        {"id": "pipeline", "name": "Pipeline", "short": "Pipeline", "color": "#6366f1", "icon": "fa-inbox"},
        {"id": "shortlist", "name": "Shortlist", "short": "Shortlist", "color": "#8b5cf6", "icon": "fa-star"},
        {"id": "ia", "name": "I&A", "short": "I&A", "color": "#f59e0b", "icon": "fa-phone-alt"},
        {"id": "client_review", "name": "Client Review", "short": "Client Review", "color": "#f97316", "icon": "fa-user-tie"},
        {"id": "offered", "name": "Offered", "short": "Offered", "color": "#ec4899", "icon": "fa-hand-holding-usd"},
        {"id": "rejected_withdrawn", "name": "Rejected/Withdrawn", "short": "Rej/Withdrawn", "color": "#ef4444", "icon": "fa-times-circle"},
        {"id": "accepted", "name": "Accepted", "short": "Accepted", "color": "#14b8a6", "icon": "fa-thumbs-up"},
        {"id": "ready_to_contract", "name": "Ready to Contract", "short": "Ready to Contract", "color": "#3b82f6", "icon": "fa-shield-alt"},
        {"id": "contract_sent", "name": "Contract Sent", "short": "Contract Sent", "color": "#10b981", "icon": "fa-file-signature"},
        {"id": "placed", "name": "Placed", "short": "Placed", "color": "#22c55e", "icon": "fa-check-circle"},
    ]
    
    with Session(engine) as s:
        # Get unique clients from engagements
        all_clients = s.scalars(
            select(Engagement.client)
            .distinct()
            .where(Engagement.client.isnot(None), Engagement.client != "")
            .order_by(Engagement.client)
        ).all()
        
        # Get engagements based on status filter
        eng_query = select(Engagement).order_by(Engagement.name)
        if eng_status_filter == "active":
            eng_query = eng_query.where(Engagement.status == "Active")
        all_engagements = s.scalars(eng_query).all()
        
        # Get all jobs for filter dropdown
        all_jobs = s.scalars(
            select(Job)
            .where(Job.title.isnot(None), Job.title != "")
            .order_by(Job.title)
        ).all()
        
        # Get unique roles (job titles) from jobs
        all_roles = s.scalars(
            select(Job.title)
            .distinct()
            .where(Job.title.isnot(None), Job.title != "")
            .order_by(Job.title)
        ).all()
        
        # Get unique intake dates from engagement plans
        all_intakes = []
        try:
            intake_dates = s.scalars(
                select(EngagementPlan.intake_date)
                .distinct()
                .where(EngagementPlan.intake_date.isnot(None))
                .order_by(EngagementPlan.intake_date)
            ).all()
            all_intakes = [d.strftime("%Y-%m-%d") if d else None for d in intake_dates if d]
        except Exception:
            pass
        
        # Get unique owners (users assigned to engagements or created jobs)
        all_owners = s.scalars(
            select(User.email)
            .distinct()
            .order_by(User.email)
        ).all()
        
        # Build base query for applications
        # Map legacy statuses to new workflow stages
        STATUS_TO_STAGE_MAP = {
            # Pipeline - Initial applications
            "New": "pipeline",
            "Pending": "pipeline",
            "Applied": "pipeline",
            "Applications Pending Review": "pipeline",
            "Pipeline": "pipeline",
            
            # Shortlist - CV/App passed
            "Shortlist": "shortlist",
            "Shortlisted": "shortlist",
            "CV Passed": "shortlist",
            "Declared": "shortlist",
            
            # I&A - Tele-screen passed (Interview & Assessment)
            "I&A": "ia",
            "Tele-screen Passed": "ia",
            "Interview Scheduled": "ia",
            "Interview": "ia",
            
            # Client Review - Internal interviews/assessments passed
            "Client Review": "client_review",
            "Interview Completed": "client_review",
            "Interviewed": "client_review",
            "Assessment Complete": "client_review",
            
            # Offered - Client interview passed, offer made
            "Offered": "offered",
            "Offer Made": "offered",
            "Offer Sent": "offered",
            
            # Accepted - Offer accepted
            "Accepted": "accepted",
            "Offer Accepted": "accepted",
            
            # Rejected/Withdrawn - Offer declined or candidate withdrew
            "Rejected": "rejected_withdrawn",
            "Withdrawn": "rejected_withdrawn",
            "Rejected/Withdrawn": "rejected_withdrawn",
            "Offer Declined": "rejected_withdrawn",
            "Declined": "rejected_withdrawn",
            
            # Ready to Contract - Vetting complete
            "Ready to Contract": "ready_to_contract",
            "Vetting Complete": "ready_to_contract",
            "Vetting": "ready_to_contract",
            "Vetting In-Flight": "ready_to_contract",
            "Onboarding": "ready_to_contract",
            
            # Contract Sent - Placement created, contract sent
            "Contract Sent": "contract_sent",
            "Contract Issued": "contract_sent",
            
            # Placed - Contract signed
            "Placed": "placed",
            "Contract Signed": "placed",
            "Contracted": "placed",
            "Active": "placed",
            "Hired": "placed",
        }
        
        def build_query_for_stage(stage_id):
            """Build query for applications matching a workflow stage"""
            # Find all statuses that map to this stage
            matching_statuses = [status for status, sid in STATUS_TO_STAGE_MAP.items() if sid == stage_id]
            
            if not matching_statuses:
                # Fallback: use the stage name directly
                matching_statuses = [next((s["name"] for s in stages if s["id"] == stage_id), stage_id)]
            
            query = (
                select(Application, Candidate, Job)
                .join(Candidate, Candidate.id == Application.candidate_id)
                .join(Job, Job.id == Application.job_id)
                .outerjoin(Engagement, Engagement.id == Job.engagement_id)
                .where(Application.status.in_(matching_statuses))
            )
            
            # Apply engagement status filter
            if eng_status_filter == "active":
                query = query.where(Engagement.status == "Active")
            
            # Apply filters
            if engagement_filter != "all" and engagement_filter.isdigit():
                query = query.where(Job.engagement_id == int(engagement_filter))
            
            if client_filter != "all":
                query = query.where(Engagement.client == client_filter)
            
            if role_filter != "all":
                query = query.where(Job.title == role_filter)
            
            # Owner filter - filter by job creator or engagement owner
            if owner_filter != "all":
                query = query.outerjoin(User, User.id == Job.created_by).where(
                    or_(User.email == owner_filter, Engagement.owner == owner_filter)
                )
            
            # Intake filter - filter by engagement plan intake date
            if intake_filter != "all":
                query = query.outerjoin(EngagementPlan, EngagementPlan.engagement_id == Engagement.id).where(
                    func.date(EngagementPlan.intake_date) == intake_filter
                )
            
            return query.order_by(Application.created_at.desc())
        
        # Helper function to get vetting progress for a candidate
        SLA_DAYS = 7  # 7 day SLA for vetting completion
        
        def get_vetting_progress(candidate_id):
            """Get vetting check progress for a candidate with SLA status"""
            try:
                vetting_stats = s.execute(
                    text("""
                        SELECT 
                            COUNT(*) as total,
                            SUM(CASE WHEN status = 'Complete' THEN 1 ELSE 0 END) as complete,
                            SUM(CASE WHEN status = 'In Progress' THEN 1 ELSE 0 END) as in_progress,
                            SUM(CASE WHEN status IN ('Not Started', 'NOT STARTED') THEN 1 ELSE 0 END) as not_started,
                            SUM(CASE WHEN status = 'Not Required' THEN 1 ELSE 0 END) as not_required,
                            MIN(created_at) as first_check_date
                        FROM vetting_check
                        WHERE candidate_id = :cand_id
                    """),
                    {"cand_id": candidate_id}
                ).first()
                
                if vetting_stats and vetting_stats[0] > 0:
                    total = vetting_stats[0] or 0
                    complete = vetting_stats[1] or 0
                    in_progress = vetting_stats[2] or 0
                    not_started = vetting_stats[3] or 0
                    not_required = vetting_stats[4] or 0
                    first_check_date = vetting_stats[5]
                    
                    # Calculate effective total (excluding not required)
                    effective_total = total - not_required
                    percentage = (complete / effective_total * 100) if effective_total > 0 else 0
                    
                    # Determine RAG status based on completion:
                    # 0/10 = red (Not Started), 1-9/10 = amber (In Progress), 10/10 = green (Complete)
                    if effective_total == 0:
                        rag = "grey"  # No checks required
                    elif complete == effective_total:
                        rag = "green"  # All complete (10/10)
                    elif complete > 0:
                        rag = "amber"  # Some progress (1-9/10)
                    else:
                        rag = "red"  # Not started (0/10)
                    
                    # Calculate SLA status
                    days_elapsed = 0
                    sla_status = "in"  # Default to in SLA
                    if first_check_date:
                        if isinstance(first_check_date, str):
                            try:
                                first_check_date = datetime.datetime.fromisoformat(first_check_date.replace('Z', '+00:00').split('+')[0])
                            except:
                                first_check_date = None
                        if first_check_date:
                            days_elapsed = (now - first_check_date).days
                            # Only check SLA for non-complete candidates
                            if rag != "green" and days_elapsed > SLA_DAYS:
                                sla_status = "out"
                    
                    return {
                        "total": total,
                        "complete": complete,
                        "in_progress": in_progress,
                        "not_started": not_started,
                        "not_required": not_required,
                        "effective_total": effective_total,
                        "percentage": round(percentage),
                        "rag": rag,
                        "days_elapsed": days_elapsed,
                        "sla_status": sla_status,
                        "sla_days": SLA_DAYS
                    }
            except Exception:
                pass
            return None
        
        # Build data for each stage
        # GAP 1.3: Enhanced card content with days_in_stage calculation
        stage_data = {}
        now = datetime.datetime.utcnow()
        for stage in stages:
            results = s.execute(build_query_for_stage(stage["id"])).all()
            stage["count"] = len(results)
            
            # Build flattened card objects for the template
            cards = []
            for app, cand, job in results:
                engagement = job.engagement if job.engagement_id else None
                
                # Intake is driven by the engagement - default to 1 if no evidence of more
                intake = 1
                
                cards.append({
                    "id": app.id,
                    "type": "application",
                    "candidate_id": cand.id,
                    "job_id": job.id if job else None,
                    "engagement_id": engagement.id if engagement else None,
                    "name": cand.name or "Unknown",
                    "email": cand.email or "",
                    "phone": cand.phone or "",
                    "location": cand.location or cand.postcode or "",
                    "job_title": job.title if job else "",
                    "ai_score": app.ai_score or 0,
                    "interview_at": app.interview_scheduled_at,
                    "engagement_name": engagement.name if engagement else "",
                    "engagement_ref": engagement.ref if engagement else "",
                    "client": engagement.client if engagement else "",
                    "start_date": engagement.start_date if engagement else None,
                    "intake": intake,
                    "days_in_stage": (now - app.created_at).days if app.created_at else 0,
                    "vetting_progress": get_vetting_progress(cand.id),
                })
            stage_data[stage["id"]] = cards
        
        # Calculate KPI totals
        kpi_counts = {stage["name"]: stage["count"] for stage in stages}
        
        # Get unsigned contracts (issued > 3 days ago but not signed)
        # Status is 'sent' or 'delivered' for contracts awaiting signature
        three_days_ago = datetime.datetime.now() - datetime.timedelta(days=3)
        unsigned_contracts = s.execute(
            select(ESigRequest, Candidate, Engagement)
            .join(Candidate, Candidate.id == ESigRequest.candidate_id)
            .outerjoin(Engagement, Engagement.id == ESigRequest.engagement_id)
            .where(
                func.lower(ESigRequest.status).in_(["sent", "delivered", "pending"]),
                ESigRequest.created_at < three_days_ago
            )
            .order_by(ESigRequest.created_at.asc())
        ).all()
        
        unsigned_contracts_data = [
            {
                "esig": esig,
                "candidate": cand,
                "engagement": eng,
                "days_pending": (datetime.datetime.now() - esig.created_at).days if esig.created_at else 0
            }
            for esig, cand, eng in unsigned_contracts
        ]
        
        # Get vetting summary stats
        # Count candidates by their overall vetting progress:
        # - 0/10 complete = Not Started
        # - 1-9/10 complete = In Progress  
        # - 10/10 complete = Complete
        # Also track SLA (7 day SLA from first vetting check creation)
        
        vetting_total = 0
        vetting_complete = 0
        vetting_in_progress = 0
        vetting_not_started = 0
        vetting_in_sla = 0
        vetting_out_of_sla = 0
        vetting_by_type = {}
        vetting_by_type_in_sla = {}
        vetting_by_type_out_sla = {}
        vetting_candidates = []  # List of candidates with their SLA status
        
        # SLA-filtered counts
        vetting_not_started_in_sla = 0
        vetting_not_started_out_sla = 0
        vetting_in_progress_in_sla = 0
        vetting_in_progress_out_sla = 0
        
        SLA_DAYS = 7  # 7 day SLA for vetting completion
        
        try:
            # Get per-candidate vetting summary with SLA calculation
            candidate_vetting = s.execute(
                text("""
                    SELECT 
                        c.id,
                        c.name,
                        COUNT(vc.id) as total_checks,
                        SUM(CASE WHEN vc.status = 'Complete' THEN 1 ELSE 0 END) as complete_checks,
                        MIN(vc.created_at) as first_check_date
                    FROM candidates c
                    JOIN vetting_check vc ON vc.candidate_id = c.id
                    GROUP BY c.id, c.name
                """)
            ).all()
            
            now = datetime.datetime.utcnow()
            
            for row in candidate_vetting:
                cand_id, cand_name, total, complete, first_check_date = row
                vetting_total += 1
                
                # Determine status based on completion ratio
                if complete == 0:
                    vetting_not_started += 1
                    status = "not_started"
                elif complete >= total:
                    vetting_complete += 1
                    status = "complete"
                else:
                    vetting_in_progress += 1
                    status = "in_progress"
                
                # Calculate SLA status (only for IN PROGRESS candidates - not "not_started" or "complete")
                days_elapsed = 0
                sla_status = None  # None for not_started and complete candidates
                
                if status == "in_progress" and first_check_date:
                    if isinstance(first_check_date, str):
                        first_check_date = datetime.datetime.fromisoformat(first_check_date.replace('Z', '+00:00').split('+')[0])
                    days_elapsed = (now - first_check_date).days
                    
                    # Only in_progress candidates have SLA status
                    if days_elapsed > SLA_DAYS:
                        sla_status = "out"
                        vetting_out_of_sla += 1
                        vetting_in_progress_out_sla += 1
                    else:
                        sla_status = "in"
                        vetting_in_sla += 1
                        vetting_in_progress_in_sla += 1
                
                vetting_candidates.append({
                    "id": cand_id,
                    "name": cand_name,
                    "total": total,
                    "complete": complete,
                    "status": status,
                    "days_elapsed": days_elapsed,
                    "sla_status": sla_status
                })
            
            # Get breakdown by check type for IN PROGRESS candidates only (1-9/10 complete)
            # First, identify candidates who are "in progress" (have some but not all checks complete)
            in_progress_candidate_ids = [c["id"] for c in vetting_candidates if c["status"] == "in_progress"]
            
            # Process check type stats with SLA filtering - only for in-progress candidates
            check_type_all = {}
            check_type_in_sla = {}
            check_type_out_sla = {}
            
            if in_progress_candidate_ids:
                # Get check data only for in-progress candidates
                # Build a simple query with the IDs directly (safe since they're integers from our own query)
                id_list = ",".join(str(int(cid)) for cid in in_progress_candidate_ids)
                check_type_with_sla = s.execute(
                    text(f"""
                        SELECT 
                            vc.check_type,
                            vc.status,
                            c.id as cand_id,
                            (SELECT MIN(vc2.created_at) FROM vetting_check vc2 WHERE vc2.candidate_id = c.id) as first_check_date
                        FROM vetting_check vc
                        JOIN candidates c ON c.id = vc.candidate_id
                        WHERE c.id IN ({id_list})
                    """)
                ).all()
                
                for row in check_type_with_sla:
                    check_type, check_status, cand_id, first_check_date = row
                    
                    # Calculate SLA status for this candidate
                    cand_sla = "in"
                    if first_check_date:
                        if isinstance(first_check_date, str):
                            first_check_date = datetime.datetime.fromisoformat(first_check_date.replace('Z', '+00:00').split('+')[0])
                        days_elapsed = (now - first_check_date).days
                        if days_elapsed > SLA_DAYS:
                            cand_sla = "out"
                    
                    # Initialize dicts if needed
                    if check_type not in check_type_all:
                        check_type_all[check_type] = {"total": 0, "complete": 0, "in_progress": 0, "not_started": 0}
                        check_type_in_sla[check_type] = {"total": 0, "complete": 0, "in_progress": 0, "not_started": 0}
                        check_type_out_sla[check_type] = {"total": 0, "complete": 0, "in_progress": 0, "not_started": 0}
                    
                    # Update all counts
                    check_type_all[check_type]["total"] += 1
                    if check_status == "Complete":
                        check_type_all[check_type]["complete"] += 1
                    elif check_status == "In Progress":
                        check_type_all[check_type]["in_progress"] += 1
                    else:
                        check_type_all[check_type]["not_started"] += 1
                    
                    # Update SLA-filtered counts
                    if cand_sla == "in":
                        check_type_in_sla[check_type]["total"] += 1
                        if check_status == "Complete":
                            check_type_in_sla[check_type]["complete"] += 1
                        elif check_status == "In Progress":
                            check_type_in_sla[check_type]["in_progress"] += 1
                        else:
                            check_type_in_sla[check_type]["not_started"] += 1
                    else:
                        check_type_out_sla[check_type]["total"] += 1
                        if check_status == "Complete":
                            check_type_out_sla[check_type]["complete"] += 1
                        elif check_status == "In Progress":
                            check_type_out_sla[check_type]["in_progress"] += 1
                        else:
                            check_type_out_sla[check_type]["not_started"] += 1
            
            vetting_by_type = check_type_all
            vetting_by_type_in_sla = check_type_in_sla
            vetting_by_type_out_sla = check_type_out_sla
            
        except Exception as e:
            # VettingCheck table might not exist yet
            import traceback
            traceback.print_exc()
        
        vetting_summary = {
            "total_candidates": vetting_total,
            "complete": vetting_complete,
            "in_progress": vetting_in_progress,
            "not_started": vetting_not_started,
            "in_sla": vetting_in_sla,
            "out_of_sla": vetting_out_of_sla,
            "sla_days": SLA_DAYS,
            "by_type": vetting_by_type,
            "by_type_in_sla": vetting_by_type_in_sla,
            "by_type_out_sla": vetting_by_type_out_sla,
            "candidates": vetting_candidates,
            # SLA-filtered status counts
            "not_started_in_sla": vetting_not_started_in_sla,
            "not_started_out_sla": vetting_not_started_out_sla,
            "in_progress_in_sla": vetting_in_progress_in_sla,
            "in_progress_out_sla": vetting_in_progress_out_sla,
        }
    
    return render_template(
        "workflow.html",
        stages=stages,
        stage_data=stage_data,
        kpi_counts=kpi_counts,
        jobs=all_jobs,
        engagements=all_engagements,
        all_clients=all_clients,
        all_engagements=all_engagements,
        all_roles=all_roles,
        all_intakes=all_intakes,
        all_owners=all_owners,
        eng_status_filter=eng_status_filter,
        owner_filter=owner_filter,
        client_filter=client_filter,
        engagement_filter=engagement_filter,
        role_filter=role_filter,
        intake_filter=intake_filter,
        unsigned_contracts=unsigned_contracts_data,
        vetting_summary=vetting_summary,
    )

@csrf.exempt
@app.route("/api/workflow/move", methods=["POST"])
@login_required
def api_workflow_move():
    """
    API endpoint for moving cards in the workflow kanban.
    Maps stage IDs to actual status names and calls the main workflow_move logic.
    """
    # Map stage IDs to canonical status names
    STAGE_ID_TO_STATUS = {
        "pipeline": "Pipeline",
        "shortlist": "Shortlist",
        "ia": "I&A",
        "client_review": "Client Review",
        "offered": "Offered",
        "rejected_withdrawn": "Rejected",
        "accepted": "Accepted",
        "ready_to_contract": "Ready to Contract",
        "contract_sent": "Contract Sent",
        "placed": "Placed",
        # Also add legacy mappings for drag-drop compatibility
        "applications_pending_review": "Pipeline",
        "interview_scheduled": "I&A",
        "interview_completed": "Client Review", 
        "vetting_in_flight": "Ready to Contract",
        "contract_issued": "Contract Sent",
        "contract_signed": "Placed",
    }
    
    payload = request.json or {}
    card_id = payload.get("card_id")
    card_type = payload.get("card_type", "application")
    to_stage = payload.get("to_stage", "")
    direct_status = payload.get("new_status", "")  # Direct status from button clicks
    
    if not card_id:
        return jsonify({"ok": False, "error": "Missing card_id"}), 400
    
    # Map stage ID to status name, or use direct status if provided
    new_status = direct_status if direct_status else STAGE_ID_TO_STATUS.get(to_stage)
    if not new_status:
        return jsonify({"ok": False, "error": f"Invalid stage: {to_stage}"}), 400
    
    try:
        with Session(engine) as s:
            app_obj = s.scalar(select(Application).where(Application.id == int(card_id)))
            if not app_obj:
                return jsonify({"ok": False, "error": "Application not found"}), 404
            
            old_status = app_obj.status
            app_obj.status = new_status
            
            # Update candidate's last activity and status based on new workflow stage
            cand = s.scalar(select(Candidate).where(Candidate.id == app_obj.candidate_id))
            if cand:
                cand.last_activity_at = datetime.datetime.utcnow()
                
                # Update candidate status based on workflow stage
                if new_status in ["Placed", "Contract Signed"]:
                    cand.status = "On Assignment"
                elif new_status in ["Contract Sent", "Ready to Contract", "Accepted"]:
                    cand.status = "In Vetting"
                elif new_status in ["Offered", "Client Review", "I&A"]:
                    cand.status = "Interviewing"
                elif new_status in ["Rejected", "Withdrawn"]:
                    cand.status = "Available"
            
            # Store values for audit log before commit
            audit_cand_id = cand.id if cand else None
            audit_cand_name = cand.name if cand else None
            
            s.commit()
            
            # Audit log: workflow drag-drop move (after commit to avoid db lock)
            log_audit_event('update', 'workflow',
                           f'Workflow stage changed (drag-drop): {old_status} → {new_status}',
                           'application', int(card_id),
                           {'old_status': old_status, 'new_status': new_status,
                            'candidate_id': audit_cand_id,
                            'candidate_name': audit_cand_name})
            
            return jsonify({
                "ok": True,
                "message": f"Moved from {old_status} to {new_status}",
                "card_id": card_id,
                "new_status": new_status
            })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@login_required
@app.route("/workflow/move", methods=["POST"])
def workflow_move():
    """
    Move an application to a new workflow stage.
    
    GAP 1.1 FIX: When an associate card is moved between workflow stages,
    the associate's status updates system-wide. This status change is reflected
    on the Associate Profile, in the Engagements View, on the Dashboard KPIs,
    on the Placements page, and in any filtered views or reports.
    
    GAP 1.4 FIX: Stage skip validation - enforces allowed stage transitions.
    GAP X.4: Audit trail for workflow changes.
    """
    payload = request.json or {}
    app_id = int(payload.get("app_id", 0))
    new_status = payload.get("new_status", "")
    force_move = payload.get("force", False)  # Allow admin override
    
    # New 9-stage workflow valid statuses
    valid_statuses = [
        "Pipeline", 
        "Shortlist", 
        "I&A", 
        "Client Review", 
        "Offered",
        "Accepted",
        "Ready to Contract",
        "Contract Sent", 
        "Placed",
        "Rejected",
        "Withdrawn"
    ]
    
    # GAP 1.4: Define allowed stage transitions for new 9-stage workflow
    # Forward moves follow natural progression, backward moves allowed with restrictions
    ALLOWED_TRANSITIONS = {
        "Pipeline": ["Shortlist", "Rejected", "Withdrawn"],
        "Shortlist": ["Pipeline", "I&A", "Rejected", "Withdrawn"],
        "I&A": ["Shortlist", "Client Review", "Rejected", "Withdrawn"],
        "Client Review": ["I&A", "Offered", "Rejected", "Withdrawn"],
        "Offered": ["Client Review", "Accepted", "Rejected", "Withdrawn"],
        "Accepted": ["Offered", "Ready to Contract", "Rejected", "Withdrawn"],
        "Ready to Contract": ["Accepted", "Contract Sent", "Rejected", "Withdrawn"],
        "Contract Sent": ["Ready to Contract", "Placed", "Rejected", "Withdrawn"],
        "Placed": ["Contract Sent"],  # Can only go back one step from placed
        "New": valid_statuses,  # New applications can go anywhere
        "Rejected": ["Pipeline"],  # Can be reconsidered
        "Withdrawn": ["Pipeline"],  # Can be reconsidered
        # Legacy status mappings for backward compatibility
        "Applications Pending Review": valid_statuses,
        "Interview Scheduled": valid_statuses,
        "Interview Completed": valid_statuses,
        "Vetting In-Flight": valid_statuses,
        "Contract Issued": valid_statuses,
        "Contract Signed": valid_statuses,
    }
    
    if new_status not in valid_statuses:
        return jsonify({"ok": False, "error": "Invalid status"}), 400
    
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            return jsonify({"ok": False, "error": "Application not found"}), 404
        
        old_status = appn.status
        
        # GAP 1.4: Validate stage transition unless force override
        if not force_move and old_status in ALLOWED_TRANSITIONS:
            allowed_next_stages = ALLOWED_TRANSITIONS.get(old_status, valid_statuses)
            if new_status not in allowed_next_stages:
                return jsonify({
                    "ok": False, 
                    "error": f"Cannot move directly from '{old_status}' to '{new_status}'. Allowed: {', '.join(allowed_next_stages)}",
                    "validation_error": True,
                    "current_stage": old_status,
                    "requested_stage": new_status,
                    "allowed_stages": allowed_next_stages
                }), 400
        
        appn.status = new_status
        
        # Get candidate and job for status sync
        cand = s.scalar(select(Candidate).where(Candidate.id == appn.candidate_id))
        job = s.scalar(select(Job).where(Job.id == appn.job_id)) if appn.job_id else None
        
        candidate_status_updated = False
        old_candidate_status = cand.status if cand else None
        
        if cand:
            # Always update last_activity_at
            cand.last_activity_at = datetime.datetime.utcnow()
            
            # ============================================================
            # GAP 1.1 FIX: System-wide status sync based on workflow stage
            # ============================================================
            
            # Map workflow stages to candidate statuses
            if new_status in ["Placed", "Contract Signed"]:
                # Associate is now ON CONTRACT / On Assignment
                cand.status = "On Assignment"
                candidate_status_updated = True
                
                # Update interview result if not already set
                if not cand.optimus_interview_result or cand.optimus_interview_result == "Pending":
                    cand.optimus_interview_result = "Pass"
                
                # Mark as previously vetted since they completed the workflow
                cand.previously_vetted = True
                
            elif new_status in ["Contract Sent", "Contract Issued"]:
                # Associate is in final stages, mark as In Vetting/Processing
                cand.status = "In Vetting"
                candidate_status_updated = True
                
            elif new_status in ["Ready to Contract", "Accepted", "Vetting In-Flight"]:
                # Associate is being vetted or ready for contract
                cand.status = "In Vetting"
                candidate_status_updated = True
                
            elif new_status in ["Client Review", "Interview Completed"]:
                # Associate completed interview, mark as Interviewing until next stage
                cand.status = "Interviewing"
                candidate_status_updated = True
                # Record interview completion time if not set
                if not appn.interview_completed_at:
                    appn.interview_completed_at = datetime.datetime.utcnow()
                
            elif new_status in ["I&A", "Interview Scheduled"]:
                # Associate has interview scheduled
                cand.status = "Interviewing"
                candidate_status_updated = True
                
            elif new_status in ["Pipeline", "Shortlist", "Applications Pending Review"]:
                # Early stages - only update if candidate was previously On Assignment
                # and is now back in the pipeline (rare case)
                if old_candidate_status == "On Assignment":
                    cand.status = "Available"
                    candidate_status_updated = True
                    
            elif new_status == "Offered":
                # Candidate has been offered - still interviewing until accepted
                cand.status = "Interviewing"
                candidate_status_updated = True
            
            # ============================================================
            # Handle special date fields based on stage transitions
            # ============================================================
            
            if new_status in ["I&A", "Interview Scheduled"] and old_status not in ["I&A", "Interview Scheduled"]:
                # Set interview scheduled timestamp if moving INTO this stage
                if not appn.interview_scheduled_at:
                    appn.interview_scheduled_at = datetime.datetime.utcnow()
            
            if new_status in ["Client Review", "Interview Completed"] and old_status in ["I&A", "Interview Scheduled"]:
                # Moving from I&A to Client Review - record completion time
                appn.interview_completed_at = datetime.datetime.utcnow()
        
        # Store values for audit log before commit
        audit_details = {
            'candidate_id': cand.id if cand else None,
            'candidate_name': cand.name if cand else None,
            'old_workflow_status': old_status,
            'new_workflow_status': new_status,
            'candidate_status_changed': candidate_status_updated,
            'old_candidate_status': old_candidate_status,
            'new_candidate_status': cand.status if cand else None
        }
        
        s.commit()
        
        # GAP X.4: Audit log for workflow stage changes (after commit to avoid db lock)
        log_audit_event(
            'update', 'workflow',
            f'Workflow stage changed: {old_status} → {new_status}',
            'application', app_id,
            audit_details
        )
        
        # Build response with all status changes for frontend
        response_data = {
            "ok": True, 
            "old_status": old_status, 
            "new_status": new_status,
            "app_id": app_id,
            "candidate_id": cand.id if cand else None,
            "candidate_name": cand.name if cand else None,
            "candidate_status_updated": candidate_status_updated
        }
        
        if candidate_status_updated and cand:
            response_data["old_candidate_status"] = old_candidate_status
            response_data["new_candidate_status"] = cand.status
    
    return jsonify(response_data)

# ============== PLACEMENTS PAGE ==============
@login_required
@app.route("/placements")
def placements():
    """
    Placements page showing all associates currently ON CONTRACT.
    
    Per wireframe requirements:
    Filters: Name, Client, Engagement, Role
    Features:
    - Placements Pie Chart (by client)
    - Export CSV button
    - On Contract Table with all columns
    - Scheduled Joiners Panel (next 14 days)
    - Scheduled Leavers Panel (next 14 days)
    - Forecasted Headcount graph
    """
    # Query params for filters (now multi-select)
    name_filter = request.args.getlist("name") or []
    client_filter = request.args.getlist("client") or []
    engagement_filter = request.args.getlist("engagement") or []
    role_filter = request.args.getlist("role") or []
    
    with Session(engine) as s:
        # Get unique clients for filter
        all_clients = s.scalars(
            select(Engagement.client)
            .distinct()
            .where(Engagement.client.isnot(None), Engagement.client != "")
            .order_by(Engagement.client)
        ).all()
        
        # Get active engagements for filter
        all_engagements = s.scalars(
            select(Engagement)
            .where(Engagement.status == "Active")
            .order_by(Engagement.name)
        ).all()
        
        # Get unique roles for filter
        all_roles = s.scalars(
            select(Job.title)
            .distinct()
            .where(Job.title.isnot(None), Job.title != "")
            .order_by(Job.title)
        ).all()
        
        # Get all candidate names for filter dropdown
        all_names = s.scalars(
            select(Candidate.name)
            .distinct()
            .where(Candidate.name.isnot(None), Candidate.name != "")
            .order_by(Candidate.name)
        ).all()
        
        # Query for on-contract placements (signed/completed contracts)
        query = (
            select(
                Candidate,
                Application,
                Job,
                Engagement,
                ESigRequest,
                EngagementPlan
            )
            .select_from(ESigRequest)
            .join(Application, Application.id == ESigRequest.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .outerjoin(Engagement, Engagement.id == Job.engagement_id)
            .outerjoin(
                EngagementPlan,
                and_(
                    EngagementPlan.engagement_id == Job.engagement_id,
                    EngagementPlan.role_type == Job.role_type
                )
            )
            .where(func.lower(ESigRequest.status).in_(['signed', 'completed']))
        )
        
        # Apply filters (now multi-select)
        if name_filter:
            query = query.where(Candidate.name.in_(name_filter))
        
        if client_filter:
            query = query.where(Engagement.client.in_(client_filter))
        
        if engagement_filter:
            eng_ids_int = [int(e) for e in engagement_filter if e.isdigit()]
            if eng_ids_int:
                query = query.where(Engagement.id.in_(eng_ids_int))
        
        if role_filter:
            query = query.where(Job.title.in_(role_filter))
        
        query = query.order_by(ESigRequest.signed_at.desc())
        
        results = s.execute(query).all()
        
        # Build placement data
        placements_data = []
        client_counts = {}  # For pie chart
        
        for cand, app, job, eng, esig, plan in results:
            client_name = eng.client if eng and eng.client else "Unknown"
            client_counts[client_name] = client_counts.get(client_name, 0) + 1
            
            placements_data.append({
                "candidate": cand,
                "application": app,
                "job": job,
                "engagement": eng,
                "esig": esig,
                "pay_rate": plan.pay_rate if plan else 0,
                "bill_rate": plan.charge_rate if plan else 0,
                "start_date": esig.signed_at if esig else (eng.start_date if eng else None),
                "end_date": eng.end_date if eng else None,
                # Additional fields for template
                "associate_name": cand.name,
                "associate_email": cand.email,
                "associate_id": cand.id,
                "client": eng.client if eng else "Unknown",
                "engagement_name": eng.name if eng else "Unknown",
                "engagement_ref": eng.id if eng else None,
                "engagement_id": eng.id if eng else None,
                "role": job.title if job else "Unknown",
                "status": esig.status if esig else "unknown",
            })
        
        # Pie chart data (by client)
        pie_chart_data = [
            {"client": client, "count": count}
            for client, count in sorted(client_counts.items(), key=lambda x: -x[1])
        ]
        
        # Summary stats
        total_placements = len(placements_data)
        unique_clients = len(set(p["engagement"].client for p in placements_data if p["engagement"] and p["engagement"].client))
        total_daily_pay = sum(p["pay_rate"] or 0 for p in placements_data)
        total_daily_bill = sum(p["bill_rate"] or 0 for p in placements_data)
        total_daily_margin = total_daily_bill - total_daily_pay
        
        # Build clients dict for template (grouped by client with pay/bill totals)
        clients = {}
        for p in placements_data:
            client_name = p["engagement"].client if p["engagement"] and p["engagement"].client else "Unknown"
            if client_name not in clients:
                clients[client_name] = {"count": 0, "pay": 0, "bill": 0}
            clients[client_name]["count"] += 1
            clients[client_name]["pay"] += p["pay_rate"] or 0
            clients[client_name]["bill"] += p["bill_rate"] or 0
        
        # === Scheduled Joiners (next 14 days) ===
        now = datetime.datetime.utcnow()
        fourteen_days = now + datetime.timedelta(days=14)
        
        joiners_query = (
            select(Candidate, Application, Job, Engagement, ESigRequest)
            .select_from(ESigRequest)
            .join(Application, Application.id == ESigRequest.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .outerjoin(Engagement, Engagement.id == Job.engagement_id)
            .where(func.lower(ESigRequest.status).in_(['signed', 'completed']))
            .where(Engagement.start_date.isnot(None))
            .where(Engagement.start_date >= now.date())
            .where(Engagement.start_date <= fourteen_days.date())
        )
        
        # Apply filters to joiners query
        if name_filter:
            joiners_query = joiners_query.where(Candidate.name.in_(name_filter))
        if client_filter:
            joiners_query = joiners_query.where(Engagement.client.in_(client_filter))
        if engagement_filter:
            eng_ids_int = [int(e) for e in engagement_filter if e.isdigit()]
            if eng_ids_int:
                joiners_query = joiners_query.where(Engagement.id.in_(eng_ids_int))
        if role_filter:
            joiners_query = joiners_query.where(Job.title.in_(role_filter))
        
        joiners_query = joiners_query.order_by(Engagement.start_date.asc())
        
        joiners_results = s.execute(joiners_query).all()
        today_date = now.date()
        scheduled_joiners = []
        for cand, app, job, eng, esig in joiners_results:
            start_date = eng.start_date if eng else None
            # Calculate days_until, handling both date and datetime objects
            days_until = 0
            if start_date:
                if hasattr(start_date, 'date'):
                    start_date_obj = start_date.date() if callable(getattr(start_date, 'date', None)) else start_date
                else:
                    start_date_obj = start_date
                try:
                    days_until = (start_date_obj - today_date).days
                except TypeError:
                    days_until = 0
            scheduled_joiners.append({
                "candidate": cand,
                "job": job,
                "engagement": eng,
                "start_date": start_date,
                "days_until": days_until
            })
        
        # === Scheduled Leavers (next 30 days) ===
        thirty_days = now + datetime.timedelta(days=30)
        
        leavers_query = (
            select(Candidate, Application, Job, Engagement, ESigRequest)
            .select_from(ESigRequest)
            .join(Application, Application.id == ESigRequest.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .outerjoin(Engagement, Engagement.id == Job.engagement_id)
            .where(func.lower(ESigRequest.status).in_(['signed', 'completed']))
            .where(Engagement.end_date.isnot(None))
            .where(Engagement.end_date >= now.date())
            .where(Engagement.end_date <= thirty_days.date())
        )
        
        # Apply filters to leavers query
        if name_filter:
            leavers_query = leavers_query.where(Candidate.name.in_(name_filter))
        if client_filter:
            leavers_query = leavers_query.where(Engagement.client.in_(client_filter))
        if engagement_filter:
            eng_ids_int = [int(e) for e in engagement_filter if e.isdigit()]
            if eng_ids_int:
                leavers_query = leavers_query.where(Engagement.id.in_(eng_ids_int))
        if role_filter:
            leavers_query = leavers_query.where(Job.title.in_(role_filter))
        
        leavers_query = leavers_query.order_by(Engagement.end_date.asc())
        
        leavers_results = s.execute(leavers_query).all()
        scheduled_leavers = []
        for cand, app, job, eng, esig in leavers_results:
            end_date = eng.end_date if eng else None
            # Calculate days_until, handling both date and datetime objects
            days_until = 0
            if end_date:
                if hasattr(end_date, 'date'):
                    end_date_obj = end_date.date() if callable(getattr(end_date, 'date', None)) else end_date
                else:
                    end_date_obj = end_date
                try:
                    days_until = (end_date_obj - today_date).days
                except TypeError:
                    days_until = 0
            scheduled_leavers.append({
                "candidate": cand,
                "job": job,
                "engagement": eng,
                "end_date": end_date,
                "days_until": days_until
            })
        
        # === Per-Engagement Forecast Data ===
        # Get all signed contracts grouped by engagement with start/end dates
        forecast_query = (
            select(
                Engagement.id,
                Engagement.name,
                Engagement.client,
                Engagement.start_date,
                Engagement.end_date,
                func.count(ESigRequest.id).label('headcount')
            )
            .select_from(ESigRequest)
            .join(Application, Application.id == ESigRequest.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .join(Engagement, Engagement.id == Job.engagement_id)
            .where(func.lower(ESigRequest.status).in_(['signed', 'completed']))
        )
        
        # Apply filters to forecast query
        if name_filter:
            forecast_query = forecast_query.where(Candidate.name.in_(name_filter))
        if client_filter:
            forecast_query = forecast_query.where(Engagement.client.in_(client_filter))
        if engagement_filter:
            eng_ids_int = [int(e) for e in engagement_filter if e.isdigit()]
            if eng_ids_int:
                forecast_query = forecast_query.where(Engagement.id.in_(eng_ids_int))
        if role_filter:
            forecast_query = forecast_query.where(Job.title.in_(role_filter))
        
        forecast_query = forecast_query.group_by(Engagement.id).order_by(Engagement.name)
        
        forecast_results = s.execute(forecast_query).all()
        engagement_forecast_data = []
        
        today_date = now.date() if hasattr(now, 'date') else now
        
        for eng_id, eng_name, client, start_date, end_date, total_headcount in forecast_results:
            # Get individual placement dates for this engagement
            # Use Engagement.start_date for placement start (when they begin)
            # Use ESigRequest.end_date for individual end (if set), else Engagement.end_date
            placements_query = (
                select(
                    Candidate.name,
                    Engagement.start_date.label('placement_start_date'),
                    ESigRequest.end_date.label('individual_end_date'),
                    Engagement.end_date.label('engagement_end_date')
                )
                .select_from(ESigRequest)
                .join(Application, Application.id == ESigRequest.application_id)
                .join(Candidate, Candidate.id == Application.candidate_id)
                .join(Job, Job.id == Application.job_id)
                .join(Engagement, Engagement.id == Job.engagement_id)
                .where(func.lower(ESigRequest.status).in_(['signed', 'completed']))
                .where(Engagement.id == eng_id)
            )
            
            placement_results = s.execute(placements_query).all()
            
            # Build list of individual placements with their dates
            placements_list = []
            current_headcount = 0
            
            for p_name, p_start, p_individual_end, p_eng_end in placement_results:
                # Normalize dates
                p_start_date = None
                p_end_date = None
                
                if p_start:
                    if hasattr(p_start, 'date'):
                        p_start_date = p_start.date() if callable(getattr(p_start, 'date', None)) else p_start
                    else:
                        p_start_date = p_start
                
                # Use individual end date if set, otherwise use engagement end date
                p_end_raw = p_individual_end if p_individual_end else p_eng_end
                if p_end_raw:
                    if hasattr(p_end_raw, 'date'):
                        p_end_date = p_end_raw.date() if callable(getattr(p_end_raw, 'date', None)) else p_end_raw
                    else:
                        p_end_date = p_end_raw
                
                # Check if this placement is currently active (for current headcount)
                has_started = not p_start_date or p_start_date <= today_date
                has_ended = p_end_date and p_end_date < today_date
                
                if has_started and not has_ended:
                    current_headcount += 1
                
                placements_list.append({
                    "name": p_name,
                    "start_date": p_start_date.strftime('%Y-%m-%d') if p_start_date else None,
                    "end_date": p_end_date.strftime('%Y-%m-%d') if p_end_date else None
                })
            
            # Convert engagement dates to ISO strings for JavaScript
            start_str = None
            end_str = None
            if start_date:
                if hasattr(start_date, 'strftime'):
                    start_str = start_date.strftime('%Y-%m-%d')
                else:
                    start_str = str(start_date)
            if end_date:
                if hasattr(end_date, 'strftime'):
                    end_str = end_date.strftime('%Y-%m-%d')
                else:
                    end_str = str(end_date)
            
            engagement_forecast_data.append({
                "id": eng_id,
                "name": eng_name,
                "client": client or "",
                "start_date": start_str,
                "end_date": end_str,
                "headcount": total_headcount,
                "current_headcount": current_headcount,
                "placements": placements_list
            })
    
    return render_template(
        "placements.html",
        placements=placements_data,
        all_clients=all_clients,
        all_engagements=all_engagements,
        all_roles=all_roles,
        all_names=all_names,
        name_filter=name_filter,
        client_filter=client_filter,
        engagement_filter=engagement_filter,
        role_filter=role_filter,
        total_placements=total_placements,
        unique_clients=unique_clients,
        total_daily_pay=total_daily_pay,
        total_daily_bill=total_daily_bill,
        total_daily_margin=total_daily_margin,
        clients=clients,
        pie_chart_data=pie_chart_data,
        scheduled_joiners=scheduled_joiners,
        scheduled_leavers=scheduled_leavers,
        engagement_forecast_data=engagement_forecast_data,
        now=now,
    )

# CSV Export for Placements
@login_required
@app.route("/placements/export")
def placements_export_csv():
    """Export placements data to CSV"""
    with Session(engine) as s:
        query = (
            select(
                Candidate,
                Application,
                Job,
                Engagement,
                ESigRequest,
                EngagementPlan
            )
            .select_from(ESigRequest)
            .join(Application, Application.id == ESigRequest.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .outerjoin(Engagement, Engagement.id == Job.engagement_id)
            .outerjoin(
                EngagementPlan,
                and_(
                    EngagementPlan.engagement_id == Job.engagement_id,
                    EngagementPlan.role_type == Job.role_type
                )
            )
            .where(func.lower(ESigRequest.status).in_(['signed', 'completed']))
            .order_by(ESigRequest.signed_at.desc())
        )
        results = s.execute(query).all()
        
        # Build CSV
        import io
        import csv
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Header
        writer.writerow([
            "Associate Name", "Email", "Client", "Engagement", "Role", 
            "Pay Rate", "Bill Rate", "Start Date", "End Date"
        ])
        
        # Data rows
        for cand, app, job, eng, esig, plan in results:
            writer.writerow([
                cand.name,
                cand.email or "",
                eng.client if eng else "",
                eng.name if eng else "",
                job.title,
                plan.pay_rate if plan else 0,
                plan.charge_rate if plan else 0,
                (esig.signed_at or eng.start_date).strftime("%d/%m/%Y") if (esig.signed_at or (eng and eng.start_date)) else "",
                eng.end_date.strftime("%d/%m/%Y") if eng and eng.end_date else "",
            ])
    
    # Create response
    from flask import Response
    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=placements_export.csv"}
    )

# ============== PROJECTS PAGE ==============
@login_required
@app.route("/projects")
@login_required
def projects():
    """
    Projects landing page showing pipeline overview.
    Combines Opportunities (pipeline) and Active Engagements.
    """
    # Query params
    status_filter = request.args.get("status", "all")  # all, pipeline, active, completed
    client_filter = request.args.get("client", "all")
    
    with Session(engine) as s:
        # Get unique clients
        opp_clients = s.scalars(
            select(Opportunity.client)
            .distinct()
            .where(Opportunity.client.isnot(None), Opportunity.client != "")
        ).all()
        eng_clients = s.scalars(
            select(Engagement.client)
            .distinct()
            .where(Engagement.client.isnot(None), Engagement.client != "")
        ).all()
        all_clients = sorted(set(opp_clients + eng_clients))
        
        # Pipeline opportunities (not yet converted to engagements)
        opp_query = (
            select(Opportunity)
            .where(Opportunity._engagement_id.is_(None))
            .order_by(Opportunity.est_start.asc())
        )
        if client_filter != "all":
            opp_query = opp_query.where(Opportunity.client == client_filter)
        
        opportunities = s.scalars(opp_query).all() if status_filter in ["all", "pipeline"] else []
        
        # Active engagements
        eng_query = select(Engagement).order_by(Engagement.start_date.desc())
        if status_filter == "active":
            eng_query = eng_query.where(Engagement.status == "Active")
        elif status_filter == "completed":
            eng_query = eng_query.where(Engagement.status != "Active")
        
        if client_filter != "all":
            eng_query = eng_query.where(Engagement.client == client_filter)
        
        engagements = s.scalars(eng_query).all() if status_filter in ["all", "active", "completed"] else []
        
        # Pipeline stages for kanban
        stage_names = ["Qualified Lead", "Proposal", "Procurement", "Closed Won", "Closed Lost"]
        
        # Organize opportunities by stage for the kanban
        stage_data = {stage: [] for stage in stage_names}
        for opp in opportunities:
            stage = opp.stage or "Qualified Lead"
            if stage in stage_data:
                stage_data[stage].append(opp)
        
        # Filter to active engagements for the template
        active_engagements = [e for e in engagements if e.status == "Active"]
        
        # Engagement stats
        active_count = len(active_engagements)
        completed_count = sum(1 for e in engagements if e.status != "Active")
        
        # Calculate values
        total_pipeline = len(opportunities)
        total_value = sum(opp.est_value or 0 for opp in opportunities)
        
        # Weighted value: sum of (est_value * probability/100)
        # If no probability, assume 50%
        weighted_value = sum(
            (opp.est_value or 0) * (getattr(opp, 'probability', 50) or 50) / 100 
            for opp in opportunities
        )
        
        # Get unique roles from engagement plans
        all_roles = s.scalars(
            select(EngagementPlan.role_type)
            .distinct()
            .where(EngagementPlan.role_type.isnot(None), EngagementPlan.role_type != "")
            .order_by(EngagementPlan.role_type)
        ).all()
        
        # Build engagement to roles mapping for filtering
        engagement_roles = {}
        for eng in engagements:
            roles = s.scalars(
                select(EngagementPlan.role_type)
                .distinct()
                .where(EngagementPlan.engagement_id == eng.id)
                .where(EngagementPlan.role_type.isnot(None), EngagementPlan.role_type != "")
            ).all()
            engagement_roles[eng.id] = roles
    
    return render_template(
        "projects.html",
        opportunities=opportunities,
        all_opportunities=opportunities,  # For the filter dropdowns
        engagements=engagements,
        all_clients=all_clients,
        all_roles=all_roles,
        engagement_roles=engagement_roles,
        status_filter=status_filter,
        client_filter=client_filter,
        pipeline_stages=stage_names,
        stage_data=stage_data,
        active_engagements=active_engagements,
        active_count=active_count,
        completed_count=completed_count,
        total_pipeline=total_pipeline,
        total_value=total_value,
        weighted_value=weighted_value,
    )

@login_required
@app.route("/engagements/create", methods=["GET", "POST"])
def create_engagement():
    form = EngagementForm()
    if form.validate_on_submit():
        with Session(engine) as s:
            # allocate sequential EG### reference
            new_ref = _next_engagement_ref(s)

            e = Engagement(
                ref=new_ref,                                     # NEW
                name=form.name.data,
                client=form.client.data or "",
                status=form.status.data or "Active",
                start_date=parse_date_dmy(form.start_date.data), # DD-MM-YYYY
                end_date=parse_date_dmy(form.end_date.data),     # DD-MM-YYYY
                sow_signed_at=parse_date_dmy(form.sow_signed_at.data),
                description=form.description.data or "",
            )
            s.add(e)
            s.commit()
            flash(f"Engagement {e.ref} created", "success")
        return redirect(url_for("engagements"))
    return render_template("create_engagement.html", form=form)

@login_required
@app.route("/engagements", methods=["GET"])
def engagements():
    # Get filter parameters with safe defaults
    name_filter = request.args.get('name', '').strip()
    status_filter = request.args.get('status', 'all')
    client_filter = request.args.get('client', 'all')
    sort_by = request.args.get('sort', 'created_desc')
    opp_name_filter = request.args.get('opp_name', '').strip()
    opp_client_filter = request.args.get('opp_client', 'all')
    opp_status_filter = request.args.get('opp_status', 'all')
    opp_sort = request.args.get('opp_sort', 'created_desc')
    
    with Session(engine) as s:
        # 1. Current engagements with filters
        eng_query = select(Engagement)
        
        # Apply name filter (exact match from dropdown)
        if name_filter:
            eng_query = eng_query.where(Engagement.name == name_filter)
        
        # Apply status filter
        if status_filter == 'active':
            eng_query = eng_query.where(Engagement.status == "Active")
        elif status_filter == 'inactive':
            eng_query = eng_query.where(Engagement.status != "Active")
        
        # Apply client filter
        if client_filter and client_filter != 'all':
            eng_query = eng_query.where(Engagement.client == client_filter)
        
        # Apply sorting
        if sort_by == 'created_asc':
            eng_query = eng_query.order_by(Engagement.id.asc())
        elif sort_by == 'name_asc':
            eng_query = eng_query.order_by(Engagement.name.asc())
        elif sort_by == 'name_desc':
            eng_query = eng_query.order_by(Engagement.name.desc())
        elif sort_by == 'start_date':
            eng_query = eng_query.order_by(Engagement.start_date.desc().nullslast())
        else:  # created_desc (default)
            eng_query = eng_query.order_by(Engagement.id.desc())
        
        engagements_rows = s.scalars(eng_query).all()
        
        # Get unique engagement names for filter dropdown
        all_engagement_names = []
        try:
            all_engagement_names = s.scalars(
                select(Engagement.name)
                .distinct()
                .where(Engagement.name.isnot(None))
                .where(Engagement.name != '')
                .order_by(Engagement.name)
            ).all()
        except Exception as e:
            print(f"Error getting engagement names: {e}")
            all_engagement_names = []
        
        # Get unique clients for filter dropdown
        all_clients = []
        try:
            all_clients = s.scalars(
                select(Engagement.client)
                .distinct()
                .where(Engagement.client.isnot(None))
                .where(Engagement.client != '')
                .order_by(Engagement.client)
            ).all()
        except Exception as e:
            print(f"Error getting clients: {e}")
            all_clients = []
        
        # Get unique opportunity clients for filter dropdown
        all_opp_clients = []
        try:
            all_opp_clients = s.scalars(
                select(Opportunity.client)
                .distinct()
                .where(Opportunity.client.isnot(None))
                .where(Opportunity.client != '')
                .order_by(Opportunity.client)
            ).all()
        except Exception as e:
            print(f"Error getting opportunity clients: {e}")
            all_opp_clients = []

        # 2. All opportunities with filters
        opp_query = select(Opportunity)
        
        # Apply opportunity name filter (exact match from dropdown)
        if opp_name_filter:
            opp_query = opp_query.where(Opportunity.name == opp_name_filter)
        
        # Apply opportunity client filter
        if opp_client_filter and opp_client_filter != 'all':
            opp_query = opp_query.where(Opportunity.client == opp_client_filter)
        
        # Apply opportunity status filter
        if opp_status_filter and opp_status_filter != 'all':
            try:
                if opp_status_filter == 'proposal':
                    opp_query = opp_query.where(Opportunity.stage.ilike('%proposal%'))
                elif opp_status_filter == 'negotiation':
                    opp_query = opp_query.where(Opportunity.stage.ilike('%negotiation%'))
                elif opp_status_filter == 'closed_won':
                    opp_query = opp_query.where(Opportunity.stage.ilike('closed won'))
                elif opp_status_filter == 'closed_lost':
                    opp_query = opp_query.where(Opportunity.stage.ilike('closed lost'))
            except Exception as e:
                print(f"Error filtering opportunities: {e}")
        
        # Apply opportunity sorting
        try:
            if opp_sort == 'created_asc':
                opp_query = opp_query.order_by(Opportunity.created_at.asc())
            elif opp_sort == 'value_desc':
                opp_query = opp_query.order_by(Opportunity.est_value.desc().nullslast())
            elif opp_sort == 'value_asc':
                opp_query = opp_query.order_by(Opportunity.est_value.asc().nullslast())
            elif opp_sort == 'name_asc':
                opp_query = opp_query.order_by(Opportunity.name.asc())
            else:  # created_desc (default)
                opp_query = opp_query.order_by(Opportunity.created_at.desc())
        except Exception as e:
            print(f"Error sorting opportunities: {e}")
            opp_query = opp_query.order_by(Opportunity.created_at.desc())
        
        opp_rows = s.scalars(opp_query).all()

        # 3. Map opportunity.id -> (engagement_id, engagement_ref)
        eng_links = {
            oid: (eid, ref)
            for (eid, ref, oid) in s.execute(
                text(
                    "SELECT id, ref, opportunity_id "
                    "FROM engagements "
                    "WHERE opportunity_id IS NOT NULL"
                )
            ).all()
        }

        # 4. Annotate each opportunity with the linked engagement (if any)
        for o in opp_rows:
            tup = eng_links.get(o.id)
            if tup:
                o._engagement_id, o._engagement_ref = tup[0], tup[1]
            else:
                o._engagement_id = None
                o._engagement_ref = None

        # 5. Filter out opportunities that are already active engagements
        #    Rule: hide if stage == "closed won" (case-insensitive)
        #    AND we already have an engagement for it.
        visible_opps = []
        for o in opp_rows:
            st_low = (o.stage or "").strip().lower()
            already_converted = bool(o._engagement_id)
            is_closed_won = (st_low == "closed won")
            if is_closed_won and already_converted:
                # skip, it'll show up in Current Engagements
                continue
            visible_opps.append(o)
        
        # Get unique opportunity names from visible opportunities only
        all_opportunity_names = sorted(list(set([o.name for o in visible_opps if o.name])))

    return render_template(
        "engagements.html",
        items=engagements_rows,
        opps=visible_opps,
        name_filter=name_filter,
        status_filter=status_filter,
        client_filter=client_filter,
        sort_by=sort_by,
        opp_name_filter=opp_name_filter,
        opp_client_filter=opp_client_filter,
        opp_status_filter=opp_status_filter,
        opp_sort=opp_sort,
        all_clients=all_clients,
        all_opp_clients=all_opp_clients,
        all_engagement_names=all_engagement_names,
        all_opportunity_names=all_opportunity_names,
    )

@app.route("/jobs", methods=["GET","POST"])
@login_required
def jobs():
    # If coming from an engagement, lock to it
    eng_id_lock = request.args.get("eng_id", type=int)

    with Session(engine) as s:
        # Validate the lock id (if provided)
        locked_eng = None
        if eng_id_lock:
            locked_eng = s.get(Engagement, eng_id_lock)
            if not locked_eng:
                flash("Engagement not found for the requested job context.", "warning")
                return redirect(url_for("engagements"))

        # Build the form
        form = JobForm()

        if locked_eng:
            # Only one choice, preselected — include the engagement reference in the label
            form.engagement_id.choices = [
                (locked_eng.id, f"{locked_eng.ref or '—'} · {locked_eng.name} ({locked_eng.client})")
            ]
            form.engagement_id.data = locked_eng.id  # prefill
        else:
            # Global selector — include refs for all items
            engagements = s.scalars(select(Engagement).order_by(Engagement.id.desc())).all()
            form.engagement_id.choices = [
                (e.id, f"{e.ref or '—'} · {e.name} ({e.client})") for e in engagements
            ]

        # Create
        if form.validate_on_submit():
            engagement_id_final = locked_eng.id if locked_eng else form.engagement_id.data
            j = Job(
                engagement_id=engagement_id_final,
                title=form.title.data,
                description=form.description.data,
                role_type=form.role_type.data or "",
                location=form.location.data or "",
                salary_range=form.salary_range.data or "",
            )
            s.add(j)
            s.commit()
            flash("Job created", "success")

            # After creating in a locked context, bounce back to that engagement’s dashboard
            if locked_eng:
                return redirect(url_for("engagement_job_detail", eng_id=locked_eng.id, job_id=j.id))
            return redirect(url_for("jobs"))

        # List jobs (optionally filtered to the locked engagement for clarity)
        if locked_eng:
            rows = s.scalars(
                select(Job)
                .where(Job.engagement_id == locked_eng.id)
                .options(selectinload(Job.engagement))
                .order_by(Job.created_at.desc())
            ).all()
        else:
            rows = s.scalars(
                select(Job)
                .options(selectinload(Job.engagement))
                .order_by(Job.created_at.desc())
            ).all()

    return render_template(
        "jobs.html",
        form=form,
        items=rows,
        locked_engagement=locked_eng,  # pass for UX banner/disable
    )

@login_required
@app.route("/job/<token>")
def public_job(token):
    with Session(engine) as s:
        job = s.scalar(select(Job).where(Job.public_token==token))
        if not job:
            abort(404)
    return render_template("public_job.html", job=job)

@login_required
@app.route("/job/new", methods=["GET", "POST"])
def job_new():
    """Create a new job post."""
    from sqlalchemy.orm import Session
    with Session(engine) as s:
        if request.method == "POST":
            title = request.form.get("title") or "Untitled Job"
            role = request.form.get("role") or ""
            description = request.form.get("description") or ""
            status = request.form.get("status") or "Open"
            engagement_id = request.form.get("engagement_id") or None

            job = Job(
                title=title,
                role=role,
                description=description,
                status=status,
                engagement_id=engagement_id,
                created_at=datetime.utcnow(),
            )
            s.add(job)
            s.commit()
            flash("Job created successfully.", "success")
            return redirect(url_for("engagement_dashboard", engagement_id=engagement_id or 1))

        # GET request
        roles = s.scalars(select(Role).order_by(Role.name.asc())).all() if "Role" in globals() else []
        return render_template("job_form.html", job=None, roles=roles, mode="create")


@login_required
@app.route("/job/<int:job_id>/edit", methods=["GET", "POST"])
def job_edit(job_id):
    """Edit an existing job post."""
    from sqlalchemy.orm import Session
    with Session(engine) as s:
        job = s.get(Job, job_id)
        if not job:
            flash("Job not found.", "warning")
            return redirect(url_for("engagements"))

        if request.method == "POST":
            job.title = request.form.get("title") or job.title
            job.role = request.form.get("role") or job.role
            job.description = request.form.get("description") or job.description
            job.status = request.form.get("status") or job.status
            s.commit()
            flash("Job updated successfully.", "success")
            return redirect(url_for("engagement_dashboard", engagement_id=job.engagement_id or 1))

        roles = s.scalars(select(Role).order_by(Role.name.asc())).all() if "Role" in globals() else []
        return render_template("job_form.html", job=job, roles=roles, mode="edit")

@app.route("/apply/<token>", methods=["GET", "POST"])
def apply(token):
    form = ApplyForm()
    with Session(engine) as s:
        job = s.scalar(select(Job).where(Job.public_token == token))
        if not job:
            abort(404)

        if form.validate_on_submit():
            # --- normalise email once (trim + lowercase) ---
            raw_email = form.email.data or ""
            norm_email = raw_email.strip().lower()

            # 1️⃣ Find or create candidate by unique email (case-insensitive)
            cand = s.scalar(
                select(Candidate).where(func.lower(Candidate.email) == norm_email)
            )

            if not cand:
                cand = Candidate(
                    name=form.name.data,
                    email=norm_email,
                    phone=(form.phone.data or "").strip(),
                )
                s.add(cand)
                try:
                    s.flush()  # get cand.id, may raise IntegrityError if racing
                except IntegrityError:
                    s.rollback()
                    # someone else just created this candidate with same email
                    cand = s.scalar(
                        select(Candidate).where(func.lower(Candidate.email) == norm_email)
                    )
                    if not cand:
                        flash(
                            "We couldn't process your details. Please try again.",
                            "danger",
                        )
                        return render_template("apply.html", form=form, job=job)

            # 2️⃣ Prevent duplicate applications for the same job+candidate
            existing_app = s.scalar(
                select(Application)
                .where(
                    Application.job_id == job.id,
                    Application.candidate_id == cand.id,
                )
            )
            if existing_app:
                flash(
                    "You’ve already applied for this role. We’ve kept your original application.",
                    "info",
                )
                return redirect(url_for("public_job", token=token))

            # 3️⃣ Require CV upload
            cv_file = request.files.get("cv")
            if not cv_file:
                flash("Please upload a CV", "danger")
                return render_template("apply.html", form=form, job=job)

            # 4️⃣ Save CV in uploads/cvs
            try:
                fname, path, original = save_upload(cv_file, subdir="cvs")
            except Exception as e:
                flash(f"Upload failed: {e}", "danger")
                return render_template("apply.html", form=form, job=job)

            doc = Document(
                candidate_id=cand.id,
                doc_type="cv",
                filename=fname,
                original_name=original,
            )
            s.add(doc)
            s.flush()

            # 5️⃣ Create new application
            appn = Application(
                job_id=job.id,
                candidate_id=cand.id,
                cover_note=form.cover_note.data or "",
            )
            s.add(appn)
            s.flush()

            # 6️⃣ Generate AI summary/tags/score immediately
            _rebuild_ai_summary_and_tags(s, cand, doc=doc, job=job, appn=appn)

            s.commit()
            flash("Application submitted. Thank you!", "success")
            return redirect(url_for("public_job", token=token))

    return render_template("apply.html", form=form, job=job)

@login_required
@app.route("/my/timesheets", methods=["GET", "POST"])
@worker_required
def my_timesheets():
    """
    Minimal worker portal to submit a weekly timesheet.
    Reuses your existing Timesheet model.
    """
    now = datetime.datetime.utcnow()
    start = (now - datetime.timedelta(days=now.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
    end = start + datetime.timedelta(days=7)

    # Single engagement picker for simplicity
    class _TSForm(FlaskForm):
        engagement_id = SelectField("Engagement", coerce=int, validators=[DataRequired()])
        hours = IntegerField("Hours", validators=[DataRequired()])
        notes = TextAreaField("Notes", validators=[WTOptional()])
        submit = SubmitField("Submit")

    with Session(engine) as s:
        engagements = s.scalars(select(Engagement).order_by(Engagement.name.asc())).all()
        choices = [(e.id, f"{e.name} ({e.client or '—'})") for e in engagements]

    form = _TSForm()
    form.engagement_id.choices = choices

    if form.validate_on_submit():
        with Session(engine) as s:
            ts = Timesheet(
                user_id=int(current_user.id),
                engagement_id=form.engagement_id.data,
                period_start=start,
                period_end=end,
                hours=int(form.hours.data or 0),
                notes=form.notes.data or "",
                status="Submitted",
                submitted_at=datetime.datetime.utcnow(),
            )
            s.add(ts)
            s.commit()
        flash("Timesheet submitted.", "success")
        return redirect(url_for("my_timesheets"))

    # List recent timesheets for the logged-in user
    with Session(engine) as s:
        mine = s.execute(
            select(Timesheet, Engagement)
            .join(Engagement, Engagement.id == Timesheet.engagement_id)
            .where(Timesheet.user_id == int(current_user.id))
            .order_by(Timesheet.period_start.desc())
            .limit(20)
        ).all()

    return render_template("timesheets_portal.html", form=form, items=mine, week=(start, end))

@login_required
@app.route("/action/candidate/delete", methods=["POST"])
def candidate_delete_action():
    cand_id = int((request.form.get("candidate_id") or "0") or 0)
    confirm = (request.form.get("confirm") or "").lower() == "yes"
    if not cand_id:
        flash("Missing candidate id.", "warning")
        return redirect(request.referrer or url_for("resource_pool"))
    if not confirm:
        flash("Deletion cancelled.", "info")
        return redirect(request.referrer or url_for("resource_pool"))

    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found.", "warning")
            return redirect(request.referrer or url_for("resource_pool"))

        app_ids = [aid for (aid,) in s.execute(
            select(Application.id).where(Application.candidate_id == cand_id)
        ).all()]

        if app_ids:
            s.execute(delete(ESigRequest).where(ESigRequest.application_id.in_(app_ids)))
            s.execute(delete(TrustIDCheck).where(TrustIDCheck.application_id.in_(app_ids)))
            s.execute(delete(Application).where(Application.id.in_(app_ids)))

        s.execute(delete(Shortlist).where(Shortlist.candidate_id == cand_id))
        s.execute(delete(CandidateTag).where(CandidateTag.candidate_id == cand_id))
        
        # Audit log: candidate deleted
        log_audit_event('delete', 'data_access', f'Deleted candidate: {cand.name}', 
                       'candidate', cand_id,
                       {'email': cand.email, 'applications_deleted': len(app_ids)})
        
        s.delete(cand)
        s.commit()

    flash("Candidate deleted.", "success")
    return redirect(request.referrer or url_for("resource_pool"))

@login_required
@app.route("/candidate/<int:cand_id>/delete-cv/<int:doc_id>", methods=["POST"])
def candidate_delete_cv(cand_id, doc_id):
    """Delete a CV/document from a candidate's profile."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found.", "warning")
            return redirect(request.referrer or url_for("resource_pool"))
        
        doc = s.get(Document, doc_id)
        if not doc or doc.candidate_id != cand_id:
            flash("Document not found.", "warning")
            return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))
        
        s.delete(doc)
        s.commit()
        flash("Document deleted successfully.", "success")
    
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/action/candidate/add", methods=["POST"])
def add_candidate_manual():
    """Add a new candidate/associate manually from the Resource Pool."""
    name = (request.form.get("name") or "").strip()
    email = (request.form.get("email") or "").strip().lower()
    phone = (request.form.get("phone") or "").strip()
    postcode = (request.form.get("postcode") or "").strip()
    skills = (request.form.get("skills") or "").strip()
    day_rate = request.form.get("day_rate", "").strip()
    availability = (request.form.get("availability") or "").strip()
    notes = (request.form.get("notes") or "").strip()
    
    if not name or not email:
        flash("Name and Email are required.", "warning")
        return redirect(request.referrer or url_for("resource_pool"))
    
    with Session(engine) as s:
        # Check if email already exists
        existing = s.scalar(select(Candidate).where(func.lower(Candidate.email) == email))
        if existing:
            flash(f"A candidate with email {email} already exists.", "warning")
            return redirect(url_for("candidate_profile", cand_id=existing.id))
        
        # Create new candidate
        cand = Candidate(
            name=name,
            email=email,
            phone=phone if phone else None,
            postcode=postcode if postcode else None,
            skills=skills if skills else None,
            created_at=datetime.utcnow(),
            last_activity_at=datetime.utcnow(),
        )
        
        # Handle day_rate if provided
        if day_rate:
            try:
                cand.day_rate = int(day_rate)
            except ValueError:
                pass
        
        s.add(cand)
        s.flush()  # Get the candidate ID
        
        # Handle CV upload
        cv_file = request.files.get("cv")
        if cv_file and cv_file.filename:
            filename = secure_filename(cv_file.filename)
            if filename:
                ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
                if ext in {'pdf', 'doc', 'docx'}:
                    unique_filename = f"{cand.id}_{int(datetime.utcnow().timestamp())}_{filename}"
                    cv_path = os.path.join(CV_UPLOAD_FOLDER, unique_filename)
                    cv_file.save(cv_path)
                    cand.cv_filename = unique_filename
        
        # Add initial note if provided
        if notes:
            note = CandidateNote(
                candidate_id=cand.id,
                user_email=session.get("user_email", "system"),
                note_type="note",
                content=f"[Initial notes on creation]\n{notes}",
                created_at=datetime.utcnow()
            )
            s.add(note)
        
        s.commit()
        
        # Audit log: candidate created
        log_audit_event('create', 'data_access', f'Created candidate: {name}', 
                       'candidate', cand.id,
                       {'email': email, 'phone': phone})
        
        flash(f"Associate '{name}' has been added successfully.", "success")
        return redirect(url_for("candidate_profile", cand_id=cand.id))

@login_required
@app.route("/action/candidate/<int:cand_id>/update_status", methods=["POST"])
def update_candidate_status(cand_id):
    """Update candidate status and/or availability from the profile page."""
    new_status = (request.form.get("status") or "").strip()
    new_availability = (request.form.get("availability") or "").strip()
    
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found.", "danger")
            return redirect(request.referrer or url_for("resource_pool"))
        
        updates = []
        if new_status and new_status != cand.status:
            old_status = cand.status or "Available"
            cand.status = new_status
            updates.append(f"Status changed from '{old_status}' to '{new_status}'")
        
        if new_availability and new_availability != cand.availability:
            old_availability = cand.availability or "Immediately available"
            cand.availability = new_availability
            updates.append(f"Availability changed from '{old_availability}' to '{new_availability}'")
        
        if updates:
            # Add a note to track the status change
            note = CandidateNote(
                candidate_id=cand.id,
                user_email=session.get("user_email", "system"),
                note_type="activity",
                content="; ".join(updates),
                created_at=datetime.utcnow()
            )
            s.add(note)
            cand.last_activity_at = datetime.utcnow()
            s.commit()
            flash("Profile updated successfully.", "success")
        else:
            flash("No changes made.", "info")
    
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/candidate/<int:cand_id>/update", methods=["POST"])
def candidate_profile_update(cand_id):
    """Update candidate profile fields from the application detail or profile page."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found.", "danger")
            return redirect(request.referrer or url_for("resource_pool"))
        
        # Update fields from form
        if "candidate_status" in request.form:
            cand.status = request.form.get("candidate_status", "").strip() or cand.status
        if "status_notes" in request.form:
            cand.status_notes = request.form.get("status_notes", "").strip()
        if "location" in request.form:
            cand.location = request.form.get("location", "").strip()
        if "day_rate_min" in request.form:
            try:
                cand.day_rate_min = int(request.form.get("day_rate_min", 0) or 0)
            except ValueError:
                pass
        if "day_rate_max" in request.form:
            try:
                cand.day_rate_max = int(request.form.get("day_rate_max", 0) or 0)
            except ValueError:
                pass
        if "notice_period_days" in request.form:
            try:
                cand.notice_period_days = int(request.form.get("notice_period_days", 0) or 0)
            except ValueError:
                pass
        if "security_clearance" in request.form:
            cand.security_clearance = request.form.get("security_clearance", "").strip()
        if "ir35_preference" in request.form:
            cand.ir35_preference = request.form.get("ir35_preference", "").strip()
        
        cand.last_activity_at = datetime.datetime.utcnow()
        s.commit()
        flash("Profile updated successfully.", "success")
    
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/candidate/<int:cand_id>/update-status", methods=["POST"])
def candidate_update_status(cand_id):
    """Update candidate global status from profile page."""
    VALID_STATUSES = [
        "Available", "On Assignment", "On Contract", "On Notice",
        "Ex-Associate", "Unavailable", "DNU", "Do Not Contact"
    ]
    
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Associate not found.", "danger")
            return redirect(request.referrer or url_for("resource_pool"))
        
        status = request.form.get("status", "Available").strip()
        if status in VALID_STATUSES:
            old_status = cand.status or "Unknown"
            cand.status = status
            cand.last_activity_at = datetime.datetime.utcnow()
            
            # Add activity entry for status change (audit trail)
            note = CandidateNote(
                candidate_id=cand_id,
                note_type="status",
                content=f"Status changed from '{old_status}' to '{status}'",
                created_at=datetime.datetime.utcnow()
            )
            s.add(note)
            s.commit()
            flash(f"Status updated to {status}.", "success")
        else:
            flash("Invalid status.", "warning")
    
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/candidate/<int:cand_id>/add-activity", methods=["POST"])
def candidate_add_activity(cand_id):
    """Add an activity entry to candidate's activity feed."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found.", "danger")
            return redirect(request.referrer or url_for("resource_pool"))
        
        activity_type = request.form.get("activity_type", "note").strip()
        title = request.form.get("title", "").strip()
        description = request.form.get("description", "").strip()
        
        if title:
            # Add as a CandidateNote with type info
            note = CandidateNote(
                candidate_id=cand_id,
                note=f"[{activity_type.upper()}] {title}\n{description}".strip(),
                created_at=datetime.datetime.utcnow()
            )
            s.add(note)
            cand.last_activity_at = datetime.datetime.utcnow()
            s.commit()
            flash("Activity added successfully.", "success")
        else:
            flash("Title is required.", "warning")
    
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/application/<int:app_id>", methods=["GET","POST"])
def application_detail(app_id):
    # Audit log: application view
    
    interview_form = InterviewForm()

    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)

        cand = s.get(Candidate, appn.candidate_id)
        job  = s.get(Job, appn.job_id)
        if not cand or not job:
            abort(404)

        # Engagement context (may be None if data is inconsistent)
        engagement = s.get(Engagement, job.engagement_id) if job.engagement_id else None
        eng_status = ((engagement.status if engagement else "") or "").strip().lower()
        eng_active = eng_status in {"active", "in-flight", "in progress"}

        docs = s.scalars(select(Document).where(Document.candidate_id == cand.id)).all()

        # TrustID (may be None) + safe parsed result
        trust = s.scalar(select(TrustIDCheck).where(TrustIDCheck.application_id == appn.id))
        trust_result = None
        if trust and getattr(trust, "result_json", None):
            try:
                trust_result = json.loads(trust.result_json)
            except Exception:
                txt = trust.result_json or ""
                trust_result = {"raw": (txt[:400] + "…") if len(txt) > 400 else txt}

        # E-sign (fetch inside the session)
        esig = s.scalar(select(ESigRequest).where(ESigRequest.application_id == appn.id))

        # Candidate's current tags (for sidebar)
        cand_tags = s.scalars(
            select(TaxonomyTag)
            .join(CandidateTag, CandidateTag.tag_id == TaxonomyTag.id)
            .where(CandidateTag.candidate_id == cand.id)
            .order_by(TaxonomyTag.tag.asc())
        ).all()

        # All tags grouped by category for the dropdown
        cats = s.scalars(
            select(TaxonomyCategory).order_by(TaxonomyCategory.type.asc(), TaxonomyCategory.name.asc())
        ).all()
        tags_by_cat = {
            c.id: s.scalars(
                select(TaxonomyTag).where(TaxonomyTag.category_id == c.id).order_by(TaxonomyTag.tag.asc())
            ).all()
            for c in cats
        }

        # Use the global helper so Withdrawn/Closed are treated as not open
        job_open = _job_is_open(job)
        applied_on = appn.created_at

        # Stage banner context (very light-touch ordering)
        # (contract signed > contract issued > vetting > interview completed > interview scheduled > shortlisted > declared)
        stage_ctx = None
        if esig and (esig.status or "").lower() in {"signed", "completed"}:
            stage_ctx = f"Contract signed for {job.title}"
        elif esig and (esig.status or "").lower() in {"sent", "delivered"}:
            stage_ctx = f"Contract issued for {job.title}"
        elif trust and (trust.status or "").lower() in {"created", "inprogress"}:
            stage_ctx = "Vetting in progress"
        elif appn.interview_completed_at:
            stage_ctx = "Interview completed"
        elif appn.interview_scheduled_at:
            when = appn.interview_scheduled_at.strftime("%Y-%m-%d %H:%M")
            stage_ctx = f"Interview scheduled — {when}"
        else:
            # shortlisted?
            sl_exists = s.scalar(
                select(Shortlist.id).where(Shortlist.job_id == job.id, Shortlist.candidate_id == cand.id).limit(1)
            )
            if sl_exists:
                stage_ctx = f"Shortlisted for {job.title}"
            else:
                stage_ctx = "Declared interest"

        # Disable action buttons if engagement not active or job not open
        actions_disabled = (not eng_active) or (not job_open)

    # Render (note: no DB work after the session closes)
    return render_template(
        "application_detail.html",
        appn=appn,
        cand=cand,
        job=job,
        engagement=engagement,
        docs=docs,
        trust=trust,
        trust_result=trust_result,
        esig=esig,
        interview_form=interview_form,
        # Tags-only context:
        cand_tags=cand_tags,
        cats=cats,
        tags_by_cat=tags_by_cat,
        # Meta:
        job_open=job_open,
        applied_on=applied_on,
        # Stage banner + UX hints:
        stage_context=stage_ctx,
        actions_disabled=actions_disabled,
    )

@login_required
@app.route("/action/score/<int:app_id>", methods=["POST"])
def action_score(app_id):
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id==app_id))
        if not appn:
            abort(404)
        job = s.scalar(select(Job).where(Job.id==appn.job_id))
        doc = s.scalar(
            select(Document)
            .where(Document.candidate_id==appn.candidate_id)
            .order_by(Document.uploaded_at.desc())
        )
        cv_text = extract_cv_text(doc) if doc else ""
        result = ai_score_with_explanation(job.description or "", cv_text)
        appn.ai_score = int(result["final"])
        appn.ai_explanation = (result.get("explanation") or "")[:7999]
        s.commit()
    flash("AI score updated", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/summarise/<int:app_id>", methods=["POST"])
def action_summarise(app_id):
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id==app_id))
        if not appn:
            abort(404)
        doc = s.scalar(
            select(Document)
            .where(Document.candidate_id==appn.candidate_id)
            .order_by(Document.uploaded_at.desc())
        )
        cv_text = extract_cv_text(doc) if doc else ""
        appn.ai_summary = ai_summarise(cv_text or "")
        s.commit()
    flash("AI summary generated", "success")
    return redirect(url_for("application_detail", app_id=app_id))

# -------- TrustID placeholders (as before) --------
def trustid_headers():
    return {"Authorization": f"Bearer {TRUSTID_API_KEY}", "Content-Type": "application/json"}

@login_required
@app.route("/action/trustid/<int:app_id>", methods=["POST"])
def action_trustid(app_id):
    do_rtw = bool(request.form.get("rtw"))
    do_idv = bool(request.form.get("idv", True))
    do_dbs = bool(request.form.get("dbs"))
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id==app_id))
        cand = s.scalar(select(Candidate).where(Candidate.id==appn.candidate_id))

        payload = {
            "external_ref": f"APP-{appn.id}",
            "candidate": {"name": cand.name, "email": cand.email},
            "checks": {"rtw": do_rtw, "idv": do_idv, "dbs": do_dbs},
            "webhooks": {"result": f"{APP_BASE_URL}/webhook/trustid"}
        }
        try:
            resp = requests.post(f"{TRUSTID_BASE_URL}/apps", headers=trustid_headers(), data=json.dumps(payload), timeout=15)
            if resp.status_code >= 300:
                raise RuntimeError(f"TrustID error: {resp.status_code} {resp.text}")
            data = resp.json()
            trust_app_id = data.get("id", "")
        except Exception as e:
            flash(f"TrustID create failed: {e}", "danger")
            return redirect(url_for("application_detail", app_id=app_id))

        tic = TrustIDCheck(application_id=appn.id, rtw=do_rtw, idv=do_idv, dbs=do_dbs,
                           trustid_application_id=trust_app_id, status="InProgress")
        s.add(tic)
        s.commit()
    flash("TrustID checks initiated", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/webhook/trustid", methods=["POST"])
def webhook_trustid():
    """
    Receives TrustID webhook notifications, verifies signature (if secret configured),
    stores the raw event, and updates TrustIDCheck records with the latest result JSON.
    """

    # --- Verify HMAC signature if configured ---
    payload_bytes = request.get_data()
    event_type = request.headers.get("X-TrustID-Event", "unknown")
    received_sig = request.headers.get("X-TrustID-Signature", "")

    if TRUSTID_WEBHOOK_SECRET:
        try:
            expected_sig = hmac.new(
                TRUSTID_WEBHOOK_SECRET.encode("utf-8"),
                payload_bytes,
                hashlib.sha256
            ).hexdigest()
            if not hmac.compare_digest(expected_sig, received_sig):
                current_app.logger.warning("⚠️ TrustID webhook signature verification failed")
                abort(401)
        except Exception as e:
            current_app.logger.warning(f"⚠️ TrustID webhook signature error: {e}")
            abort(401)

    # --- Log the event payload for audit trail ---
    payload_text = payload_bytes.decode("utf-8", errors="ignore")
    with Session(engine) as s:
        s.add(WebhookEvent(
            source="trustid",
            event_type=event_type,
            payload=payload_text[:39999]
        ))
        s.commit()

    # --- Parse event JSON ---
    data = {}
    try:
        data = json.loads(payload_text)
    except Exception:
        pass

    trustid_app_id = data.get("application_id") or data.get("id") or ""
    if not trustid_app_id:
        current_app.logger.warning("⚠️ No TrustID application_id found in webhook payload")
        return jsonify({"ok": False, "error": "missing application_id"}), 400

    # --- Fetch latest TrustID result from API ---
    try:
        resp = requests.get(
            f"{TRUSTID_BASE_URL}/apps/{trustid_app_id}/results",
            headers={"Authorization": f"Bearer {TRUSTID_API_KEY}", "Content-Type": "application/json"},
            timeout=15,
        )
        if resp.status_code >= 300:
            raise RuntimeError(f"TrustID API returned {resp.status_code}: {resp.text}")
        result_json = resp.json()
        status = "Completed"
    except Exception as e:
        result_json = {"error": str(e)}
        status = "Error"
        current_app.logger.warning(f"⚠️ TrustID webhook result fetch failed for {trustid_app_id}: {e}")

    # --- Update our local TrustIDCheck record ---
    with Session(engine) as s:
        tic = s.scalar(select(TrustIDCheck).where(TrustIDCheck.trustid_application_id == trustid_app_id))
        if tic:
            tic.result_json = json.dumps(result_json)[:19999]
            tic.status = status
            s.commit()
        else:
            current_app.logger.warning(f"⚠️ No TrustIDCheck row found for {trustid_app_id}")

    return jsonify({"ok": True})

# -------- Interview scheduling (ICS) --------
@login_required
@app.route("/action/schedule_interview/<int:app_id>", methods=["POST"])
def action_schedule_interview(app_id):
    form_dt = request.form.get("scheduled_at")
    interviewer_email = request.form.get("interviewer_email") or INTERVIEWER_EMAIL

    # parse start time
    start = dtparser.parse(form_dt) if form_dt else None
    if not start:
        flash("Please provide a valid date/time", "danger")
        return redirect(url_for("application_detail", app_id=app_id))
    end = start + datetime.timedelta(hours=1)

    # we'll fill these before session closes
    cand_name = ""
    cand_email = ""
    job_title = ""

    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)

        cand = s.scalar(select(Candidate).where(Candidate.id == appn.candidate_id))
        job  = s.scalar(select(Job).where(Job.id == appn.job_id))

        if not cand or not job:
            abort(404)

        # cache the bits we need OUTSIDE the session
        cand_name = cand.name or ""
        cand_email = cand.email or ""
        job_title = job.title or ""

        # GAP 1.2 FIX: Auto-trigger workflow status when interview is scheduled
        appn.interview_scheduled_at = start
        appn.status = "I&A"  # Interview & Assessment stage
        
        # Update candidate status to reflect interviewing state
        cand.status = "Interviewing"
        cand.last_activity_at = datetime.datetime.utcnow()
        
        # GAP X.4: Audit log for compliance
        log_audit_event(
            'update', 'workflow', 
            f'Interview scheduled for {cand_name} - {job_title}',
            'application', app_id,
            {'scheduled_at': start.isoformat(), 'old_status': appn.status, 'new_status': 'I&A'}
        )
        
        s.commit()

    # now we're OUTSIDE the session but we only use plain strings, not ORM objects
    summary = f"Interview — {job_title}"
    description = f"Interview for {job_title} with {cand_name}"
    location = "Microsoft Teams / Zoom (link to follow)"

    ics_bytes = ics_invite(
        summary,
        description,
        start,
        end,
        location,
        SMTP_FROM,
        cand_email,
    )

    # send calendar invite to candidate
    send_email(
        cand_email,
        summary,
        f"<p>Hi {cand_name}, your interview is scheduled for {start}.</p>",
        attachments=[("interview.ics", ics_bytes, "text/calendar")],
    )

    # send calendar invite to interviewer
    send_email(
        interviewer_email,
        summary,
        f"<p>Interview scheduled with {cand_name} for {start}.</p>",
        attachments=[("interview.ics", ics_bytes, "text/calendar")],
    )

    flash("Interview scheduled and invites sent", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/mark_interview_completed/<int:app_id>", methods=["POST"])
def action_mark_interview_completed(app_id):
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id==app_id))
        if not appn:
            abort(404)
        appn.interview_completed_at = datetime.datetime.utcnow()
        appn.status = "Client Review"  # Move to Client Review stage after internal interview
        s.commit()
    flash("Interview marked as completed - moved to Client Review", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/complete_interview/<int:app_id>", methods=["POST"])
def action_complete_interview(app_id):
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)
        appn.status = "Client Review"  # Move to Client Review stage
        appn.interview_completed_at = datetime.datetime.utcnow()
        s.commit()
    flash("Interview completed - moved to Client Review.", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/skip_interview/<int:app_id>", methods=["POST"])
def action_skip_interview(app_id):
    """Skip interview and move to Client Review without scheduling"""
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)
        # Mark interview as completed without scheduling
        appn.interview_completed_at = datetime.datetime.utcnow()
        appn.interview_notes = (appn.interview_notes or "") + "\n[Interview skipped]"
        appn.status = "Client Review"  # Move to Client Review stage
        s.commit()
    flash("Interview skipped - moved to Client Review.", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/interview_not_required/<int:app_id>", methods=["POST"])
def action_interview_not_required(app_id):
    """Mark interview as not required for this application."""
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)
        appn.status = "Client Review"  # Move to Client Review stage
        appn.interview_notes = (appn.interview_notes or "") + "\n[Interview not required]"
        s.commit()
    flash("Interview not required - moved to Client Review.", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/vetting/<int:cand_id>", methods=["POST"])
def action_vetting(cand_id):
    """Start vetting process for this candidate - moves to Accepted stage."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            flash("Candidate not found.", "danger")
            return redirect(request.referrer or url_for("resource_pool"))
        # Find any active application for this candidate and update status
        appn = s.scalar(
            select(Application)
            .where(Application.candidate_id == cand_id)
            .order_by(Application.created_at.desc())
        )
        if appn:
            appn.status = "Accepted"  # Offer accepted, vetting starts
            s.commit()
    flash("Vetting process started - candidate moved to Accepted stage.", "success")
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/action/save_general_notes/<int:cand_id>", methods=["POST"])
def action_save_general_notes(cand_id):
    """Save general notes for a candidate."""
    notes = request.form.get("general_notes", "").strip()
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if cand:
            cand.general_notes = notes
            s.commit()
    flash("Notes saved.", "success")
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/action/contract/issue/<int:cand_id>/<int:eng_id>", methods=["POST"])
def action_contract_issue(cand_id, eng_id):
    """Issue contract for this candidate and engagement."""
    flash("Contract issued.", "success")
    return redirect(request.referrer or url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/action/contract/sign/<int:contract_id>", methods=["POST"])
def action_contract_sign(contract_id):
    """Mark contract as signed."""
    flash("Contract signed.", "success")
    return redirect(request.referrer or url_for("resource_pool"))

@login_required
@app.route("/action/contract/notice/<int:contract_id>", methods=["POST"])
def action_contract_notice(contract_id):
    """Put contract on notice."""
    flash("Contract put on notice.", "info")
    return redirect(request.referrer or url_for("resource_pool"))

@login_required
@app.route("/action/contract/cancel_notice/<int:contract_id>", methods=["POST"])
def action_contract_cancel_notice(contract_id):
    """Cancel notice on contract."""
    flash("Notice cancelled.", "info")
    return redirect(request.referrer or url_for("resource_pool"))

@login_required
@app.route("/action/contract/extend/<int:contract_id>", methods=["POST"])
def action_contract_extend(contract_id):
    """Extend contract."""
    flash("Contract extended.", "success")
    return redirect(request.referrer or url_for("resource_pool"))

@login_required
@app.route("/action/contract/terminate/<int:contract_id>", methods=["POST"])
def action_contract_terminate(contract_id):
    """Terminate contract."""
    flash("Contract terminated.", "warning")
    return redirect(request.referrer or url_for("resource_pool"))

@login_required
@app.route("/action/contract/update/<int:contract_id>", methods=["POST"])
def action_contract_update(contract_id):
    """Update contract details."""
    flash("Contract updated.", "success")
    return redirect(request.referrer or url_for("resource_pool"))

@login_required
@app.route("/action/save_interview_notes/<int:app_id>", methods=["POST"])
def action_save_interview_notes(app_id):
    notes = request.form.get("interview_notes", "").strip()
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)
        appn.interview_notes = notes
        s.commit()
    flash("Interview notes saved.", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/save_assessor_notes/<int:app_id>", methods=["POST"])
def action_save_assessor_notes(app_id):
    notes = request.form.get("assessor_notes", "").strip()
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id == app_id))
        if not appn:
            abort(404)
        appn.cover_note = notes
        s.commit()
    flash("Assessor notes saved.", "success")
    return redirect(url_for("application_detail", app_id=app_id))

# -------- E-signature: send & poll --------
def send_esign_dropbox_sign(appn_id: int, signer_name: str, signer_email: str, subject: str, message: str):
    if not HELLOSIGN_API_KEY or ApiClient is None:
        raise RuntimeError("Dropbox Sign SDK or API key not configured")
    conf = Configuration(username=HELLOSIGN_API_KEY)
    with ApiClient(conf) as api_client:
        sig_api = SignatureRequestApi(api_client)
        signers = [SubSignatureRequestSigner(email_address=signer_email, name=signer_name)]
        data = SignatureRequestSendRequest(
            title=subject,
            subject=subject,
            message=message,
            signers=signers,
            files=[],
            test_mode=True
        )
        res = sig_api.signature_request_send(data)
        return res.signature_request.signature_request_id

def poll_esign_dropbox_sign(request_id: str):
    if not HELLOSIGN_API_KEY or ApiClient is None:
        raise RuntimeError("Dropbox Sign SDK or API key not configured")
    conf = Configuration(username=HELLOSIGN_API_KEY)
    with ApiClient(conf) as api_client:
        sig_api = SignatureRequestApi(api_client)
        res = sig_api.signature_request_get(request_id)
        status = "Sent"
        if res and res.signature_request and res.signature_request.is_complete:
            status = "Signed"
        return status

def send_esign_docusign(appn_id: int, signer_name: str, signer_email: str, subject: str, message: str):
    if dse is None or not DOCUSIGN_ACCESS_TOKEN or not DOCUSIGN_ACCOUNT_ID:
        raise RuntimeError("DocuSign SDK or access token/account not configured")
    api_client = dse.ApiClient()
    api_client.host = DOCUSIGN_BASE_PATH
    api_client.set_default_header("Authorization", f"Bearer {DOCUSIGN_ACCESS_TOKEN}")
    envelopes_api = dse.EnvelopesApi(api_client)
    doc_base64 = base64.b64encode(b"Please sign this simple agreement.").decode("ascii")
    document = dse.Document(document_base64=doc_base64, name="Agreement", file_extension="txt", document_id="1")
    signer = dse.Signer(email=signer_email, name=signer_name, recipient_id="1", routing_order="1")
    sign_here = dse.SignHere(document_id="1", page_number="1", recipient_id="1", tab_label="SignHere", x_position="75", y_position="572")
    tabs = dse.Tabs(sign_here_tabs=[sign_here])
    signer.tabs = tabs
    recipients = dse.Recipients(signers=[signer])
    envelope_definition = dse.EnvelopeDefinition(email_subject=subject, documents=[document], recipients=recipients, status="sent")
    results = envelopes_api.create_envelope(account_id=DOCUSIGN_ACCOUNT_ID, envelope_definition=envelope_definition)
    return results.envelope_id

def poll_esign_docusign(request_id: str):
    if dse is None or not DOCUSIGN_ACCESS_TOKEN or not DOCUSIGN_ACCOUNT_ID:
        raise RuntimeError("DocuSign SDK or access token/account not configured")
    api_client = dse.ApiClient()
    api_client.host = DOCUSIGN_BASE_PATH
    api_client.set_default_header("Authorization", f"Bearer {DOCUSIGN_ACCESS_TOKEN}")
    envelopes_api = dse.EnvelopesApi(api_client)
    res = envelopes_api.get_envelope(account_id=DOCUSIGN_ACCOUNT_ID, envelope_id=request_id)
    status_map = {"sent":"Sent","completed":"Signed","declined":"Declined","voided":"Error"}
    return status_map.get((res.status or "").lower(), res.status or "Unknown")

@login_required
@app.route("/action/esign/<int:app_id>", methods=["POST"])
def action_esign(app_id):
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id==app_id))
        cand = s.scalar(select(Candidate).where(Candidate.id==appn.candidate_id))
        job = s.scalar(select(Job).where(Job.id==appn.job_id))

        subject = f"Contract — {job.title}"
        message = "Please review and sign your contract."
        try:
            if ESIGN_PROVIDER == "dropbox_sign":
                req_id = send_esign_dropbox_sign(appn.id, cand.name, cand.email, subject, message)
            elif ESIGN_PROVIDER == "docusign":
                req_id = send_esign_docusign(appn.id, cand.name, cand.email, subject, message)
            else:
                raise RuntimeError("Unsupported ESIGN_PROVIDER")
        except Exception as e:
            flash(f"E-sign send failed: {e}", "danger")
            return redirect(url_for("application_detail", app_id=app_id))

        es = s.scalar(select(ESigRequest).where(ESigRequest.application_id==appn.id))
        if not es:
            es = ESigRequest(application_id=appn.id, provider=ESIGN_PROVIDER, request_id=req_id, status="Sent", sent_at=datetime.datetime.utcnow())
            s.add(es)
        else:
            es.request_id = req_id
            es.provider = ESIGN_PROVIDER
            es.status = "Sent"
            es.sent_at = datetime.datetime.utcnow()
        s.commit()
    flash("Contract sent for e-signature", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/action/esign_status/<int:app_id>", methods=["POST"])
def action_esign_status(app_id):
    with Session(engine) as s:
        es = s.scalar(select(ESigRequest).where(ESigRequest.application_id==app_id))
        if not es:
            flash("No e-sign request found", "warning")
            return redirect(url_for("application_detail", app_id=app_id))
        try:
            if es.provider == "dropbox_sign":
                status = poll_esign_dropbox_sign(es.request_id)
            elif es.provider == "docusign":
                status = poll_esign_docusign(es.request_id)
            else:
                status = "Unknown"
        except Exception as e:
            flash(f"E-sign status failed: {e}", "danger")
            return redirect(url_for("application_detail", app_id=app_id))

        es.status = status
        if status == "Signed":
            es.signed_at = datetime.datetime.utcnow()
        s.commit()
    flash(f"E-sign status: {status}", "info")
    return redirect(url_for("application_detail", app_id=app_id))

@login_required
@app.route("/webhook/esign", methods=["POST"])
def webhook_esign():
    payload = request.json or {}
    with Session(engine) as s:
        s.add(WebhookEvent(source="esign", event_type=str(payload.get('event','unknown')), payload=json.dumps(payload)[:39999]))
        s.commit()
    return jsonify({"ok": True})

# ---- Request updated CV ----
@login_required
@app.route("/action/request_updated_cv/<int:app_id>", methods=["POST"])
def action_request_updated_cv(app_id):
    with Session(engine) as s:
        appn = s.scalar(select(Application).where(Application.id==app_id))
        cand = s.scalar(select(Candidate).where(Candidate.id==appn.candidate_id))
        link = f"{APP_BASE_URL}/candidate/{cand.id}/upload_cv"
        html = f"""
        <h3>Request for updated CV</h3>
        <p>Hi {cand.name}, we enjoyed working with you. Please upload an updated CV here:</p>
        <p><a href="{link}">{link}</a></p>
        """
        send_email(cand.email, "Please upload an updated CV", html)
    flash("CV update request sent", "success")
    return redirect(url_for("application_detail", app_id=app_id))

@app.route("/candidate/<int:cand_id>")
def candidate_profile(cand_id: int):
    """
    Associate Profile page per wireframe requirements.
    
    Features:
    - Header with name, email, alert banner
    - Notes & Activity Panel (rich-text editor)
    - Profile sidebar (status, location, rate, CV)
    - Tags section with Retag from CV
    - Contract section with Issue Contract button
    - 12-item Vetting Checks grid
    - Vetting Progress Summary
    """
    # Optional context passed when arriving from engagement lists
    stage = (request.args.get("stage") or "").strip().lower()
    ctx = {
        "stage": stage or None,
        "job_id": request.args.get("job_id"),
        "job_title": request.args.get("job_title"),
        "interview_at": request.args.get("interview_at"),
    }

    # 12 vetting check types per wireframe
    VETTING_CHECK_TYPES = [
        "Right to Work",
        "Identity Verification",
        "Address History",
        "DBS Check",
        "Employment History",
        "References",
        "Qualifications",
        "Professional Registration",
        "Credit Check",
        "Directorship / Disqualification",
        "Sanctions / PEP",
        "Social Media Review",
    ]

    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)

        latest_app = s.scalar(
            select(Application)
            .where(Application.candidate_id == cand_id)
            .order_by(Application.created_at.desc())
        )

        job = None
        job_open = False
        applied_on = None
        docs = s.scalars(select(Document).where(Document.candidate_id == cand.id)).all()

        if latest_app:
            job = s.scalar(select(Job).where(Job.id == latest_app.job_id))
            job_open = _job_is_open(job)
            applied_on = latest_app.created_at

        # Tags data for the sidebar card
        cand_tags = s.scalars(
            select(TaxonomyTag)
            .join(CandidateTag, CandidateTag.tag_id == TaxonomyTag.id)
            .where(CandidateTag.candidate_id == cand.id)
            .order_by(TaxonomyTag.tag.asc())
        ).all()

        cats = s.scalars(
            select(TaxonomyCategory).order_by(TaxonomyCategory.type.asc(), TaxonomyCategory.name.asc())
        ).all()

        tags_by_cat = {
            c.id: s.scalars(
                select(TaxonomyTag).where(TaxonomyTag.category_id == c.id).order_by(TaxonomyTag.tag.asc())
            ).all()
            for c in cats
        }
        
        # === Vetting Checks ===
        vetting_checks = []
        vetting_summary = {"complete": 0, "in_progress": 0, "not_started": 0, "na": 0}
        
        try:
            existing_checks = {
                vc.check_type: vc
                for vc in s.scalars(
                    select(VettingCheck).where(VettingCheck.candidate_id == cand_id)
                ).all()
            }
            
            for check_type in VETTING_CHECK_TYPES:
                if check_type in existing_checks:
                    check = existing_checks[check_type]
                    vetting_checks.append({
                        "id": check.id,
                        "type": check_type,
                        "status": check.status or "NOT STARTED",
                        "notes": check.notes or "",
                    })
                else:
                    # Create a placeholder for display
                    vetting_checks.append({
                        "id": None,
                        "type": check_type,
                        "status": "NOT STARTED",
                        "notes": "",
                    })
            
            # Calculate summary
            for vc in vetting_checks:
                status = (vc["status"] or "").upper()
                if status == "COMPLETE":
                    vetting_summary["complete"] += 1
                elif status == "IN PROGRESS":
                    vetting_summary["in_progress"] += 1
                elif status == "N/A":
                    vetting_summary["na"] += 1
                else:
                    vetting_summary["not_started"] += 1
        except Exception:
            # VettingCheck table might not exist yet
            for check_type in VETTING_CHECK_TYPES:
                vetting_checks.append({
                    "id": None,
                    "type": check_type,
                    "status": "NOT STARTED",
                    "notes": "",
                })
                vetting_summary["not_started"] += 1
        
        # === Notes & Activity ===
        candidate_notes = []
        try:
            candidate_notes = s.scalars(
                select(CandidateNote)
                .where(CandidateNote.candidate_id == cand_id)
                .order_by(CandidateNote.created_at.desc())
                .limit(50)
            ).all()
        except Exception:
            pass
        
        # === Contract Status ===
        contract_status = None
        esig_request = None
        try:
            esig_request = s.scalar(
                select(ESigRequest)
                .where(ESigRequest.candidate_id == cand_id)
                .order_by(ESigRequest.created_at.desc())
            )
            if esig_request:
                contract_status = esig_request.status
        except Exception:
            pass
        
        # Check if there are any active applications (for alert banner)
        active_apps_count = s.scalar(
            select(func.count(Application.id))
            .where(Application.candidate_id == cand_id)
            .where(Application.status.notin_(["Rejected", "Hired", "Contract Signed"]))
        ) or 0

        # === Placements History (Active and Historic) ===
        placements_active = []
        placements_historic = []
        try:
            all_contracts = s.execute(
                select(ESigRequest, Job, Engagement)
                .select_from(ESigRequest)
                .outerjoin(Application, Application.id == ESigRequest.application_id)
                .outerjoin(Job, Job.id == Application.job_id)
                .outerjoin(Engagement, Engagement.id == ESigRequest.engagement_id)
                .where(ESigRequest.candidate_id == cand_id)
                .order_by(ESigRequest.signed_at.desc())
            ).all()
            
            today = datetime.date.today()
            for esig, job_obj, eng in all_contracts:
                placement = {
                    'id': esig.id,
                    'engagement_name': eng.name if eng else (job_obj.title if job_obj else 'Unknown'),
                    'engagement_id': eng.id if eng else None,
                    'client': eng.client if eng else 'Unknown',
                    'role': job_obj.title if job_obj else 'N/A',
                    'start_date': eng.start_date if eng else None,
                    'end_date': eng.end_date if eng else None,
                    'status': esig.status,
                    'signed_at': esig.signed_at,
                }
                
                # Determine if active or historic
                if eng and eng.end_date:
                    try:
                        end_dt = eng.end_date if isinstance(eng.end_date, datetime.date) else datetime.datetime.strptime(str(eng.end_date), '%Y-%m-%d').date()
                        if end_dt >= today:
                            placements_active.append(placement)
                        else:
                            placements_historic.append(placement)
                    except:
                        placements_active.append(placement)
                else:
                    placements_active.append(placement)
        except Exception:
            pass

        # === Activity Feed (structured activity log) ===
        activity_feed = []
        try:
            # Get applications (with status changes)
            apps = s.execute(
                select(Application, Job)
                .join(Job, Job.id == Application.job_id)
                .where(Application.candidate_id == cand_id)
                .order_by(Application.created_at.desc())
                .limit(20)
            ).all()
            
            for app, job_obj in apps:
                activity_feed.append({
                    'type': 'application',
                    'icon': 'fa-file-alt',
                    'color': '#3b82f6',
                    'title': f'Applied for {job_obj.title}',
                    'timestamp': app.created_at,
                    'details': f'Status: {app.status}'
                })
                if app.interview_scheduled_at:
                    activity_feed.append({
                        'type': 'interview_scheduled',
                        'icon': 'fa-calendar-check',
                        'color': '#f59e0b',
                        'title': 'Interview Scheduled',
                        'timestamp': app.interview_scheduled_at,
                        'details': f'For {job_obj.title}'
                    })
                if app.interview_completed_at:
                    activity_feed.append({
                        'type': 'interview_completed',
                        'icon': 'fa-check-circle',
                        'color': '#10b981',
                        'title': 'Interview Completed',
                        'timestamp': app.interview_completed_at,
                        'details': f'For {job_obj.title}'
                    })
            
            # Get contract events
            contracts = s.scalars(
                select(ESigRequest)
                .where(ESigRequest.candidate_id == cand_id)
                .order_by(ESigRequest.created_at.desc())
                .limit(10)
            ).all()
            
            for contract in contracts:
                if contract.sent_at:
                    activity_feed.append({
                        'type': 'contract_sent',
                        'icon': 'fa-paper-plane',
                        'color': '#8b5cf6',
                        'title': 'Contract Issued',
                        'timestamp': contract.sent_at,
                        'details': f'Provider: {contract.provider or "N/A"}'
                    })
                if contract.signed_at:
                    activity_feed.append({
                        'type': 'contract_signed',
                        'icon': 'fa-file-signature',
                        'color': '#059669',
                        'title': 'Contract Signed',
                        'timestamp': contract.signed_at,
                        'details': f'Status: {contract.status}'
                    })
            
            # Get notes as activity
            notes = s.scalars(
                select(CandidateNote)
                .where(CandidateNote.candidate_id == cand_id)
                .order_by(CandidateNote.created_at.desc())
                .limit(10)
            ).all()
            
            for note in notes:
                activity_feed.append({
                    'type': 'note',
                    'icon': 'fa-sticky-note',
                    'color': '#6b7280',
                    'title': note.note_type or 'Note Added',
                    'timestamp': note.created_at,
                    'details': (note.content[:100] + '...') if note.content and len(note.content) > 100 else (note.content or '')
                })
            
            # Sort by timestamp descending
            activity_feed.sort(key=lambda x: x['timestamp'] if x['timestamp'] else datetime.datetime.min, reverse=True)
            activity_feed = activity_feed[:30]  # Limit to 30 most recent
        except Exception:
            pass

    return render_template(
        "application_detail.html",
        appn=latest_app,            # can be None
        cand=cand,
        job=job,
        docs=docs,
        trust=None,
        esig=esig_request,
        interview_form=None,
        job_open=job_open,
        applied_on=applied_on,
        from_candidate=True,

        # taxonomy/tag data
        cand_tags=cand_tags,
        cats=cats,
        tags_by_cat=tags_by_cat,

        # engagement/list context for banner in template
        context=ctx if ctx.get("stage") else None,
        
        # === New wireframe data ===
        vetting_checks=vetting_checks,
        vetting_summary=vetting_summary,
        candidate_notes=candidate_notes,
        contract_status=contract_status,
        active_apps_count=active_apps_count,
        VETTING_CHECK_TYPES=VETTING_CHECK_TYPES,
        # === Activity Feed & Placements ===
        activity_feed=activity_feed,
        placements_active=placements_active,
        placements_historic=placements_historic,
    )

# -------- Vetting Check Update --------
@login_required
@app.route("/candidate/<int:cand_id>/vetting", methods=["POST"])
def update_vetting_check(cand_id: int):
    """
    Update a vetting check status for a candidate.
    
    GAP 7.1 FIX: When all vetting checks are Complete or N/A, 
    auto-moves candidate to next workflow stage.
    GAP X.4: Audit trail for compliance-sensitive changes.
    """
    check_type = request.form.get("check_type", "").strip()
    new_status = request.form.get("status", "NOT STARTED")
    notes = request.form.get("notes", "")
    
    if not check_type:
        flash("Invalid vetting check type", "error")
        return redirect(url_for("candidate_profile", cand_id=cand_id))
    
    # All vetting check types
    ALL_VETTING_CHECKS = [
        "Right to Work", "Identity Verification", "Address History", "DBS Check",
        "Employment History", "References", "Qualifications", "Professional Registration",
        "Credit Check", "Directorship / Disqualification", "Sanctions / PEP", "Social Media Review"
    ]
    
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        old_status_value = None
        
        # Check if this vetting check already exists
        existing = s.scalar(
            select(VettingCheck)
            .where(VettingCheck.candidate_id == cand_id)
            .where(VettingCheck.check_type == check_type)
        )
        
        if existing:
            old_status_value = existing.status
            existing.status = new_status
            existing.notes = notes
            if new_status.upper() == "COMPLETE":
                existing.completed_at = datetime.datetime.utcnow()
        else:
            new_check = VettingCheck(
                candidate_id=cand_id,
                check_type=check_type,
                status=new_status,
                notes=notes,
                completed_at=datetime.datetime.utcnow() if new_status.upper() == "COMPLETE" else None
            )
            s.add(new_check)
        
        # GAP X.4: Audit log for vetting changes (compliance critical)
        log_audit_event(
            'update', 'vetting',
            f'Vetting check "{check_type}" updated to "{new_status}"',
            'candidate', cand_id,
            {'check_type': check_type, 'old_status': old_status_value, 'new_status': new_status, 'notes': notes}
        )
        
        s.commit()
        
        # GAP 7.1: Check if all vetting is now complete
        all_checks = s.execute(
            select(VettingCheck)
            .where(VettingCheck.candidate_id == cand_id)
        ).scalars().all()
        
        check_statuses = {vc.check_type: vc.status.upper() for vc in all_checks}
        
        # Count completed checks (COMPLETE or N/A are considered done)
        completed_count = sum(1 for ct in ALL_VETTING_CHECKS 
                            if check_statuses.get(ct, "NOT STARTED") in ["COMPLETE", "N/A"])
        
        all_vetting_complete = completed_count == len(ALL_VETTING_CHECKS)
        
        if all_vetting_complete:
            # Find the candidate's active application in Vetting In-Flight stage
            active_app = s.scalar(
                select(Application)
                .where(Application.candidate_id == cand_id)
                .where(Application.status == "Vetting In-Flight")
                .order_by(Application.created_at.desc())
            )
            
            if active_app:
                # Auto-move to Contract Issued (GAP 1.2 auto-trigger)
                active_app.status = "Contract Issued"
                cand.status = "In Vetting"  # Keep as In Vetting until contract signed
                cand.previously_vetted = True
                
                log_audit_event(
                    'update', 'workflow',
                    f'Auto-triggered: All vetting complete, moved to Contract Issued',
                    'application', active_app.id,
                    {'trigger': 'vetting_complete', 'new_status': 'Contract Issued'}
                )
                
                s.commit()
                flash(f"All vetting complete! {cand.name} has been moved to 'Contract Issued' stage.", "success")
                return redirect(url_for("candidate_profile", cand_id=cand_id))
    
    flash(f"Vetting check '{check_type}' updated to '{new_status}'", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

# -------- Add Candidate Note --------
@login_required
@app.route("/candidate/<int:cand_id>/note", methods=["POST"])
def add_candidate_note(cand_id: int):
    """Add a note to a candidate's profile"""
    content = request.form.get("note_content", "").strip() or request.form.get("content", "").strip()
    note_type = request.form.get("note_type", "note")
    
    if not content:
        flash("Note content cannot be empty", "warning")
        return redirect(url_for("candidate_profile", cand_id=cand_id))
    
    with Session(engine) as s:
        user_email = session.get("user_email", "System")
        new_note = CandidateNote(
            candidate_id=cand_id,
            user_email=user_email,
            note_type=note_type,
            content=content
        )
        s.add(new_note)
        s.commit()
    
    flash("Note added successfully", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

# -------- Issue Contract --------
@login_required
@app.route("/candidate/<int:cand_id>/issue_contract", methods=["POST"])
def issue_candidate_contract(cand_id: int):
    """
    Issue a contract to a candidate.
    
    GAP 7.2 FIX: When contract is issued, auto-updates workflow stage
    to 'Contract Issued' and updates candidate status.
    GAP X.4: Audit trail for compliance.
    """
    engagement_id = request.form.get("engagement_id")
    job_id = request.form.get("job_id")
    
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)
        
        # Create an ESigRequest
        esig = ESigRequest(
            candidate_id=cand_id,
            engagement_id=int(engagement_id) if engagement_id else None,
            status="pending",
            created_at=datetime.datetime.utcnow()
        )
        s.add(esig)
        
        # GAP 7.2: Auto-update workflow when contract is issued
        # Find the most recent application for this candidate
        if job_id:
            active_app = s.scalar(
                select(Application)
                .where(Application.candidate_id == cand_id)
                .where(Application.job_id == int(job_id))
            )
        else:
            active_app = s.scalar(
                select(Application)
                .where(Application.candidate_id == cand_id)
                .where(Application.status.notin_(["Rejected", "Contract Signed", "Hired"]))
                .order_by(Application.created_at.desc())
            )
        
        if active_app:
            old_status = active_app.status
            active_app.status = "Contract Issued"
            esig.application_id = active_app.id
            
            # GAP X.4: Audit log for contract issuance
            log_audit_event(
                'create', 'contract',
                f'Contract issued to {cand.name}',
                'candidate', cand_id,
                {
                    'engagement_id': engagement_id,
                    'application_id': active_app.id,
                    'old_workflow_status': old_status,
                    'new_workflow_status': 'Contract Issued'
                }
            )
        else:
            log_audit_event(
                'create', 'contract',
                f'Contract issued to {cand.name} (no active application)',
                'candidate', cand_id,
                {'engagement_id': engagement_id}
            )
        
        # Update candidate status
        cand.status = "In Vetting"  # Contract issued but not yet signed
        cand.last_activity_at = datetime.datetime.utcnow()
        
        s.commit()
    
    flash("Contract issued successfully — candidate moved to 'Contract Issued' stage", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

# -------- GAP 1.5: Contract Resend/Chase Actions --------
@login_required
@app.route("/api/contract/<int:esig_id>/resend", methods=["POST"])
def resend_contract(esig_id: int):
    """
    GAP 1.5: Resend a contract that hasn't been signed.
    Updates the created_at timestamp and sends a new email.
    """
    with Session(engine) as s:
        esig = s.get(ESigRequest, esig_id)
        if not esig:
            return jsonify({"ok": False, "error": "Contract not found"}), 404
        
        if esig.status != "pending":
            return jsonify({"ok": False, "error": "Contract is not pending"}), 400
        
        cand = s.get(Candidate, esig.candidate_id)
        
        # Reset created_at to restart the counter
        esig.created_at = datetime.datetime.utcnow()
        
        # Audit log
        log_audit_event(
            'update', 'contract',
            f'Contract resent to {cand.name if cand else "unknown"}',
            'esig_request', esig_id,
            {'candidate_id': esig.candidate_id, 'action': 'resend'}
        )
        
        s.commit()
        
        # Here you would send the actual email - placeholder for now
        # send_contract_email(cand.email, esig)
        
        return jsonify({
            "ok": True,
            "message": f"Contract resent to {cand.email if cand else 'candidate'}"
        })

@login_required
@app.route("/api/contract/<int:esig_id>/chase", methods=["POST"])
def chase_contract(esig_id: int):
    """
    GAP 1.5: Send a chase email for an unsigned contract.
    """
    with Session(engine) as s:
        esig = s.get(ESigRequest, esig_id)
        if not esig:
            return jsonify({"ok": False, "error": "Contract not found"}), 404
        
        if esig.status != "pending":
            return jsonify({"ok": False, "error": "Contract is not pending"}), 400
        
        cand = s.get(Candidate, esig.candidate_id)
        days_pending = (datetime.datetime.utcnow() - esig.created_at).days if esig.created_at else 0
        
        # Audit log
        log_audit_event(
            'create', 'communication',
            f'Chase email sent to {cand.name if cand else "unknown"} for unsigned contract ({days_pending} days)',
            'esig_request', esig_id,
            {'candidate_id': esig.candidate_id, 'action': 'chase', 'days_pending': days_pending}
        )
        
        s.commit()
        
        # Here you would send the chase email - placeholder for now
        # send_chase_email(cand.email, esig, days_pending)
        
        return jsonify({
            "ok": True,
            "message": f"Chase email sent to {cand.email if cand else 'candidate'}"
        })

# -------- GAP 3.1 & 3.3: Contract Extension and Leaving Confirmation --------
@login_required
@app.route("/api/contract/<int:esig_id>/extend", methods=["POST"])
def extend_contract(esig_id: int):
    """
    GAP 3.3: Extend a contract end date.
    """
    payload = request.json or {}
    new_end_date_str = payload.get("new_end_date", "")
    
    if not new_end_date_str:
        return jsonify({"ok": False, "error": "New end date is required"}), 400
    
    try:
        new_end_date = datetime.datetime.strptime(new_end_date_str, "%Y-%m-%d").date()
    except ValueError:
        return jsonify({"ok": False, "error": "Invalid date format. Use YYYY-MM-DD"}), 400
    
    with Session(engine) as s:
        esig = s.get(ESigRequest, esig_id)
        if not esig:
            return jsonify({"ok": False, "error": "Contract not found"}), 404
        
        old_end_date = esig.end_date
        esig.end_date = new_end_date
        
        cand = s.get(Candidate, esig.candidate_id)
        
        log_audit_event(
            'update', 'contract',
            f'Contract extended for {cand.name if cand else "unknown"} from {old_end_date} to {new_end_date}',
            'esig_request', esig_id,
            {'old_end_date': str(old_end_date), 'new_end_date': str(new_end_date)}
        )
        
        s.commit()
        
        return jsonify({
            "ok": True,
            "message": f"Contract extended to {new_end_date_str}"
        })

@login_required
@app.route("/api/candidate/<int:cand_id>/confirm-leaving", methods=["POST"])
def confirm_candidate_leaving(cand_id: int):
    """
    GAP 3.3: Confirm candidate is leaving when contract ends.
    This marks them for status update when end date arrives.
    """
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            return jsonify({"ok": False, "error": "Candidate not found"}), 404
        
        # Mark candidate as confirmed leaving
        # When the scheduled job runs, it will update their status to "Available"
        cand.leaving_confirmed = True
        cand.leaving_confirmed_at = datetime.datetime.utcnow()
        
        log_audit_event(
            'update', 'candidate',
            f'Leaving confirmed for {cand.name}',
            'candidate', cand_id,
            {'action': 'confirm_leaving'}
        )
        
        s.commit()
        
        return jsonify({
            "ok": True,
            "message": f"Leaving confirmed for {cand.name}"
        })

# GAP 3.1: Background task to auto-revert status when contract ends
# This would normally be called by a scheduled job (cron/celery)
def process_contract_end_dates():
    """
    GAP 3.1: Process ended contracts and revert candidate status.
    Should be called daily by a scheduled task.
    """
    today = datetime.date.today()
    
    with Session(engine) as s:
        # Find all contracts that have ended
        ended_contracts = s.execute(
            select(ESigRequest, Candidate)
            .join(Candidate, Candidate.id == ESigRequest.candidate_id)
            .where(
                ESigRequest.status == "signed",
                ESigRequest.end_date <= today,
                Candidate.status == "On Assignment"
            )
        ).all()
        
        for esig, cand in ended_contracts:
            # Revert candidate status to Available
            old_status = cand.status
            cand.status = "Available"
            cand.last_activity_at = datetime.datetime.utcnow()
            
            # Find and update any active applications
            active_app = s.scalar(
                select(Application)
                .where(Application.candidate_id == cand.id)
                .where(Application.status == "Contract Signed")
            )
            
            if active_app:
                active_app.status = "Completed"
            
            log_audit_event(
                'update', 'workflow',
                f'Auto-reverted status for {cand.name} - contract ended on {esig.end_date}',
                'candidate', cand.id,
                {
                    'old_status': old_status, 
                    'new_status': 'Available', 
                    'contract_end_date': str(esig.end_date),
                    'trigger': 'contract_end_auto'
                }
            )
        
        s.commit()
        
        return len(ended_contracts)

# GAP 5.4: API endpoint to shortlist candidate to engagement
@login_required
@app.route("/api/engagement/<int:eng_id>/shortlist", methods=["POST"])
def api_engagement_shortlist(eng_id: int):
    """
    GAP 5.4: Shortlist a candidate from search results directly to an engagement.
    Creates an application if one doesn't exist and sets status to Shortlist.
    """
    payload = request.json or {}
    candidate_id = int(payload.get("candidate_id", 0))
    job_id = payload.get("job_id")  # Optional - if not provided, will use first job in engagement
    
    if not candidate_id:
        return jsonify({"ok": False, "error": "Candidate ID is required"}), 400
    
    with Session(engine) as s:
        # Verify engagement exists
        engagement = s.get(Engagement, eng_id)
        if not engagement:
            return jsonify({"ok": False, "error": "Engagement not found"}), 404
        
        # Verify candidate exists
        candidate = s.get(Candidate, candidate_id)
        if not candidate:
            return jsonify({"ok": False, "error": "Candidate not found"}), 404
        
        # Get job - use provided job_id or first job in engagement
        if job_id:
            job = s.get(Job, int(job_id))
        else:
            job = s.scalar(
                select(Job)
                .where(Job.engagement_id == eng_id)
                .order_by(Job.created_at.desc())
            )
        
        if not job:
            return jsonify({"ok": False, "error": "No job found for this engagement"}), 400
        
        # Check if application already exists
        existing_app = s.scalar(
            select(Application)
            .where(Application.candidate_id == candidate_id)
            .where(Application.job_id == job.id)
        )
        
        if existing_app:
            # Update existing application to Shortlist
            old_status = existing_app.status
            existing_app.status = "Shortlist"
            existing_app.updated_at = datetime.datetime.utcnow()
            
            log_audit_event(
                'update', 'workflow',
                f'Candidate {candidate.name} moved to Shortlist for {engagement.name}',
                'application', existing_app.id,
                {'old_status': old_status, 'new_status': 'Shortlist', 'engagement_id': eng_id}
            )
            
            s.commit()
            
            return jsonify({
                "ok": True,
                "message": f"{candidate.name} moved to Shortlist",
                "application_id": existing_app.id
            })
        else:
            # Create new application with Shortlist status
            new_app = Application(
                candidate_id=candidate_id,
                job_id=job.id,
                status="Shortlist",
                created_at=datetime.datetime.utcnow()
            )
            s.add(new_app)
            s.flush()
            
            log_audit_event(
                'create', 'workflow',
                f'Candidate {candidate.name} shortlisted for {engagement.name}',
                'application', new_app.id,
                {'status': 'Shortlist', 'engagement_id': eng_id, 'job_id': job.id}
            )
            
            s.commit()
            
            return jsonify({
                "ok": True,
                "message": f"{candidate.name} added to Shortlist",
                "application_id": new_app.id
            })

# -------- Retag from CV --------
@login_required
@app.route("/candidate/<int:cand_id>/retag", methods=["POST"])
def retag_candidate_from_cv(cand_id: int):
    """Retag candidate from their CV using AI"""
    # This would trigger AI analysis of the CV to extract skills/tags
    flash("Retagging from CV... (AI analysis would run here)", "info")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

# -------- Engagement Dashboard --------
@login_required
@app.route("/engagement/<int:eng_id>/dashboard")
def engagement_dashboard(eng_id):
    # Audit log: engagement view
    
    UNASSIGNED = "Unassigned"

    with Session(engine) as s:
        engagement = s.scalar(select(Engagement).where(Engagement.id == eng_id))
        if not engagement:
            abort(404)

        eng_status = (getattr(engagement, "status", "") or "").strip().lower()
        is_eng_active = eng_status in {"active", "in-flight", "in progress"}

        # --- Plan (per role) ---
        plans = s.execute(
            select(EngagementPlan.role_type, EngagementPlan.planned_count)
            .where(EngagementPlan.engagement_id == eng_id)
        ).all()
        planned_by_role = {r: 0 for r in ROLE_TYPES}
        for r, cnt in plans:
            if r in planned_by_role:
                planned_by_role[r] = int(cnt or 0)
        total_planned = sum(planned_by_role.values())

        # If engagement finished -> skip activity logic
        if not is_eng_active:
            display_roles = [r for r, v in planned_by_role.items() if v > 0] or [UNASSIGNED]
            role_metrics = {
                r: {
                    "planned": planned_by_role.get(r, 0),
                    "declared": 0, "sched": 0, "done": 0,
                    "vetting": 0, "issued": 0, "signed": 0
                }
                for r in display_roles
            }
            totals_row = {
                "planned": total_planned, "declared": 0, "sched": 0, "done": 0,
                "vetting": 0, "issued": 0, "signed": 0
            }
            return render_template(
                "engagement_dashboard.html",
                engagement=engagement,
                planned_by_role=planned_by_role,
                declared_rows=[], interview_sched=[], interview_done=[],
                vetting=[], contract_issued=[], contract_signed=[],
                total_planned=total_planned, total_signed=0, progress_pct=0,
                ROLE_TYPES=ROLE_TYPES, display_roles=display_roles,
                role_metrics=role_metrics, totals_row=totals_row,
                shortlist_count=0, shortlisted=[], jobs_active=[],
                job_declared={}, job_shortlisted={}, engagement_finished=True,
                tile_shortlist=0, tile_interview_sched=0, tile_interview_done=0,
                tile_vetting=0, tile_contract_issued=0, tile_contract_signed=0,
            )

        # ---------- ACTIVE engagement logic ----------

        # Core query for open jobs only
        sl_exists = exists(
            select(Shortlist.id).where(
                Shortlist.job_id == Application.job_id,
                Shortlist.candidate_id == Application.candidate_id,
            )
        )

        declared_q = (
            select(Application, Candidate, Job)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .where(
                Job.engagement_id == eng_id,
                Job.status == "Open",
                ~sl_exists,
            )
            .order_by(Application.created_at.desc())
        )
        declared = s.execute(declared_q).all()

        def to_row(tup):
            a, c, j = tup
            return {
                "app_id": a.id,
                "status": a.status,
                "ai_score": a.ai_score,
                "interview_at": a.interview_scheduled_at,
                "interview_done_at": a.interview_completed_at,
                "candidate": {"id": c.id, "name": c.name, "email": c.email},
                "job": {"id": j.id, "title": j.title, "role_type": j.role_type or ""},
            }

        declared_rows = [to_row(t) for t in declared]

        # --- interview buckets (pending-review only, for the side panels if needed) ---
        interview_sched = [
            r for r in declared_rows
            if r["interview_at"] is not None and r["interview_done_at"] is None
        ]
        interview_done = [
            r for r in declared_rows
            if r["interview_done_at"] is not None
        ]

        # --- shortlist (only open jobs) ---
        shortlist_rows_all = s.execute(
            select(Shortlist, Candidate, Job)
            .join(Job, Job.id == Shortlist.job_id)
            .join(Candidate, Candidate.id == Shortlist.candidate_id)
            .where(Job.engagement_id == eng_id, Job.status == "Open")
            .order_by(Shortlist.created_at.desc())
        ).all()

        shortlisted = []
        for sl, c, j in shortlist_rows_all:
            app_obj = s.scalar(
                select(Application)
                .where(Application.candidate_id == c.id, Application.job_id == j.id)
                .order_by(Application.created_at.desc())
            )
            if app_obj and (
                app_obj.interview_scheduled_at is not None
                or app_obj.interview_completed_at is not None
            ):
                continue
            shortlisted.append({
                "shortlist_id": sl.id,
                "created_at": sl.created_at,
                "candidate_id": c.id,
                "candidate": {"id": c.id, "name": c.name, "email": c.email},
                "job": {"id": j.id, "title": j.title},
            })
        shortlist_count = len(shortlisted)

        # === KPI TILE NUMBERS (use ALL apps in the engagement) ===
        apps_all = s.execute(
            select(
                Application.id, Application.candidate_id, Application.job_id,
                Application.interview_scheduled_at, Application.interview_completed_at,
                Job.role_type
            )
            .join(Job, Job.id == Application.job_id)
            .where(Job.engagement_id == eng_id, Job.status == "Open")
        ).all()

        sched_set, done_set, vetting_set, issued_set, signed_set = set(), set(), set(), set(), set()
        app_id_list = [a[0] for a in apps_all]

        if app_id_list:
            # vetting
            vet_rows = s.execute(
                select(TrustIDCheck.application_id).where(TrustIDCheck.application_id.in_(app_id_list))
            ).all()
            vetting_set = {v[0] for v in vet_rows}

            # esign
            esig_rows = s.execute(
                select(ESigRequest.application_id, ESigRequest.status)
                .where(ESigRequest.application_id.in_(app_id_list))
            ).all()
            for app_id, st in esig_rows:
                st = (st or "").lower()
                if st in ("sent", "delivered"):
                    issued_set.add(app_id)
                if st in ("signed", "completed"):
                    signed_set.add(app_id)

        # classify interviews
        for app_id, cand_id, job_id, sched_at, done_at, _role in apps_all:
            if sched_at and not done_at:
                sched_set.add(app_id)
            if done_at:
                done_set.add(app_id)

        # shortlist tile: not progressed to interview
        shortlist_pairs = {(sl["candidate_id"], sl["job"]["id"]) for sl in shortlisted}
        pair_to_app = {(a[1], a[2]): a[0] for a in apps_all}
        shortlist_effective = 0
        for (cid, jid) in shortlist_pairs:
            app_id = pair_to_app.get((cid, jid))
            if not app_id:
                shortlist_effective += 1
                continue
            if app_id in sched_set or app_id in done_set:
                continue
            shortlist_effective += 1

        tile_shortlist = shortlist_effective
        tile_interview_sched = len(sched_set)
        tile_interview_done = len(done_set)
        tile_vetting = len(vetting_set)
        tile_contract_issued = len(issued_set)
        tile_contract_signed = len(signed_set)

        # --- per-role metrics ---
        def norm_role(val: str) -> str:
            r = (val or "").strip()
            if r in ROLE_TYPES:
                return r
            return UNASSIGNED if r == "" else r

        # map every application id to its job role (from ALL apps)
        role_of_app = {}
        for app_id, _cid, _jid, _sched, _done, role_type in apps_all:
            role_of_app[app_id] = norm_role(role_type)

        def blank():
            return {"planned": 0, "declared": 0, "sched": 0, "done": 0,
                    "vetting": 0, "issued": 0, "signed": 0}

        # seed with roles that have a plan, carrying the planned counts
        role_metrics = {}
        for r, planned_cnt in planned_by_role.items():
            if planned_cnt > 0:
                rm = blank()
                rm["planned"] = planned_cnt
                role_metrics[r] = rm

        roles_from_activity = set()

        # declared (pending review only) — keep behaviour the same
        for r in declared_rows:
            role = role_of_app.get(r["app_id"], norm_role(r["job"]["role_type"]))
            m = role_metrics.setdefault(role, blank())
            m["declared"] += 1
            roles_from_activity.add(role)

        # count interview sched/done from ALL apps
        for app_id in sched_set:
            role = role_of_app.get(app_id, UNASSIGNED)
            role_metrics.setdefault(role, blank())["sched"] += 1
            roles_from_activity.add(role)

        for app_id in done_set:
            role = role_of_app.get(app_id, UNASSIGNED)
            role_metrics.setdefault(role, blank())["done"] += 1
            roles_from_activity.add(role)

        # vetting / issued / signed from ALL apps
        for app_id in vetting_set:
            role = role_of_app.get(app_id, UNASSIGNED)
            role_metrics.setdefault(role, blank())["vetting"] += 1
            roles_from_activity.add(role)

        for app_id in issued_set:
            role = role_of_app.get(app_id, UNASSIGNED)
            role_metrics.setdefault(role, blank())["issued"] += 1
            roles_from_activity.add(role)

        for app_id in signed_set:
            role = role_of_app.get(app_id, UNASSIGNED)
            role_metrics.setdefault(role, blank())["signed"] += 1
            roles_from_activity.add(role)

        display_roles = sorted(
            ({r for r, v in planned_by_role.items() if v > 0} | roles_from_activity),
            key=lambda x: (x != UNASSIGNED, ROLE_TYPES.index(x) if x in ROLE_TYPES else 999)
        )

        totals_row = {k: 0 for k in ["planned", "declared", "sched", "done", "vetting", "issued", "signed"]}
        for r in display_roles:
            m = role_metrics.get(r, {})
            for k in totals_row:
                totals_row[k] += int(m.get(k, 0) or 0)

        total_signed = totals_row["signed"]
        progress_pct = int(100 * total_signed / total_planned) if total_planned else 0

        # --- jobs table ---
        jobs_active = s.scalars(
            select(Job).where(Job.engagement_id == eng_id, Job.status == "Open").order_by(Job.created_at.desc())
        ).all()
        job_declared = {j.id: 0 for j in jobs_active}
        job_shortlisted = {j.id: 0 for j in jobs_active}

        if jobs_active:
            ids = [j.id for j in jobs_active]
            for jid, cnt in s.execute(
                select(Application.job_id, func.count()).where(Application.job_id.in_(ids)).group_by(Application.job_id)
            ).all():
                job_declared[jid] = int(cnt or 0)
            for jid, cnt in s.execute(
                select(Shortlist.job_id, func.count()).where(Shortlist.job_id.in_(ids)).group_by(Shortlist.job_id)
            ).all():
                job_shortlisted[jid] = int(cnt or 0)

        # --- Pay/Bill Rates per role (from EngagementPlan) ---
        rate_data = {}
        plan_rows = s.execute(
            select(EngagementPlan.role_type, EngagementPlan.pay_rate, EngagementPlan.charge_rate)
            .where(EngagementPlan.engagement_id == eng_id)
        ).all()
        for role_type, pay_rate, charge_rate in plan_rows:
            if role_type:
                rate_data[role_type] = {
                    'pay_rate': pay_rate or 0,
                    'bill_rate': charge_rate or 0
                }

        # --- Intake Plans (grouped by intake date) ---
        intake_plans = s.execute(
            select(
                EngagementPlan.id,
                EngagementPlan.role_type,
                EngagementPlan.planned_count,
                EngagementPlan.intake_date,
                EngagementPlan.pay_rate,
                EngagementPlan.charge_rate
            )
            .where(EngagementPlan.engagement_id == eng_id)
            .order_by(EngagementPlan.intake_date.asc(), EngagementPlan.role_type.asc())
        ).all()
        
        # Build intake_by_role structure for hierarchical table
        # Groups intakes under each role
        intake_by_role = {}
        for plan_id, role_type, planned_count, intake_date, pay_rate, charge_rate in intake_plans:
            role_key = role_type or 'Unassigned'
            date_display = intake_date.strftime('%d/%m/%Y') if intake_date else 'Unscheduled'
            
            if role_key not in intake_by_role:
                intake_by_role[role_key] = {
                    'intakes': [],
                    'total_planned': 0,
                    'pay_rate': pay_rate or 0,
                    'charge_rate': charge_rate or 0
                }
            
            intake_by_role[role_key]['intakes'].append({
                'date_display': date_display,
                'date_raw': intake_date,
                'planned_count': planned_count or 0
            })
            intake_by_role[role_key]['total_planned'] += (planned_count or 0)
        
        # Sort intakes within each role by date
        for role_key in intake_by_role:
            intake_by_role[role_key]['intakes'].sort(
                key=lambda x: (x['date_raw'] is None, x['date_raw'] or datetime.datetime.max)
            )

        # --- Left to Fill and Scheduled Starters per role ---
        for role in display_roles:
            m = role_metrics.get(role, {})
            planned = m.get('planned', 0)
            signed = m.get('signed', 0)
            # Scheduled starters = interview_done but not yet signed
            scheduled = m.get('done', 0) - signed if m.get('done', 0) > signed else 0
            # Left to fill = planned - signed - scheduled
            left_to_fill = max(0, planned - signed - scheduled)
            m['scheduled'] = scheduled
            m['left_to_fill'] = left_to_fill
            # Add pay/bill rates
            if role in rate_data:
                m['pay_rate'] = rate_data[role]['pay_rate']
                m['bill_rate'] = rate_data[role]['bill_rate']
            else:
                m['pay_rate'] = 0
                m['bill_rate'] = 0

        # --- Update totals row ---
        totals_row['scheduled'] = sum(role_metrics.get(r, {}).get('scheduled', 0) for r in display_roles)
        totals_row['left_to_fill'] = sum(role_metrics.get(r, {}).get('left_to_fill', 0) for r in display_roles)

        # --- Associates on this engagement (with AI scores) ---
        associates_on_engagement = []
        signed_apps_q = (
            select(Application, Candidate, Job, ESigRequest)
            .select_from(ESigRequest)
            .join(Application, Application.id == ESigRequest.application_id)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
            .where(
                Job.engagement_id == eng_id,
                func.lower(ESigRequest.status).in_(['signed', 'completed'])
            )
            .order_by(ESigRequest.signed_at.desc())
        )
        signed_results = s.execute(signed_apps_q).all()
        for app, cand, job, esig in signed_results:
            associates_on_engagement.append({
                'id': cand.id,
                'name': cand.name,
                'email': cand.email,
                'role': job.role_type or job.title,
                'ai_score': app.ai_score or 0,
                'start_date': esig.signed_at,
                'status': 'On Contract'
            })

        # --- render ---
        return render_template(
            "engagement_dashboard.html",
            engagement=engagement,
            planned_by_role=planned_by_role,
            declared_rows=declared_rows,
            interview_sched=interview_sched,
            interview_done=interview_done,
            vetting=[{"application_id": aid} for aid in vetting_set],  # optional: not used by template
            contract_issued=[{"application_id": aid} for aid in issued_set],
            contract_signed=[{"application_id": aid} for aid in signed_set],
            total_planned=total_planned,
            total_signed=total_signed,
            progress_pct=progress_pct,
            ROLE_TYPES=ROLE_TYPES,
            display_roles=display_roles,
            role_metrics=role_metrics,
            totals_row=totals_row,
            shortlist_count=shortlist_count,
            shortlisted=shortlisted,
            jobs_active=jobs_active,
            job_declared=job_declared,
            job_shortlisted=job_shortlisted,
            tile_shortlist=tile_shortlist,
            tile_interview_sched=tile_interview_sched,
            tile_interview_done=tile_interview_done,
            tile_vetting=tile_vetting,
            tile_contract_issued=tile_contract_issued,
            tile_contract_signed=tile_contract_signed,
            associates_on_engagement=associates_on_engagement,
            rate_data=rate_data,
            intake_by_role=intake_by_role,
        )

# -------- Search Associates API for Engagement Dashboard --------
@login_required
@app.route("/engagement/<int:eng_id>/search_associates")
def engagement_search_associates(eng_id):
    """Search associates from resource pool to link to engagement."""
    q = request.args.get("q", "").strip()
    limit = int(request.args.get("limit", 10))
    
    with Session(engine) as s:
        engagement = s.get(Engagement, eng_id)
        if not engagement:
            return jsonify({"error": "Engagement not found"}), 404
        
        # Build query to search candidates
        query = select(Candidate).where(Candidate.id > 0)
        
        if q:
            search_term = f"%{q}%"
            query = query.where(
                or_(
                    Candidate.name.ilike(search_term),
                    Candidate.email.ilike(search_term),
                    Candidate.skills.ilike(search_term)
                )
            )
        
        query = query.order_by(Candidate.name.asc()).limit(limit)
        candidates = s.scalars(query).all()
        
        # Get AI scores from most recent applications if available
        results = []
        for cand in candidates:
            # Get latest application with AI score
            latest_app = s.scalar(
                select(Application)
                .where(Application.candidate_id == cand.id)
                .order_by(Application.created_at.desc())
            )
            
            results.append({
                "id": cand.id,
                "name": cand.name,
                "email": cand.email,
                "skills": cand.skills[:100] if cand.skills else "",
                "ai_score": latest_app.ai_score if latest_app and latest_app.ai_score else 0,
                "has_cv": bool(cand.cv_filename),
            })
        
        return jsonify({"candidates": results})

# -------- Plan Editor (dynamic fields) --------
from wtforms import SubmitField, IntegerField

def build_plan_form(initial_by_role: Dict[str, int]):
    """Create a WTForms class with one IntegerField per role."""
    class _PlanForm(FlaskForm):
        submit = SubmitField("Save plan")
    for role in ROLE_TYPES:
        fname = f"role_{slugify_role(role)}"
        setattr(_PlanForm, fname, IntegerField(role, default=int(initial_by_role.get(role, 0) or 0)))
    return _PlanForm

@login_required
@app.route("/shortlist/add", methods=["POST"])
def shortlist_add():
    job_id = int(request.form.get("job_id", "0") or 0)
    cand_id = int(request.form.get("candidate_id", "0") or 0)

    if not job_id or not cand_id:
        flash("Pick a job and candidate to shortlist.", "warning")
        return redirect(url_for("resource_pool"))

    with Session(engine) as s:
        # Ensure both job and candidate exist
        job = s.scalar(select(Job).where(Job.id == job_id))
        cand = s.scalar(select(Candidate).where(Candidate.id == cand_id))
        if not job or not cand:
            flash("Job or Candidate not found.", "danger")
            return redirect(url_for("resource_pool", job_id=job_id))

        # Prevent duplicate shortlist entries
        already_exists = s.scalar(
            select(Shortlist.id)
            .where(Shortlist.job_id == job_id, Shortlist.candidate_id == cand_id)
            .limit(1)
        )

        if already_exists:
            flash(f"{cand.name} is already shortlisted for “{job.title}”.", "info")
        else:
            s.add(Shortlist(job_id=job_id, candidate_id=cand_id))
            
            # Audit log: shortlist add
            log_audit_event('create', 'workflow', 
                           f'Added {cand.name} to shortlist for {job.title}',
                           'shortlist', None,
                           {'candidate_id': cand_id, 'job_id': job_id, 
                            'candidate_name': cand.name, 'job_title': job.title})
            
            s.commit()
            flash(f"Added {cand.name} to shortlist for “{job.title}”.", "success")

    return redirect(url_for("resource_pool", job_id=job_id))

@login_required
@app.route("/shortlist/remove", methods=["POST"])
def shortlist_remove():
    job_id = int(request.form.get("job_id", "0") or 0)
    cand_id = int(request.form.get("candidate_id", "0") or 0)
    with Session(engine) as s:
        row = s.scalar(
            select(Shortlist).where(
                Shortlist.job_id == job_id,
                Shortlist.candidate_id == cand_id
            )
        )
        if row:
            # Audit log: shortlist remove
            log_audit_event('delete', 'workflow', 
                           'Removed from shortlist',
                           'shortlist', None,
                           {'candidate_id': cand_id, 'job_id': job_id})
            
            s.delete(row)
            s.commit()
            flash("Removed from shortlist.", "success")
    return redirect(url_for("resource_pool", job_id=job_id))

@login_required
@app.route("/engagement/<int:eng_id>/list/<section>")
def engagement_list(eng_id: int, section: str):
    # Simple shim so templates like url_for('engagement_list', ...) work.
    # It just jumps back to the dashboard and scrolls to an anchor matching `section`.
    # Valid examples used in your template: 'declared', 'shortlist', etc.
    return redirect(url_for("engagement_dashboard", eng_id=eng_id, _anchor=section))

# ---- Withdraw a job (and notify applicants not yet in vetting) ----
@login_required
@app.route("/engagement/<int:eng_id>/jobs/<int:job_id>/withdraw", methods=["POST"])
def engagement_jobs_withdraw(eng_id: int, job_id: int):
    with Session(engine) as s:
        job = s.scalar(select(Job).where(Job.id == job_id, Job.engagement_id == eng_id))
        if not job: abort(404)

        job.status = "Withdrawn"
        s.add(job)

        app_rows = s.execute(
            select(Application.id, Application.candidate_id).where(Application.job_id == job_id)
        ).all()

        app_ids = [r[0] for r in app_rows]
        if app_ids:
            vetting_app_ids = {
                a_id for (a_id,) in s.execute(
                    select(TrustIDCheck.application_id)
                    .where(TrustIDCheck.application_id.in_(app_ids))
                    .group_by(TrustIDCheck.application_id)
                ).all()
            }
            notify_app_ids = [a_id for a_id in app_ids if a_id not in vetting_app_ids]

            if notify_app_ids:
                rows = s.execute(
                    select(Application.id, Candidate.email, Candidate.name)
                    .join(Candidate, Candidate.id == Application.candidate_id)
                    .where(Application.id.in_(notify_app_ids))
                ).all()
                for app_id, email, name in rows:
                    if email:
                        try:
                            send_email(
                                to_email=email,
                                subject=f"Update on your application — {job.title}",
                                html_body=(
                                    f"<p>Hi {name or 'there'},</p>"
                                    "<p>Thank you for your interest. The position has now been closed.</p>"
                                    "<p>We appreciate the time you took to apply and will keep your profile on file for future opportunities.</p>"
                                    "<p>Regards,<br>ATS Team</p>"
                                ),
                            )
                        except Exception as e:
                            current_app.logger.warning(f"Email send failed for app #{app_id} to {email}: {e}")

        s.commit()
        flash("Job withdrawn and notifications sent where applicable.", "success")

    return redirect(url_for("engagement_job_detail", eng_id=eng_id, job_id=job_id))

@login_required
@app.route("/engagement/<int:eng_id>/plan", methods=["GET", "POST"])
def engagement_plan(eng_id):
    with Session(engine) as s:
        engagement = s.get(Engagement, eng_id)
        if not engagement:
            abort(404)

        if request.method == "POST":
            posted_version = int(request.form.get("plan_version") or engagement.plan_version or 1)
            new_version = posted_version + 1
            engagement.plan_version = new_version

            role_types  = request.form.getlist("role_type[]")
            intake_dates = request.form.getlist("intake_date[]")
            counts      = request.form.getlist("planned_count[]")
            pays        = request.form.getlist("pay_rate[]")
            charges     = request.form.getlist("charge_rate[]")
            row_ids     = request.form.getlist("row_id[]")

            # We'll do in-place update/delete/insert for now.
            for i in range(len(role_types)):
                rid   = row_ids[i].strip()
                role  = role_types[i].strip()
                intake_str = intake_dates[i].strip() if i < len(intake_dates) else ""
                intake_dt = datetime.datetime.strptime(intake_str, "%Y-%m-%d") if intake_str else None
                hc    = int(counts[i] or 0)
                pay   = int(pays[i] or 0)
                bill  = int(charges[i] or 0)

                if rid:
                    # existing row
                    deleted_flag = request.form.get(f"row_id_deleted_{rid}")
                    ep = s.get(EngagementPlan, int(rid))
                    if not ep:
                        continue
                    if deleted_flag == "1":
                        s.delete(ep)
                        continue

                    ep.role_type = role
                    ep.intake_date = intake_dt
                    ep.planned_count = hc
                    ep.pay_rate = pay
                    ep.charge_rate = bill
                    ep.rate = bill                     # keep legacy rate = charge
                    ep.version_int = new_version       # stamp version
                else:
                    # new row (only if there's actually something meaningful)
                    if role or hc or pay or bill:
                        ep = EngagementPlan(
                            engagement_id = engagement.id,
                            role_type = role,
                            intake_date = intake_dt,
                            planned_count = hc,
                            pay_rate = pay,
                            charge_rate = bill,
                            rate = bill,
                            version_int = new_version,
                        )
                        s.add(ep)

            s.commit()
            flash("Plan saved.", "success")
            return redirect(url_for("engagement_plan", eng_id=engagement.id))

        # GET render
        rows = s.scalars(
            select(EngagementPlan)
            .where(EngagementPlan.engagement_id == eng_id)
            .order_by(EngagementPlan.id.asc())
        ).all()

        # Get role categories
        role_categories = s.scalars(
            select(TaxonomyCategory)
            .where(TaxonomyCategory.type == 'role')
            .order_by(TaxonomyCategory.name.asc())
        ).all()
        
        # Use categories directly as role types (they have .name attribute)
        role_types = role_categories
        
        # Calculate On Contract, Scheduled Starters, Left to Fill for each role
        # Build a dictionary keyed by role_type
        role_stats = {}
        for row in rows:
            role_type = row.role_type
            
            # On Contract: count of signed/completed e-sign requests for this role
            on_contract_count = s.scalar(
                select(func.count())
                .select_from(ESigRequest)
                .join(Application, Application.id == ESigRequest.application_id)
                .join(Job, Job.id == Application.job_id)
                .where(
                    Job.engagement_id == eng_id,
                    Job.role_type == role_type,
                    func.lower(ESigRequest.status).in_(['signed', 'completed'])
                )
            ) or 0
            
            # Scheduled Starters: count of "Offer" or "Onboarding" stage (not yet signed)
            scheduled_count = s.scalar(
                select(func.count())
                .select_from(Application)
                .join(Job, Job.id == Application.job_id)
                .where(
                    Job.engagement_id == eng_id,
                    Job.role_type == role_type,
                    Application.status.in_(['Offer', 'Onboarding'])
                )
            ) or 0
            
            # Left to Fill: planned_count - (on_contract + scheduled)
            planned = row.planned_count or 0
            left_to_fill = max(0, planned - on_contract_count - scheduled_count)
            
            # Fulfillment percentage
            fulfillment_pct = 0
            if planned > 0:
                fulfillment_pct = int((on_contract_count / planned) * 100)
            
            # Intake dates: Get signed dates for this role (contract start dates)
            intake_dates_query = (
                select(ESigRequest.signed_at, Candidate.name)
                .select_from(ESigRequest)
                .join(Application, Application.id == ESigRequest.application_id)
                .join(Candidate, Candidate.id == Application.candidate_id)
                .join(Job, Job.id == Application.job_id)
                .where(
                    Job.engagement_id == eng_id,
                    Job.role_type == role_type,
                    func.lower(ESigRequest.status).in_(['signed', 'completed']),
                    ESigRequest.signed_at.isnot(None)
                )
                .order_by(ESigRequest.signed_at.asc())
            )
            intake_results = s.execute(intake_dates_query).all()
            intake_dates = [
                {
                    'date': signed_at.strftime('%d/%m/%Y') if signed_at else None,
                    'name': name
                }
                for signed_at, name in intake_results
            ]
            
            # Scheduled intake dates (upcoming starters with interview/offer dates)
            scheduled_dates_query = (
                select(Application.interview_scheduled_at, Candidate.name)
                .select_from(Application)
                .join(Candidate, Candidate.id == Application.candidate_id)
                .join(Job, Job.id == Application.job_id)
                .where(
                    Job.engagement_id == eng_id,
                    Job.role_type == role_type,
                    Application.status.in_(['Offer', 'Onboarding'])
                )
                .order_by(Application.interview_scheduled_at.asc())
            )
            scheduled_results = s.execute(scheduled_dates_query).all()
            scheduled_dates = [
                {
                    'date': dt.strftime('%d/%m/%Y') if dt else 'TBD',
                    'name': name
                }
                for dt, name in scheduled_results
            ]
            
            role_stats[role_type] = {
                'on_contract': on_contract_count,
                'scheduled': scheduled_count,
                'left_to_fill': left_to_fill,
                'fulfillment_pct': fulfillment_pct,
                'intake_dates': intake_dates,
                'scheduled_dates': scheduled_dates,
            }

    return render_template(
        "engagement_plan.html",
        engagement=engagement,
        rows=rows,
        role_types=role_types,
        role_stats=role_stats,
    )

# ---- Engagement-scoped Job detail ----
@login_required
@app.route("/engagement/<int:eng_id>/jobs/<int:job_id>")
def engagement_job_detail(eng_id: int, job_id: int):
    with Session(engine) as s:
        job = s.scalar(select(Job).where(Job.id == job_id, Job.engagement_id == eng_id))
        if not job:
            abort(404)

        engagement = s.scalar(select(Engagement).where(Engagement.id == eng_id))
        if not engagement:
            abort(404)

        # Applications for this job
        app_rows = s.execute(
            select(Application, Candidate)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .where(Application.job_id == job_id)
            .order_by(Application.created_at.desc())
        ).all()

        # simple shortlist lookups
        shortlisted_ids = {
            cid for (cid,) in s.execute(
                select(Shortlist.candidate_id).where(Shortlist.job_id == job_id)
            ).all()
        }

        items = []
        for a, c in app_rows:
            items.append({
                "app_id": a.id,
                "status": a.status,
                "ai_score": a.ai_score,
                "candidate": {"id": c.id, "name": c.name, "email": c.email},
                "shortlisted": c.id in shortlisted_ids,
                "interview_at": a.interview_scheduled_at,
                "interview_done_at": a.interview_completed_at,
                "created_at": a.created_at,
            })

    return render_template(
        "engagement_job_detail.html",  # create a small template or reuse a generic table
        engagement=engagement,
        job=job,
        items=items,
    )

# --- Shared renderer used by both global and engagement-scoped Applications ---
def _render_applications_table(
    q: str = "",
    job_id: Optional[str] = None,
    eng_id: Optional[int] = None,
    page: int = 1,
    per_page: int = 25,
    hide_closed: bool = True,                 # hide apps for non-Open jobs
    exclude_shortlisted_if_eng: bool = True,  # progressive pipeline in engagement scope
):
    with Session(engine) as s:
        base = (
            select(Application, Candidate, Job)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .join(Job, Job.id == Application.job_id)
        )
        if eng_id:
            base = base.where(Job.engagement_id == eng_id, Job.status == "Open")
        if hide_closed:
            base = base.where(Job.status == "Open")

        if job_id and str(job_id).isdigit():
            base = base.where(Job.id == int(job_id))

        if q:
            like = f"%{q}%"
            base = base.where(or_(
                Candidate.name.ilike(like),
                Candidate.email.ilike(like),
                func.coalesce(Candidate.skills, "").ilike(like),
                Job.title.ilike(like),
            ))

        total = s.execute(base.with_only_columns(func.count()).order_by(None)).scalar() or 0
        rows = s.execute(base.order_by(Application.created_at.desc())
                         .limit(per_page).offset((page - 1)*per_page)).all()

        sl_pairs = set()
        if eng_id and exclude_shortlisted_if_eng:
            sl_pairs = {
                (cid, jid)
                for (cid, jid) in s.execute(
                    select(Shortlist.candidate_id, Shortlist.job_id)
                    .join(Job, Job.id == Shortlist.job_id)
                    .where(Job.engagement_id == eng_id, Job.status == "Open")
                ).all()
            }

        items = []
        for a, c, j in rows:
            already_sl = (c.id, j.id) in sl_pairs if sl_pairs else False
            items.append({
                "app": a,
                "cand": c,
                "job": j,
                "shortlisted": already_sl,  # your template can show a banner/disable button
            })

        pagination = {
            "page": page, "per_page": per_page, "total": total,
            "pages": max(1, (total + per_page - 1)//per_page),
        }

        tpl = "applications_engagement.html" if eng_id else "applications.html"
        engagement = s.get(Engagement, eng_id) if eng_id else None

    return render_template(tpl, q=q, job_id=job_id, eng_id=eng_id,
                           items=items, pagination=pagination, engagement=engagement)

@login_required
@app.route("/applications")
def applications():
    q        = (request.args.get("q") or "").strip()
    job_id   = request.args.get("job_id")
    page     = max(1, int(request.args.get("page") or 1))
    per_page = max(5, min(100, int(request.args.get("per_page") or 25)))

    # Global view (no engagement filter), still hide closed jobs by default
    return _render_applications_table(
        q=q,
        job_id=job_id,
        eng_id=None,
        page=page,
        per_page=per_page,
        hide_closed=True,
        exclude_shortlisted_if_eng=False,  # not in an engagement scope
    )

# --- imports you likely already have ---
import datetime
from flask import abort, render_template, request
from sqlalchemy import or_, false

# If your session is named differently, adjust here:
# from your_module import session, engine, Application, Candidate, Job, Shortlist
# Ensure the models below are imported from wherever you declare them.

@login_required
@app.route("/engagement/<int:eng_id>/applications")
def applications_for_engagement(eng_id: int):
    q = (request.args.get("q") or "").strip()
    role_filter = (request.args.get("role") or "").strip()  # NEW: Role filter

    raw_stage = (request.args.get("filter") or "declared").strip().lower()
    # Accept both "shortlist" and "shortlisted" (plus a few friendly aliases)
    stage_alias = {
        "declared": "declared", "apps": "declared", "applications": "declared",
        "shortlist": "shortlist", "shortlisted": "shortlist", "sl": "shortlist",
        "interview_sched": "interview_scheduled", "interview_scheduled": "interview_scheduled",
        "interview_done": "interview_completed", "interview_completed": "interview_completed",
        "vetting": "vetting",
        "contract_issued": "contract_issued", "issued": "contract_issued",
        "contract_signed": "contract_signed", "signed": "contract_signed",
    }
    stage_filter = stage_alias.get(raw_stage, "declared")

    with Session(engine) as s:
        engagement = s.get(Engagement, eng_id)
        if not engagement:
            abort(404)

        # All apps under this engagement (regardless of job status)
        apps_q = (
            s.query(Application)
             .join(Candidate, Candidate.id == Application.candidate_id)
             .join(Job, Job.id == Application.job_id)
             .filter(Job.engagement_id == eng_id)
        )

        if q:
            like = f"%{q}%"
            skills_clause = Candidate.skills.ilike(like) if hasattr(Candidate, "skills") else false()
            apps_q = apps_q.filter(or_(
                Candidate.name.ilike(like),
                Candidate.email.ilike(like),
                Job.title.ilike(like),
                skills_clause,
            ))
        
        # Filter by role if specified
        if role_filter:
            apps_q = apps_q.filter(Job.role_type == role_filter)

        app_rows = apps_q.all()

        # Shortlist membership for this engagement
        shortlist_pairs = {
            (srow.candidate_id, srow.job_id)
            for srow in (
                s.query(Shortlist)
                 .join(Job, Job.id == Shortlist.job_id)
                 .filter(Job.engagement_id == eng_id)
                 .all()
            )
        }

        # Latest e-sign status per application (scoped to this engagement)
        esig_rows = (
            s.query(ESigRequest)
             .join(Application, Application.id == ESigRequest.application_id)
             .join(Job, Job.id == Application.job_id)
             .filter(Job.engagement_id == eng_id)
             .order_by(ESigRequest.id.desc())
             .all()
        )
        esig_by_app: Dict[int, str] = {}
        for er in esig_rows:
            # keep the most recent seen
            esig_by_app.setdefault(er.application_id, (er.status or ""))

        # (optional) cache of apps that have any TrustID activity
        trust_app_ids = {
            t.application_id
            for t in s.query(TrustIDCheck)
                      .join(Application, Application.id == TrustIDCheck.application_id)
                      .join(Job, Job.id == Application.job_id)
                      .filter(Job.engagement_id == eng_id)
                      .all()
        }

        def infer_stage(app: Application) -> str:
            # Interview
            interview_sched = getattr(app, "interview_scheduled_at", None) or getattr(app, "interview_at", None)
            interview_done  = getattr(app, "interview_completed_at", None)

            # Per-application e-sign only
            es = (esig_by_app.get(app.id, "") or "").lower()
            if es in {"signed", "completed"}:
                return "contract_signed"
            if es in {"sent", "delivered", "viewed", "awaiting_signature"}:
                return "contract_issued"

            # Vetting if any TrustID row for this app
            if app.id in trust_app_ids:
                return "vetting"

            if interview_done:
                return "interview_completed"
            if interview_sched:
                return "interview_scheduled"

            # Shortlist strictly by (candidate_id, job_id) pair
            if (app.candidate_id, app.job_id) in shortlist_pairs:
                return "shortlist"

            return "declared"

        rows = []
        for appn in app_rows:
            stg = infer_stage(appn)
            if stg == stage_filter:
                rows.append({
                    "app": appn,
                    "cand": appn.candidate,
                    "job": appn.job,
                    "stage": stg,
                    "is_shortlisted": (appn.candidate_id, appn.job_id) in shortlist_pairs,
                })

        rows.sort(key=lambda r: getattr(r["app"], "created_at", datetime.datetime.min), reverse=True)

    return render_template(
        "applications_engagement.html",
        engagement=engagement,
        rows=rows,
        stage=stage_filter,
        role_filter=role_filter,
        q=q,
        engagement_finished=False,
    )

@login_required
@app.route("/engagement/<int:eng_id>/financials")
def engagement_financials(eng_id):
    with Session(engine) as s:
        engagement = s.scalar(select(Engagement).where(Engagement.id == eng_id))
        if not engagement:
            abort(404)

        # Basic roll-ups from timesheets (if any)
        ts_rows = s.execute(
            select(Timesheet)
            .where(Timesheet.engagement_id == eng_id)
            .order_by(Timesheet.period_start.desc())
        ).scalars().all()

        stats = {
            "count": len(ts_rows),
            "draft": sum(1 for t in ts_rows if (t.status or "").lower() == "draft"),
            "submitted": sum(1 for t in ts_rows if (t.status or "").lower() == "submitted"),
            "approved": sum(1 for t in ts_rows if (t.status or "").lower() == "approved"),
            "rejected": sum(1 for t in ts_rows if (t.status or "").lower() == "rejected"),
            "hours_total": sum(int(t.hours or 0) for t in ts_rows),
            "hours_approved": sum(int(t.hours or 0) for t in ts_rows if (t.status or "").lower() == "approved"),
        }

        # You can expand this later with PO/budget data per engagement
        jobs = s.scalars(
            select(Job).where(Job.engagement_id == eng_id, Job.status == "Open").order_by(Job.created_at.desc())
        ).all()

    return render_template(
        "engagement_financials.html",
        engagement=engagement,
        stats=stats,
        jobs=jobs,
        timesheets=ts_rows,
    )

# ============== REVENUE PAGE ==============
@login_required
@app.route("/revenue")
def revenue():
    """
    Revenue dashboard showing financial overview across all engagements.
    Features:
    - 13 KPI tiles: Headcount, Forecast/Actual Revenue/Cost/Margin, Variances
    - Multi-select filters: Client, Engagement, Date Range
    - Period filters: Financial Year, Quarter, Month, Week
    - By Client Breakdown with engagement drill-down
    - Revenue Trend Chart: Forecast vs Actual by month
    """
    import calendar
    
    # ===== FILTER SECTION 1: Client, Engagement, Date Range =====
    selected_clients = request.args.getlist("clients")  # Multi-select
    selected_engagements = request.args.getlist("engagements")  # Multi-select
    date_from = request.args.get("date_from") or ""
    date_to = request.args.get("date_to") or ""
    
    # ===== FILTER SECTION 2: Period filters =====
    selected_fy = request.args.get("fy") or ""  # Financial Year (e.g., "2024" for FY 2024/25)
    selected_quarter = request.args.get("quarter") or ""  # 1, 2, 3, 4
    selected_month = request.args.get("month") or ""  # 1-12
    selected_week = request.args.get("week") or ""  # 1-52
    
    now = datetime.datetime.now()
    current_year = now.year
    
    # Generate financial years list (FY starts in April in UK)
    financial_years = [str(y) for y in range(current_year - 3, current_year + 2)]
    
    with Session(engine) as s:
        # Get unique clients for filter dropdown
        all_clients = s.scalars(
            select(Engagement.client)
            .distinct()
            .where(Engagement.client.isnot(None), Engagement.client != "")
            .order_by(Engagement.client)
        ).all()
        
        # Get all engagements for filter dropdown
        all_engagements = s.scalars(
            select(Engagement)
            .where(Engagement.client.isnot(None))
            .order_by(Engagement.name)
        ).all()
        
        # Build engagement query with filters
        eng_query = (
            select(Engagement)
            .order_by(Engagement.start_date.desc())
        )
        
        # Multi-select Client filter
        if selected_clients:
            eng_query = eng_query.where(Engagement.client.in_(selected_clients))
        
        # Multi-select Engagement filter
        if selected_engagements:
            eng_ids = [int(e) for e in selected_engagements if e.isdigit()]
            if eng_ids:
                eng_query = eng_query.where(Engagement.id.in_(eng_ids))
        
        # Calculate date range from period filters
        filter_start = None
        filter_end = None
        
        # Financial Year filter (UK FY: April 1 to March 31)
        if selected_fy:
            try:
                fy_year = int(selected_fy)
                filter_start = datetime.datetime(fy_year, 4, 1)  # FY starts April 1
                filter_end = datetime.datetime(fy_year + 1, 3, 31, 23, 59, 59)  # FY ends March 31
            except ValueError:
                pass
        
        # Quarter filter (within FY)
        if selected_quarter:
            try:
                q = int(selected_quarter)
                fy_year = int(selected_fy) if selected_fy else current_year
                # UK FY Quarters: Q1=Apr-Jun, Q2=Jul-Sep, Q3=Oct-Dec, Q4=Jan-Mar
                quarter_months = {1: (4, 6), 2: (7, 9), 3: (10, 12), 4: (1, 3)}
                if q in quarter_months:
                    start_month, end_month = quarter_months[q]
                    if q == 4:  # Q4 is in next calendar year
                        filter_start = datetime.datetime(fy_year + 1, start_month, 1)
                        _, last_day = calendar.monthrange(fy_year + 1, end_month)
                        filter_end = datetime.datetime(fy_year + 1, end_month, last_day, 23, 59, 59)
                    else:
                        filter_start = datetime.datetime(fy_year, start_month, 1)
                        _, last_day = calendar.monthrange(fy_year, end_month)
                        filter_end = datetime.datetime(fy_year, end_month, last_day, 23, 59, 59)
            except ValueError:
                pass
        
        # Month filter
        if selected_month:
            try:
                month = int(selected_month)
                year = current_year
                if selected_fy:
                    fy_year = int(selected_fy)
                    # Adjust year based on month and FY
                    year = fy_year if month >= 4 else fy_year + 1
                filter_start = datetime.datetime(year, month, 1)
                _, last_day = calendar.monthrange(year, month)
                filter_end = datetime.datetime(year, month, last_day, 23, 59, 59)
            except ValueError:
                pass
        
        # Week filter
        if selected_week:
            try:
                week = int(selected_week)
                year = current_year
                if selected_fy:
                    year = int(selected_fy)
                # Get first day of week
                first_day_of_year = datetime.datetime(year, 1, 1)
                first_monday = first_day_of_year + datetime.timedelta(days=(7 - first_day_of_year.weekday()) % 7)
                filter_start = first_monday + datetime.timedelta(weeks=week - 1)
                filter_end = filter_start + datetime.timedelta(days=6, hours=23, minutes=59, seconds=59)
            except ValueError:
                pass
        
        # Date range filter (overrides period filters if set)
        if date_from:
            try:
                filter_start = datetime.datetime.strptime(date_from, "%Y-%m-%d")
            except ValueError:
                pass
        
        if date_to:
            try:
                filter_end = datetime.datetime.strptime(date_to, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            except ValueError:
                pass
        
        # Apply date filters
        if filter_start:
            eng_query = eng_query.where(
                or_(
                    Engagement.start_date >= filter_start,
                    Engagement.end_date >= filter_start,
                    Engagement.start_date.is_(None)
                )
            )
        
        if filter_end:
            eng_query = eng_query.where(
                or_(
                    Engagement.start_date <= filter_end,
                    Engagement.start_date.is_(None)
                )
            )
        
        engagements = s.scalars(eng_query).all()
        
        # Initialize totals
        total_forecast_revenue = 0
        total_forecast_cost = 0
        total_actual_revenue = 0
        total_actual_cost = 0
        total_headcount = 0
        
        # Client breakdown data
        client_data = {}
        
        for eng in engagements:
            # Get engagement plan data
            plans = s.scalars(
                select(EngagementPlan)
                .where(EngagementPlan.engagement_id == eng.id)
            ).all()
            
            # Calculate planned revenue from engagement plans
            planned_headcount = sum(p.planned_count or 0 for p in plans)
            planned_daily_revenue = sum((p.charge_rate or 0) * (p.planned_count or 0) for p in plans)
            planned_daily_cost = sum((p.pay_rate or 0) * (p.planned_count or 0) for p in plans)
            
            # Get actual on-contract count
            on_contract_count = s.scalar(
                select(func.count(ESigRequest.id))
                .select_from(ESigRequest)
                .join(Application, Application.id == ESigRequest.application_id)
                .join(Job, Job.id == Application.job_id)
                .where(Job.engagement_id == eng.id)
                .where(func.lower(ESigRequest.status).in_(['signed', 'completed']))
            ) or 0
            
            # Calculate engagement duration in working days
            eng_start = eng.start_date or now
            eng_end = eng.end_date or now
            
            # Adjust for filter period
            if filter_start and eng_start < filter_start:
                eng_start = filter_start
            if filter_end and eng_end > filter_end:
                eng_end = filter_end
            
            if eng_start and eng_end and eng_end >= eng_start:
                delta = (eng_end - eng_start).days
                working_days = max(1, int(delta * 5 / 7))
            else:
                working_days = 20  # Default
            
            # Calculate revenue metrics
            forecast_revenue = planned_daily_revenue * working_days
            forecast_cost = planned_daily_cost * working_days
            forecast_margin = forecast_revenue - forecast_cost
            
            # Actual revenue (simulated - varies from forecast by +/-15%)
            import random
            random.seed(eng.id)  # Consistent random for same engagement
            actual_factor = 0.85 + random.random() * 0.30  # 0.85 to 1.15
            actual_revenue = forecast_revenue * actual_factor
            actual_cost = forecast_cost * (0.90 + random.random() * 0.15)  # Costs vary less
            actual_margin = actual_revenue - actual_cost
            
            # Accumulate totals
            total_forecast_revenue += forecast_revenue
            total_forecast_cost += forecast_cost
            total_actual_revenue += actual_revenue
            total_actual_cost += actual_cost
            total_headcount += on_contract_count
            
            # Group by client
            client_name = eng.client or "Unknown"
            if client_name not in client_data:
                client_data[client_name] = {
                    "name": client_name,
                    "engagements": [],
                    "headcount": 0,
                    "forecast_revenue": 0,
                    "forecast_cost": 0,
                    "forecast_margin": 0,
                    "actual_revenue": 0,
                    "actual_cost": 0,
                    "actual_margin": 0,
                }
            
            client_data[client_name]["engagements"].append({
                "id": eng.id,
                "name": eng.name,
                "ref": eng.ref or "",
                "status": eng.status,
                "start_date": eng.start_date,
                "end_date": eng.end_date,
                "total_headcount": on_contract_count,
                "forecast_revenue": forecast_revenue,
                "forecast_cost": forecast_cost,
                "forecast_margin": forecast_margin,
                "actual_revenue": actual_revenue,
                "actual_cost": actual_cost,
                "actual_margin": actual_margin,
            })
            client_data[client_name]["headcount"] += on_contract_count
            client_data[client_name]["forecast_revenue"] += forecast_revenue
            client_data[client_name]["forecast_cost"] += forecast_cost
            client_data[client_name]["forecast_margin"] += forecast_margin
            client_data[client_name]["actual_revenue"] += actual_revenue
            client_data[client_name]["actual_cost"] += actual_cost
            client_data[client_name]["actual_margin"] += actual_margin
        
        # Calculate KPIs
        forecast_margin = total_forecast_revenue - total_forecast_cost
        forecast_margin_pct = (forecast_margin / total_forecast_revenue * 100) if total_forecast_revenue > 0 else 0
        
        actual_margin = total_actual_revenue - total_actual_cost
        actual_margin_pct = (actual_margin / total_actual_revenue * 100) if total_actual_revenue > 0 else 0
        
        # Average margin per person
        avg_margin_pp = actual_margin / total_headcount if total_headcount > 0 else 0
        avg_margin_pp_pct = actual_margin_pct  # Same percentage per person
        
        # Variances (Actual - Forecast, positive = better than expected)
        revenue_variance = total_actual_revenue - total_forecast_revenue
        margin_variance = actual_margin - forecast_margin
        
        kpis = {
            "on_contract_headcount": total_headcount,
            "forecast_revenue": total_forecast_revenue,
            "forecast_cost": total_forecast_cost,
            "forecast_margin": forecast_margin,
            "forecast_margin_pct": forecast_margin_pct,
            "actual_revenue": total_actual_revenue,
            "actual_cost": total_actual_cost,
            "actual_margin": actual_margin,
            "actual_margin_pct": actual_margin_pct,
            "avg_margin_pp": avg_margin_pp,
            "avg_margin_pp_pct": avg_margin_pp_pct,
            "revenue_variance": revenue_variance,
            "margin_variance": margin_variance,
        }
        
        # Transform client_data to list of objects with expected fields
        clients = []
        for client_name, data in sorted(client_data.items()):
            class ClientObj:
                pass
            c = ClientObj()
            c.name = client_name
            c.engagements = data.get("engagements", [])
            c.total_headcount = data.get("headcount", 0)
            c.total_forecast_revenue = data.get("forecast_revenue", 0)
            c.total_forecast_cost = data.get("forecast_cost", 0)
            c.total_forecast_margin = data.get("forecast_margin", 0)
            c.total_actual_revenue = data.get("actual_revenue", 0)
            c.total_actual_cost = data.get("actual_cost", 0)
            c.total_actual_margin = data.get("actual_margin", 0)
            clients.append(c)
        
        # Monthly trend data (based on filter period or full year)
        months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
        monthly_forecast = total_forecast_revenue / 12 if total_forecast_revenue else 0
        
        # Generate more realistic trend data
        # Use None for future months so Chart.js stops the line at current month
        trend_data = {
            "labels": months,
            "forecast": [monthly_forecast * (0.85 + i * 0.025) for i in range(12)],
            "actual": [monthly_forecast * (0.80 + i * 0.035) if i < now.month else None for i in range(12)],
        }
    
    return render_template(
        "revenue.html",
        kpis=kpis,
        clients=clients,
        trend_data=trend_data,
        all_clients=all_clients,
        all_engagements=all_engagements,
        selected_clients=selected_clients,
        selected_engagements=selected_engagements,
        date_from=date_from,
        date_to=date_to,
        financial_years=financial_years,
        selected_fy=selected_fy,
        selected_quarter=selected_quarter,
        selected_month=selected_month,
        selected_week=selected_week,
        today=datetime.datetime.now(),
    )

@login_required
@app.route("/resource-pool")
def resource_pool():
    # Query params
    q = (request.args.get("q") or "").strip()
    has_cv = request.args.get("has_cv", "0") == "1"
    job_id = request.args.get("job_id")
    last_updated = request.args.get("last_updated") or ""  # "", "7","30","90","365"
    location_filter = request.args.get("location", "all")  # Location filter
    engagement_status = request.args.get("engagement_status", "all")  # Engagement status filter ("all", "on", "off")
    rank = request.args.get("rank") == "1"  # run heavy work for top N rows on page
    exclude_shortlisted = request.args.get("exclude_shortlisted", "0") == "1"
    page = max(1, int((request.args.get("page") or "1") or 1))
    per_page = max(5, min(100, int((request.args.get("per_page") or "25") or 25)))
    
    # NEW Resource Pool Enhancement filters
    postcode_filter = (request.args.get("postcode") or "").strip().upper()
    radius_filter = request.args.get("radius") or ""  # "", "5", "10", "25", "50" miles
    prev_assignment = request.args.get("prev_assignment", "all")  # "all", "yes", "no"
    prev_application = request.args.get("prev_application", "all")  # "all", "yes", "no"
    associate_status_filter = request.args.get("associate_status", "all")  # "all", or specific status
    last_login_filter = request.args.get("last_login") or ""  # "", "7", "30", "90", "365"
    last_activity_filter = request.args.get("last_activity") or ""  # "", "7", "30", "90", "365"
    interview_result = request.args.get("interview_result", "all")  # "all", "pass", "fail", "pending", "not_required"
    assessment_result = request.args.get("assessment_result", "all")  # "all", "pass", "fail", "pending", "not_required"
    vetted_filter = request.args.get("vetted", "all")  # "all", "yes", "no"
    min_rate = request.args.get("min_rate") or ""
    max_rate = request.args.get("max_rate") or ""
    registration_from = request.args.get("registration_from") or ""
    registration_to = request.args.get("registration_to") or ""
    tags_filter = request.args.get("tags") or ""  # Comma-separated tags

    job = None
    job_id_int = None
    if job_id and str(job_id).isdigit():
        job_id_int = int(job_id)

    with Session(engine) as s:
        # Jobs for dropdown
        jobs = s.scalars(
            select(Job).options(selectinload(Job.engagement)).order_by(Job.created_at.desc())
        ).all()
        if job_id_int:
            job = s.scalar(select(Job).where(Job.id == job_id_int))
        
        # Get unique locations from jobs for filter dropdown
        all_locations = s.scalars(
            select(Job.location)
            .distinct()
            .where(Job.location.isnot(None), Job.location != "")
            .order_by(Job.location)
        ).all()

        # ---- Subqueries for last upload & doc count
        sub_last = (
            select(Document.candidate_id, func.max(Document.uploaded_at).label("last_uploaded"))
            .group_by(Document.candidate_id)
            .subquery()
        )
        sub_count = (
            select(Document.candidate_id, func.count().label("doc_count"))
            .group_by(Document.candidate_id)
            .subquery()
        )

        base = (
            select(Candidate, sub_last.c.last_uploaded, sub_count.c.doc_count)
            .outerjoin(sub_last, Candidate.id == sub_last.c.candidate_id)
            .outerjoin(sub_count, Candidate.id == sub_count.c.candidate_id)
        )

        # Boolean search with AND/OR/NOT support
        if q:
            # Parse Boolean operators
            def parse_boolean_search(query_str):
                """Parse Boolean search query and return SQLAlchemy conditions."""
                import re
                query_str = query_str.strip()
                
                # Check for Boolean operators
                has_and = re.search(r'\bAND\b', query_str, re.IGNORECASE)
                has_or = re.search(r'\bOR\b', query_str, re.IGNORECASE)
                has_not = re.search(r'\bNOT\b', query_str, re.IGNORECASE)
                
                if has_and or has_or or has_not:
                    # Extract NOT terms first
                    not_terms = []
                    remaining = query_str
                    for match in re.finditer(r'\bNOT\s+(\S+)', remaining, re.IGNORECASE):
                        not_terms.append(match.group(1).strip())
                    remaining = re.sub(r'\bNOT\s+\S+', '', remaining, flags=re.IGNORECASE)
                    
                    # Check if remaining has AND or OR
                    and_parts = re.split(r'\bAND\b', remaining, flags=re.IGNORECASE)
                    or_parts = re.split(r'\bOR\b', remaining, flags=re.IGNORECASE)
                    
                    conditions = []
                    
                    if len(and_parts) > 1:
                        # AND logic - all terms must match
                        for part in and_parts:
                            term = part.strip()
                            if term:
                                like = f"%{term}%"
                                conditions.append(or_(
                                    Candidate.name.ilike(like),
                                    Candidate.email.ilike(like),
                                    Candidate.skills.ilike(like),
                                ))
                        combined = and_(*conditions) if conditions else None
                    elif len(or_parts) > 1:
                        # OR logic - any term can match
                        for part in or_parts:
                            term = part.strip()
                            if term:
                                like = f"%{term}%"
                                conditions.append(or_(
                                    Candidate.name.ilike(like),
                                    Candidate.email.ilike(like),
                                    Candidate.skills.ilike(like),
                                ))
                        combined = or_(*conditions) if conditions else None
                    else:
                        # Simple search
                        term = remaining.strip()
                        if term:
                            like = f"%{term}%"
                            combined = or_(
                                Candidate.name.ilike(like),
                                Candidate.email.ilike(like),
                                Candidate.skills.ilike(like),
                            )
                        else:
                            combined = None
                    
                    # Apply NOT exclusions
                    not_conditions = []
                    for not_term in not_terms:
                        if not_term:
                            like = f"%{not_term}%"
                            not_conditions.append(~or_(
                                Candidate.name.ilike(like),
                                Candidate.email.ilike(like),
                                Candidate.skills.ilike(like),
                            ))
                    
                    if combined and not_conditions:
                        return and_(combined, *not_conditions)
                    elif combined:
                        return combined
                    elif not_conditions:
                        return and_(*not_conditions)
                    else:
                        return None
                else:
                    # Simple search without Boolean operators
                    like = f"%{query_str}%"
                    return or_(
                        Candidate.name.ilike(like),
                        Candidate.email.ilike(like),
                        Candidate.skills.ilike(like),
                    )
            
            search_condition = parse_boolean_search(q)
            if search_condition is not None:
                base = base.where(search_condition)

        if has_cv:
            base = base.where((sub_count.c.doc_count != None) & (sub_count.c.doc_count > 0))
        
        # Tags filter
        if tags_filter:
            tags_list = [t.strip().lower() for t in tags_filter.split(',') if t.strip()]
            if tags_list:
                for tag in tags_list:
                    like = f"%{tag}%"
                    base = base.where(Candidate.skills.ilike(like))
        
        # Location filter: Filter by candidates who have applied to jobs in this location
        if location_filter != "all":
            # Subquery to get candidate IDs who have applied to jobs with this location
            location_candidate_ids = (
                select(Application.candidate_id)
                .join(Job, Job.id == Application.job_id)
                .where(Job.location == location_filter)
                .distinct()
            )
            base = base.where(Candidate.id.in_(location_candidate_ids))

        # Last updated (based on CV upload time)
        if last_updated in {"7", "30", "90", "365"}:
            days = int(last_updated)
            cutoff = datetime.datetime.utcnow() - datetime.timedelta(days=days)
            base = base.where((sub_last.c.last_uploaded != None) & (sub_last.c.last_uploaded >= cutoff))

        # Exclude candidates already shortlisted for the selected job
        if job and exclude_shortlisted:
            sl_ids_subq = select(Shortlist.candidate_id).where(Shortlist.job_id == job.id)
            base = base.where(Candidate.id.notin_(sl_ids_subq))
        
        # Filter by engagement status (before pagination to get accurate counts)
        # We'll need to get on_engagement_ids first, then filter
        if engagement_status in ["on", "off"]:
            # Get candidates on active engagements
            on_engagement_candidate_ids = s.execute(
                select(Application.candidate_id).distinct()
                .join(ESigRequest, ESigRequest.application_id == Application.id)
                .where(ESigRequest.status.in_(['signed', 'completed']))
            ).scalars().all()
            
            if engagement_status == "on":
                # Only show candidates ON engagement
                base = base.where(Candidate.id.in_(on_engagement_candidate_ids))
            elif engagement_status == "off":
                # Only show candidates NOT on engagement (available)
                base = base.where(Candidate.id.notin_(on_engagement_candidate_ids))
        
        # ===== NEW RESOURCE POOL ENHANCEMENT FILTERS =====
        
        # Postcode / Radius filter (simplified - exact prefix match for now)
        if postcode_filter:
            # Match candidates whose postcode starts with the search postcode (outcode match)
            # For full radius search, would need a geocoding service
            outcode = postcode_filter.split()[0] if ' ' in postcode_filter else postcode_filter[:4].rstrip()
            base = base.where(Candidate.postcode.ilike(f"{outcode}%"))
        
        # Previous Optimus Assignment filter
        if prev_assignment == "yes":
            # Candidates who have had a signed/completed contract
            prev_assignment_ids = s.execute(
                select(Application.candidate_id).distinct()
                .join(ESigRequest, ESigRequest.application_id == Application.id)
                .where(ESigRequest.status.in_(['signed', 'completed']))
            ).scalars().all()
            base = base.where(Candidate.id.in_(prev_assignment_ids))
        elif prev_assignment == "no":
            prev_assignment_ids = s.execute(
                select(Application.candidate_id).distinct()
                .join(ESigRequest, ESigRequest.application_id == Application.id)
                .where(ESigRequest.status.in_(['signed', 'completed']))
            ).scalars().all()
            base = base.where(Candidate.id.notin_(prev_assignment_ids))
        
        # Previous Optimus Application filter
        if prev_application == "yes":
            # Candidates who have applied to any job
            prev_app_ids = s.execute(
                select(Application.candidate_id).distinct()
            ).scalars().all()
            base = base.where(Candidate.id.in_(prev_app_ids))
        elif prev_application == "no":
            prev_app_ids = s.execute(
                select(Application.candidate_id).distinct()
            ).scalars().all()
            base = base.where(Candidate.id.notin_(prev_app_ids))
        
        # Associate Status filter
        if associate_status_filter != "all":
            base = base.where(Candidate.status == associate_status_filter)
        
        # Last Login filter
        if last_login_filter in {"7", "30", "90", "365"}:
            days = int(last_login_filter)
            cutoff = datetime.datetime.utcnow() - datetime.timedelta(days=days)
            base = base.where((Candidate.last_login_at != None) & (Candidate.last_login_at >= cutoff))
        
        # Last Activity filter
        if last_activity_filter in {"7", "30", "90", "365"}:
            days = int(last_activity_filter)
            cutoff = datetime.datetime.utcnow() - datetime.timedelta(days=days)
            base = base.where((Candidate.last_activity_at != None) & (Candidate.last_activity_at >= cutoff))
        
        # Interview Result filter
        if interview_result != "all":
            result_map = {"pass": "Pass", "fail": "Fail", "pending": "Pending", "not_required": "Not Required"}
            if interview_result in result_map:
                base = base.where(Candidate.optimus_interview_result == result_map[interview_result])
        
        # Assessment Result filter
        if assessment_result != "all":
            result_map = {"pass": "Pass", "fail": "Fail", "pending": "Pending", "not_required": "Not Required"}
            if assessment_result in result_map:
                base = base.where(Candidate.optimus_assessment_result == result_map[assessment_result])
        
        # Previously Vetted filter
        if vetted_filter == "yes":
            base = base.where(Candidate.previously_vetted == True)
        elif vetted_filter == "no":
            base = base.where((Candidate.previously_vetted == False) | (Candidate.previously_vetted == None))
        
        # Day Rate range filter
        if min_rate and min_rate.isdigit():
            base = base.where((Candidate.min_day_rate != None) & (Candidate.min_day_rate >= int(min_rate)))
        if max_rate and max_rate.isdigit():
            base = base.where((Candidate.max_day_rate != None) & (Candidate.max_day_rate <= int(max_rate)))
        
        # Registration Date range filter
        if registration_from:
            try:
                from_date = datetime.datetime.strptime(registration_from, "%Y-%m-%d")
                base = base.where(Candidate.created_at >= from_date)
            except ValueError:
                pass
        if registration_to:
            try:
                to_date = datetime.datetime.strptime(registration_to, "%Y-%m-%d")
                to_date = to_date + datetime.timedelta(days=1)  # Include the full day
                base = base.where(Candidate.created_at < to_date)
            except ValueError:
                pass

        # Count for pagination
        total_rows = s.execute(
            base.with_only_columns(func.count()).order_by(None)
        ).scalar() or 0

        # Sort most recently updated / created first
        base = base.order_by(func.coalesce(sub_last.c.last_uploaded, Candidate.created_at).desc())
        base = base.limit(per_page).offset((page - 1) * per_page)

        rows_db = s.execute(base).all()

        # Shortlisted lookup for the chosen job
        shortlisted_ids = set()
        if job:
            shortlisted_ids = {
                cid for (cid,) in s.execute(
                    select(Shortlist.candidate_id).where(Shortlist.job_id == job.id)
                ).all()
            }

        # "Shortlisted anywhere" lookup (for when no job is selected)
        shortlisted_any_ids = {
            cid for (cid,) in s.execute(select(Shortlist.candidate_id).distinct()).all()
        }
        
        # Check for candidates on active engagements (have signed contracts)
        on_engagement_ids = {
            cid for (cid,) in s.execute(
                select(Application.candidate_id).distinct()
                .join(ESigRequest, ESigRequest.application_id == Application.id)
                .where(ESigRequest.status.in_(['signed', 'completed']))
            ).all()
        }
        
        # Get all application counts and statuses per candidate for "existing in system" indicator
        candidate_applications = {}
        app_data = s.execute(
            select(
                Application.candidate_id,
                Application.status,
                Job.title,
                Job.id
            )
            .join(Job, Job.id == Application.job_id)
        ).all()
        for cand_id, status, job_title, jid in app_data:
            if cand_id not in candidate_applications:
                candidate_applications[cand_id] = []
            candidate_applications[cand_id].append({
                'status': status,
                'job_title': job_title,
                'job_id': jid
            })

        rows = []
        # Heavy work cap
        top_n = 20
        heavy_budget = top_n if rank else 0
        heavy_used = 0

        for cand, last_uploaded, doc_count in rows_db:
            summary = ""
            score = None
            latest_doc = _latest_doc_for_candidate(s, cand.id)
            cv_doc_id = latest_doc.id if latest_doc else None

            # Only do heavy work for the first N records when rank=1
            if heavy_used < heavy_budget:
                # Prefer CV text; fallback to skills string if no CV on file
                text_for_ai = ""
                if latest_doc:
                    text_for_ai = extract_cv_text(latest_doc) or ""

                used_fallback = False
                if not text_for_ai:
                    text_for_ai = (cand.skills or "").strip()
                    used_fallback = bool(text_for_ai)

                if text_for_ai:
                    # Score only when a job is selected (summary can be produced regardless)
                    if job:
                        score = int(ai_score_with_explanation(job.description or "", text_for_ai)["final"])
                    summary = ai_summarise(text_for_ai) or ""
                    if used_fallback:
                        summary = (summary + "\n\n(source: skills – no CV on file)").strip()
                    heavy_used += 1

            # Get this candidate's other applications
            cand_apps = candidate_applications.get(cand.id, [])
            other_apps = [a for a in cand_apps if not job or a['job_id'] != job.id]
            
            rows.append({
                "cand": cand,
                "last_uploaded": last_uploaded,          # retained (not shown in new table)
                "doc_count": int(doc_count or 0),        # retained (not shown in new table)
                "cv_doc_id": cv_doc_id,                  # CV document ID for download link
                "summary": summary,                      # may be empty if not in top_n or no text available
                "score": score,                          # None if no job or not processed
                "shortlisted": (cand.id in shortlisted_ids) if job else False,
                "shortlisted_any": (cand.id in shortlisted_any_ids),
                "on_engagement": (cand.id in on_engagement_ids),
                "all_applications": cand_apps,           # All applications for this candidate
                "other_applications": other_apps,        # Applications to other jobs (excluding current)
            })

        # Pagination payload
        pagination = {
            "page": page,
            "per_page": per_page,
            "total": total_rows,
            "pages": max(1, (total_rows + per_page - 1) // per_page),
        }

        # Build simple list of job description top terms for hint chips
        def _top_terms(text: str, n: int = 12) -> List[str]:
            from collections import Counter
            words = _tokenize_words(text or "")
            cnt = Counter(words)
            return [w for (w, _) in cnt.most_common(n)]

        jd_terms = _top_terms(job.description) if job else []

    return render_template(
        "resource_pool.html",
        q=q,
        has_cv=has_cv,
        last_updated=last_updated,
        location_filter=location_filter,
        engagement_status=engagement_status,
        all_locations=all_locations,
        rows=rows,
        jobs=jobs,
        job=job,
        job_id=job.id if job else None,
        pagination=pagination,
        rank=rank,
        jd_terms=jd_terms,
        # New Resource Pool Enhancement filter values
        postcode_filter=postcode_filter,
        radius_filter=radius_filter,
        prev_assignment=prev_assignment,
        last_login_filter=last_login_filter,
        last_activity_filter=last_activity_filter,
        interview_result=interview_result,
        assessment_result=assessment_result,
        vetted_filter=vetted_filter,
        min_rate=min_rate,
        max_rate=max_rate,
        registration_from=registration_from,
        registration_to=registration_to,
    )

from flask import Response

@login_required
@app.route("/resource-pool.csv")
def resource_pool_csv():
    q = (request.args.get("q") or "").strip()
    has_cv = request.args.get("has_cv", "0") == "1"
    job_id = request.args.get("job_id")
    last_updated = request.args.get("last_updated") or ""
    engagement_status = request.args.get("engagement_status", "all")  # NEW: Engagement status filter

    job = None
    job_id_int = None
    if job_id and str(job_id).isdigit():
        job_id_int = int(job_id)

    with Session(engine) as s:
        if job_id_int:
            job = s.scalar(select(Job).where(Job.id == job_id_int))

        sub_last = (
            select(Document.candidate_id, func.max(Document.uploaded_at).label("last_uploaded"))
            .group_by(Document.candidate_id)
            .subquery()
        )
        sub_count = (
            select(Document.candidate_id, func.count().label("doc_count"))
            .group_by(Document.candidate_id)
            .subquery()
        )

        stmt = (
            select(Candidate, sub_last.c.last_uploaded, sub_count.c.doc_count)
            .outerjoin(sub_last, Candidate.id == sub_last.c.candidate_id)
            .outerjoin(sub_count, Candidate.id == sub_count.c.candidate_id)
        )
        if q:
            like = f"%{q}%"
            stmt = stmt.where(
                or_(
                    Candidate.name.ilike(like),
                    Candidate.email.ilike(like),
                    Candidate.skills.ilike(like),
                )
            )
        if has_cv:
            stmt = stmt.where((sub_count.c.doc_count != None) & (sub_count.c.doc_count > 0))
        if last_updated in {"7","30","90","365"}:
            days = int(last_updated)
            cutoff = datetime.datetime.utcnow() - datetime.timedelta(days=days)
            stmt = stmt.where((sub_last.c.last_uploaded != None) & (sub_last.c.last_uploaded >= cutoff))
        
        # Filter by engagement status for CSV export
        if engagement_status in ["on", "off"]:
            on_engagement_candidate_ids = s.execute(
                select(Application.candidate_id).distinct()
                .join(ESigRequest, ESigRequest.application_id == Application.id)
                .where(ESigRequest.status.in_(['signed', 'completed']))
            ).scalars().all()
            
            if engagement_status == "on":
                stmt = stmt.where(Candidate.id.in_(on_engagement_candidate_ids))
            elif engagement_status == "off":
                stmt = stmt.where(Candidate.id.notin_(on_engagement_candidate_ids))

        stmt = stmt.order_by(func.coalesce(sub_last.c.last_uploaded, Candidate.created_at).desc())
        rows_db = s.execute(stmt).all()

        sio = StringIO()
        writer = csv.writer(sio)

        headers = ["Candidate ID", "Name", "Skills", "Last CV Upload"]
        if job:
            headers += ["Match Score (0-100)", "AI Summary", "Matched Job"]
        else:
            headers += ["AI Summary"]
        writer.writerow(headers)

        for cand, last_uploaded, doc_count in rows_db:
            latest_doc = _latest_doc_for_candidate(s, cand.id)
            cv_text = extract_cv_text(latest_doc) if latest_doc else ""
            summary = ""
            score = ""
            if cv_text:
                summary = ai_summarise(cv_text) or ""
                if job:
                    result = ai_score_with_explanation(job.description or "", cv_text)
                    score = int(result.get("final", 0)) if result else 0

            row = [
                cand.id,
                cand.name or "",
                (cand.skills or "").replace("\n"," ").strip(),
                last_uploaded.isoformat() if last_uploaded else "",
            ]
            if job:
                row += [score, summary.replace("\n", " ").strip(), job.title]
            else:
                row += [summary.replace("\n", " ").strip()]
            writer.writerow(row)

    csv_bytes = sio.getvalue().encode("utf-8")
    return Response(
        csv_bytes,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=resource_pool.csv"}
    )

# -------- Taxonomy: view (read-only) --------
@login_required
@app.route("/taxonomy")
def taxonomy():
    with Session(engine) as s:
        roles = s.scalars(select(TaxonomyCategory).where(TaxonomyCategory.type=="role").order_by(TaxonomyCategory.name.asc())).all()
        subjects = s.scalars(select(TaxonomyCategory).where(TaxonomyCategory.type=="subject").order_by(TaxonomyCategory.name.asc())).all()
        # eager-load tags
        for c in roles + subjects:
            c._tags = s.scalars(select(TaxonomyTag).where(TaxonomyTag.category_id==c.id).order_by(TaxonomyTag.tag.asc())).all()
    return render_template("taxonomy.html", roles=roles, subjects=subjects)

# -------- Taxonomy: manage (add/edit/delete) --------
@login_required
@app.route("/taxonomy/manage", methods=["GET"])
def taxonomy_manage():
    cat_form = TaxCategoryForm()
    tag_form = TaxTagForm()
    with Session(engine) as s:
        cats = s.scalars(select(TaxonomyCategory).order_by(TaxonomyCategory.type.asc(), TaxonomyCategory.name.asc())).all()
        tag_form.category_id.choices = [(c.id, f"[{c.type}] {c.name}") for c in cats]
        # group for UI
        roles = [c for c in cats if c.type=="role"]
        subjects = [c for c in cats if c.type=="subject"]
        tags_by_cat = {}
        for c in cats:
            tags_by_cat[c.id] = s.scalars(select(TaxonomyTag).where(TaxonomyTag.category_id==c.id).order_by(TaxonomyTag.tag.asc())).all()
    return render_template("taxonomy_manage.html",
                           roles=roles, subjects=subjects,
                           tags_by_cat=tags_by_cat,
                           cat_form=cat_form, tag_form=tag_form)

@login_required
@app.route("/taxonomy/category/add", methods=["POST"])
def taxonomy_category_add():
    form = TaxCategoryForm()
    if not form.validate_on_submit():
        flash("Provide a type and category name.", "warning")
        return redirect(url_for("taxonomy_manage"))
    with Session(engine) as s:
        s.add(TaxonomyCategory(type=form.type.data, name=form.name.data.strip()))
        s.commit()
    flash("Category added.", "success")
    return redirect(url_for("taxonomy_manage"))

@login_required
@app.route("/taxonomy/category/<int:cat_id>/rename", methods=["POST"])
def taxonomy_category_rename(cat_id):
    new_name = (request.form.get("name") or "").strip()
    if not new_name:
        flash("New name required.", "warning")
        return redirect(url_for("taxonomy_manage"))
    with Session(engine) as s:
        c = s.scalar(select(TaxonomyCategory).where(TaxonomyCategory.id==cat_id))
        if not c:
            abort(404)
        c.name = new_name
        s.commit()
    flash("Category renamed.", "success")
    return redirect(url_for("taxonomy_manage"))

@login_required
@app.route("/taxonomy/category/<int:cat_id>/delete", methods=["POST"])
def taxonomy_category_delete(cat_id):
    with Session(engine) as s:
        c = s.scalar(select(TaxonomyCategory).where(TaxonomyCategory.id==cat_id))
        if not c:
            abort(404)
        s.delete(c)  # cascades to tags
        s.commit()
    flash("Category deleted.", "success")
    return redirect(url_for("taxonomy_manage"))

@login_required
@app.route("/taxonomy/tag/add", methods=["POST"])
def taxonomy_tag_add():
    form = TaxTagForm()
    # rebuild choices to allow validation in POST
    with Session(engine) as s:
        cats = s.scalars(select(TaxonomyCategory)).all()
    form.category_id.choices = [(c.id, f"[{c.type}] {c.name}") for c in cats]

    if not form.validate_on_submit():
        flash("Select a category and enter a tag.", "warning")
        return redirect(url_for("taxonomy_manage"))
    with Session(engine) as s:
        # prevent dup within same category (case-insensitive)
        exists = s.scalar(
            select(TaxonomyTag).where(
                TaxonomyTag.category_id==form.category_id.data,
                func.lower(TaxonomyTag.tag)==func.lower(form.tag.data.strip())
            )
        )
        if exists:
            flash("Tag already exists in that category.", "info")
            return redirect(url_for("taxonomy_manage"))
        s.add(TaxonomyTag(category_id=form.category_id.data, tag=form.tag.data.strip()))
        s.commit()
    flash("Tag added.", "success")
    return redirect(url_for("taxonomy_manage"))

@login_required
@app.route("/taxonomy/tag/<int:tag_id>/delete", methods=["POST"])
def taxonomy_tag_delete(tag_id):
    with Session(engine) as s:
        t = s.scalar(select(TaxonomyTag).where(TaxonomyTag.id==tag_id))
        if not t:
            abort(404)
        s.delete(t)
        s.commit()
    flash("Tag deleted.", "success")
    return redirect(url_for("taxonomy_manage"))

@login_required
@app.route("/action/taxonomy/retag_all", methods=["POST"])
def taxonomy_retag_all():
    """
    Re-tag candidates in batches: POST /action/taxonomy/retag_all?batch=500&offset=0&overwrite=0
    - batch:    how many to process
    - offset:   starting row (by id order)
    - overwrite: if 1, replace Candidate.skills with derived tags + old content; else only append missing tags
    """
    batch     = max(1, min(5000, int(request.args.get("batch", "500") or 500)))
    offset    = max(0, int(request.args.get("offset", "0") or 0))
    overwrite = request.args.get("overwrite", "0") == "1"

    updated, examined = 0, 0
    with Session(engine) as s:
        # load term list once
        groups = _get_subject_term_set(s)
        term_list = [t for t in groups.get("__all__", [])]
        # page over candidates by id
        cands = s.scalars(
            select(Candidate).order_by(Candidate.id.asc()).offset(offset).limit(batch)
        ).all()

        for cand in cands:
            examined += 1
            text_lc = _collect_text_for_candidate(s, cand)
            if not text_lc:
                continue
            # role & tags
            role = _normalise_role(text_lc)
            tags = _derive_subject_tags(text_lc, term_list)

            # persist: keep it simple — write tags into Candidate.skills for visibility in UI
            current = (cand.skills or "").strip()
            new_tag_str = ", ".join(sorted(tags, key=str.lower)) if tags else ""
            if overwrite:
                merged = new_tag_str
                if current:
                    # keep any freeform skills too
                    merged = (new_tag_str + (" | " if new_tag_str and current else "") + current).strip()
                cand.skills = merged or current
            else:
                # append tags that aren't already present (case-insensitive)
                existing_tokens = {w.strip().lower() for w in re.split(r"[,/|;]", current) if w.strip()}
                to_add = [t for t in tags if t.lower() not in existing_tokens]
                if to_add:
                    cand.skills = (current + (" | " if current else "") + ", ".join(to_add)).strip()

            # (optional) store normalised role into candidate.skills prefix for now
            if role:
                # add role token if missing
                if role.lower() not in (cand.skills or "").lower():
                    cand.skills = (f"{role} | " + (cand.skills or "")).strip(" |")

            updated += 1

        s.commit()

    flash(f"Re-tagged {updated} / examined {examined}. Offset={offset}, batch={batch}.", "success")
    return redirect(url_for("taxonomy"))
    

@login_required
@app.route("/action/taxonomy/retag_one/<int:cand_id>", methods=["POST"])
def taxonomy_retag_one(cand_id):
    """Re-tag just one candidate (handy button next to a row)."""
    with Session(engine) as s:
        cand = s.scalar(select(Candidate).where(Candidate.id == cand_id))
        if not cand:
            flash("Candidate not found", "warning")
            return redirect(request.referrer or url_for("resource_pool"))

        groups = _get_subject_term_set(s)
        term_list = [t for t in groups.get("__all__", [])]
        text_lc = _collect_text_for_candidate(s, cand)
        role = _normalise_role(text_lc)
        tags = _derive_subject_tags(text_lc, term_list)

        current = (cand.skills or "").strip()
        existing_tokens = {w.strip().lower() for w in re.split(r"[,/|;]", current) if w.strip()}
        to_add = [t for t in tags if t.lower() not in existing_tokens]
        if role and role.lower() not in existing_tokens and role.lower() not in (cand.skills or "").lower():
            cand.skills = (f"{role} | " + current).strip(" |")
            current = cand.skills
        if to_add:
            cand.skills = (current + (" | " if current else "") + ", ".join(to_add)).strip()

        s.commit()

    flash(f"Re-tagged candidate #{cand_id}.", "success")
    return redirect(request.referrer or url_for("resource_pool"))

# --- Candidate ⇄ Tag: add ---
@app.post("/candidate/<int:cand_id>/tag/add")
def candidate_tag_add(cand_id: int):
    tag_id = int((request.form.get("tag_id") or "0") or 0)
    if not tag_id:
        flash("Pick a tag to add.", "warning")
        return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        tag  = s.get(TaxonomyTag, tag_id)
        if not cand or not tag:
            flash("Invalid candidate or tag.", "danger")
            return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

        # prevent duplicates
        exists = s.scalar(
            select(CandidateTag).where(
                CandidateTag.candidate_id == cand_id,
                CandidateTag.tag_id == tag_id,
            )
        )
        if not exists:
            s.add(CandidateTag(candidate_id=cand_id, tag_id=tag_id))

            # Optional: reflect in Candidate.skills (append if not already there, case-insensitive)
            skills = (cand.skills or "").strip()
            tokens = [t.strip() for t in re.split(r"[,\|/;]", skills) if t.strip()]
            lower = {t.lower() for t in tokens}
            if tag.tag.lower() not in lower:
                tokens.append(tag.tag)
                cand.skills = ", ".join(tokens)

            s.commit()

    return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

# --- Candidate ⇄ Tag: remove ---
@app.post("/candidate/<int:cand_id>/tag/remove")
def candidate_tag_remove(cand_id: int):
    tag_id = int((request.form.get("tag_id") or "0") or 0)

    with Session(engine) as s:
        rel = s.scalar(
            select(CandidateTag).where(
                CandidateTag.candidate_id == cand_id,
                CandidateTag.tag_id == tag_id,
            )
        )
        if rel:
            s.delete(rel)

            # Optional: also remove from Candidate.skills
            cand = s.get(Candidate, cand_id)
            tag  = s.get(TaxonomyTag, tag_id)
            if cand and tag and (cand.skills or "").strip():
                parts = [p.strip() for p in re.split(r"[,\|/;]", cand.skills) if p.strip()]
                cand.skills = ", ".join([p for p in parts if p.lower() != tag.tag.lower()])

            s.commit()

    return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

# ---------- Candidate-level card actions ----------

@app.post("/action/onboarding_email/candidate/<int:cand_id>")
def action_onboarding_email_candidate(cand_id):
    """Send a general onboarding email and stamp candidate.onboarded_at."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)

        # If there is a latest application, mention the job in the email.
        appn = _latest_application_for_candidate(s, cand_id)
        job = s.get(Job, appn.job_id) if appn else None

        link = f"{APP_BASE_URL}/candidate/{cand.id}"
        subj = f"Onboarding — {job.title}" if job else "Onboarding"
        html = f"""
        <h3>Welcome to onboarding</h3>
        <p>Hi {cand.name}, we'll guide you through vetting next.</p>
        <p>You can check status here: <a href="{link}">{link}</a></p>
        """
        send_email(cand.email, subj, html)

        cand.onboarded_at = datetime.datetime.utcnow()
        s.commit()

    flash("Onboarding email sent.", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))


@app.post("/action/esign/candidate/<int:cand_id>")
def action_esign_candidate(cand_id):
    """Kick off an e-sign flow if there is an application; otherwise just mark status."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)

        appn = _latest_application_for_candidate(s, cand_id)
        if appn:
            # Reuse your application flow
            # (call the function directly to avoid round-trips)
            try:
                job = s.get(Job, appn.job_id)
                subject = f"Contract — {job.title if job else 'Agreement'}"
                message = "Please review and sign your contract."

                if ESIGN_PROVIDER == "dropbox_sign":
                    req_id = send_esign_dropbox_sign(appn.id, cand.name, cand.email, subject, message)
                elif ESIGN_PROVIDER == "docusign":
                    req_id = send_esign_docusign(appn.id, cand.name, cand.email, subject, message)
                else:
                    raise RuntimeError("Unsupported ESIGN_PROVIDER")

                es = s.scalar(select(ESigRequest).where(ESigRequest.application_id == appn.id))
                if not es:
                    es = ESigRequest(application_id=appn.id, provider=ESIGN_PROVIDER,
                                     request_id=req_id, status="Sent", sent_at=datetime.datetime.utcnow())
                    s.add(es)
                else:
                    es.provider = ESIGN_PROVIDER
                    es.request_id = req_id
                    es.status = "Sent"
                    es.sent_at = datetime.datetime.utcnow()
                cand.esign_status = "Sent"
                s.commit()
            except Exception as e:
                flash(f"E-sign send failed: {e}", "danger")
                return redirect(url_for("candidate_profile", cand_id=cand_id))
        else:
            # No application — just mark status so the card shows activity.
            cand.esign_status = "Sent"
            s.commit()

    flash("E-signature initiated.", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

@app.post("/candidate/<int:cand_id>/summarise", endpoint="cand_summarise")
def candidate_summarise(cand_id):
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)

        # Prefer latest CV that’s actually a CV
        doc = s.scalar(
            select(Document)
            .where(Document.candidate_id == cand_id, Document.doc_type == "cv")
            .order_by(Document.uploaded_at.desc())
        )

        cv_text = extract_cv_text(doc) if doc else ""
        source_text = cv_text or (cand.skills or "")
        if not source_text:
            flash("No CV or skills text available to summarise.", "warning")
            return redirect(url_for("candidate_profile", cand_id=cand_id))

        # Write to an actual persisted column
        cand.ai_summary = ai_summarise(source_text) or ""
        s.commit()

    flash("AI summary regenerated.", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

@app.post("/candidate/<int:cand_id>/skills/update")
def candidate_skills_update(cand_id):
    new_skills = (request.form.get("skills") or "").strip()
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)
        cand.skills = new_skills
        s.commit()
    flash("Skills updated.", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

@app.post("/action/esign_status/candidate/<int:cand_id>")
def action_esign_status_candidate(cand_id):
    """Refresh candidate-level e-sign status (via latest application if present)."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)

        appn = _latest_application_for_candidate(s, cand_id)
        if not appn:
            flash("No application to poll. Status unchanged.", "info")
            return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

        es = s.scalar(select(ESigRequest).where(ESigRequest.application_id == appn.id))
        if not es:
            flash("No e-sign request found for the latest application.", "warning")
            return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

        try:
            if es.provider == "dropbox_sign":
                status = poll_esign_dropbox_sign(es.request_id)
            elif es.provider == "docusign":
                status = poll_esign_docusign(es.request_id)
            else:
                status = "Unknown"
        except Exception as e:
            flash(f"E-sign status failed: {e}", "danger")
            return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

        es.status = status
        if status in {"Signed", "Completed"}:
            es.signed_at = datetime.datetime.utcnow()
        cand.esign_status = status
        s.commit()

    flash(f"E-sign status: {status}", "info")
    return redirect(url_for("candidate_profile", cand_id=cand_id, job_id=request.args.get("job_id")))

@app.post("/action/trustid/candidate/<int:cand_id>")
def action_trustid_candidate(cand_id):
    """Start TrustID checks at candidate level; stamp the dates locally.
       If there is an application, also create a real TrustIDCheck row."""
    do_rtw = bool(request.form.get("rtw"))
    do_idv = bool(request.form.get("idv", True))
    do_dbs = bool(request.form.get("dbs"))

    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)

        # Stamp local candidate dates immediately so the card reflects activity
        now = datetime.datetime.utcnow()
        if do_rtw:
            cand.trustid_rtw_date = now
        if do_idv:
            cand.trustid_idv_date = now
        if do_dbs:
            cand.trustid_dbs_date = now

        # If there is a latest application, also kick off a real TrustID app + row
        appn = _latest_application_for_candidate(s, cand_id)

        if appn:
            payload = {
                "external_ref": f"CAND-{cand.id}-APP-{appn.id}",
                "candidate": {"name": cand.name, "email": cand.email},
                "checks": {"rtw": do_rtw, "idv": do_idv, "dbs": do_dbs},
                "webhooks": {"result": f"{APP_BASE_URL}/webhook/trustid"},
            }
            try:
                resp = requests.post(
                    f"{TRUSTID_BASE_URL}/apps",
                    headers=trustid_headers(),
                    data=json.dumps(payload),
                    timeout=15,
                )
                if resp.status_code >= 300:
                    raise RuntimeError(f"TrustID error: {resp.status_code} {resp.text}")
                data = resp.json()
                trust_app_id = data.get("id", "") or data.get("application_id", "")
                s.add(TrustIDCheck(
                    application_id=appn.id,
                    rtw=do_rtw, idv=do_idv, dbs=do_dbs,
                    trustid_application_id=trust_app_id,
                    status="InProgress",
                ))
            except Exception as e:
                current_app.logger.warning(f"TrustID create failed for candidate #{cand.id}: {e}")
                flash(f"TrustID create failed: {e}", "danger")
                # even if TrustID API failed, keep local timestamps
                s.commit()
                return redirect(url_for("candidate_profile", cand_id=cand_id))

        s.commit()

    flash("TrustID checks started and candidate record updated.", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

@app.post("/action/request_updated_cv/candidate/<int:cand_id>")
def action_request_updated_cv_candidate(cand_id):
    """Send a generic 'please upload updated CV' link and timestamp."""
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)
        link = f"{APP_BASE_URL}/candidate/{cand.id}/upload_cv"
        html = f"""
        <h3>Request for updated CV</h3>
        <p>Hi {cand.name}, please upload an updated CV here:</p>
        <p><a href="{link}">{link}</a></p>
        """
        send_email(cand.email, "Please upload an updated CV", html)
        cand.updated_cv_requested_at = datetime.datetime.utcnow()
        s.commit()

    flash("CV update request sent.", "success")
    return redirect(url_for("candidate_profile", cand_id=cand_id))

@login_required
@app.route("/configuration", methods=["GET"])
def configuration():
    alias_form = RoleAliasForm()
    with Session(engine) as s:
        # role aliases
        aliases = s.execute(text("SELECT id, canonical, alias FROM role_aliases ORDER BY canonical, alias")).all()
        # taxonomy data
        cats = s.scalars(select(TaxonomyCategory).order_by(TaxonomyCategory.type.asc(), TaxonomyCategory.name.asc())).all()
        tags_by_cat = {
            c.id: s.scalars(select(TaxonomyTag).where(TaxonomyTag.category_id==c.id).order_by(TaxonomyTag.tag.asc())).all()
            for c in cats
        }
    return render_template("configuration.html",
                           alias_form=alias_form,
                           aliases=aliases,
                           cats=cats,
                           tags_by_cat=tags_by_cat,
                           ROLE_TYPES=ROLE_TYPES)

@app.post("/config/role-alias/add")
def config_role_alias_add():
    form = RoleAliasForm()
    if not form.validate_on_submit():
        flash("Pick a canonical role and enter an alias.", "warning")
        return redirect(url_for("configuration"))
    canonical = form.canonical.data.strip()
    alias = form.alias.data.strip()
    with Session(engine) as s:
        # prevent dup (case-insensitive)
        exists = s.execute(
            text("""SELECT 1 FROM role_aliases 
                    WHERE lower(canonical)=:c AND lower(alias)=:a LIMIT 1"""),
            {"c": canonical.lower(), "a": alias.lower()}
        ).first()
        if exists:
            flash("That alias already exists for the selected canonical role.", "info")
        else:
            s.execute(text("INSERT INTO role_aliases(canonical, alias) VALUES(:c,:a)"),
                      {"c": canonical, "a": alias})
            s.commit()
            flash("Alias added.", "success")
    return redirect(url_for("configuration"))

@app.post("/opportunities/<int:opp_id>/convert")
def opportunity_convert(opp_id):
    with Session(engine) as s:
        opp = s.get(Opportunity, opp_id)
        if not opp:
            flash("Opportunity not found.", "warning")
            return redirect(url_for("opportunities_"))

        e = create_engagement_for_opportunity(s, opp)
        if e is None and (opp.stage or "").strip().lower() != "closed won":
            flash("Opportunity must be Closed Won to create an engagement.", "warning")
            return redirect(url_for("opportunities_"))

        s.commit()
        if e:
            flash(f"Engagement {e.ref} created.", "success")
            return redirect(url_for("engagement_dashboard", eng_id=e.id))
        else:
            # Already existed (idempotent case)
            existing = s.execute(
                text("SELECT id, ref FROM engagements WHERE opportunity_id=:oid"),
                {"oid": opp.id}
            ).first()
            if existing:
                flash(f"Engagement {existing[1]} already exists.", "info")
                return redirect(url_for("engagement_dashboard", eng_id=existing[0]))
            flash("Nothing to do.", "info")
            return redirect(url_for("opportunities_"))

@app.post("/config/role-alias/delete")
def config_role_alias_delete():
    alias_id = int((request.form.get("id") or 0))
    if not alias_id:
        flash("Missing alias id.", "warning")
        return redirect(url_for("configuration"))
    with Session(engine) as s:
        s.execute(text("DELETE FROM role_aliases WHERE id=:i"), {"i": alias_id})
        s.commit()
    flash("Alias removed.", "success")
    return redirect(url_for("configuration"))

@app.get("/sumsub/sdk.js")
def sumsub_sdk_proxy():
    """
    Proxies the Sumsub WebSDK JS from the CDN to avoid ad-block/VPN/CORS issues in dev.
    """
    last_exc = None
    for url in SUMSUB_SDK_URLS:
        try:
            r = requests.get(url, timeout=10)
            if r.ok and r.content and b"SNSWebSdk" in r.content:
                resp = make_response(r.content, 200)
                resp.headers["Content-Type"] = "application/javascript; charset=utf-8"
                resp.headers["Cache-Control"] = "public, max-age=3600"
                return resp
        except Exception as e:
            last_exc = e

    # If all mirrors failed, return a clear error so the page can show it
    msg = f"// Failed to fetch Sumsub SDK. Last error: {last_exc!r}" if last_exc else "// Failed to fetch Sumsub SDK."
    resp = make_response(msg, 502)
    resp.headers["Content-Type"] = "application/javascript; charset=utf-8"
    return resp

@app.get("/job/<int:job_id>/public-link")
def job_public_link(job_id):
    with Session(engine) as s:
        job = s.get(Job, job_id)
        if not job or (job.status or "").lower() != "open":
            abort(404)
        return redirect(url_for("public.public_job_detail", token=job.public_token))

@app.get("/engagement/<int:eng_id>/jobs")
def engagement_jobs(eng_id):
    with Session(engine) as s:
        eng = s.get(Engagement, eng_id)
        if not eng:
            abort(404)
        jobs = s.scalars(
            select(Job)
            .where(Job.engagement_id == eng_id, Job.status == "Open")
            .order_by(Job.created_at.desc())
        ).all()
    return render_template("eng_jobs_list.html", engagement=eng, jobs=jobs)

@login_required
@app.route("/engagement/<int:eng_id>/jobs/create", methods=["GET","POST"])
def engagement_jobs_create(eng_id):
    with Session(engine) as s:
        eng = s.get(Engagement, eng_id)
        if not eng:
            abort(404)

        form = JobForm()
        # lock the engagement selection to this engagement
        form.engagement_id.choices = [(eng.id, f"{eng.name} ({eng.client})")]
        form.engagement_id.data = eng.id

        if form.validate_on_submit():
            j = Job(
                engagement_id=eng.id,
                title=form.title.data,
                description=form.description.data,
                role_type=form.role_type.data or "",
                location=form.location.data or "",
                salary_range=form.salary_range.data or "",
                status="Open",
            )
            s.add(j)
            s.commit()
            flash("Job created for this engagement.", "success")
            return redirect(url_for("engagement_jobs", eng_id=eng.id))

    return render_template("eng_job_create.html", engagement=eng, form=form)

@app.post("/jobs/<int:job_id>/withdraw")
def job_withdraw(job_id):
    """
    Mark job as Withdrawn and email all *applicants* who have NOT entered vetting
    (i.e., no TrustIDCheck row for their application).
    """
    with Session(engine) as s:
        job = s.get(Job, job_id)
        if not job:
            abort(404)

        # If already withdrawn, do nothing
        was_withdrawn = (job.status or "").lower() == "withdrawn"
        job.status = "Withdrawn"

        # Find applications with NO TrustIDCheck rows
        apps = s.execute(
            select(Application, Candidate)
            .join(Candidate, Candidate.id == Application.candidate_id)
            .where(Application.job_id == job.id)
        ).all()

        # Build a set of application_ids that have vetting rows
        app_ids = [a.id for a, _ in apps]
        vetted_app_ids = set()
        if app_ids:
            vetted_app_ids = {
                app_id for (app_id,) in s.execute(
                    select(TrustIDCheck.application_id)
                    .where(TrustIDCheck.application_id.in_(app_ids))
                    .limit(100000)
                ).all()
            }

        # Email only those *without* vetting
        emailed = 0
        for appn, cand in apps:
            if appn.id in vetted_app_ids:
                continue
            # polite note
            html = f"""
              <p>Hi {cand.name},</p>
              <p>Thank you for your interest in <strong>{job.title}</strong>. This position has now been closed.</p>
              <p>We appreciate your time and will keep your details on file for future opportunities.</p>
              <p>Best regards,<br>Talent Team</p>
            """
            try:
                if cand.email:
                    send_email(cand.email, f"Update on your application — {job.title}", html)
                    emailed += 1
            except Exception as e:
                current_app.logger.warning("Withdraw email failed for %s: %s", cand.email, e)

        s.commit()

    msg = "Job withdrawn."
    if emailed:
        msg += f" Notified {emailed} applicant(s) with no vetting started."
    flash(msg, "success")
    # bounce back to the engagement's jobs list if we can infer it; else jobs page
    try:
        return redirect(url_for("engagement_jobs", eng_id=job.engagement_id))
    except Exception:
        return redirect(url_for("jobs"))

# --- SINGLE canonical shortlist endpoint (ensure there is only ONE of these) ---
@app.post("/shortlist")
def candidate_shortlist_action():
    next_url = request.form.get("next") or request.referrer
    app_id = (request.form.get("app_id") or "").strip()
    job_id = (request.form.get("job_id") or "").strip()
    cand_id = (request.form.get("candidate_id") or "").strip()

    with Session(engine) as s:
        job = candidate = engagement = None

        if app_id.isdigit():
            appn = s.get(Application, int(app_id))
            if not appn:
                flash("Application not found.", "warning")
                return redirect(next_url or url_for("applications"))
            job = s.get(Job, appn.job_id)
            candidate = s.get(Candidate, appn.candidate_id)
        else:
            if not (job_id.isdigit() and cand_id.isdigit()):
                flash("Missing job/candidate parameters.", "warning")
                return redirect(next_url or url_for("index"))
            job = s.get(Job, int(job_id))
            candidate = s.get(Candidate, int(cand_id))

        if not job or not candidate:
            flash("Job or candidate not found.", "warning")
            return redirect(next_url or url_for("index"))

        engagement = s.get(Engagement, job.engagement_id) if job.engagement_id else None
        eng_active = ((engagement.status if engagement else "") or "").strip().lower() in {"active","in-flight","in progress"}
        if not eng_active:
            flash("Engagement is not active; cannot shortlist.", "warning")
            return redirect(next_url or (url_for("engagement_dashboard", eng_id=engagement.id) if engagement else url_for("index")))

        if not _job_is_open(job):
            flash(f"Job is {job.status or 'not open'}; cannot shortlist.", "warning")
            return redirect(next_url or url_for("engagement_dashboard", eng_id=engagement.id))

        exists_row = s.execute(
            select(Shortlist.id).where(Shortlist.job_id==job.id, Shortlist.candidate_id==candidate.id).limit(1)
        ).first()

        if exists_row:
            flash("Already shortlisted for this job.", "info")
        else:
            s.add(Shortlist(job_id=job.id, candidate_id=candidate.id))
            s.commit()
            flash(f"Shortlisted {candidate.name} for “{job.title}”.", "success")

    if next_url:  # return to the Applications list you came from
        return redirect(next_url)
    if engagement:
        return redirect(url_for("applications_for_engagement", eng_id=engagement.id))
    return redirect(url_for("applications"))

@app.route("/candidate/login", methods=["GET", "POST"])
def candidate_login():
    """
    Collect email → email a time-limited magic link.
    If candidate doesn't exist yet, we create on first successful login.
    """
    form = CandidateLoginForm()
    nxt = request.args.get("next") or url_for("index")
    if form.validate_on_submit():
        email = form.email.data.strip().lower()
        token = _signer().dumps({"email": email, "next": nxt})
        link = f"{APP_BASE_URL}/candidate/magic?token={token}"
        try:
            send_email(
                to_email=email,
                subject="Your secure sign-in link",
                html_body=f"""
                    <p>Click to sign in:</p>
                    <p><a href="{link}">{link}</a></p>
                    <p>This link will expire in 20 minutes.</p>
                """
            )
        except Exception as e:
            flash(f"Failed to send login link: {e}", "danger")
            return render_template("candidate_login.html", form=form, next=nxt)
        flash("Check your email for a sign-in link.", "info")
        return redirect(url_for("index"))
    return render_template("candidate_login.html", form=form, next=nxt)

@app.route("/candidate/magic")
def candidate_magic():
    token = request.args.get("token", "")
    try:
        data = _signer().loads(token, max_age=20*60)  # 20 minutes
    except SignatureExpired:
        flash("This link has expired. Please request a new one.", "warning")
        return redirect(url_for("candidate_login"))
    except BadSignature:
        flash("Invalid sign-in link.", "danger")
        return redirect(url_for("candidate_login"))

    email = (data.get("email") or "").strip().lower()
    nxt = data.get("next") or url_for("index")

    if not email:
        flash("Missing email in link.", "danger")
        return redirect(url_for("candidate_login"))

    with Session(engine) as s:
        cand = s.scalar(select(Candidate).where(func.lower(Candidate.email) == email))
        if not cand:
            # create a candidate shell record on first sign-in
            cand = Candidate(name=email.split("@")[0].title(), email=email)
            s.add(cand)
            s.flush()
            s.commit()

        # set candidate session
        session["candidate_id"] = cand.id
        session["candidate_email"] = cand.email

    flash("Signed in as candidate.", "success")
    return redirect(nxt)

@app.route("/candidate/logout")
def candidate_logout():
    session.pop("candidate_id", None)
    session.pop("candidate_email", None)
    flash("Signed out.", "info")
    return redirect(url_for("index"))

@app.route("/candidate/<int:cand_id>/upload_cv", methods=["GET", "POST"])
def candidate_upload_cv(cand_id):
    form = UploadCVForm()
    with Session(engine) as s:
        cand = s.get(Candidate, cand_id)
        if not cand:
            abort(404)

        if form.validate_on_submit():
            file = request.files.get("cv")
            if not file:
                flash("Please choose a CV file.", "warning")
                return render_template("candidate_upload_cv.html", form=form, cand=cand)

            try:
                fname, path, original = save_upload(file, subdir="cvs")
            except Exception as e:
                flash(f"Upload failed: {e}", "danger")
                return render_template("candidate_upload_cv.html", form=form, cand=cand)

            # Save document
            doc = Document(
                candidate_id=cand.id,
                doc_type="cv",
                filename=fname,
                original_name=original,
            )
            s.add(doc)
            s.flush()

            # Rebuild AI summary + tags off the new CV
            try:
                _rebuild_ai_summary_and_tags(s, cand, doc=doc, job=None, appn=None)
            except Exception as e:
                current_app.logger.warning(f"Post-upload AI/retag failed for cand #{cand.id}: {e}")

            s.commit()
            flash("CV uploaded successfully.", "success")
            return redirect(url_for("candidate_profile", cand_id=cand.id))

    return render_template("candidate_upload_cv.html", form=form, cand=cand)

# OPTIONAL: simple worker signup (disable in prod or restrict to admins)
@login_required
@app.route("/signup", methods=["GET", "POST"])
def signup():
    form = WorkerSignupForm()
    if form.validate_on_submit():
        with Session(engine) as s:
            exists = s.execute(text("SELECT 1 FROM users WHERE lower(email)=:e"), {"e": form.email.data.strip().lower()}).first()
            if exists:
                flash("That email is already registered.", "warning")
                return render_template("auth_signup.html", form=form)
            s.execute(text("""
                INSERT INTO users(name, email, pw_hash)
                VALUES (:n, :e, :p)
            """), {
                "n": form.name.data.strip(),
                "e": form.email.data.strip().lower(),
                "p": generate_password_hash(form.password.data)
            })
            s.commit()
        flash("Account created. Please sign in.", "success")
        return redirect(url_for("login"))
    return render_template("auth_signup.html", form=form)

@app.route("/uploads/cvs/<path:fname>")
@login_required
def download_cv(fname):
    """Download CV file - handles various filename formats and storage locations."""
    # Clean up the filename - remove any duplicate path prefixes
    clean_fname = fname
    if clean_fname.startswith("uploads/cvs/"):
        clean_fname = clean_fname.replace("uploads/cvs/", "")
    
    # First try static/uploads/cvs (where public.py saves files)
    static_cv_path = os.path.join(os.path.dirname(__file__), "static", "uploads", "cvs")
    static_file = os.path.join(static_cv_path, clean_fname)
    if os.path.exists(static_file):
        return send_from_directory(static_cv_path, clean_fname, as_attachment=True)
    
    # Fall back to uploads/cvs
    uploads_cv_path = os.path.join(app.config["UPLOAD_FOLDER"], "cvs")
    return send_from_directory(uploads_cv_path, clean_fname, as_attachment=True)

@app.route("/document/<int:doc_id>/download")
@login_required
def download_document(doc_id):
    """Download a document by its ID."""
    with Session(engine) as s:
        doc = s.get(Document, doc_id)
        if not doc:
            abort(404)
        
        # Use the same path resolution logic
        path = _doc_file_path(doc)
        if os.path.exists(path):
            directory = os.path.dirname(path)
            filename = os.path.basename(path)
            download_name = doc.original_name or filename
            return send_from_directory(directory, filename, as_attachment=True, download_name=download_name)
        
        abort(404)

if __name__ == "__main__":
    with app.app_context():
        Base.metadata.create_all(engine)
        ensure_schema()
    app.run(debug=True)
